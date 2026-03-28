from __future__ import annotations

import base64
import zipfile

from app.service_shared import *


class StudentAnalyticsServiceMixin:
    def _list_filtered_student_practice_questions(
        self,
        filters: Dict[str, str],
        actor: Actor,
        *,
        include_chapter_progress: bool = True,
    ) -> List[Dict[str, str]]:
        only_personal_bank = str(filters.get("onlyPersonalBank", "")).strip().lower() == "true"
        student_record_map: Dict[str, Dict[str, object]] = {}
        if only_personal_bank:
            student_record_map = {
                str(record.get("questionId", "")).strip(): record
                for record in self.repository.list_student_question_record_rows(actor.user_id)
                if str(record.get("questionId", "")).strip()
            }
        db_filters = self._pick_filters(
            filters,
            ("knowledgeId", "subjectId", "chapter", "type", "difficulty", "chapterCode", "pointCode"),
        )
        db_filters["chapterCode"] = str(filters.get("chapterCode") or "").strip()
        db_filters["pointCode"] = str(filters.get("pointCode") or "").strip()
        questions = self.repository.list_visible_published_questions(db_filters, ROLE_SUPER_ADMIN, actor.user_id)
        questions = self._filter_questions_for_student_scope(questions, actor.user_id, filters)
        if not only_personal_bank and str(filters.get("module", "")).strip().lower() != "free":
            questions = self._filter_unlocked_chapter_questions(questions, actor.user_id)
        questions = self._filter_student_question_views(
            questions,
            actor.user_id,
            filters,
            student_record_map=student_record_map,
        )
        return [
            self._enrich_question_for_student(
                question,
                actor.user_id,
                include_chapter_progress=include_chapter_progress,
                student_record=student_record_map.get(str(question.get("id", "")).strip()),
            )
            for question in questions
        ]

    def _resolve_wrong_book_path_node_descendants(
        self,
        knowledge_path_node_id: str,
        subject_code: str = "",
    ) -> set[str]:
        normalized_path_node_id = str(knowledge_path_node_id or "").strip()
        if not normalized_path_node_id:
            return set()

        repository_filters: Dict[str, str] = {}
        normalized_subject_code = str(subject_code or "").strip()
        if normalized_subject_code:
            repository_filters["subjectCode"] = normalized_subject_code
            repository_filters["subject_code"] = normalized_subject_code

        knowledge_rows = self.repository.list_knowledge("", repository_filters)
        if not any(str(item.get("id", "")).strip() == normalized_path_node_id for item in knowledge_rows):
            return set()

        children_by_parent: Dict[str, List[str]] = {}
        for item in knowledge_rows:
            knowledge_id = str(item.get("id", "")).strip()
            parent_id = str(item.get("parentId", "") or "").strip()
            if not knowledge_id:
                continue
            children_by_parent.setdefault(parent_id, []).append(knowledge_id)

        descendant_ids: set[str] = set()
        queue: List[str] = [normalized_path_node_id]
        while queue:
            current_id = queue.pop(0)
            if current_id in descendant_ids:
                continue
            descendant_ids.add(current_id)
            queue.extend(children_by_parent.get(current_id, []))
        return descendant_ids

    def list_student_practice_questions(
        self,
        filters: Dict[str, str],
        page: int,
        size: int,
        actor: Actor,
    ) -> Tuple[List[Dict[str, str]], int]:
        questions = self._list_filtered_student_practice_questions(filters, actor)
        return self._paginate_questions(questions, page, size)

    def generate_adaptive_practice(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        request = AdaptivePracticeRequest.model_validate(dict(payload or {}))
        requested_count = int(request.count)
        preferred_knowledge_id = str(request.knowledge_id or "").strip()
        profile = self._get_student_profile(actor.user_id)
        scope_filters = {
            "examCategoryCode": str(profile.get("examCategoryCode", "")).strip(),
            "jointExamGroupCode": str(profile.get("jointExamGroupCode", "")).strip(),
            "subjectCode": "",
        }

        question_pool = self.repository.list_visible_published_questions(
            {"subjectId": "", "chapter": "", "type": "", "difficulty": ""},
            ROLE_SUPER_ADMIN,
            actor.user_id,
        )
        question_pool = self._filter_questions_for_student_scope(question_pool, actor.user_id, scope_filters)
        if not question_pool:
            return {"questionIds": []}

        history_rows = self._collect_adaptive_history_rows(actor.user_id, scope_filters)
        mastery_snapshot = self._build_adaptive_mastery_snapshot(question_pool, history_rows)
        target_knowledge_ids = self._select_adaptive_target_knowledge_ids(mastery_snapshot, top_n=3, mastery_threshold=0.4)

        if preferred_knowledge_id:
            if not any(str(question.get("knowledgeId", "")).strip() == preferred_knowledge_id for question in question_pool):
                raise validation_failed("knowledgeId 不存在或当前学生不可见。")
            target_knowledge_ids = [preferred_knowledge_id] + [
                item
                for item in target_knowledge_ids
                if str(item).strip() and str(item).strip() != preferred_knowledge_id
            ]

        target_knowledge_quota = self._allocate_adaptive_question_counts(
            target_knowledge_ids,
            mastery_snapshot,
            requested_count,
        )
        if preferred_knowledge_id:
            target_knowledge_quota[preferred_knowledge_id] = max(1, int(target_knowledge_quota.get(preferred_knowledge_id, 0)))
            while sum(int(value) for value in target_knowledge_quota.values()) > requested_count:
                reduced = False
                for knowledge_id in list(target_knowledge_quota.keys()):
                    if knowledge_id == preferred_knowledge_id:
                        continue
                    if int(target_knowledge_quota.get(knowledge_id, 0)) <= 0:
                        continue
                    target_knowledge_quota[knowledge_id] = int(target_knowledge_quota[knowledge_id]) - 1
                    reduced = True
                    if sum(int(value) for value in target_knowledge_quota.values()) <= requested_count:
                        break
                if not reduced:
                    break

        difficulty_cycle = self._resolve_adaptive_difficulty_cycle(history_rows)
        question_ids = self._select_adaptive_question_ids(
            question_pool,
            target_knowledge_ids,
            difficulty_cycle,
            requested_count,
            target_knowledge_quota,
        )
        return {"questionIds": question_ids}

    def get_student_challenge_point_summary(self, subject_code: str, actor: Actor) -> Dict[str, object]:
        normalized_subject_code = str(subject_code or "").strip()
        if not normalized_subject_code:
            raise validation_failed("subjectCode 不能为空。")
        summary = self._build_challenge_point_summary(actor.user_id, normalized_subject_code, leaderboard_limit=10)
        summary["studentUserId"] = actor.user_id
        return summary

    def list_student_practice_chapters(self, subject_id: str, actor: Actor) -> List[Dict[str, object]]:
        normalized_subject_id = subject_id.strip()
        if not normalized_subject_id:
            raise validation_failed("subjectId 不能为空。")
        subject_questions = self.repository.list_visible_published_questions(
            {"subjectId": normalized_subject_id, "chapter": "", "type": "", "difficulty": ""},
            ROLE_SUPER_ADMIN,
            actor.user_id,
        )
        subject_questions = self._filter_questions_for_student_scope(subject_questions, actor.user_id, {})
        chapter_order = self._chapter_order_for_subject(normalized_subject_id, subject_questions)
        rows: List[Dict[str, object]] = []
        for chapter in chapter_order:
            progress = self._build_chapter_progress(normalized_subject_id, chapter, actor.user_id)
            rows.append(
                {
                    "subjectId": normalized_subject_id,
                    "chapter": chapter,
                    "answered": int(progress.get("answered", 0)),
                    "total": int(progress.get("total", 0)),
                    "accuracy": float(progress.get("accuracy", 0.0)),
                    "isUnlocked": bool(progress.get("isUnlocked", False)),
                }
            )

        current_chapter = ""
        for row in rows:
            if row["isUnlocked"] and row["answered"] < row["total"]:
                current_chapter = str(row["chapter"])
                break
        if not current_chapter:
            unlocked_rows = [row for row in rows if row["isUnlocked"]]
            if unlocked_rows:
                current_chapter = str(unlocked_rows[-1]["chapter"])

        for row in rows:
            row["isCurrent"] = str(row["chapter"]) == current_chapter
            if row["isCurrent"]:
                row["statusLabel"] = "正在闯关"
            elif row["isUnlocked"]:
                row["statusLabel"] = "已解锁"
            else:
                row["statusLabel"] = "未解锁"
        return rows

    def list_personal_bank_questions(
        self,
        filters: Dict[str, str],
        page: int,
        size: int,
        actor: Actor,
    ) -> Tuple[List[Dict[str, str]], int]:
        scoped_filters = dict(filters)
        scoped_filters["onlyPersonalBank"] = "true"
        return self.list_student_practice_questions(scoped_filters, page, size, actor)

    def list_all_personal_bank_questions(self, filters: Dict[str, str], actor: Actor) -> List[Dict[str, str]]:
        scoped_filters = dict(filters)
        scoped_filters["onlyPersonalBank"] = "true"
        return self._list_filtered_student_practice_questions(
            scoped_filters,
            actor,
            include_chapter_progress=False,
        )

    def get_personal_bank_summary(self, filters: Dict[str, str], actor: Actor) -> Dict[str, object]:
        items = self.list_all_personal_bank_questions(filters, actor)
        synced_plans = self._sync_student_review_plans(actor)
        filtered_question_ids = {
            str(item.get("id", "")).strip()
            for item in items
            if str(item.get("id", "")).strip()
        }
        return self._summarize_personal_bank_items(
            items,
            review_plan_rows=self._project_review_plan_rows(synced_plans, filtered_question_ids),
        )

    def export_personal_bank(self, filters: Dict[str, str], actor: Actor) -> Dict[str, str]:
        export_format = self._normalize_export_format(filters.get("format", "csv"), PERSONAL_BANK_EXPORT_FORMATS, "沉淀题库导出", "csv")
        items = self.list_all_personal_bank_questions(filters, actor)
        summary = self.get_personal_bank_summary(filters, actor)
        if export_format == "pdf":
            lines = [
                "沉淀题库导出",
                f"studentUserId: {actor.user_id}",
                f"totalCount: {summary['totalCount']}",
                f"accuracy: {summary['accuracy']}",
                f"recentCollectedAt: {summary['recentCollectedAt'] or '-'}",
            ]
            for question in items:
                lines.extend(
                    [
                        "",
                        f"{self._question_subject_id(question)} / {self._question_chapter(question)} / {question['id']}",
                        question["stem"],
                        f"answer: {question['answer']}",
                    ]
                )
            return {"format": "pdf", "content": "\n".join(lines)}
        lines = ["questionId,subjectId,chapter,type,difficulty,stem,answer"]
        for question in items:
            lines.append(
                ",".join(
                    [
                        question["id"],
                        self._question_subject_id(question),
                        self._question_chapter(question),
                        question["type"],
                        self._question_difficulty(question),
                        self._csv_cell(question["stem"]),
                        self._csv_cell(question["answer"]),
                    ]
                )
            )
        return {"format": "csv", "content": "\n".join(lines)}

    def list_student_personal_bank_review_plans(self, filters: Dict[str, str], actor: Actor) -> List[Dict[str, object]]:
        synced_plans = self._sync_student_review_plans(actor)
        filtered_question_ids = {
            str(item.get("id", "")).strip()
            for item in self.list_all_personal_bank_questions(filters, actor)
            if str(item.get("id", "")).strip()
        }
        return self._project_review_plan_rows(
            synced_plans,
            filtered_question_ids,
        )

    def _resolve_personal_bank_filtered_question_ids(self, filters: Dict[str, str], actor: Actor) -> set[str]:
        return {
            str(item.get("id", "")).strip()
            for item in self.list_all_personal_bank_questions(filters, actor)
            if str(item.get("id", "")).strip()
        }

    def get_student_personal_bank_review_plan(self, plan_id: str, filters: Dict[str, str], actor: Actor) -> Dict[str, object]:
        self._sync_student_review_plans(actor)
        bundle = self._load_student_review_plan_bundle(actor.user_id, plan_id)
        if not bundle:
            raise not_found("复习计划不存在。")
        filtered_question_ids = self._resolve_personal_bank_filtered_question_ids(filters, actor)
        question_ids = [
            str(item.get("questionId", "")).strip()
            for item in bundle.get("items", [])
            if (
                isinstance(item, dict)
                and str(item.get("questionId", "")).strip()
                and str(item.get("questionId", "")).strip() in filtered_question_ids
            )
        ]
        questions_by_id = {
            str(question.get("id", "")).strip(): self._public_question(self._enrich_question_for_student(question, actor.user_id, include_chapter_progress=False))
            for question in self.repository.list_questions_by_ids(question_ids)
        }
        projected = self._project_review_plan_rows([bundle], filtered_question_ids)[0]
        projected["items"] = [
            {
                "itemId": str(item.get("id", "")).strip(),
                "questionId": str(item.get("questionId", "")).strip(),
                "status": str(item.get("status", "")).strip() or "PENDING",
                "sort": self._safe_int(item.get("sort", 0), 0),
                "completedAt": str(item.get("completedAt", "")).strip(),
                "question": questions_by_id.get(str(item.get("questionId", "")).strip()),
            }
            for item in bundle.get("items", [])
            if isinstance(item, dict) and str(item.get("questionId", "")).strip() in filtered_question_ids
        ]
        return projected

    def start_student_personal_bank_review_plan(self, plan_id: str, filters: Dict[str, str], actor: Actor) -> Dict[str, object]:
        self._sync_student_review_plans(actor)
        bundle = self._load_student_review_plan_bundle(actor.user_id, plan_id)
        if not bundle:
            raise not_found("复习计划不存在。")
        now_iso = self._now_iso()
        ext_json = self._load_json_object(bundle.get("extJson", "{}"))
        if not isinstance(ext_json, dict):
            ext_json = {}
        ext_json["questionCount"] = len(bundle.get("items", []))
        ext_json["completedCount"] = sum(
            1
            for item in bundle.get("items", [])
            if isinstance(item, dict) and str(item.get("status", "")).strip() == "COMPLETED"
        )
        self.repository.replace_student_review_plan(
            {
                "id": str(bundle.get("id", "")).strip(),
                "studentUserId": actor.user_id,
                "planType": str(bundle.get("planType", "")).strip(),
                "planName": str(bundle.get("planName", "")).strip(),
                "status": "IN_PROGRESS" if bundle.get("items") else "PENDING",
                "generatedAt": str(bundle.get("generatedAt", "")).strip() or now_iso,
                "startedAt": str(bundle.get("startedAt", "")).strip() or now_iso,
                "completedAt": "",
                "lastExecutedAt": now_iso,
                "extJson": self._dump_json(ext_json),
                "createTime": str(bundle.get("createTime", "")).strip() or now_iso,
                "updateTime": now_iso,
            },
            [
                {
                    "id": str(item.get("id", "")).strip(),
                    "planId": str(bundle.get("id", "")).strip(),
                    "studentUserId": actor.user_id,
                    "questionId": str(item.get("questionId", "")).strip(),
                    "status": str(item.get("status", "")).strip() or "PENDING",
                    "sort": self._safe_int(item.get("sort", 0), 0),
                    "completedAt": str(item.get("completedAt", "")).strip(),
                    "extJson": str(item.get("extJson", "{}")),
                    "createTime": str(item.get("createTime", "")).strip() or now_iso,
                    "updateTime": now_iso,
                }
                for item in bundle.get("items", [])
                if isinstance(item, dict)
            ],
        )
        return self.get_student_personal_bank_review_plan(plan_id, filters, actor)

    def complete_student_personal_bank_review_plan_question(
        self,
        plan_id: str,
        question_id: str,
        filters: Dict[str, str],
        actor: Actor,
    ) -> Dict[str, object]:
        self._sync_student_review_plans(actor)
        bundle = self._load_student_review_plan_bundle(actor.user_id, plan_id)
        if not bundle:
            raise not_found("复习计划不存在。")
        normalized_question_id = str(question_id or "").strip()
        if not normalized_question_id:
            raise validation_failed("questionId 不能为空。")
        now_iso = self._now_iso()
        matched = False
        next_items: List[Dict[str, object]] = []
        for item in bundle.get("items", []):
            if not isinstance(item, dict):
                continue
            item_question_id = str(item.get("questionId", "")).strip()
            item_status = str(item.get("status", "")).strip() or "PENDING"
            completed_at = str(item.get("completedAt", "")).strip()
            if item_question_id == normalized_question_id:
                item_status = "COMPLETED"
                completed_at = completed_at or now_iso
                matched = True
            next_items.append(
                {
                    "id": str(item.get("id", "")).strip(),
                    "planId": str(bundle.get("id", "")).strip(),
                    "studentUserId": actor.user_id,
                    "questionId": item_question_id,
                    "status": item_status,
                    "sort": self._safe_int(item.get("sort", 0), 0),
                    "completedAt": completed_at,
                    "extJson": str(item.get("extJson", "{}")),
                    "createTime": str(item.get("createTime", "")).strip() or now_iso,
                    "updateTime": now_iso,
                }
            )
        if not matched:
            raise not_found("复习计划题目不存在。")
        completed_count = sum(1 for item in next_items if str(item.get("status", "")).strip() == "COMPLETED")
        total_count = len(next_items)
        plan_status = "COMPLETED" if total_count and completed_count >= total_count else "IN_PROGRESS"
        ext_json = self._load_json_object(bundle.get("extJson", "{}"))
        if not isinstance(ext_json, dict):
            ext_json = {}
        ext_json["questionCount"] = total_count
        ext_json["completedCount"] = completed_count
        self.repository.replace_student_review_plan(
            {
                "id": str(bundle.get("id", "")).strip(),
                "studentUserId": actor.user_id,
                "planType": str(bundle.get("planType", "")).strip(),
                "planName": str(bundle.get("planName", "")).strip(),
                "status": plan_status,
                "generatedAt": str(bundle.get("generatedAt", "")).strip() or now_iso,
                "startedAt": str(bundle.get("startedAt", "")).strip() or now_iso,
                "completedAt": now_iso if plan_status == "COMPLETED" else "",
                "lastExecutedAt": now_iso,
                "extJson": self._dump_json(ext_json),
                "createTime": str(bundle.get("createTime", "")).strip() or now_iso,
                "updateTime": now_iso,
            },
            next_items,
        )
        return self.get_student_personal_bank_review_plan(plan_id, filters, actor)

    def get_student_dashboard(self, actor: Actor) -> Dict[str, object]:
        _, _, profile = self._load_student_profile_bundle(actor.user_id)
        self._assert_student_profile_complete(profile)
        return self._build_student_dashboard_payload(profile, actor.user_id)

    def save_student_profile(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        profile_input = parse_student_profile_model(payload)
        exam_category_code = str(profile_input.examCategoryCode).strip()
        joint_exam_group_code = str(profile_input.jointExamGroupCode).strip()
        self._validate_student_profile_selection(exam_category_code, joint_exam_group_code)
        record, record_ext, profile = self._load_student_profile_bundle(actor.user_id)
        profile = self._apply_student_profile_selection(
            profile,
            exam_category_code,
            joint_exam_group_code,
        )
        self.repository.set_student_profile_selection(
            actor.user_id,
            exam_category_code,
            joint_exam_group_code,
            self._now_iso(),
        )
        self._sync_student_directory_scope_projection(
            actor.user_id,
            exam_category_code,
            joint_exam_group_code,
        )
        self._save_student_profile_snapshot(record, record_ext, profile, actor.user_id)
        return self.get_student_dashboard(actor)

    def submit_student_session(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        submission = parse_student_submit_model(payload)
        _, _, profile = self._load_student_profile_bundle(actor.user_id)
        self._assert_student_profile_complete(profile)
        current_session = profile.get("examSession", {})
        if not isinstance(current_session, dict):
            current_session = {}
        answered_count = max(int(current_session.get("answeredCount", 0)), int(submission.answeredCount))
        elapsed_sec = max(int(current_session.get("elapsedSec", 0)), int(submission.elapsedSec))
        exam_session = {
            "answeredCount": answered_count,
            "elapsedSec": elapsed_sec,
            "updateTime": self._now_iso(),
        }
        self.repository.set_student_profile_exam_session(actor.user_id, exam_session, self._now_iso())
        return exam_session

    def student_daily_check_in(self, actor: Actor) -> Dict[str, object]:
        settings = self._load_system_state()["systemSettings"]
        _, _, profile = self._load_student_profile_bundle(actor.user_id)
        self._assert_student_profile_complete(profile)
        today = datetime.now(timezone.utc).date().isoformat()
        streak_days = self._apply_student_daily_check_in(profile, settings, today, actor.user_id)
        self._notify_check_in_success(actor.user_id, int(profile["points"]))
        return {
            "points": profile["points"],
            "title": profile["title"],
            "checkInDates": profile["checkInDates"],
            "streakDays": streak_days,
        }

    def list_wrong_book_questions(
        self,
        page: int,
        size: int,
        actor: Actor,
        filters: Optional[Dict[str, str]] = None,
    ) -> Tuple[List[Dict[str, str]], int]:
        normalized_filters = {key: str(value or "").strip() for key, value in dict(filters or {}).items()}
        target_knowledge_id = str(normalized_filters.get("knowledgeId", "")).strip()
        target_path_node_id = str(normalized_filters.get("knowledgePathNodeId", "")).strip()
        target_chapter = str(normalized_filters.get("chapter", "")).strip()
        target_chapter_code = str(normalized_filters.get("chapterCode", "")).strip()
        target_point_code = str(normalized_filters.get("pointCode", "")).strip()
        target_subject_code = str(normalized_filters.get("subjectCode", "")).strip()
        include_archived = str(normalized_filters.get("includeArchived", "")).strip().lower() == "true"
        descendant_knowledge_ids = self._resolve_wrong_book_path_node_descendants(
            target_path_node_id,
            subject_code=target_subject_code,
        )
        records = self.repository.list_student_records_by_user(actor.user_id)
        matched: List[Dict[str, str]] = []
        for question in records:
            if not self._is_question_visible_to_student(question, actor.user_id):
                continue
            if target_knowledge_id and str(question.get("knowledgeId", "")).strip() != target_knowledge_id:
                continue
            if descendant_knowledge_ids and str(question.get("knowledgeId", "")).strip() not in descendant_knowledge_ids:
                continue
            ext_json = self._load_json_object(question["extJson"])
            if target_chapter and str(ext_json.get("chapter", "")).strip() != target_chapter:
                continue
            if target_chapter_code and str(ext_json.get("chapterCode", "")).strip() != target_chapter_code:
                continue
            if target_point_code and str(ext_json.get("pointCode", "")).strip() != target_point_code:
                continue
            if target_subject_code and str(ext_json.get("subjectCode", "")).strip() != target_subject_code:
                continue
            student_record = ext_json.get("studentRecord", {})
            record_ext = self._load_json_object(student_record.get("extJson", "{}"))
            formal_state = student_record.get("formalState", {}) if isinstance(student_record.get("formalState", {}), dict) else {}
            chapter_practice = record_ext.get("chapterPractice", {})
            wrong_book = dict(record_ext.get("wrongBook", {}))
            if int(formal_state.get("wrongBookFlag", 0) or 0) > 0:
                wrong_book["isCollected"] = True
            if int(formal_state.get("wrongBookArchivedFlag", 0) or 0) > 0:
                wrong_book["isArchived"] = True
            if str(formal_state.get("wrongBookCollectedAt", "")).strip():
                wrong_book["collectedAt"] = str(formal_state.get("wrongBookCollectedAt", "")).strip()
            if str(formal_state.get("wrongBookLastWrongAt", "")).strip():
                wrong_book["lastWrongAt"] = str(formal_state.get("wrongBookLastWrongAt", "")).strip()
            if str(formal_state.get("wrongBookReviewedAt", "")).strip():
                wrong_book["reviewedAt"] = str(formal_state.get("wrongBookReviewedAt", "")).strip()
            if str(formal_state.get("wrongBookArchivedAt", "")).strip():
                wrong_book["archivedAt"] = str(formal_state.get("wrongBookArchivedAt", "")).strip()
            if str(formal_state.get("wrongBookRestoredAt", "")).strip():
                wrong_book["restoredAt"] = str(formal_state.get("wrongBookRestoredAt", "")).strip()
            if str(formal_state.get("wrongBookLastReasonCode", "")).strip():
                wrong_book["lastReasonCode"] = str(formal_state.get("wrongBookLastReasonCode", "")).strip()
            if str(formal_state.get("wrongBookLastReasonLabel", "")).strip():
                wrong_book["lastReasonLabel"] = str(formal_state.get("wrongBookLastReasonLabel", "")).strip()
            if int(formal_state.get("wrongBookReviewCount", 0) or 0) > 0:
                wrong_book["reviewCount"] = int(formal_state.get("wrongBookReviewCount", 0) or 0)
            if int(formal_state.get("wrongBookPostWrongAttemptCount", 0) or 0) > 0:
                wrong_book["postWrongAttemptCount"] = int(formal_state.get("wrongBookPostWrongAttemptCount", 0) or 0)
            if int(formal_state.get("wrongBookPostWrongCorrectCount", 0) or 0) > 0:
                wrong_book["postWrongCorrectCount"] = int(formal_state.get("wrongBookPostWrongCorrectCount", 0) or 0)
            if int(formal_state.get("wrongCount", 0) or 0) > 0:
                wrong_book["wrongCount"] = int(formal_state.get("wrongCount", 0) or 0)
            if not include_archived and bool(wrong_book.get("isArchived")):
                continue
            should_include = bool(wrong_book.get("isCollected")) or (bool(chapter_practice) and not chapter_practice.get("isCorrect", True))
            if should_include:
                wrong_book["isCollected"] = True
                wrong_book["collectedAt"] = str(wrong_book.get("collectedAt", "")) or self._now_iso()
                if not isinstance(wrong_book.get("reasonStats"), list) or not wrong_book.get("reasonStats"):
                    fallback = self._resolve_wrong_reason_payload(
                        question,
                        str(chapter_practice.get("normalizedAnswer", "")),
                        bool(chapter_practice.get("isTimeout", False)),
                    )
                    wrong_book["reasonStats"] = self._upsert_wrong_reason_stats(
                        [],
                        str(fallback.get("reasonCode", "KNOWLEDGE_GAP")),
                        str(fallback.get("reasonLabel", "知识点掌握不牢")),
                    )
                if int(wrong_book.get("wrongCount", 0)) <= 0:
                    wrong_book["wrongCount"] = sum(int(item.get("count", 0)) for item in wrong_book.get("reasonStats", []))
                record_ext["wrongBook"] = wrong_book
                ext_json["studentState"] = {
                    "wrongBook": wrong_book,
                    "chapterPractice": chapter_practice,
                    "aiMarking": record_ext.get("aiMarking", {}),
                    "aiTutor": record_ext.get("aiTutor", {}),
                }
                question["extJson"] = self._dump_json(ext_json)
                matched.append(question)
        return self._paginate_questions(matched, page, size)

    def archive_wrong_book_questions(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        raw_question_ids = payload.get("questionIds")
        if raw_question_ids is None:
            raw_question_ids = payload.get("questionIds", [])
        question_ids: List[str] = []
        if isinstance(raw_question_ids, list):
            for item in raw_question_ids:
                question_id = str(item or "").strip()
                if question_id and question_id not in question_ids:
                    question_ids.append(question_id)
        if not question_ids:
            subject_code = str(payload.get("subjectCode") or "").strip()
            wrong_questions, _ = self.list_wrong_book_questions(
                1,
                500,
                actor,
                {"subjectCode": subject_code, "subject_code": subject_code},
            )
            summary = self.get_student_error_book_summary({"subjectCode": subject_code, "subject_code": subject_code}, actor)
            mastered_ids = {
                str(item.get("questionId", "")).strip()
                for item in summary.get("questionInsights", [])
                if str(item.get("reviewStatusKey", "")).strip() == "mastered"
            }
            question_ids = [
                str(item.get("id", "")).strip()
                for item in wrong_questions
                if str(item.get("id", "")).strip() in mastered_ids
            ]
        archived_count = 0
        for question_id in question_ids:
            record = self.repository.get_student_question_bank(question_id, actor.user_id)
            if not record:
                continue
            record_ext = self._load_json_object(record["extJson"])
            wrong_book = record_ext.get("wrongBook", {})
            if not isinstance(wrong_book, dict) or not wrong_book.get("isCollected"):
                continue
            wrong_book["isArchived"] = True
            wrong_book["archivedAt"] = self._now_iso()
            record_ext["wrongBook"] = wrong_book
            personal_bank = record_ext.get("personalBank", {})
            if not isinstance(personal_bank, dict):
                personal_bank = {}
            personal_bank["isCollected"] = True
            personal_bank["collectedAt"] = str(personal_bank.get("collectedAt", "")) or self._now_iso()
            personal_bank["sourceType"] = "HARVESTED_ARCHIVE"
            personal_bank["sourceLabel"] = "已斩获归档"
            record_ext["personalBank"] = personal_bank
            self._save_student_question_record_bundle(record, record_ext, actor.user_id)
            archived_count += 1
        return {"archivedCount": archived_count, "questionIds": question_ids}

    def restore_archived_wrong_book_questions(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        raw_question_ids = payload.get("questionIds")
        if raw_question_ids is None:
            raw_question_ids = payload.get("questionIds", [])
        question_ids: List[str] = []
        if isinstance(raw_question_ids, list):
            for item in raw_question_ids:
                question_id = str(item or "").strip()
                if question_id and question_id not in question_ids:
                    question_ids.append(question_id)
        restored_count = 0
        for question_id in question_ids:
            record = self.repository.get_student_question_bank(question_id, actor.user_id)
            if not record:
                continue
            record_ext = self._load_json_object(record["extJson"])
            wrong_book = record_ext.get("wrongBook", {})
            if not isinstance(wrong_book, dict) or not bool(wrong_book.get("isArchived")):
                continue
            wrong_book["isArchived"] = False
            wrong_book["restoredAt"] = self._now_iso()
            record_ext["wrongBook"] = wrong_book
            self._save_student_question_record_bundle(record, record_ext, actor.user_id)
            restored_count += 1
        return {"restoredCount": restored_count, "questionIds": question_ids}

    def get_student_error_book_summary(self, filters: Dict[str, str], actor: Actor) -> Dict[str, object]:
        payload = self._build_student_error_book_summary_payload(actor, filters)
        return {
            key: value
            for key, value in payload.items()
            if not str(key).startswith("_")
        }

    def list_teacher_wrong_book_questions(
        self,
        filters: Dict[str, object],
        page: int,
        size: int,
        actor: Actor,
    ) -> Tuple[List[Dict[str, str]], int]:
        target_student_user_id = str(filters.get("studentUserId") or "").strip()
        target_actor = self._build_error_book_target_actor(actor, target_student_user_id)
        selected_subject_codes = set(self._normalize_subject_codes_filter(filters))
        normalized_filters = {
            "knowledgeId": str(filters.get("knowledgeId") or "").strip(),
            "knowledgePathNodeId": str(filters.get("knowledgePathNodeId") or "").strip(),
            "chapter": str(filters.get("chapter") or "").strip(),
            "chapterCode": str(filters.get("chapterCode") or "").strip(),
            "pointCode": str(filters.get("pointCode") or "").strip(),
            "subjectCode": str(filters.get("subjectCode") or "").strip(),
        }
        if len(selected_subject_codes) == 1 and not normalized_filters["subjectCode"]:
            only_subject_code = next(iter(selected_subject_codes))
            normalized_filters["subjectCode"] = only_subject_code
        items, total = self.list_wrong_book_questions(1, 500, target_actor, normalized_filters)
        if selected_subject_codes:
            filtered_items = []
            for item in items:
                ext_json = self._load_json_object(item.get("extJson", "{}"))
                if str(ext_json.get("subjectCode", "")).strip() not in selected_subject_codes:
                    continue
                filtered_items.append(item)
            items = filtered_items
            total = len(items)
        return self._paginate_questions(items, page, size)

    def list_teacher_error_book_students(self, actor: Actor) -> List[Dict[str, object]]:
        students, _ = self.list_managed_users({"role": ROLE_STUDENT, "keyword": ""}, 1, 500, actor)
        return [
            {
                "studentUserId": str(item.get("userId", "")).strip(),
                "studentName": str(item.get("name", "")).strip() or str(item.get("userId", "")).strip(),
                "examCategoryCode": str(item.get("examCategoryCode", "")).strip(),
                "jointExamGroupCode": str(item.get("jointExamGroupCode", "")).strip(),
                "classId": str(item.get("jointExamGroupCode", "")).strip(),
                "className": str((get_joint_exam_group(str(item.get("jointExamGroupCode", "")).strip()) or {}).get("jointExamGroupName", "")).strip() or str(item.get("jointExamGroupCode", "")).strip(),
                "vocationalMajor": str(item.get("vocationalMajor", "")).strip(),
                "prepStage": str(item.get("prepStage", "")).strip(),
            }
            for item in students
            if str(item.get("userId", "")).strip()
        ]

    def _list_teacher_students_for_class(self, class_id: str, actor: Actor) -> List[Dict[str, object]]:
        normalized_class_id = str(class_id or "").strip()
        students = self.list_teacher_error_book_students(actor)
        if normalized_class_id:
            students = [item for item in students if str(item.get("classId", "")).strip() == normalized_class_id]
        return students

    def get_teacher_error_book_class_overview(self, filters: Dict[str, object], actor: Actor) -> Dict[str, object]:
        class_id = str(filters.get("classId") or "").strip()
        selected_subject_codes = self._normalize_subject_codes_filter(filters)
        student_rows = self._list_teacher_students_for_class(class_id, actor)
        if not student_rows:
            return {
                "classMeta": {
                    "classId": class_id,
                    "className": class_id or "当前班级",
                    "studentCount": 0,
                },
                "summaryCards": {
                    "studentCount": 0,
                    "analyzedStudentCount": 0,
                    "averageCoverageRate": 0.0,
                    "averageAccuracy": 0.0,
                    "atRiskStudentCount": 0,
                    "overdueQuestionCount": 0,
                },
                "topChapters": [],
                "topStudents": [],
            }

        if not class_id:
            class_id = str(student_rows[0].get("classId", "")).strip()
        class_name = str(student_rows[0].get("className", "")).strip() or class_id or "当前班级"

        total_coverage = 0.0
        total_accuracy = 0.0
        total_practice_points = 0
        analyzed_student_count = 0
        practice_point_student_count = 0
        at_risk_student_count = 0
        overdue_question_count = 0
        top_chapter_map: Dict[str, Dict[str, object]] = {}
        top_students: List[Dict[str, object]] = []

        for student in student_rows:
            student_user_id = str(student.get("studentUserId", "")).strip()
            summary_payload = self._build_teacher_error_book_summary_payload(
                {
                    "studentUserId": student_user_id,
                    "subjectCodes": ",".join(selected_subject_codes),
                },
                actor,
            )
            current_subject = summary_payload.get("currentSubject", {}) if isinstance(summary_payload.get("currentSubject", {}), dict) else {}
            practice_points = self._build_teacher_practice_point_summary(student_user_id, selected_subject_codes)
            wrong_count = int(current_subject.get("wrongCount", 0))
            if wrong_count > 0:
                analyzed_student_count += 1
            coverage_rate = float(current_subject.get("knowledgeCoverageRate", 0.0))
            analytics_overview = current_subject.get("analyticsOverview", {}) if isinstance(current_subject.get("analyticsOverview", {}), dict) else {}
            average_accuracy = float(analytics_overview.get("averageAccuracy", 0.0))
            total_coverage += coverage_rate
            total_accuracy += average_accuracy
            total_practice_points += int(practice_points.get("total", 0) or 0)
            if int(practice_points.get("total", 0) or 0) > 0:
                practice_point_student_count += 1
            overdue_question_count += int(current_subject.get("overdueQuestionCount", 0))
            if float(current_subject.get("lowestMastery", 0.0)) < 0.4 or wrong_count >= 3:
                at_risk_student_count += 1
            for chapter_row in current_subject.get("topChapters", []):
                chapter_key = str(chapter_row.get("chapterCode", "")).strip() or str(chapter_row.get("chapterName", "")).strip()
                target = top_chapter_map.setdefault(
                    chapter_key,
                    {
                        "chapterCode": str(chapter_row.get("chapterCode", "")).strip(),
                        "chapterName": str(chapter_row.get("chapterName", "")).strip() or chapter_key or "未标注章节",
                        "wrongCount": 0,
                    },
                )
                target["wrongCount"] = int(target.get("wrongCount", 0)) + int(chapter_row.get("wrongCount", 0))
            top_students.append(
                {
                    "studentUserId": str(student.get("studentUserId", "")).strip(),
                    "studentName": str(student.get("studentName", "")).strip(),
                    "wrongCount": wrong_count,
                    "lowestMastery": float(current_subject.get("lowestMastery", 0.0)),
                    "coverageRate": coverage_rate,
                    "practicePointTotal": int(practice_points.get("total", 0) or 0),
                    "practicePointTodayDelta": int(practice_points.get("todayDelta", 0) or 0),
                }
            )

        student_count = len(student_rows)
        return {
            "classMeta": {
                "classId": class_id,
                "className": class_name,
                "studentCount": student_count,
                "selectedSubjectCodes": selected_subject_codes,
            },
            "summaryCards": {
                "studentCount": student_count,
                "analyzedStudentCount": analyzed_student_count,
                "averagePracticePoints": round(total_practice_points / student_count, 2) if student_count else 0.0,
                "practicePointStudentCount": practice_point_student_count,
                "averageCoverageRate": round(total_coverage / student_count, 4) if student_count else 0.0,
                "averageAccuracy": round(total_accuracy / student_count, 4) if student_count else 0.0,
                "atRiskStudentCount": at_risk_student_count,
                "overdueQuestionCount": overdue_question_count,
            },
            "topChapters": sorted(
                top_chapter_map.values(),
                key=lambda item: (-int(item.get("wrongCount", 0)), str(item.get("chapterCode", "")), str(item.get("chapterName", ""))),
            )[:5],
            "topStudents": sorted(
                top_students,
                key=lambda item: (-int(item.get("wrongCount", 0)), float(item.get("lowestMastery", 0.0)), str(item.get("studentUserId", ""))),
            )[:8],
        }

    def _build_teacher_practice_point_summary(
        self,
        student_user_id: str,
        selected_subject_codes: List[str],
    ) -> Dict[str, object]:
        normalized_student_user_id = str(student_user_id or "").strip()
        normalized_subject_codes = [
            str(item or "").strip()
            for item in (selected_subject_codes or [])
            if str(item or "").strip()
        ]
        if not normalized_student_user_id:
            return {
                "subjectCode": "",
                "total": 0,
                "todayDelta": 0,
                "rank": 0,
                "leaderboard": [],
                "awardUnlocked": False,
                "awardProgress": 0,
                "awardThreshold": int(self.CHALLENGE_POINT_AWARD_THRESHOLD),
                "award": {},
                "selectedSubjectCount": 0,
                **self._build_challenge_point_level(0),
            }

        if len(normalized_subject_codes) == 1:
            summary = dict(self._build_challenge_point_summary(normalized_student_user_id, normalized_subject_codes[0]))
            summary["selectedSubjectCount"] = 1
            return summary

        subject_code_set = set(normalized_subject_codes)
        repository_rows = self.repository.list_challenge_point_subjects_by_student(normalized_student_user_id)
        filtered_rows = [
            row
            for row in repository_rows
            if not subject_code_set or str(row.get("subjectCode", "")).strip() in subject_code_set
        ]
        today = datetime.now(timezone.utc).date().isoformat()
        total_points = sum(int(row.get("totalPoints", 0) or 0) for row in filtered_rows)
        today_delta = sum(
            self.repository.count_today_challenge_points(
                normalized_student_user_id,
                str(row.get("subjectCode", "")).strip(),
                today,
            )
            for row in filtered_rows
            if str(row.get("subjectCode", "")).strip()
        )
        level_payload = self._build_challenge_point_level(total_points)
        return {
            "subjectCode": "__MULTI__" if len(filtered_rows) != 1 else str(filtered_rows[0].get("subjectCode", "")).strip(),
            "total": total_points,
            "todayDelta": today_delta,
            "rank": 0,
            "leaderboard": [],
            "awardUnlocked": total_points >= int(self.CHALLENGE_POINT_AWARD_THRESHOLD),
            "awardProgress": min(total_points, int(self.CHALLENGE_POINT_AWARD_THRESHOLD)),
            "awardThreshold": int(self.CHALLENGE_POINT_AWARD_THRESHOLD),
            "award": {},
            "selectedSubjectCount": len(filtered_rows),
            **level_payload,
        }

    def _docx_style_or_none(self, document: object, style_name: str) -> object | None:
        try:
            return document.styles[style_name]
        except (KeyError, AttributeError, TypeError):
            return None

    def _ensure_docx_child(self, parent: object, tag: str, qn: object, OxmlElement: object) -> object:
        child = parent.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            parent.append(child)
        return child

    def _set_docx_rfonts(
        self,
        parent: object,
        ascii_font: str,
        east_asia_font: str,
        qn: object,
        OxmlElement: object,
    ) -> None:
        r_fonts = self._ensure_docx_child(parent, "w:rFonts", qn, OxmlElement)
        if ascii_font:
            r_fonts.set(qn("w:ascii"), ascii_font)
            r_fonts.set(qn("w:hAnsi"), ascii_font)
            r_fonts.set(qn("w:cs"), ascii_font)
        if east_asia_font:
            r_fonts.set(qn("w:eastAsia"), east_asia_font)

    def _set_docx_style_font(
        self,
        style: object | None,
        ascii_font: str,
        east_asia_font: str,
        qn: object,
        OxmlElement: object,
        *,
        size: object | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        if style is None:
            return
        if ascii_font:
            style.font.name = ascii_font
        if size is not None:
            style.font.size = size
        if bold is not None:
            style.font.bold = bold
        if italic is not None:
            style.font.italic = italic
        r_pr = self._ensure_docx_child(style._element, "w:rPr", qn, OxmlElement)
        self._set_docx_rfonts(r_pr, ascii_font, east_asia_font, qn, OxmlElement)

    def _apply_docx_run_font(
        self,
        run: object | None,
        ascii_font: str,
        east_asia_font: str,
        qn: object,
        OxmlElement: object,
    ) -> None:
        if run is None:
            return
        if ascii_font:
            run.font.name = ascii_font
        r_pr = run._element.get_or_add_rPr()
        self._set_docx_rfonts(r_pr, ascii_font, east_asia_font, qn, OxmlElement)

    def _configure_docx_export_theme(
        self,
        document: object,
        Pt: object,
        qn: object,
        OxmlElement: object,
    ) -> None:
        body_ascii_font = "Microsoft YaHei"
        body_east_asia_font = "Microsoft YaHei"
        heading_ascii_font = "Microsoft YaHei"
        heading_east_asia_font = "Microsoft YaHei"
        ui_ascii_font = "Microsoft YaHei"
        ui_east_asia_font = "Microsoft YaHei"

        styles_root = document.styles.element
        doc_defaults = self._ensure_docx_child(styles_root, "w:docDefaults", qn, OxmlElement)
        r_pr_default = self._ensure_docx_child(doc_defaults, "w:rPrDefault", qn, OxmlElement)
        r_pr = self._ensure_docx_child(r_pr_default, "w:rPr", qn, OxmlElement)
        self._set_docx_rfonts(r_pr, body_ascii_font, body_east_asia_font, qn, OxmlElement)
        lang = self._ensure_docx_child(r_pr, "w:lang", qn, OxmlElement)
        lang.set(qn("w:val"), "en-US")
        lang.set(qn("w:eastAsia"), "zh-CN")
        lang.set(qn("w:bidi"), "ar-SA")

        normal_style = self._docx_style_or_none(document, "Normal")
        self._set_docx_style_font(
            normal_style,
            body_ascii_font,
            body_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(11),
        )
        if normal_style is not None:
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.35

        heading_1 = self._docx_style_or_none(document, "Heading 1")
        self._set_docx_style_font(
            heading_1,
            heading_ascii_font,
            heading_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(16),
            bold=True,
        )
        if heading_1 is not None:
            heading_1.paragraph_format.space_before = Pt(12)
            heading_1.paragraph_format.space_after = Pt(8)
            heading_1.paragraph_format.line_spacing = 1.2

        heading_2 = self._docx_style_or_none(document, "Heading 2")
        self._set_docx_style_font(
            heading_2,
            heading_ascii_font,
            heading_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(13),
            bold=True,
        )
        if heading_2 is not None:
            heading_2.paragraph_format.space_before = Pt(10)
            heading_2.paragraph_format.space_after = Pt(6)
            heading_2.paragraph_format.line_spacing = 1.2

        heading_3 = self._docx_style_or_none(document, "Heading 3")
        self._set_docx_style_font(
            heading_3,
            heading_ascii_font,
            heading_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(12),
            bold=True,
        )
        if heading_3 is not None:
            heading_3.paragraph_format.space_before = Pt(8)
            heading_3.paragraph_format.space_after = Pt(4)
            heading_3.paragraph_format.line_spacing = 1.2

        list_bullet = self._docx_style_or_none(document, "List Bullet")
        self._set_docx_style_font(
            list_bullet,
            body_ascii_font,
            body_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(10.5),
        )
        if list_bullet is not None:
            list_bullet.paragraph_format.space_before = Pt(0)
            list_bullet.paragraph_format.space_after = Pt(2)
            list_bullet.paragraph_format.line_spacing = 1.25

        header_style = self._docx_style_or_none(document, "Header")
        self._set_docx_style_font(
            header_style,
            ui_ascii_font,
            ui_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(9),
        )

        footer_style = self._docx_style_or_none(document, "Footer")
        self._set_docx_style_font(
            footer_style,
            ui_ascii_font,
            ui_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(9),
        )

        hyperlink_style = self._docx_style_or_none(document, "Hyperlink")
        self._set_docx_style_font(
            hyperlink_style,
            ui_ascii_font,
            ui_east_asia_font,
            qn,
            OxmlElement,
            size=Pt(10),
        )

    def _shade_docx_cell(self, cell: object, fill: str, qn: object, OxmlElement: object) -> None:
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = cell_properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            cell_properties.append(shading)
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)

    def _format_docx_table(
        self,
        table: object,
        Pt: object,
        qn: object,
        OxmlElement: object,
        WD_ALIGN_PARAGRAPH: object,
        WD_TABLE_ALIGNMENT: object,
        WD_CELL_VERTICAL_ALIGNMENT: object,
        *,
        center: bool = False,
        header_rows: int = 0,
        emphasize_first_column: bool = False,
    ) -> None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            is_header_row = row_index < header_rows
            for column_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if is_header_row:
                    self._shade_docx_cell(cell, "EEF2F7", qn, OxmlElement)
                elif emphasize_first_column and column_index == 0:
                    self._shade_docx_cell(cell, "F7F7F7", qn, OxmlElement)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center or is_header_row else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        self._apply_docx_run_font(run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)
                        run.font.size = Pt(10)
                        if is_header_row or (emphasize_first_column and column_index == 0):
                            run.bold = True

    def export_teacher_error_book_class_report(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        normalized_payload = dict(payload or {}) if isinstance(payload, dict) else {}
        class_id = str(normalized_payload.get("classId") or "").strip()
        if not class_id:
            raise validation_failed("classId 不能为空。")
        class_overview = self.get_teacher_error_book_class_overview(normalized_payload, actor)
        class_meta = class_overview.get("classMeta", {}) if isinstance(class_overview.get("classMeta", {}), dict) else {}
        summary_cards = class_overview.get("summaryCards", {}) if isinstance(class_overview.get("summaryCards", {}), dict) else {}
        top_chapters = class_overview.get("topChapters", []) if isinstance(class_overview.get("topChapters", []), list) else []
        top_students = class_overview.get("topStudents", []) if isinstance(class_overview.get("topStudents", []), list) else []
        selected_subject_codes = self._normalize_subject_codes_filter(normalized_payload)
        student_rows = self._list_teacher_students_for_class(class_id, actor)
        reason_counter: Dict[str, int] = {}
        chapter_reason_counter: Dict[str, Dict[str, int]] = {}
        for student in student_rows:
            summary_payload = self._build_teacher_error_book_summary_payload(
                {
                    "studentUserId": str(student.get("studentUserId", "")).strip(),
                    "subjectCodes": ",".join(selected_subject_codes),
                },
                actor,
            )
            for insight in summary_payload.get("questionInsights", []):
                question_id = str(insight.get("questionId", "")).strip()
                if not question_id:
                    continue
                target_actor = self._build_error_book_target_actor(actor, str(student.get("studentUserId", "")).strip())
                wrong_questions, _ = self.list_wrong_book_questions(1, 500, target_actor, {})
                matched_question = next((item for item in wrong_questions if str(item.get("id", "")).strip() == question_id), None)
                if not matched_question:
                    continue
                ext_json = self._load_json_object(matched_question.get("extJson", "{}"))
                student_state = ext_json.get("studentState", {}) if isinstance(ext_json.get("studentState", {}), dict) else {}
                wrong_book = student_state.get("wrongBook", {}) if isinstance(student_state.get("wrongBook", {}), dict) else {}
                reason_stats = wrong_book.get("reasonStats", []) if isinstance(wrong_book.get("reasonStats", []), list) else []
                chapter_key = str(insight.get("chapterCode", "")).strip() or str(insight.get("chapterName", "")).strip() or "未标注章节"
                for row in reason_stats:
                    if not isinstance(row, dict):
                        continue
                    reason_label = str(row.get("reasonLabel", "")).strip()
                    reason_count = max(1, int(row.get("count", 0)))
                    if not reason_label:
                        continue
                    reason_counter[reason_label] = reason_counter.get(reason_label, 0) + reason_count
                    chapter_map = chapter_reason_counter.setdefault(chapter_key, {})
                    chapter_map[reason_label] = chapter_map.get(reason_label, 0) + reason_count
        if not self.is_docx_available():
            raise failed_dependency("当前环境缺少 python-docx 依赖，无法导出 Word 文档。")
        try:
            from docx import Document as _Document  # type: ignore
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            from docx.oxml import OxmlElement  # type: ignore
            from docx.oxml.ns import qn  # type: ignore
            from docx.shared import Inches, Pt  # type: ignore
        except Exception as exc:
            raise failed_dependency("当前环境缺少 python-docx 依赖，无法导出 Word 文档。") from exc

        document = _Document()
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        self._configure_docx_export_theme(document, Pt, qn, OxmlElement)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f"{str(class_meta.get('className', '')).strip() or class_id} 班级学情汇总总报告")
        title_run.bold = True
        title_run.font.size = Pt(22)
        self._apply_docx_run_font(title_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("教师智能学情修复中心 / 班级总览导出版")
        subtitle_run.font.size = Pt(12)
        self._apply_docx_run_font(subtitle_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        class_summary_sentence = document.add_paragraph()
        class_summary_sentence.alignment = WD_ALIGN_PARAGRAPH.CENTER
        analyzed_student_count = int(summary_cards.get("analyzedStudentCount", 0))
        student_count = int(summary_cards.get("studentCount", 0))
        average_accuracy = float(summary_cards.get("averageAccuracy", 0.0))
        at_risk_student_count = int(summary_cards.get("atRiskStudentCount", 0))
        if at_risk_student_count >= max(1, student_count // 2):
            summary_sentence = "班级整体处于高关注区，建议优先组织高频错因讲评并分层跟进重点学生。"
        elif average_accuracy < 0.6:
            summary_sentence = "班级整体正在爬坡，建议围绕高频章节组织二轮复盘并跟踪主错因清理进度。"
        else:
            summary_sentence = "班级整体状态可控，建议保持章节复盘节奏并重点压缩剩余高频错因。"
        class_summary_sentence.add_run(
            f"首页摘要结论：当前班级共 {student_count} 人，已分析 {analyzed_student_count} 人，平均正确率约 {int(round(average_accuracy * 100))}% 。{summary_sentence}"
        ).bold = True

        risk_card = document.add_paragraph()
        risk_card.alignment = WD_ALIGN_PARAGRAPH.CENTER
        risk_label = "高风险班级" if at_risk_student_count >= max(1, student_count // 2) else "常规关注班级"
        risk_advice = "优先集中讲评主错因并逐个跟进重点学生。" if risk_label == "高风险班级" else "保持章节复盘节奏并按风险学生定向提醒。"
        risk_card.add_run(f"风险等级结论卡：{risk_label} / 建议：{risk_advice}").bold = True

        class_advice = document.add_paragraph()
        class_advice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if average_accuracy < 0.5:
            advice_text = "班级一句话建议卡：先统一讲评高频错因，再按重点学生名单做分层复盘。"
        elif int(summary_cards.get("overdueQuestionCount", 0)) > 0:
            advice_text = "班级一句话建议卡：优先清理遗忘预警题，再组织章节复盘和随堂检测。"
        else:
            advice_text = "班级一句话建议卡：保持章节复盘节奏，重点压缩剩余高频错因并强化易失分考点。"
        class_advice.add_run(advice_text).bold = True

        summary_table = document.add_table(rows=2, cols=3)
        summary_table.style = "Table Grid"
        summary_table.cell(0, 0).text = "班级学生数"
        summary_table.cell(0, 1).text = "有错题学生"
        summary_table.cell(0, 2).text = "风险学生数"
        summary_table.cell(1, 0).text = str(summary_cards.get("studentCount", 0))
        summary_table.cell(1, 1).text = str(summary_cards.get("analyzedStudentCount", 0))
        summary_table.cell(1, 2).text = str(summary_cards.get("atRiskStudentCount", 0))
        self._format_docx_table(
            summary_table,
            Pt,
            qn,
            OxmlElement,
            WD_ALIGN_PARAGRAPH,
            WD_TABLE_ALIGNMENT,
            WD_CELL_VERTICAL_ALIGNMENT,
            center=True,
            header_rows=1,
        )

        summary_table_b = document.add_table(rows=2, cols=3)
        summary_table_b.style = "Table Grid"
        summary_table_b.cell(0, 0).text = "平均覆盖率"
        summary_table_b.cell(0, 1).text = "平均正确率"
        summary_table_b.cell(0, 2).text = "遗忘预警题"
        summary_table_b.cell(1, 0).text = f"{int(round(float(summary_cards.get('averageCoverageRate', 0.0)) * 100))}%"
        summary_table_b.cell(1, 1).text = f"{int(round(float(summary_cards.get('averageAccuracy', 0.0)) * 100))}%"
        summary_table_b.cell(1, 2).text = str(summary_cards.get("overdueQuestionCount", 0))
        self._format_docx_table(
            summary_table_b,
            Pt,
            qn,
            OxmlElement,
            WD_ALIGN_PARAGRAPH,
            WD_TABLE_ALIGNMENT,
            WD_CELL_VERTICAL_ALIGNMENT,
            center=True,
            header_rows=1,
        )

        chapters_heading = document.add_heading("班级高频章节", level=2)
        chapters_heading.paragraph_format.space_after = Pt(8)
        if top_chapters:
            chapter_table = document.add_table(rows=1, cols=2)
            chapter_table.style = "Table Grid"
            chapter_table.cell(0, 0).text = "章节"
            chapter_table.cell(0, 1).text = "错题量"
            for row in top_chapters:
                row_cells = chapter_table.add_row().cells
                row_cells[0].text = str(row.get("chapterName", "")).strip() or str(row.get("chapterCode", "")).strip() or "未标注章节"
                row_cells[1].text = str(row.get("wrongCount", 0))
            self._format_docx_table(
                chapter_table,
                Pt,
                qn,
                OxmlElement,
                WD_ALIGN_PARAGRAPH,
                WD_TABLE_ALIGNMENT,
                WD_CELL_VERTICAL_ALIGNMENT,
                header_rows=1,
            )
        else:
            document.add_paragraph("当前班级暂无高频章节数据。")

        students_heading = document.add_heading("重点学生", level=2)
        students_heading.paragraph_format.space_after = Pt(8)
        if top_students:
            student_table = document.add_table(rows=1, cols=4)
            student_table.style = "Table Grid"
            student_table.cell(0, 0).text = "学生"
            student_table.cell(0, 1).text = "错题量"
            student_table.cell(0, 2).text = "最低掌握度"
            student_table.cell(0, 3).text = "覆盖率"
            for row in top_students:
                row_cells = student_table.add_row().cells
                row_cells[0].text = str(row.get("studentName", "")).strip() or str(row.get("studentUserId", "")).strip()
                row_cells[1].text = str(row.get("wrongCount", 0))
                row_cells[2].text = f"{int(round(float(row.get('lowestMastery', 0.0)) * 100))}%"
                row_cells[3].text = f"{int(round(float(row.get('coverageRate', 0.0)) * 100))}%"
            self._format_docx_table(
                student_table,
                Pt,
                qn,
                OxmlElement,
                WD_ALIGN_PARAGRAPH,
                WD_TABLE_ALIGNMENT,
                WD_CELL_VERTICAL_ALIGNMENT,
                header_rows=1,
            )
        else:
            document.add_paragraph("当前班级暂无重点学生数据。")

        reason_heading = document.add_heading("班级错因统计图表化文本", level=2)
        reason_heading.paragraph_format.space_after = Pt(8)
        if reason_counter:
            total_reason_count = sum(int(item) for item in reason_counter.values())
            for reason_label, count in sorted(reason_counter.items(), key=lambda item: (-int(item[1]), item[0])):
                ratio = (int(count) / total_reason_count) if total_reason_count > 0 else 0.0
                bars = "■" * max(1, min(10, int(round(ratio * 10))))
                document.add_paragraph(f"{reason_label} | {bars} | {count}次 | {ratio:.0%}")
        else:
            document.add_paragraph("当前班级暂无错因统计文本。")

        chapter_reason_heading = document.add_heading("章节错因图表化文本", level=2)
        chapter_reason_heading.paragraph_format.space_after = Pt(8)
        if chapter_reason_counter:
            for chapter_key, reason_map in sorted(chapter_reason_counter.items(), key=lambda item: item[0]):
                chapter_label = str((next((item.get("chapterName", "") for item in top_chapters if str(item.get("chapterCode", "")).strip() == chapter_key or str(item.get("chapterName", "")).strip() == chapter_key), "")) or chapter_key)
                document.add_paragraph(f"{chapter_label}")
                chapter_total = sum(int(item) for item in reason_map.values())
                for reason_label, count in sorted(reason_map.items(), key=lambda item: (-int(item[1]), item[0]))[:3]:
                    ratio = (int(count) / chapter_total) if chapter_total > 0 else 0.0
                    bars = "■" * max(1, min(10, int(round(ratio * 10))))
                    document.add_paragraph(f"  {reason_label} | {bars} | {count}次 | {ratio:.0%}")
        else:
            document.add_paragraph("当前班级暂无章节错因图表化文本。")

        export_time = datetime.now(timezone(timedelta(hours=8))).strftime("%Y-%m-%d %H:%M:%S")
        footer = document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run(f"导出时间：{export_time}")
        self._apply_docx_run_font(footer_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        content = io.BytesIO()
        document.save(content)
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
        return {
            "format": "docx",
            "fileName": f"class-error-book-report-{class_id}-{timestamp}.docx",
            "mediaType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "contentBase64": base64.b64encode(content.getvalue()).decode("ascii"),
            "classId": class_id,
        }

    def export_teacher_error_book_class_package(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        normalized_payload = dict(payload or {}) if isinstance(payload, dict) else {}
        class_id = str(normalized_payload.get("classId") or "").strip()
        if not class_id:
            raise validation_failed("classId 不能为空。")
        student_rows = self._list_teacher_students_for_class(class_id, actor)
        if not student_rows:
            raise validation_failed("当前班级暂无可导出的学生。")
        selected_subject_codes = self._normalize_subject_codes_filter(normalized_payload)
        class_report = self.export_teacher_error_book_class_report(normalized_payload, actor)
        zip_buffer = io.BytesIO()
        export_items: List[Dict[str, object]] = []
        with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zip_file:
            readme_lines = [
                "教师端全班分析包 README",
                f"classId: {class_id}",
                f"selectedSubjectCodes: {','.join(selected_subject_codes) if selected_subject_codes else 'ALL'}",
                "使用说明:",
                "1. 先阅读班级总报告，定位班级高频章节与重点学生。",
                "2. 再按学生姓名顺序查看个人分析卷，逐个安排复盘。",
                "3. 若需课堂讲评，请优先处理班级总报告中的高频错因与高频章节。",
                "",
                "推荐阅读顺序:",
                "1. 班级总报告首页摘要结论",
                "2. 班级高频章节与重点学生",
                "3. 逐个学生个人分析卷",
                "",
                "目录树说明:",
                f"{str(class_report.get('fileName', f'{class_id}.docx'))}",
                "文件清单:",
                f"- {str(class_report.get('fileName', f'{class_id}.docx'))}",
            ]
            zip_file.writestr(
                str(class_report.get("fileName", f"{class_id}.docx")),
                base64.b64decode(str(class_report.get("contentBase64", ""))),
            )
            sorted_students = sorted(
                student_rows,
                key=lambda item: (str(item.get("studentName", "")).strip() or str(item.get("studentUserId", "")).strip(), str(item.get("studentUserId", "")).strip()),
            )
            for student in sorted_students:
                student_user_id = str(student.get("studentUserId", "")).strip()
                if not student_user_id:
                    continue
                export_result = self.export_teacher_error_book_word(
                    {
                        "studentUserId": student_user_id,
                        "subjectCodes": ",".join(selected_subject_codes),
                    },
                    actor,
                )
                student_name = str(student.get("studentName", "")).strip() or student_user_id
                file_name = f"{student_name}-{student_user_id}.docx"
                zip_file.writestr(file_name, base64.b64decode(str(export_result.get("contentBase64", ""))))
                export_items.append(
                    {
                        "studentUserId": student_user_id,
                        "studentName": student_name,
                        "fileName": file_name,
                        "questionCount": int(export_result.get("questionCount", 0)),
                    }
                )
                readme_lines.append(f"  |- {file_name}")
            zip_file.writestr("README.txt", "\n".join(readme_lines))
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
        return {
            "format": "zip",
            "fileName": f"class-error-book-package-{class_id}-{timestamp}.zip",
            "mediaType": "application/zip",
            "contentBase64": base64.b64encode(zip_buffer.getvalue()).decode("ascii"),
            "classId": class_id,
            "selectedSubjectCodes": selected_subject_codes,
            "items": export_items,
        }

    def _build_teacher_error_book_summary_payload(self, filters: Dict[str, object], actor: Actor) -> Dict[str, object]:
        target_student_user_id = str(filters.get("studentUserId") or "").strip()
        if not target_student_user_id:
            raise validation_failed("studentUserId 不能为空。")
        target_actor = self._build_error_book_target_actor(actor, target_student_user_id)
        summary_payload = self._build_student_error_book_summary_payload(target_actor, {})
        selected_subject_codes = self._normalize_subject_codes_filter(filters)
        question_insights = [
            item
            for item in summary_payload.get("questionInsights", [])
            if not selected_subject_codes or str(item.get("subjectCode", "")).strip() in set(selected_subject_codes)
        ]
        subject_rows = [
            item
            for item in summary_payload.get("subjectRows", [])
            if not selected_subject_codes or str(item.get("subjectCode", "")).strip() in set(selected_subject_codes)
        ]
        analytics_rows = self._extract_teacher_error_book_analytics_rows(target_actor, question_insights)
        average_accuracy = (
            sum(float(item.get("isCorrect", False)) for item in analytics_rows) / len(analytics_rows)
            if analytics_rows
            else 0.0
        )
        average_duration = (
            sum(float(item.get("answerDurationSec", 0)) for item in analytics_rows) / len(analytics_rows)
            if analytics_rows
            else 0.0
        )
        current_subject = {}
        if len(subject_rows) == 1:
            current_subject = dict(subject_rows[0])
            current_subject["analyticsOverview"] = {
                "averageAccuracy": round(average_accuracy, 4),
                "averageAnswerDurationSec": round(average_duration, 2),
                "answerCount": len(analytics_rows),
                "selectedSubjectCount": len(subject_rows),
            }
            current_subject["practicePoints"] = self._build_teacher_practice_point_summary(
                target_student_user_id,
                [str(current_subject.get("subjectCode", "")).strip()],
            )
        elif subject_rows:
            aggregated_top_chapters: Dict[str, Dict[str, object]] = {}
            repair_suggestions: List[Dict[str, object]] = []
            review_warnings: List[Dict[str, object]] = []
            weakest_question_ids: List[str] = []
            total_point_count = 0
            practiced_point_count = 0
            wrong_point_count = 0
            wrong_count = 0
            lowest_mastery = 1.0
            for row in subject_rows:
                wrong_count += int(row.get("wrongCount", 0))
                total_point_count += int(row.get("totalPointCount", 0))
                practiced_point_count += int(row.get("practicedPointCount", 0))
                wrong_point_count += int(row.get("wrongPointCount", 0))
                lowest_mastery = min(lowest_mastery, float(row.get("lowestMastery", 0.0)))
                weakest_question_ids.extend(
                    [
                        str(item).strip()
                        for item in row.get("weakestQuestionIds", [])
                        if str(item).strip()
                    ]
                )
                for chapter_row in row.get("topChapters", []):
                    chapter_key = str(chapter_row.get("chapterCode", "")).strip() or str(chapter_row.get("chapterName", "")).strip()
                    target = aggregated_top_chapters.setdefault(
                        chapter_key,
                        {
                            "chapterCode": str(chapter_row.get("chapterCode", "")).strip(),
                            "chapterName": str(chapter_row.get("chapterName", "")).strip() or chapter_key or "未标注章节",
                            "wrongCount": 0,
                        },
                    )
                    target["wrongCount"] = int(target.get("wrongCount", 0)) + int(chapter_row.get("wrongCount", 0))
                ai_suggestions = row.get("aiSuggestions", {}) if isinstance(row.get("aiSuggestions", {}), dict) else {}
                repair_suggestions.extend(ai_suggestions.get("repairSuggestions", []) if isinstance(ai_suggestions.get("repairSuggestions", []), list) else [])
                review_warnings.extend(ai_suggestions.get("reviewWarnings", []) if isinstance(ai_suggestions.get("reviewWarnings", []), list) else [])
            top_chapters = sorted(
                aggregated_top_chapters.values(),
                key=lambda item: (-int(item.get("wrongCount", 0)), str(item.get("chapterCode", "")), str(item.get("chapterName", ""))),
            )[:3]
            current_subject = {
                "subjectCode": "__MULTI__",
                "subjectName": "多科联合分析",
                "wrongCount": wrong_count,
                "redDotCount": wrong_count,
                "totalPointCount": total_point_count,
                "practicedPointCount": practiced_point_count,
                "wrongPointCount": wrong_point_count,
                "knowledgeCoverageRate": round(practiced_point_count / total_point_count, 4) if total_point_count else 0.0,
                "errorCoverageRate": round(wrong_point_count / total_point_count, 4) if total_point_count else 0.0,
                "topChapters": top_chapters,
                "overdueQuestionCount": len(review_warnings),
                "lowestMastery": 0.0 if lowest_mastery == 1.0 else round(lowest_mastery, 4),
                "weakestQuestionIds": weakest_question_ids[:10],
                "aiSuggestions": {
                    "repairSuggestions": repair_suggestions[:5],
                    "practiceSuggestion": {
                        "questionIds": weakest_question_ids[:10],
                        "questionCount": len(weakest_question_ids[:10]),
                        "lowestMastery": 0.0 if lowest_mastery == 1.0 else round(lowest_mastery, 4),
                    },
                    "reviewWarnings": review_warnings[:8],
                },
                "analyticsOverview": {
                    "averageAccuracy": round(average_accuracy, 4),
                    "averageAnswerDurationSec": round(average_duration, 2),
                    "answerCount": len(analytics_rows),
                    "selectedSubjectCount": len(subject_rows),
                },
                "practicePoints": self._build_teacher_practice_point_summary(
                    target_student_user_id,
                    [str(item.get("subjectCode", "")).strip() for item in subject_rows],
                ),
            }
        practice_point_summary = self._build_teacher_practice_point_summary(target_student_user_id, selected_subject_codes)
        student_profile = self._get_student_profile(target_student_user_id)
        return {
            "studentUserId": target_student_user_id,
            "selectedSubjectCodes": selected_subject_codes,
            "selectedSubjectCode": selected_subject_codes[0] if len(selected_subject_codes) == 1 else "",
            "subjectRows": subject_rows,
            "questionInsights": question_insights,
            "currentSubject": current_subject,
            "studentMeta": {
                "studentUserId": target_student_user_id,
                "studentName": str((self._get_managed_user(target_student_user_id) or {}).get("name", "")).strip() or target_student_user_id,
                "examCategoryCode": str(student_profile.get("examCategoryCode", "")).strip(),
                "jointExamGroupCode": str(student_profile.get("jointExamGroupCode", "")).strip(),
                "vocationalMajor": str(student_profile.get("vocationalMajor", "")).strip(),
                "prepStage": str(student_profile.get("prepStage", "")).strip(),
                "practicePoints": practice_point_summary,
            },
        }

    def _extract_teacher_error_book_analytics_rows(
        self,
        target_actor: Actor,
        question_insights: List[Dict[str, object]],
    ) -> List[Dict[str, object]]:
        question_id_set = {
            str(item.get("questionId", "")).strip()
            for item in question_insights
            if str(item.get("questionId", "")).strip()
        }
        if not question_id_set:
            return []
        wrong_questions, _ = self.list_wrong_book_questions(1, 500, target_actor, {})
        analytics_rows: List[Dict[str, object]] = []
        for question in wrong_questions:
            question_id = str(question.get("id", "")).strip()
            if question_id not in question_id_set:
                continue
            ext_json = self._load_json_object(question.get("extJson", "{}"))
            student_state = ext_json.get("studentState", {}) if isinstance(ext_json.get("studentState", {}), dict) else {}
            chapter_practice = student_state.get("chapterPractice", {}) if isinstance(student_state.get("chapterPractice", {}), dict) else {}
            analytics_rows.append(
                {
                    "questionId": question_id,
                    "isCorrect": bool(chapter_practice.get("isCorrect", False)),
                    "answerDurationSec": float(chapter_practice.get("answerDurationSec", 0)),
                }
            )
        return analytics_rows

    def get_teacher_error_book_summary(self, filters: Dict[str, object], actor: Actor) -> Dict[str, object]:
        return self._build_teacher_error_book_summary_payload(filters, actor)

    def list_teacher_similar_wrong_book_questions(self, question_id: str, filters: Dict[str, object], actor: Actor) -> Dict[str, object]:
        target_student_user_id = str(filters.get("studentUserId") or "").strip()
        target_actor = self._build_error_book_target_actor(actor, target_student_user_id)
        return self.list_similar_wrong_book_questions(question_id, target_actor)

    def export_teacher_error_book_word(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        normalized_payload = dict(payload or {}) if isinstance(payload, dict) else {}
        raw_student_user_ids = normalized_payload.get("studentUserIds")
        normalized_student_user_ids: List[str] = []
        if isinstance(raw_student_user_ids, list):
            for item in raw_student_user_ids:
                student_user_id = str(item or "").strip()
                if student_user_id and student_user_id not in normalized_student_user_ids:
                    normalized_student_user_ids.append(student_user_id)
        target_student_user_id = str(
            normalized_payload.get("studentUserId")
            or (normalized_student_user_ids[0] if normalized_student_user_ids else "")
        ).strip()
        if target_student_user_id and target_student_user_id not in normalized_student_user_ids:
            normalized_student_user_ids.insert(0, target_student_user_id)
        selected_subject_codes = self._normalize_subject_codes_filter(normalized_payload)
        raw_question_ids = normalized_payload.get("questionIds")
        question_ids = [str(item or "").strip() for item in raw_question_ids] if isinstance(raw_question_ids, list) else []
        if not normalized_student_user_ids:
            raise validation_failed("studentUserId 不能为空。")

        if len(normalized_student_user_ids) == 1:
            target_actor = self._build_error_book_target_actor(actor, normalized_student_user_ids[0])
            if not question_ids and selected_subject_codes:
                summary_payload = self._build_teacher_error_book_summary_payload(
                    {
                        "studentUserId": normalized_student_user_ids[0],
                        "subjectCodes": ",".join(selected_subject_codes),
                    },
                    actor,
                )
                question_ids = [
                    str(item.get("questionId", "")).strip()
                    for item in summary_payload.get("questionInsights", [])
                    if str(item.get("questionId", "")).strip()
                ]
            export_payload = {
                "questionIds": [item for item in question_ids if item],
            }
            export_result = self.export_wrong_book_word(export_payload, target_actor)
            export_result["studentUserId"] = normalized_student_user_ids[0]
            export_result["studentUserIds"] = normalized_student_user_ids
            export_result["selectedSubjectCodes"] = selected_subject_codes
            return export_result

        zip_buffer = io.BytesIO()
        export_items: List[Dict[str, object]] = []
        with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zip_file:
            for student_user_id in normalized_student_user_ids:
                target_actor = self._build_error_book_target_actor(actor, student_user_id)
                summary_payload = self._build_teacher_error_book_summary_payload(
                    {
                        "studentUserId": student_user_id,
                        "subjectCodes": ",".join(selected_subject_codes),
                    },
                    actor,
                )
                export_question_ids = [
                    str(item.get("questionId", "")).strip()
                    for item in summary_payload.get("questionInsights", [])
                    if str(item.get("questionId", "")).strip()
                ]
                export_result = self.export_wrong_book_word(
                    {"questionIds": export_question_ids},
                    target_actor,
                )
                file_name = str(export_result.get("fileName", f"{student_user_id}.docx")).strip() or f"{student_user_id}.docx"
                zip_file.writestr(file_name, base64.b64decode(str(export_result.get("contentBase64", ""))))
                export_items.append(
                    {
                        "studentUserId": student_user_id,
                        "fileName": file_name,
                        "questionCount": int(export_result.get("questionCount", 0)),
                    }
                )
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
        return {
            "format": "zip",
            "fileName": f"teacher-error-book-batch-{timestamp}.zip",
            "mediaType": "application/zip",
            "contentBase64": base64.b64encode(zip_buffer.getvalue()).decode("ascii"),
            "studentUserIds": normalized_student_user_ids,
            "selectedSubjectCodes": selected_subject_codes,
            "items": export_items,
        }

    def list_similar_wrong_book_questions(self, question_id: str, actor: Actor) -> Dict[str, object]:
        summary_payload = self._build_student_error_book_summary_payload(actor, {})
        question_map = summary_payload.get("_questionMap", {})
        wrong_book_question_ids = summary_payload.get("_wrongBookQuestionIds", set())
        insight_map = {
            str(item.get("questionId", "")).strip(): item
            for item in summary_payload.get("questionInsights", [])
            if str(item.get("questionId", "")).strip()
        }
        current_question = question_map.get(question_id)
        if not current_question:
            raise not_found("错题记录不存在。")

        current_ext = self._load_json_object(current_question.get("extJson", "{}"))
        subject_code = str(current_ext.get("subjectCode", "")).strip()
        point_code = str(current_ext.get("pointCode", "")).strip()
        if not point_code:
            raise validation_failed("当前错题缺少 point_code，暂无法生成举一反三题单。")

        scope_filters = self._build_student_error_book_scope_filters(actor, {"subjectCode": subject_code})
        candidate_questions = self.repository.list_visible_published_questions(
            {
                "subjectCode": subject_code,
                "subject_code": subject_code,
                "pointCode": point_code,
                "pointCode": point_code,
                "examCategoryCode": str(scope_filters.get("examCategoryCode", "")).strip(),
                "exam_category_code": str(scope_filters.get("examCategoryCode", "")).strip(),
                "jointExamGroupCode": str(scope_filters.get("jointExamGroupCode", "")).strip(),
                "joint_exam_group_code": str(scope_filters.get("jointExamGroupCode", "")).strip(),
                "policyVersionCode": POLICY_VERSION_CODE,
                "policyVersion": POLICY_VERSION_CODE,
                "policy_version": POLICY_VERSION_CODE,
            },
            ROLE_SUPER_ADMIN,
            actor.user_id,
        )
        candidate_questions = self._filter_questions_for_student_scope(candidate_questions, actor.user_id, scope_filters)

        selected_items: List[Dict[str, str]] = []
        selected_ids: set[str] = set()
        for allow_existing_wrong_book in (False, True):
            for question in candidate_questions:
                candidate_id = str(question.get("id", "")).strip()
                if not candidate_id or candidate_id == question_id or candidate_id in selected_ids:
                    continue
                if not allow_existing_wrong_book and candidate_id in wrong_book_question_ids:
                    continue
                selected_items.append(self._public_question(self._enrich_question_for_student(question, actor.user_id)))
                selected_ids.add(candidate_id)
                if len(selected_items) >= 3:
                    break
            if len(selected_items) >= 3:
                break

        return {
            "questionId": question_id,
            "subjectCode": subject_code,
            "pointCode": point_code,
            "pointName": str((insight_map.get(question_id) or {}).get("pointName", "")).strip() or str(current_question.get("knowledgeId", "")).strip(),
            "items": selected_items[:3],
        }

    def export_wrong_book_word(self, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        normalized_payload = dict(payload or {}) if isinstance(payload, dict) else {}
        requested_subject_code = str(normalized_payload.get("subjectCode") or "").strip()
        raw_question_ids = normalized_payload.get("questionIds")
        if raw_question_ids is None:
            raw_question_ids = []
        if not isinstance(raw_question_ids, list):
            raise validation_failed("questionIds 必须为数组。")

        question_ids: List[str] = []
        seen_question_ids: set[str] = set()
        for item in raw_question_ids:
            question_id = str(item or "").strip()
            if not question_id or question_id in seen_question_ids:
                continue
            seen_question_ids.add(question_id)
            question_ids.append(question_id)

        summary_payload = self._build_student_error_book_summary_payload(
            actor,
            {"subjectCode": requested_subject_code},
        )
        question_map = summary_payload.get("_questionMap", {})
        insight_map = {
            str(item.get("questionId", "")).strip(): item
            for item in summary_payload.get("questionInsights", [])
            if str(item.get("questionId", "")).strip()
        }
        current_subject = summary_payload.get("currentSubject", {}) if isinstance(summary_payload.get("currentSubject", {}), dict) else {}
        selected_subject_code = requested_subject_code or str(summary_payload.get("selectedSubjectCode", "")).strip()
        if not question_ids:
            current_subject_insights = [
                item
                for item in summary_payload.get("questionInsights", [])
                if str(item.get("subjectCode", "")).strip() == selected_subject_code
            ]
            question_ids = [
                str(item.get("questionId", "")).strip()
                for item in current_subject_insights
                if str(item.get("questionId", "")).strip()
            ]
        if not question_ids:
            raise validation_failed("当前没有可打印的错题，请先选择题目。")

        export_rows: List[Tuple[Dict[str, str], Dict[str, object]]] = []
        filtered_mastered_count = 0
        for question_id in question_ids:
            question = question_map.get(question_id)
            insight = insight_map.get(question_id)
            if not question or not insight:
                continue
            if selected_subject_code and str(insight.get("subjectCode", "")).strip() != selected_subject_code:
                continue
            if float(insight.get("mastery", 0.0)) >= 0.8:
                filtered_mastered_count += 1
                continue
            export_rows.append((question, insight))
        if not export_rows:
            raise validation_failed("当前选中题目已全部掌握，暂无待攻克题目可导出。")
        if not self.is_docx_available():
            raise failed_dependency("当前环境缺少 python-docx 依赖，无法导出 Word 文档。")

        try:
            from docx import Document as _Document  # type: ignore
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # type: ignore
            from docx.oxml import OxmlElement  # type: ignore
            from docx.oxml.ns import qn  # type: ignore
            from docx.shared import Inches, Pt  # type: ignore
        except Exception as exc:
            raise failed_dependency("当前环境缺少 python-docx 依赖，无法导出 Word 文档。") from exc

        def append_field_run(paragraph: object, field_code: str) -> None:
            run = paragraph.add_run()
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = field_code
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run._r.append(begin)
            run._r.append(instruction)
            run._r.append(end)

        bookmark_counter = [0]

        def append_bookmark(paragraph: object, bookmark_name: str) -> None:
            bookmark_counter[0] += 1
            bookmark_id = str(bookmark_counter[0])
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), bookmark_id)
            start.set(qn("w:name"), bookmark_name)
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), bookmark_id)
            paragraph._p.insert(0, start)
            paragraph._p.append(end)

        def append_internal_hyperlink(paragraph: object, text: str, anchor: str) -> None:
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("w:anchor"), anchor)
            hyperlink.set(qn("w:history"), "1")
            run = OxmlElement("w:r")
            run_properties = OxmlElement("w:rPr")
            run_style = OxmlElement("w:rStyle")
            run_style.set(qn("w:val"), "Hyperlink")
            run_properties.append(run_style)
            run.append(run_properties)
            text_node = OxmlElement("w:t")
            text_node.text = text
            run.append(text_node)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)

        student_profile = self._get_student_profile(actor.user_id)
        actor_user = self.repository.get_user_by_id(actor.user_id)
        actor_user_ext = self._load_json_object((actor_user or {}).get("extJson", "{}"))
        student_name = str(actor_user_ext.get("name", "")).strip() or actor.user_id
        exam_category_code = str(student_profile.get("examCategoryCode", "")).strip()
        exam_category = get_exam_category(exam_category_code)
        exam_category_name = str((exam_category or {}).get("examCategoryName", "")).strip() or exam_category_code or "-"
        joint_group_code = str(student_profile.get("jointExamGroupCode", "")).strip()
        joint_group = get_joint_exam_group(joint_group_code)
        joint_group_name = str((joint_group or {}).get("jointExamGroupName", "")).strip() or joint_group_code or "-"
        vocational_major = str(student_profile.get("vocationalMajor", "")).strip() or "-"
        prep_stage = str(student_profile.get("prepStage", "")).strip() or "-"
        export_time_display = datetime.now(timezone(timedelta(hours=8))).strftime("%Y-%m-%d %H:%M:%S")
        subject_name = str(current_subject.get("subjectName", "")).strip() or selected_subject_code or "错题中心"
        difficulty_label_map = {
            "easy": "基础",
            "medium": "提升",
            "hard": "冲刺",
        }
        chapter_index_map: Dict[str, Dict[str, object]] = {}
        question_bookmark_map: Dict[int, str] = {}
        appendix_chapter_groups: List[Dict[str, object]] = []
        reason_counter: Dict[str, int] = {}
        reason_labels_by_question: Dict[int, List[str]] = {}
        reason_details_by_question: Dict[int, List[Tuple[str, int]]] = {}
        chapter_reason_counter: Dict[str, Dict[str, int]] = {}
        for _, insight in export_rows:
            chapter_key = str(insight.get("chapterCode", "")).strip() or str(insight.get("chapterName", "")).strip() or "未标注章节"
            chapter_row = chapter_index_map.setdefault(
                chapter_key,
                {
                    "chapterCode": str(insight.get("chapterCode", "")).strip(),
                    "chapterName": str(insight.get("chapterName", "")).strip() or chapter_key,
                    "questionCount": 0,
                },
            )
            chapter_row["questionCount"] = int(chapter_row.get("questionCount", 0)) + 1
        appendix_group_map: Dict[str, Dict[str, object]] = {}
        chapter_summary_map: Dict[str, Dict[str, object]] = {}
        for index, row in enumerate(export_rows, start=1):
            question_bookmark_map[index] = f"qb_q_{index}"
            question, insight = row
            chapter_key = str(insight.get("chapterCode", "")).strip() or str(insight.get("chapterName", "")).strip() or "未标注章节"
            question_ext_json = self._load_json_object(question.get("extJson", "{}"))
            question_student_state = question_ext_json.get("studentState", {}) if isinstance(question_ext_json.get("studentState", {}), dict) else {}
            question_wrong_book = question_student_state.get("wrongBook", {}) if isinstance(question_student_state.get("wrongBook", {}), dict) else {}
            question_reason_stats = question_wrong_book.get("reasonStats", []) if isinstance(question_wrong_book.get("reasonStats", []), list) else []
            question_reason_labels: List[str] = []
            question_reason_rows: List[Tuple[str, int]] = []
            for item in question_reason_stats:
                if not isinstance(item, dict):
                    continue
                reason_label = str(item.get("reasonLabel", "")).strip()
                reason_count = max(1, int(item.get("count", 0)))
                if not reason_label:
                    continue
                question_reason_labels.append(reason_label)
                question_reason_rows.append((reason_label, reason_count))
                reason_counter[reason_label] = reason_counter.get(reason_label, 0) + reason_count
                chapter_reason_map = chapter_reason_counter.setdefault(chapter_key, {})
                chapter_reason_map[reason_label] = chapter_reason_map.get(reason_label, 0) + reason_count
            question_reason_rows.sort(key=lambda item: (-int(item[1]), item[0]))
            reason_labels_by_question[index] = [item[0] for item in question_reason_rows]
            reason_details_by_question[index] = question_reason_rows
            group = appendix_group_map.setdefault(
                chapter_key,
                {
                    "chapterCode": str(insight.get("chapterCode", "")).strip(),
                    "chapterName": str(insight.get("chapterName", "")).strip() or chapter_key,
                    "rows": [],
                },
            )
            group["rows"].append((index, question, insight))
            chapter_summary = chapter_summary_map.setdefault(
                chapter_key,
                {
                    "lowestMastery": 1.0,
                    "focusPoint": "",
                    "highestWrongCount": 0,
                    "masteryTotal": 0.0,
                    "questionCount": 0,
                    "reviewCandidates": [],
                },
            )
            mastery = float(insight.get("mastery", 0.0))
            wrong_count = int(insight.get("wrongCount", 0))
            chapter_summary["masteryTotal"] = float(chapter_summary.get("masteryTotal", 0.0)) + mastery
            chapter_summary["questionCount"] = int(chapter_summary.get("questionCount", 0)) + 1
            review_candidates = chapter_summary.get("reviewCandidates", [])
            if isinstance(review_candidates, list):
                review_candidates.append(
                    {
                        "pointName": str(insight.get("pointName", "")).strip() or str(question.get("knowledgeId", "")).strip(),
                        "mastery": mastery,
                        "wrongCount": wrong_count,
                    }
                )
            if mastery <= float(chapter_summary.get("lowestMastery", 1.0)):
                chapter_summary["lowestMastery"] = mastery
                chapter_summary["focusPoint"] = str(insight.get("pointName", "")).strip() or str(question.get("knowledgeId", "")).strip()
            chapter_summary["highestWrongCount"] = max(int(chapter_summary.get("highestWrongCount", 0)), wrong_count)
        appendix_chapter_groups = sorted(
            appendix_group_map.values(),
            key=lambda item: (str(item.get("chapterCode", "")), str(item.get("chapterName", ""))),
        )
        chapter_question_count_map = {
            str(item.get("chapterCode", "")).strip() or str(item.get("chapterName", "")).strip() or "未标注章节": int(item.get("questionCount", 0))
            for item in chapter_index_map.values()
        }
        sorted_reason_rows = sorted(reason_counter.items(), key=lambda item: (-int(item[1]), item[0]))
        total_reason_count = sum(int(count) for _, count in sorted_reason_rows)
        average_mastery_all = (
            sum(float(item.get("mastery", 0.0)) for _, item in export_rows) / len(export_rows)
            if export_rows
            else 0.0
        )
        if average_mastery_all < 0.35:
            risk_level_label = "高风险待突破"
        elif average_mastery_all < 0.6:
            risk_level_label = "中风险需强化"
        else:
            risk_level_label = "低风险稳步复盘"
        if average_mastery_all < 0.35:
            overall_summary = "当前学情处于高风险待突破区，建议先按章节完成整页复盘，再集中攻克主错因。"
        elif average_mastery_all < 0.6:
            overall_summary = "当前学情处于强化提升区，建议优先处理主错因高频章节，逐步抬升掌握度。"
        else:
            overall_summary = "当前学情总体可控，建议保持章节复盘节奏，重点清理残留高频错因。"
        profile_summary_line = (
            f"{exam_category_name} / {vocational_major} / {prep_stage}，"
            f"当前整体掌握度约 {int(round(average_mastery_all * 100))}% ，"
            f"建议采用“先章节突破，再错因复盘”的修复节奏。"
        )
        document = _Document()
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        self._configure_docx_export_theme(document, Pt, qn, OxmlElement)

        header_paragraph = section.header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_paragraph.add_run(f"{subject_name} 智能学情修复中心")
        header_run.font.size = Pt(10)
        self._apply_docx_run_font(header_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        footer_paragraph = section.footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_prefix = footer_paragraph.add_run("第 ")
        footer_prefix.font.size = Pt(10)
        self._apply_docx_run_font(footer_prefix, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)
        append_field_run(footer_paragraph, "PAGE")
        footer_suffix = footer_paragraph.add_run(" 页")
        footer_suffix.font.size = Pt(10)
        self._apply_docx_run_font(footer_suffix, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_eyebrow = document.add_paragraph()
        cover_eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_eyebrow_run = cover_eyebrow.add_run("SMART REPAIR WORKBOOK")
        cover_eyebrow_run.bold = True
        cover_eyebrow_run.font.size = Pt(11)
        self._apply_docx_run_font(cover_eyebrow_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_title = document.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title_run = cover_title.add_run(f"{subject_name} 个人线下提分卷")
        cover_title_run.bold = True
        cover_title_run.font.size = Pt(24)
        self._apply_docx_run_font(cover_title_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_subtitle = document.add_paragraph()
        cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_subtitle_run = cover_subtitle.add_run("错题修复中心线下背诵版")
        cover_subtitle_run.font.size = Pt(13)
        self._apply_docx_run_font(cover_subtitle_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_rule = document.add_paragraph()
        cover_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_rule_run = cover_rule.add_run("".join(["=", "=", "=", "=", "=", "=", "=", "=", "=", "="]))
        cover_rule_run.font.size = Pt(12)
        self._apply_docx_run_font(cover_rule_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_summary_table = document.add_table(rows=1, cols=3)
        cover_summary_table.style = "Table Grid"
        cover_summary_table.cell(0, 0).text = f"待攻克题数\n{len(export_rows)} 题"
        cover_summary_table.cell(0, 1).text = f"章节数量\n{len(chapter_index_map)} 章"
        cover_summary_table.cell(0, 2).text = "本卷结构\n目录 / 正文 / 附录"
        self._format_docx_table(
            cover_summary_table,
            Pt,
            qn,
            OxmlElement,
            WD_ALIGN_PARAGRAPH,
            WD_TABLE_ALIGNMENT,
            WD_CELL_VERTICAL_ALIGNMENT,
            center=True,
        )

        cover_risk_table = document.add_table(rows=1, cols=3)
        cover_risk_table.style = "Table Grid"
        cover_risk_table.cell(0, 0).text = f"整体风险分层\n{risk_level_label}"
        cover_risk_table.cell(0, 1).text = f"平均掌握度\n{int(round(average_mastery_all * 100))}%"
        cover_risk_table.cell(0, 2).text = f"高频错因数\n{len(sorted_reason_rows)} 类"
        self._format_docx_table(
            cover_risk_table,
            Pt,
            qn,
            OxmlElement,
            WD_ALIGN_PARAGRAPH,
            WD_TABLE_ALIGNMENT,
            WD_CELL_VERTICAL_ALIGNMENT,
            center=True,
        )

        cover_table = document.add_table(rows=4, cols=2)
        cover_table.style = "Table Grid"
        cover_table.cell(0, 0).text = "学生姓名"
        cover_table.cell(0, 1).text = student_name
        cover_table.cell(1, 0).text = "专业组"
        cover_table.cell(1, 1).text = f"{joint_group_name}（{joint_group_code or '-'}）"
        cover_table.cell(2, 0).text = "备考阶段"
        cover_table.cell(2, 1).text = prep_stage
        cover_table.cell(3, 0).text = "导出时间"
        cover_table.cell(3, 1).text = export_time_display
        self._format_docx_table(
            cover_table,
            Pt,
            qn,
            OxmlElement,
            WD_ALIGN_PARAGRAPH,
            WD_TABLE_ALIGNMENT,
            WD_CELL_VERTICAL_ALIGNMENT,
            emphasize_first_column=True,
        )

        cover_profile_table = document.add_table(rows=1, cols=3)
        cover_profile_table.style = "Table Grid"
        cover_profile_table.cell(0, 0).text = f"学科门类\n{exam_category_name}"
        cover_profile_table.cell(0, 1).text = f"专业方向\n{vocational_major}"
        cover_profile_table.cell(0, 2).text = f"学习画像\n{prep_stage}"
        self._format_docx_table(
            cover_profile_table,
            Pt,
            qn,
            OxmlElement,
            WD_ALIGN_PARAGRAPH,
            WD_TABLE_ALIGNMENT,
            WD_CELL_VERTICAL_ALIGNMENT,
            center=True,
        )

        cover_profile_summary = document.add_paragraph()
        cover_profile_summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_profile_summary_run = cover_profile_summary.add_run(f"专业画像摘要：{profile_summary_line}")
        cover_profile_summary_run.font.size = Pt(10)
        self._apply_docx_run_font(cover_profile_summary_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_summary_sentence = document.add_paragraph()
        cover_summary_sentence.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_summary_sentence_run = cover_summary_sentence.add_run(f"学情一句话结论：{overall_summary}")
        cover_summary_sentence_run.bold = True
        cover_summary_sentence_run.font.size = Pt(10)
        self._apply_docx_run_font(cover_summary_sentence_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_teacher_note = document.add_paragraph()
        cover_teacher_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_teacher_note_run = cover_teacher_note.add_run(
            "老师寄语：先稳住主错因，再逐章清空错题，今天的每一次复盘都会在考场上回报你。"
        )
        cover_teacher_note_run.italic = True
        cover_teacher_note_run.font.size = Pt(10)
        self._apply_docx_run_font(cover_teacher_note_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_meta_rows = [
            f"导出学生：{actor.user_id}",
            f"科目：{subject_name}",
            f"策略版本：{POLICY_VERSION_CODE}",
            f"待攻克题数：{len(export_rows)} 题",
            f"已过滤已掌握题：{filtered_mastered_count} 题",
        ]
        for row in cover_meta_rows:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(row)
            run.font.size = Pt(12)
            self._apply_docx_run_font(run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        document.add_paragraph("")
        cover_note = document.add_paragraph()
        cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_note_run = cover_note.add_run("使用建议：先遮住答案与解析，完成默写，再对照订正。")
        cover_note_run.italic = True
        cover_note_run.font.size = Pt(11)
        self._apply_docx_run_font(cover_note_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_structure = document.add_paragraph()
        cover_structure.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_structure_run = cover_structure.add_run("本卷结构：目录导览 -> 题目正文 -> 答案附录")
        cover_structure_run.bold = True
        cover_structure_run.font.size = Pt(10)
        self._apply_docx_run_font(cover_structure_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_badge = document.add_paragraph()
        cover_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_badge_run = cover_badge.add_run("学习建议：按章节完成，再回到附录集中核对")
        cover_badge_run.italic = True
        cover_badge_run.font.size = Pt(10)
        self._apply_docx_run_font(cover_badge_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

        cover_break = document.add_paragraph()
        cover_break.add_run().add_break(WD_BREAK.PAGE)

        directory_heading = document.add_heading("目录导览", level=1)
        directory_heading.paragraph_format.space_after = Pt(10)
        directory_note = document.add_paragraph("建议先按题号完成整卷，再对照后附答案附录统一复盘。")
        directory_note.paragraph_format.space_after = Pt(8)

        for index, row in enumerate(export_rows, start=1):
            question, insight = row
            difficulty = difficulty_label_map.get(
                str(self._question_difficulty(question)).strip().lower(),
                "综合",
            )
            directory_item = document.add_paragraph()
            directory_item.add_run(f"{index}. ").bold = True
            directory_item.add_run(f"[{difficulty}] ")
            directory_item.add_run(str(insight.get("pointName", "")).strip() or str(question.get("knowledgeId", "")).strip())
            directory_item.add_run(
                f"  ({str(insight.get('chapterName', '')).strip() or str(insight.get('chapterCode', '')).strip() or '未标注章节'})"
            )
            directory_item.add_run("  ·  跳转 ")
            append_internal_hyperlink(directory_item, "题目页", question_bookmark_map.get(index, f"qb_q_{index}"))
            directory_item.add_run("  ·  第 ")
            append_field_run(directory_item, f"PAGEREF {question_bookmark_map.get(index, '')} \\h")
            directory_item.add_run(" 页")

        document.add_paragraph("")
        chapter_index_heading = document.add_heading("章节索引", level=2)
        chapter_index_heading.paragraph_format.space_after = Pt(8)
        for chapter_row in sorted(
            chapter_index_map.values(),
            key=lambda item: (str(item.get("chapterCode", "")), str(item.get("chapterName", ""))),
        ):
            chapter_paragraph = document.add_paragraph(style="List Bullet")
            chapter_paragraph.add_run(str(chapter_row.get("chapterName", "")).strip() or "未标注章节")
            chapter_paragraph.add_run(f"  ·  {int(chapter_row.get('questionCount', 0))} 题")

        document.add_paragraph("")
        reason_summary_heading = document.add_heading("错因统计汇总", level=2)
        reason_summary_heading.paragraph_format.space_after = Pt(8)
        if sorted_reason_rows:
            reason_summary_table = document.add_table(rows=1, cols=4)
            reason_summary_table.style = "Table Grid"
            reason_summary_table.cell(0, 0).text = "错因标签"
            reason_summary_table.cell(0, 1).text = "累计次数"
            reason_summary_table.cell(0, 2).text = "占比"
            reason_summary_table.cell(0, 3).text = "等级提示"
            for reason_label, count in sorted_reason_rows:
                row_cells = reason_summary_table.add_row().cells
                row_cells[0].text = reason_label
                row_cells[1].text = str(count)
                ratio = (int(count) / total_reason_count) if total_reason_count > 0 else 0.0
                row_cells[2].text = f"{ratio:.0%}"
                if ratio >= 0.5:
                    level = "重灾区"
                elif ratio >= 0.25:
                    level = "重点关注"
                else:
                    level = "常规复盘"
                row_cells[3].text = level
            self._format_docx_table(
                reason_summary_table,
                Pt,
                qn,
                OxmlElement,
                WD_ALIGN_PARAGRAPH,
                WD_TABLE_ALIGNMENT,
                WD_CELL_VERTICAL_ALIGNMENT,
                header_rows=1,
            )
        else:
            reason_summary_empty = document.add_paragraph("当前导出题单暂无错因统计。")
            reason_summary_empty.paragraph_format.space_after = Pt(8)

        document.add_paragraph("")
        chapter_reason_heading = document.add_heading("章节错因分布", level=2)
        chapter_reason_heading.paragraph_format.space_after = Pt(8)
        if chapter_reason_counter:
            chapter_reason_table = document.add_table(rows=1, cols=4)
            chapter_reason_table.style = "Table Grid"
            chapter_reason_table.cell(0, 0).text = "章节"
            chapter_reason_table.cell(0, 1).text = "错因分布"
            chapter_reason_table.cell(0, 2).text = "章节主错因"
            chapter_reason_table.cell(0, 3).text = "章节提示"
            for chapter_key, reason_map in sorted(chapter_reason_counter.items(), key=lambda item: item[0]):
                chapter_row_cells = chapter_reason_table.add_row().cells
                chapter_label = chapter_index_map.get(chapter_key, {}).get("chapterName", chapter_key)
                chapter_total = sum(int(count) for count in reason_map.values())
                ordered_reasons = sorted(reason_map.items(), key=lambda item: (-int(item[1]), item[0]))
                chapter_row_cells[0].text = str(chapter_label or chapter_key)
                chapter_row_cells[1].text = " / ".join(
                    [
                        f"{reason_label} {count}次 {((int(count) / chapter_total) if chapter_total > 0 else 0.0):.0%}"
                        for reason_label, count in ordered_reasons
                    ]
                )
                chapter_row_cells[2].text = ordered_reasons[0][0] if ordered_reasons else "-"
                dominant_ratio = (int(ordered_reasons[0][1]) / chapter_total) if ordered_reasons and chapter_total > 0 else 0.0
                if dominant_ratio >= 0.5:
                    chapter_row_cells[3].text = "先攻主错因"
                elif dominant_ratio >= 0.25:
                    chapter_row_cells[3].text = "优先复盘主错因"
                else:
                    chapter_row_cells[3].text = "均衡复盘"
            self._format_docx_table(
                chapter_reason_table,
                Pt,
                qn,
                OxmlElement,
                WD_ALIGN_PARAGRAPH,
                WD_TABLE_ALIGNMENT,
                WD_CELL_VERTICAL_ALIGNMENT,
                header_rows=1,
            )
        else:
            chapter_reason_empty = document.add_paragraph("当前导出题单暂无章节级错因分布。")
            chapter_reason_empty.paragraph_format.space_after = Pt(8)

        directory_break = document.add_paragraph()
        directory_break.add_run().add_break(WD_BREAK.PAGE)

        current_chapter_key = ""
        for index, row in enumerate(export_rows, start=1):
            question, insight = row
            ext_json = self._load_json_object(question.get("extJson", "{}"))
            difficulty = difficulty_label_map.get(
                str(self._question_difficulty(question)).strip().lower(),
                "综合",
            )
            chapter_key = str(insight.get("chapterCode", "")).strip() or str(insight.get("chapterName", "")).strip() or "未标注章节"
            if current_chapter_key != chapter_key:
                if index > 1:
                    page_break = document.add_paragraph()
                    page_break.add_run().add_break(WD_BREAK.PAGE)
                chapter_divider = document.add_paragraph()
                chapter_divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_divider_run = chapter_divider.add_run("CHAPTER")
                chapter_divider_run.bold = True
                chapter_divider_run.font.size = Pt(12)
                self._apply_docx_run_font(chapter_divider_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

                chapter_title = document.add_paragraph()
                chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_title_run = chapter_title.add_run(str(insight.get("chapterName", "")).strip() or str(insight.get("chapterCode", "")).strip() or "未标注章节")
                chapter_title_run.bold = True
                chapter_title_run.font.size = Pt(20)
                self._apply_docx_run_font(chapter_title_run, "Microsoft YaHei", "Microsoft YaHei", qn, OxmlElement)

                chapter_meta = document.add_paragraph()
                chapter_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_meta_run = chapter_meta.add_run(
                    f"本章共 {chapter_question_count_map.get(chapter_key, 0)} 题待攻克"
                )
                chapter_meta_run.font.size = Pt(11)

                chapter_focus = document.add_paragraph()
                chapter_focus.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_focus_summary = chapter_summary_map.get(chapter_key, {})
                chapter_focus_run = chapter_focus.add_run(
                    f"本章重点：{str(chapter_focus_summary.get('focusPoint', '')).strip() or '优先处理高频错题'}"
                )
                chapter_focus_run.bold = True
                chapter_focus_run.font.size = Pt(11)

                chapter_hint = document.add_paragraph()
                chapter_hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
                average_mastery = 0.0
                if int(chapter_focus_summary.get("questionCount", 0)) > 0:
                    average_mastery = float(chapter_focus_summary.get("masteryTotal", 0.0)) / int(chapter_focus_summary.get("questionCount", 0))
                chapter_hint_run = chapter_hint.add_run(
                    f"复习建议：当前章平均掌握度约 {int(round(average_mastery * 100))}% ，建议先完成本章整页练习，再翻到附录统一核对答案。"
                )
                chapter_hint_run.italic = True
                chapter_hint_run.font.size = Pt(10)

                chapter_review_candidates = chapter_focus_summary.get("reviewCandidates", [])
                chapter_review_order_rows = (
                    sorted(
                        chapter_review_candidates,
                        key=lambda item: (float(item.get("mastery", 0.0)), -int(item.get("wrongCount", 0)), str(item.get("pointName", ""))),
                    )[:3]
                    if isinstance(chapter_review_candidates, list)
                    else []
                )
                chapter_review_order = document.add_paragraph()
                chapter_review_order.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_review_order_run = chapter_review_order.add_run(
                    "本章推荐复习顺序："
                    + (
                        " -> ".join(
                            [
                                str(item.get("pointName", "")).strip() or "未命名考点"
                                for item in chapter_review_order_rows
                            ]
                        )
                        if chapter_review_order_rows
                        else "暂无"
                    )
                )
                chapter_review_order_run.font.size = Pt(10)

                target_mastery = min(95, max(int(round(average_mastery * 100)) + 15, 60))
                chapter_goal = document.add_paragraph()
                chapter_goal.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_goal_run = chapter_goal.add_run(
                    f"本章完成目标：完成 {chapter_question_count_map.get(chapter_key, 0)} 题复盘，并将本章平均掌握度提升到 {target_mastery}% 以上。"
                )
                chapter_goal_run.bold = True
                chapter_goal_run.font.size = Pt(10)

                chapter_pass_standard = document.add_paragraph()
                chapter_pass_standard.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_pass_standard_run = chapter_pass_standard.add_run(
                    f"本章自测通过标准：复盘后正确率达到 80% 以上，且主错因对应题目至少完成 2 轮回看。"
                )
                chapter_pass_standard_run.bold = True
                chapter_pass_standard_run.font.size = Pt(10)

                estimated_minutes = max(15, int(chapter_question_count_map.get(chapter_key, 0)) * 8)
                chapter_eta = document.add_paragraph()
                chapter_eta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_eta_run = chapter_eta.add_run(
                    f"预计完成时长：约 {estimated_minutes} 分钟，可拆成 2 轮完成。"
                )
                chapter_eta_run.font.size = Pt(10)

                chapter_reason_top_rows = sorted(
                    chapter_reason_counter.get(chapter_key, {}).items(),
                    key=lambda item: (-int(item[1]), item[0]),
                )[:3]
                chapter_reason_top = document.add_paragraph()
                chapter_reason_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_reason_top_run = chapter_reason_top.add_run(
                    "本章错因 TOP3："
                    + (
                        " / ".join([f"{reason_label}({count})" for reason_label, count in chapter_reason_top_rows])
                        if chapter_reason_top_rows
                        else "暂无"
                    )
                )
                chapter_reason_top_run.font.size = Pt(10)

                chapter_summary_table = document.add_table(rows=1, cols=3)
                chapter_summary_table.style = "Table Grid"
                chapter_summary_table.cell(0, 0).text = f"本章小结\n{chapter_question_count_map.get(chapter_key, 0)} 题"
                chapter_summary_table.cell(0, 1).text = f"平均掌握度\n{int(round(average_mastery * 100))}%"
                chapter_summary_table.cell(0, 2).text = f"最高错题次数\n{int(chapter_focus_summary.get('highestWrongCount', 0))} 次"
                self._format_docx_table(
                    chapter_summary_table,
                    Pt,
                    qn,
                    OxmlElement,
                    WD_ALIGN_PARAGRAPH,
                    WD_TABLE_ALIGNMENT,
                    WD_CELL_VERTICAL_ALIGNMENT,
                    center=True,
                )
                current_chapter_key = chapter_key
            else:
                page_break = document.add_paragraph()
                page_break.add_run().add_break(WD_BREAK.PAGE)

            question_heading = document.add_heading(
                f"{index}. [{difficulty}] {str(insight.get('pointName', '')).strip() or question.get('knowledgeId', '')}",
                level=2,
            )
            question_heading.paragraph_format.space_after = Pt(8)
            append_bookmark(question_heading, question_bookmark_map.get(index, f"qb_q_{index}"))

            path_paragraph = document.add_paragraph(
                " / ".join(
                    [
                        str(insight.get("subjectName", "")).strip() or str(insight.get("subjectCode", "")).strip() or "-",
                        str(insight.get("chapterName", "")).strip() or str(insight.get("chapterCode", "")).strip() or "未标注章节",
                        str(insight.get("pointCode", "")).strip() or str(question.get("knowledgeId", "")).strip() or "-",
                    ]
                )
            )
            path_paragraph.paragraph_format.space_after = Pt(6)

            metric_paragraph = document.add_paragraph(
                f"掌握度：{int(insight.get('masteryScore', 0))}%    错题次数：{int(insight.get('wrongCount', 0))}    难度角标：{difficulty}"
            )
            metric_paragraph.paragraph_format.space_after = Pt(6)

            body_reason_labels = reason_labels_by_question.get(index, [])
            body_reason_details = reason_details_by_question.get(index, [])
            reason_badge_table = document.add_table(rows=1, cols=3)
            reason_badge_table.style = "Table Grid"
            primary_reason = body_reason_details[0][0] if body_reason_details else "待补充"
            reason_badge_count = int(body_reason_details[0][1]) if body_reason_details else 0
            secondary_reasons = " / ".join(
                [
                    f"{reason_label}({count})"
                    for reason_label, count in body_reason_details[1:]
                ]
            ) if len(body_reason_details) > 1 else "无"
            reason_badge_table.cell(0, 0).text = f"主错因\n{primary_reason}"
            reason_badge_table.cell(0, 1).text = f"出现次数\n{reason_badge_count} 次"
            reason_badge_table.cell(0, 2).text = f"次级错因\n{secondary_reasons}"
            self._format_docx_table(
                reason_badge_table,
                Pt,
                qn,
                OxmlElement,
                WD_ALIGN_PARAGRAPH,
                WD_TABLE_ALIGNMENT,
                WD_CELL_VERTICAL_ALIGNMENT,
                center=True,
            )

            stem_title = document.add_paragraph()
            stem_title.add_run("题干").bold = True
            stem_body = document.add_paragraph(str(question.get('stem', '')).strip())
            stem_body.paragraph_format.space_after = Pt(10)

            options = question.get("optionsJson", "[]")
            try:
                parsed_options = json.loads(options) if isinstance(options, str) else options
            except (TypeError, ValueError):
                parsed_options = []
            if isinstance(parsed_options, list) and parsed_options:
                option_title = document.add_paragraph()
                option_title.add_run("选项").bold = True
                for option in parsed_options:
                    if not isinstance(option, dict):
                        continue
                    option_key = str(option.get("key", "")).strip()
                    option_content = str(option.get("content", "")).strip()
                    document.add_paragraph(f"{option_key}. {option_content}", style="List Bullet")

            appendix_hint = document.add_paragraph()
            appendix_hint.add_run("答案与解析").bold = True
            appendix_hint.add_run("：见文末附录，先独立回忆再对照订正。")
            appendix_hint.paragraph_format.space_after = Pt(10)

            memory_title = document.add_paragraph()
            memory_title.add_run("背诵默写区").bold = True
            memory_hint = document.add_paragraph("可先合上答案页，按知识点口述、默写定义和关键结论。")
            memory_hint.paragraph_format.space_after = Pt(6)
            for _ in range(8):
                blank_line = document.add_paragraph("____________________________________________________________")
                blank_line.paragraph_format.space_after = Pt(8)

            review_prompt = document.add_paragraph()
            review_prompt.add_run("复盘提醒：").bold = True
            review_prompt.add_run("写下本题再次出错的原因，并补一个最容易混淆的反例。")

        appendix_break = document.add_paragraph()
        appendix_break.add_run().add_break(WD_BREAK.PAGE)

        appendix_heading = document.add_heading("答案附录", level=1)
        appendix_heading.paragraph_format.space_after = Pt(10)
        appendix_note = document.add_paragraph("建议做完整份题单后，再统一翻到附录核对答案与解析。")
        appendix_note.paragraph_format.space_after = Pt(10)

        for group_index, group in enumerate(appendix_chapter_groups, start=1):
            if group_index > 1:
                appendix_group_break = document.add_paragraph()
                appendix_group_break.add_run().add_break(WD_BREAK.PAGE)

            chapter_heading = document.add_heading(
                f"附录章节：{str(group.get('chapterName', '')).strip() or str(group.get('chapterCode', '')).strip() or '未标注章节'}",
                level=2,
            )
            chapter_heading.paragraph_format.space_after = Pt(8)

            for row_index, question, insight in group.get("rows", []):
                ext_json = self._load_json_object(question.get("extJson", "{}"))
                student_state = ext_json.get("studentState", {}) if isinstance(ext_json.get("studentState", {}), dict) else {}
                wrong_book = student_state.get("wrongBook", {}) if isinstance(student_state.get("wrongBook", {}), dict) else {}
                reason_stats = wrong_book.get("reasonStats", []) if isinstance(wrong_book.get("reasonStats", []), list) else []
                reason_labels = [
                    str(item.get("reasonLabel", "")).strip()
                    for item in reason_stats
                    if isinstance(item, dict) and str(item.get("reasonLabel", "")).strip()
                ]
                appendix_item_heading = document.add_heading(
                    f"附录 {row_index}. {str(insight.get('pointName', '')).strip() or question.get('knowledgeId', '')}",
                    level=3,
                )
                appendix_item_heading.paragraph_format.space_after = Pt(8)

                appendix_metric = document.add_paragraph(
                    f"章节：{str(insight.get('chapterName', '')).strip() or str(insight.get('chapterCode', '')).strip() or '未标注章节'}"
                )
                appendix_metric.paragraph_format.space_after = Pt(4)

                appendix_tagline = document.add_paragraph(
                    f"标签：掌握度 {int(insight.get('masteryScore', 0))}%  /  错题次数 {int(insight.get('wrongCount', 0))}  /  原题难度 {difficulty_label_map.get(str(self._question_difficulty(question)).strip().lower(), '综合')}"
                )
                appendix_tagline.paragraph_format.space_after = Pt(6)

                appendix_reason_line = document.add_paragraph(
                    f"错因标签：{' / '.join(reason_labels) if reason_labels else '待补充'}"
                )
                appendix_reason_line.paragraph_format.space_after = Pt(6)

                answer_title = document.add_paragraph()
                answer_title.add_run("答案").bold = True
                answer_body = document.add_paragraph(str(question.get('answer', '')).strip() or '-')
                answer_body.paragraph_format.space_after = Pt(8)

                analysis_title = document.add_paragraph()
                analysis_title.add_run("解析").bold = True
                analysis_body = document.add_paragraph(str(ext_json.get('analysis', '')).strip() or '-')
                analysis_body.paragraph_format.space_after = Pt(10)

                appendix_backlink = document.add_paragraph()
                appendix_backlink.add_run("返回原题页：").bold = True
                append_internal_hyperlink(appendix_backlink, "回到正文", question_bookmark_map.get(row_index, f"qb_q_{row_index}"))
                appendix_backlink.add_run("  ·  原题第 ")
                append_field_run(appendix_backlink, f"PAGEREF {question_bookmark_map.get(row_index, '')} \\h")
                appendix_backlink.add_run(" 页")

                appendix_prompt = document.add_paragraph()
                appendix_prompt.add_run("对照复盘：").bold = True
                appendix_prompt.add_run("标出自己漏掉的关键词，并写下下一次遇到同类题的判断锚点。")

        buffer = io.BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
        return {
            "format": "docx",
            "fileName": f"error-book-{selected_subject_code or actor.user_id}-{timestamp}.docx",
            "mediaType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "contentBase64": base64.b64encode(content).decode("ascii"),
            "questionCount": len(export_rows),
            "filteredMasteredCount": filtered_mastered_count,
        }

    def collect_wrong_book_question(self, question_id: str, actor: Actor) -> Dict[str, str]:
        question = self.repository.get_question(question_id)
        if not question:
            raise not_found("题目不存在。")
        record, record_ext = self._load_student_question_record_bundle(question_id, actor.user_id, create_if_missing=True)
        wrong_book = dict(record_ext.get("wrongBook", {}))
        wrong_book["isCollected"] = True
        wrong_book["collectedAt"] = str(wrong_book.get("collectedAt", "")) or self._now_iso()
        wrong_book["wrongCount"] = int(wrong_book.get("wrongCount", 0))
        wrong_book["reasonStats"] = self._upsert_wrong_reason_stats(
            wrong_book.get("reasonStats"),
            str(wrong_book.get("lastReasonCode", "MANUAL_COLLECT")) or "MANUAL_COLLECT",
            str(wrong_book.get("lastReasonLabel", "手动加入错题本")) or "手动加入错题本",
        )
        record_ext["wrongBook"] = wrong_book
        self._save_student_question_record_bundle(record, record_ext, actor.user_id)
        return self._public_question(self._enrich_question_for_student(question, actor.user_id))

    def toggle_personal_bank_question(self, question_id: str, payload: Dict[str, object], actor: Actor) -> Dict[str, str]:
        question = self.repository.get_question(question_id)
        if not question or question["status"] != "PUBLISHED":
            raise not_found("题目不存在或尚未发布。")
        if not self._is_question_visible_to_student(question, actor.user_id):
            raise forbidden("当前学生不可收藏该题目。")
        record, record_ext = self._load_student_question_record_bundle(question_id, actor.user_id, create_if_missing=True)
        is_collected = bool(payload.get("isCollected", True))
        record_ext["personalBank"] = {
            "isCollected": is_collected,
            "collectedAt": self._now_iso() if is_collected else "",
        }
        self._save_student_question_record_bundle(record, record_ext, actor.user_id)
        return self._public_question(self._enrich_question_for_student(question, actor.user_id))

    def review_wrong_book_question(self, question_id: str, actor: Actor) -> Dict[str, object]:
        record = self.repository.get_student_question_bank(question_id, actor.user_id)
        if not record:
            raise not_found("错题记录不存在。")
        record_ext = self._load_json_object(record["extJson"])
        wrong_book = record_ext.get("wrongBook", {})
        if not wrong_book.get("isCollected"):
            raise validation_failed("当前题目尚未加入错题本。")
        wrong_book["reviewedAt"] = self._now_iso()
        wrong_book["reviewCount"] = int(wrong_book.get("reviewCount", 0)) + 1
        record_ext["wrongBook"] = wrong_book
        self._save_student_question_record_bundle(record, record_ext, actor.user_id)
        self._append_student_daily_metric(actor.user_id, "wrongBookReviewed", 1)
        self._grant_daily_task_points(actor.user_id)
        return {"questionId": question_id, "reviewCount": wrong_book["reviewCount"]}

    def generate_personalized_wrong_book_paper(
        self,
        actor: Actor,
        filters: Optional[Dict[str, str]] = None,
    ) -> Dict[str, object]:
        normalized_filters = {key: str(value or "").strip() for key, value in dict(filters or {}).items()}
        wrong_book_questions, _ = self.list_wrong_book_questions(1, 500, actor, normalized_filters)
        if not wrong_book_questions:
            raise validation_failed("当前错题量不足/薄弱知识点不明确，无法生成专属同类卷，请多刷题积累错题。")
        wrong_book_question_ids = {
            str(item.get("id", "")).strip()
            for item in wrong_book_questions
            if str(item.get("id", "")).strip()
        }
        target_knowledge_id = str(normalized_filters.get("knowledgeId") or "").strip()
        target_path_node_id = str(normalized_filters.get("knowledgePathNodeId") or "").strip()
        target_chapter = str(normalized_filters.get("chapter", "")).strip()
        target_chapter_code = str(normalized_filters.get("chapterCode") or "").strip()
        target_point_code = str(normalized_filters.get("pointCode") or "").strip()
        target_subject_code = str(normalized_filters.get("subjectCode") or "").strip()
        descendant_knowledge_ids = self._resolve_wrong_book_path_node_descendants(
            target_path_node_id,
            subject_code=target_subject_code,
        )
        knowledge_tag_pool: List[str] = []
        for question in wrong_book_questions:
            question_ext = self._load_json_object(question["extJson"])
            knowledge_tag_pool.extend(question_ext.get("knowledgeTags", []))
        question_scope_filters = self._build_student_error_book_scope_filters(actor, normalized_filters)
        available_questions = self.repository.list_visible_published_questions(
            {
                "subjectCode": target_subject_code,
                "subject_code": target_subject_code,
                "pointCode": target_point_code,
                "pointCode": target_point_code,
                "examCategoryCode": str(question_scope_filters.get("examCategoryCode", "")).strip(),
                "exam_category_code": str(question_scope_filters.get("examCategoryCode", "")).strip(),
                "jointExamGroupCode": str(question_scope_filters.get("jointExamGroupCode", "")).strip(),
                "joint_exam_group_code": str(question_scope_filters.get("jointExamGroupCode", "")).strip(),
                "policyVersionCode": POLICY_VERSION_CODE,
                "policyVersion": POLICY_VERSION_CODE,
                "policy_version": POLICY_VERSION_CODE,
            },
            ROLE_SUPER_ADMIN,
            actor.user_id,
        )
        available_questions = self._filter_questions_for_student_scope(available_questions, actor.user_id, question_scope_filters)
        selected: List[Dict[str, str]] = []
        for question in available_questions:
            question_id = str(question.get("id", "")).strip()
            if not question_id or question_id in wrong_book_question_ids:
                continue
            if target_knowledge_id and str(question.get("knowledgeId", "")).strip() != target_knowledge_id:
                continue
            if descendant_knowledge_ids and str(question.get("knowledgeId", "")).strip() not in descendant_knowledge_ids:
                continue
            question_ext = self._load_json_object(question.get("extJson", "{}"))
            if target_chapter and str(question_ext.get("chapter", "")).strip() != target_chapter:
                continue
            if target_chapter_code and str(question_ext.get("chapterCode", "")).strip() != target_chapter_code:
                continue
            if target_point_code and str(question_ext.get("pointCode", "")).strip() != target_point_code:
                continue
            if target_subject_code and str(question_ext.get("subjectCode", "")).strip() != target_subject_code:
                continue
            tags = self._question_knowledge_tags(question)
            if any(tag in knowledge_tag_pool for tag in tags):
                selected.append(question)
        selected = selected[:15]
        if not selected:
            raise validation_failed("当前错题量不足/薄弱知识点不明确，无法生成专属同类卷，请多刷题积累错题。")
        paper_id = f"paper-personal-{uuid.uuid4().hex[:8]}"
        question_ids = [item["id"] for item in selected]
        question_score = 10
        subject_ids = self._paper_subject_ids(selected)
        self._save_paper_bindings(
            paper_id=paper_id,
            paper_name="AI专属同类卷",
            subject_id=subject_ids[0] if subject_ids else "",
            paper_type="wrongBook",
            paper_status="PUBLISHED",
            duration_minutes=30,
            total_score=question_score * len(question_ids),
            visible_to_students=True,
            target_class_ids=[],
            question_ids=question_ids,
            question_score_map={question_id: question_score for question_id in question_ids},
            rule_mode="personalized",
            actor=actor,
            owner_student_user_id=actor.user_id,
        )
        return {
            "paperId": paper_id,
            "questionIds": question_ids,
            "scope": {
                "subjectCode": target_subject_code,
                "knowledgeId": target_knowledge_id,
                "knowledgePathNodeId": target_path_node_id,
                "chapter": target_chapter,
                "chapterCode": target_chapter_code,
                "pointCode": target_point_code,
            },
        }

    def generate_reasoned_wrong_book_paper(self, actor: Actor, reason_code: str = "", question_count: int = 12) -> Dict[str, object]:
        wrong_book_questions, _ = self.list_wrong_book_questions(1, 200, actor)
        if not wrong_book_questions:
            raise validation_failed("当前错题量不足，无法生成专项巩固卷。")
        normalized_reason_code = reason_code.strip().upper()
        target_count = max(5, min(30, int(question_count)))
        scoped_wrong_questions = wrong_book_questions
        if normalized_reason_code:
            scoped_wrong_questions = [
                question
                for question in wrong_book_questions
                if self._question_matches_wrong_reason(question, normalized_reason_code)
            ]
            if not scoped_wrong_questions:
                raise validation_failed("当前错题本中暂无该错因记录，请先完成对应错因的练习。")
        knowledge_tag_pool: List[str] = []
        for question in scoped_wrong_questions:
            question_ext = self._load_json_object(question["extJson"])
            knowledge_tag_pool.extend(question_ext.get("knowledgeTags", []))
        available_questions = self.repository.list_visible_published_questions({}, ROLE_SUPER_ADMIN, actor.user_id)
        available_questions = self._filter_questions_for_student_scope(available_questions, actor.user_id, {})
        available_by_id = {item["id"]: item for item in available_questions}
        selected: List[Dict[str, str]] = []
        selected_ids: set[str] = set()

        def append_candidate(question: Dict[str, str]) -> None:
            question_id = str(question.get("id", ""))
            if not question_id or question_id in selected_ids:
                return
            selected.append(question)
            selected_ids.add(question_id)

        # 先保留错因命中的错题，再扩展到同知识点题目，最后兜底补足题量。
        for question in scoped_wrong_questions:
            matched = available_by_id.get(str(question["id"]))
            if matched:
                append_candidate(matched)
            if len(selected) >= target_count:
                break
        if len(selected) < target_count and knowledge_tag_pool:
            for question in available_questions:
                if any(tag in knowledge_tag_pool for tag in self._question_knowledge_tags(question)):
                    append_candidate(question)
                if len(selected) >= target_count:
                    break
        if len(selected) < target_count:
            for question in wrong_book_questions:
                matched = available_by_id.get(str(question["id"]))
                if matched:
                    append_candidate(matched)
                if len(selected) >= target_count:
                    break
        if len(selected) < target_count:
            for question in available_questions:
                append_candidate(question)
                if len(selected) >= target_count:
                    break
        if not selected:
            raise validation_failed("当前错题样本不足，暂无法生成专项巩固卷。")
        paper_id = f"paper-reason-{uuid.uuid4().hex[:8]}"
        question_ids = [item["id"] for item in selected]
        question_score = 10
        subject_ids = self._paper_subject_ids(selected)
        label_map = {
            "TIMEOUT": "超时专项",
            "BLANK": "漏答专项",
            "EXPRESSION": "表达专项",
            "KNOWLEDGE_GAP": "知识点专项",
        }
        label = label_map.get(normalized_reason_code, "错因专项")
        self._save_paper_bindings(
            paper_id=paper_id,
            paper_name=f"{label}巩固卷",
            subject_id=subject_ids[0] if subject_ids else "",
            paper_type="wrongBook",
            paper_status="PUBLISHED",
            duration_minutes=35,
            total_score=question_score * len(question_ids),
            visible_to_students=True,
            target_class_ids=[],
            question_ids=question_ids,
            question_score_map={question_id: question_score for question_id in question_ids},
            rule_mode="reasoned",
            actor=actor,
            owner_student_user_id=actor.user_id,
        )
        return {
            "paperId": paper_id,
            "questionIds": question_ids,
            "reasonCode": normalized_reason_code,
            "reasonLabel": label,
        }

    def submit_ai_marking(self, question_id: str, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        question = self.repository.get_question(question_id)
        if not question or question["status"] != "PUBLISHED":
            raise not_found("题目不存在或尚未发布。")
        if not self._is_question_visible_to_student(question, actor.user_id):
            raise forbidden("当前学生不可提交该题的 AI 批改任务。")
        if question["type"] != "subjective":
            raise validation_failed("仅主观题支持 AI 批改。")
        submission = parse_ai_marking_submit_model(payload)
        return self._create_task(
            actor.user_id,
            "AI_MARKING",
            {
                "questionId": question_id,
                "requestPayload": submission.model_dump(),
                "result": {},
                "resultSummary": "",
                "errorMessage": "",
                "queue": {
                    "queueName": "question-bank-ai",
                    "requestedAt": self._now_iso(),
                    "startedAt": "",
                    "completedAt": "",
                    "pollCount": 0,
                },
            },
        )

    def ask_ai_tutor(self, question_id: str, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        question = self.repository.get_question(question_id)
        if not question or question["status"] != "PUBLISHED":
            raise not_found("题目不存在或尚未发布。")
        if not self._is_question_visible_to_student(question, actor.user_id):
            raise forbidden("当前学生不可提交该题的 AI 答疑任务。")
        submission = parse_ai_tutor_ask_model(payload)
        _, _, profile = self._load_student_profile_bundle(actor.user_id)
        ai_quota = self._current_ai_quota(profile)
        if int(ai_quota.get("usedCount", 0)) >= int(ai_quota.get("dailyLimit", 20)):
            raise validation_failed("今日 AI 对话配额已用完，可通过每日打卡或分享好友获取额外配额。")
        ai_quota["usedCount"] = int(ai_quota.get("usedCount", 0)) + 1
        self.repository.set_student_profile_ai_quota(actor.user_id, ai_quota, self._now_iso())
        return self._create_task(
            actor.user_id,
            "AI_TUTOR",
            {
                "questionId": question_id,
                "requestPayload": submission.model_dump(),
                "result": {},
                "resultSummary": "",
                "errorMessage": "",
                "queue": {
                    "queueName": "question-bank-ai",
                    "requestedAt": self._now_iso(),
                    "startedAt": "",
                    "completedAt": "",
                    "pollCount": 0,
                },
            },
        )

    def submit_practice_answer(self, question_id: str, payload: Dict[str, object], actor: Actor) -> Dict[str, str]:
        question = self.repository.get_question(question_id)
        if not question or question["status"] != "PUBLISHED":
            raise not_found("题目不存在或尚未发布。")
        if not self._is_question_visible_to_student(question, actor.user_id):
            raise forbidden("当前学生不可作答该题目。")
        source_type = self._challenge_point_source_type(payload.get("sourceType") or "")
        assignment_id = str(payload.get("assignmentId") or "").strip()
        if source_type != "FREE_PRACTICE" and not self._is_question_chapter_unlocked(question, actor.user_id):
            raise validation_failed("当前章节尚未解锁，请先完成上一章节并达到 80% 正确率。")
        parse_payload = dict(payload or {})
        raw_attempt_key = (
            parse_payload.get("attemptKey")
            or parse_payload.get("attemptId")
            or parse_payload.get("submissionId")
        )
        parse_payload.pop("sourceType", None)
        parse_payload.pop("attemptKey", None)
        parse_payload.pop("attemptId", None)
        parse_payload.pop("submissionId", None)
        submission = parse_practice_submit_model(parse_payload)
        challenge_attempt_key = self._resolve_challenge_point_attempt_key(raw_attempt_key)
        normalized_answer = self._normalize_answer(question["type"], submission.answer)
        time_limit_sec = self._question_time_limit_sec(question)
        is_timeout = submission.elapsedSec > time_limit_sec
        is_correct = False if is_timeout else self._judge_answer(question, normalized_answer)
        record, record_ext = self._load_student_question_record_bundle(question_id, actor.user_id, create_if_missing=True)
        chapter_practice = record_ext.get("chapterPractice", {})
        had_wrong_book = bool((record_ext.get("wrongBook", {}) or {}).get("isCollected"))
        answer_history = chapter_practice.get("answerHistory", [])
        if not isinstance(answer_history, list):
            answer_history = []
        answer_history.append(
            {
                "submittedAt": self._now_iso(),
                "isCorrect": is_correct,
                "answerDurationSec": submission.elapsedSec,
            }
        )
        chapter_practice.update(
            {
                "lastAnswer": submission.answer,
                "normalizedAnswer": normalized_answer,
                "isCorrect": is_correct,
                "answerDurationSec": submission.elapsedSec,
                "timeLimitSec": time_limit_sec,
                "isTimeout": is_timeout,
                "submitCount": int(chapter_practice.get("submitCount", 0)) + 1,
                "correctCount": int(chapter_practice.get("correctCount", 0)) + (1 if is_correct else 0),
                "submittedAt": self._now_iso(),
                "subjectId": self._question_subject_id(question),
                "chapter": self._question_chapter(question),
                "answerHistory": answer_history[-12:],
            }
        )
        record_ext["chapterPractice"] = chapter_practice
        wrong_book = dict(record_ext.get("wrongBook", {}))
        if had_wrong_book:
            wrong_book["postWrongAttemptCount"] = int(wrong_book.get("postWrongAttemptCount", 0)) + 1
            if is_correct:
                wrong_book["postWrongCorrectCount"] = int(wrong_book.get("postWrongCorrectCount", 0)) + 1
        if not is_correct:
            reason_payload = self._resolve_wrong_reason_payload(question, normalized_answer, is_timeout)
            reason_code = str(reason_payload.get("reasonCode", "KNOWLEDGE_GAP"))
            reason_label = str(reason_payload.get("reasonLabel", "知识点掌握不牢"))
            wrong_book["isCollected"] = True
            wrong_book["collectedAt"] = str(wrong_book.get("collectedAt", "")) or self._now_iso()
            wrong_book["lastWrongAt"] = self._now_iso()
            wrong_book["wrongCount"] = int(wrong_book.get("wrongCount", 0)) + 1
            wrong_book["lastReasonCode"] = reason_code
            wrong_book["lastReasonLabel"] = reason_label
            wrong_book["reasonStats"] = self._upsert_wrong_reason_stats(wrong_book.get("reasonStats"), reason_code, reason_label)
        if wrong_book:
            record_ext["wrongBook"] = wrong_book
        self._save_student_question_record_bundle(record, record_ext, actor.user_id)
        self._append_student_daily_metric(actor.user_id, "practiceAnswers", 1)
        self._grant_daily_task_points(actor.user_id)
        if assignment_id and str(question.get("type", "")).strip() != "subjective" and hasattr(self, "record_exam_task_practice_progress"):
            self.record_exam_task_practice_progress(
                assignment_id,
                question_id,
                actor,
                pending_review=False,
            )
        challenge_payload = {
            "challengePointDelta": 0,
            "challengePointGranted": False,
            "challengePointTotal": 0,
            "subjectRank": 0,
            "challengePoints": self._build_challenge_point_summary(actor.user_id, self._question_subject_code(question)),
            "challengeAwardGranted": False,
        }
        if is_correct:
            challenge_payload = self._grant_challenge_point_for_question(
                question,
                actor.user_id,
                source_type,
                challenge_attempt_key,
            )
        result = self._public_question(self._enrich_question_for_student(question, actor.user_id))
        result["answer"] = submission.answer
        result["isCorrect"] = is_correct
        result.update(challenge_payload)
        return result

    def list_student_paper_questions(
        self,
        paper_id: str,
        page: int,
        size: int,
        actor: Actor,
    ) -> Tuple[List[Dict[str, str]], int]:
        questions = self._list_questions_for_paper(paper_id, actor, visible_to_students=True)
        questions = [self._enrich_question_for_student(question, actor.user_id, paper_id=paper_id) for question in questions]
        return self._paginate_questions(questions, page, size)

    def submit_student_paper_question(
        self,
        paper_id: str,
        question_id: str,
        payload: Dict[str, object],
        actor: Actor,
    ) -> Dict[str, object]:
        paper_questions = self._list_questions_for_paper(paper_id, actor, visible_to_students=True)
        question = next((item for item in paper_questions if str(item.get("id", "")).strip() == question_id), None)
        if not question:
            raise not_found("试卷题目不存在或未开放给学生。")
        if str(question.get("type", "")).strip() == "subjective":
            raise validation_failed("主观题不支持即时判分，请在交卷后统一批改。")
        parse_payload = dict(payload or {})
        raw_attempt_key = (
            parse_payload.get("attemptKey")
            or parse_payload.get("attemptId")
            or parse_payload.get("submissionId")
        )
        parse_payload.pop("sourceType", None)
        parse_payload.pop("attemptKey", None)
        parse_payload.pop("attemptId", None)
        parse_payload.pop("submissionId", None)
        submission = parse_practice_submit_model(parse_payload)
        challenge_attempt_key = self._resolve_challenge_point_attempt_key(raw_attempt_key)
        normalized_answer = self._normalize_answer(question["type"], submission.answer)
        is_correct = self._judge_answer(question, normalized_answer)
        self._save_student_paper_attempt(
            question_id,
            actor.user_id,
            paper_id,
            submission.answer,
            normalized_answer,
            is_correct,
            submission.elapsedSec,
            False,
            is_pending_ai_marking=False,
            ai_marking_task_id="",
        )
        challenge_payload = {
            "challengePointDelta": 0,
            "challengePointGranted": False,
            "challengePointTotal": 0,
            "subjectRank": 0,
            "challengePoints": self._build_challenge_point_summary(actor.user_id, self._question_subject_code(question)),
            "challengeAwardGranted": False,
        }
        if is_correct:
            challenge_payload = self._grant_challenge_point_for_question(
                question,
                actor.user_id,
                "MOCK_EXAM",
                challenge_attempt_key,
            )
        result = self._public_question(self._enrich_question_for_student(question, actor.user_id, paper_id=paper_id))
        result["answer"] = submission.answer
        result["isCorrect"] = is_correct
        result["paperId"] = paper_id
        result.update(challenge_payload)
        return result

    def list_student_available_paper_questions(self, page: int, size: int, actor: Actor) -> Tuple[List[Dict[str, str]], int]:
        questions = self.repository.list_visible_published_questions({}, ROLE_SUPER_ADMIN, actor.user_id)
        questions = self._filter_questions_for_student_scope(questions, actor.user_id, {})
        matched: List[Dict[str, str]] = []
        for question in questions:
            paper_bindings = self._visible_student_paper_bindings(question, actor.user_id)
            if not paper_bindings:
                continue
            question = self._with_paper_bindings(question, paper_bindings)
            matched.append(self._enrich_question_for_student(question, actor.user_id))
        return self._paginate_questions(matched, page, size)

    def list_student_paper_reports(self, page: int, size: int, actor: Actor) -> Tuple[List[Dict[str, object]], int]:
        reports = [item for item in self._paper_reports() if str(item.get("studentUserId", "")) == actor.user_id]
        reports.sort(key=lambda item: (str(item.get("submittedAt", "")), str(item.get("paperId", ""))), reverse=True)
        normalized = [self._enrich_student_paper_report_item(item, actor.user_id) for item in reports]
        paged, total = self._paginate_items(normalized, page, size)
        return paged, total

    def get_student_paper_report_detail(self, report_id: str, actor: Actor) -> Dict[str, object]:
        report = self._find_student_paper_report_by_id(actor.user_id, report_id)
        if not report:
            raise not_found("模拟报告不存在。")
        normalized = self._enrich_student_paper_report_item(report, actor.user_id)
        detail = report.get("reportDetail", {})
        if not isinstance(detail, dict):
            detail = {}
        type_accuracy = detail.get("typeAccuracy")
        if not isinstance(type_accuracy, list):
            type_accuracy = report.get("typeAccuracy")
        question_results = detail.get("questionResults")
        normalized["typeAccuracy"] = self._normalize_type_accuracy_rows(type_accuracy if isinstance(type_accuracy, list) else [])
        normalized["questionResults"] = self._refresh_report_question_results(
            self._normalize_question_results(question_results if isinstance(question_results, list) else []),
            actor.user_id,
        )
        if normalized["questionResults"]:
            normalized["score"] = sum(int(item.get("score", 0)) for item in normalized["questionResults"])
            normalized["totalScore"] = sum(int(item.get("totalScore", 0)) for item in normalized["questionResults"])
            normalized["scoreRate"] = (
                round(float(normalized["score"]) / float(normalized["totalScore"]), 4) if int(normalized["totalScore"]) else 0.0
            )
            normalized["wrongQuestionIds"] = self._derive_wrong_question_ids(normalized["questionResults"])
            normalized["typeAccuracy"] = self._build_synced_report_type_accuracy(normalized["questionResults"])
        normalized["summary"] = self._build_report_detail_summary(normalized["questionResults"], normalized["wrongQuestionIds"])
        normalized["summary"]["pendingSubjectiveCount"] = int(normalized["pendingSubjectiveCount"])
        return normalized

    def submit_student_paper(self, paper_id: str, payload: Dict[str, object], actor: Actor) -> Dict[str, object]:
        paper_questions = self._list_questions_for_paper(paper_id, actor, visible_to_students=True)
        if not paper_questions:
            raise not_found("试卷不存在或未开放给学生。")
        today = datetime.now(timezone.utc).date().isoformat()
        existing_report = self._find_student_paper_report(actor.user_id, paper_id, today)
        if existing_report:
            return self._paper_submit_response_from_report(existing_report)
        subject_ids = self._paper_subject_ids(paper_questions)
        self._assert_daily_paper_subject_limit(actor.user_id, subject_ids, today)
        submission = parse_paper_submit_model(payload)
        result = self._evaluate_student_paper_submission(paper_questions, submission.answers, paper_id, actor.user_id)
        subject_id = subject_ids[0] if subject_ids else ""
        submitted_at = self._now_iso()
        report_id = f"paper-report-{uuid.uuid4().hex[:10]}"
        self._append_student_daily_metric(actor.user_id, "papersCompleted", 1)
        self._save_paper_report(
            actor.user_id,
            paper_id,
            {
                "reportId": report_id,
                "paperId": paper_id,
                "subjectId": subject_id,
                "subjectIds": subject_ids,
                "score": result["score"],
                "totalScore": result["totalScore"],
                "scoreRate": round(result["score"] / result["totalScore"], 4) if result["totalScore"] else 0.0,
                "totalElapsedSec": submission.totalElapsedSec,
                "submittedAt": submitted_at,
                "typeAccuracy": result["typeAccuracy"],
                "wrongQuestionIds": result["wrongQuestionIds"],
                "pendingSubjectiveTaskIds": result["pendingSubjectiveTaskIds"],
                "reportDetail": {
                    "typeAccuracy": result["typeAccuracy"],
                    "questionResults": result["questionResults"],
                },
            },
        )
        mock_exam_session = self.repository.find_mock_exam_session_by_paper(actor.user_id, paper_id)
        if mock_exam_session:
            self.repository.upsert_mock_exam_session(
                {
                    **mock_exam_session,
                    "status": "SUBMITTED",
                    "submittedAt": submitted_at,
                    "updateTime": submitted_at,
                }
            )
        self._grant_daily_task_points(actor.user_id)
        return self._paper_submit_response_from_result(
            report_id,
            paper_id,
            subject_id,
            subject_ids,
            result["score"],
            result["totalScore"],
            submission.totalElapsedSec,
            result["typeAccuracy"],
            result["wrongQuestionIds"],
            result["pendingSubjectiveTaskIds"],
            submitted_at,
        )

    def list_analytics_records(
        self,
        filters: Dict[str, str],
        page: int,
        size: int,
        actor: Actor,
    ) -> Tuple[List[Dict[str, str]], int]:
        scoped_filters = self._apply_required_list_scope(filters, actor.role, actor.user_id)
        self._validate_date_filters(scoped_filters)
        student_user_id = str(scoped_filters.get("studentUserId", "")).strip() or None
        records, total = self.repository.list_student_records(student_user_id, scoped_filters, page, size)
        enriched = [self._enrich_analytics_record(record) for record in records]
        return enriched, total

    def build_analytics_summary(self, filters: Dict[str, str], actor: Actor) -> Dict[str, object]:
        scoped_filters = self._apply_required_list_scope(filters, actor.role, actor.user_id)
        self._validate_date_filters(scoped_filters)
        rollup_rows = self.repository.aggregate_student_analytics_rollups(scoped_filters)
        point_rows, chapter_rows, student_subject_rows, student_activity_rows, tag_rows = self._partition_analytics_rollups(rollup_rows)
        if not point_rows and not student_subject_rows:
            return self._empty_analytics_summary(scoped_filters)

        question_count_by_subject = {
            str(item.get("subjectId", "")).strip(): int(item.get("questionCount", 0) or 0)
            for item in self.repository.count_visible_published_questions_by_subject(scoped_filters, actor.role, actor.user_id)
            if str(item.get("subjectId", "")).strip()
        }
        total_answer_duration = sum(float(item.get("totalAnswerDurationSec", 0) or 0) for item in student_subject_rows)
        total_question_count = sum(int(item.get("answerCount", 0) or 0) for item in student_subject_rows)
        average_duration = total_answer_duration / total_question_count if total_question_count else 0.0
        mastery_rows = self._build_mastery_rows_from_rollups(student_subject_rows, question_count_by_subject, average_duration)

        ordered_rankings: List[Dict[str, object]] = []
        for row in chapter_rows:
            answer_count = max(1, int(row.get("answerCount", 0) or 0))
            wrong_count = max(0, int(row.get("wrongCount", 0) or 0))
            ordered_rankings.append(
                {
                    "subjectId": str(row.get("subjectId", "")).strip(),
                    "chapter": str(row.get("chapter", "")).strip(),
                    "wrongCount": wrong_count,
                    "answerCount": int(row.get("answerCount", 0) or 0),
                    "wrongRate": round(wrong_count / answer_count, 4),
                }
            )
        ordered_rankings.sort(key=lambda item: (-int(item["wrongCount"]), item["subjectId"], item["chapter"]))

        latest_submission_by_student = {
            str(item.get("studentUserId", "")).strip(): str(item.get("latestSubmittedAt", "")).strip()
            for item in student_activity_rows
            if str(item.get("studentUserId", "")).strip()
        }
        student_activity = {
            str(item.get("studentUserId", "")).strip(): int(item.get("activityCount", 0) or 0)
            for item in student_activity_rows
            if str(item.get("studentUserId", "")).strip()
        }
        weak_knowledge_tags = [
            {
                "knowledgeTag": str(item.get("knowledgeTag", "")).strip(),
                "wrongCount": int(item.get("wrongCount", 0) or 0),
            }
            for item in sorted(
                tag_rows,
                key=lambda item: (-int(item.get("wrongCount", 0) or 0), str(item.get("knowledgeTag", ""))),
            )[:8]
            if int(item.get("wrongCount", 0) or 0) > 0 and str(item.get("knowledgeTag", "")).strip()
        ]

        total_point_rows = self.repository.count_l5_knowledge_points_by_subject(scoped_filters)
        subject_coverage = self._build_subject_coverage_rows(point_rows, total_point_rows)
        student_ids = sorted(student_activity.keys())
        average_accuracy = sum(item["accuracy"] for item in mastery_rows) / len(mastery_rows) if mastery_rows else 0.0
        student_rankings, mastered_student_count, at_risk_student_count = self._build_student_rankings(
            mastery_rows,
            latest_submission_by_student,
        )
        low_activity_students = self._build_low_activity_students(student_activity, latest_submission_by_student)
        active_student_count = len([item for item in student_activity.values() if item > 0])
        total_student_count = max(len(self._student_directory_ids()), len(student_ids))
        return {
            "timeRangeLabel": self._format_time_range_label(scoped_filters),
            "studentCount": len(student_ids),
            "activeStudentCount": active_student_count,
            "inactiveStudentCount": max(0, len(self._student_directory_ids()) - active_student_count),
            "coverageRate": round(active_student_count / total_student_count, 4) if total_student_count else 0.0,
            "questionCount": total_question_count,
            "averageAccuracy": round(average_accuracy, 4),
            "averageAnswerDurationSec": round(average_duration, 2),
            "masteredStudentCount": mastered_student_count,
            "atRiskStudentCount": at_risk_student_count,
            "chapterRankings": ordered_rankings,
            "weakKnowledgeTags": weak_knowledge_tags,
            "lowActivityStudents": low_activity_students,
            "mastery": mastery_rows,
            "studentRankings": student_rankings[:8],
            "subjectCoverage": subject_coverage,
            "aiReport": self._resolve_ai_analytics_report_async(
                ordered_rankings,
                mastery_rows,
                weak_knowledge_tags,
                low_activity_students,
            ),
        }

    def export_analytics(self, filters: Dict[str, str], actor: Actor) -> Dict[str, str]:
        scoped_filters = self._apply_required_list_scope(filters, actor.role, actor.user_id)
        summary = self.build_analytics_summary(scoped_filters, actor)
        export_format = self._normalize_export_format(scoped_filters.get("format", "csv"), ANALYTICS_EXPORT_FORMATS, "学情导出", "csv")
        return {"format": export_format, "content": "\n".join(self._build_analytics_export_lines(summary, export_format))}
