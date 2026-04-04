from __future__ import annotations

import base64
import io
import json
import re
import sqlite3
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document

import app.repository as repository_module
import app.service_shared as service_shared
from app.auth import Actor
from app.db import (
    backfill_student_profile_hot_state_tables,
    ensure_student_profile_state_defaults,
    get_connection,
    migrate_legacy_passwords,
)
from app.content_baseline import all_joint_exam_group_codes
from app.contracts import QUESTION_FIELDS, TASK_FIELDS
from tests.support import (
    ADVANCED_MATH_POINT_ID,
    ENGLISH_POINT_ID,
    INFO_TECH_CHAPTER_ID,
    INFO_TECH_POINT_ID,
    INFO_TECH_ROOT_ID,
    INFO_TECH_SECTION_ID,
    POLITICS_CHAPTER_ID,
    POLITICS_POINT_ID,
    POLITICS_ROOT_ID,
    POLITICS_SECTION_ID,
    knowledge_payload,
    make_client,
    payload,
    poll_task,
    student_headers,
    super_admin_headers,
    teacher_headers,
)


def auth_headers(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


def super_admin_auth_headers(client: object) -> dict[str, str]:
    login_response = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800000001", "password": "seed-password-admin-001"},
    )
    assert login_response.status_code == 200
    if hasattr(client, "cookies"):
        client.cookies.clear()
    return auth_headers(login_response.json()["data"]["accessToken"])


def teacher_auth_headers(
    client: object,
    phone: str = "13800000002",
    password: str = "seed-password-teacher-001",
) -> dict[str, str]:
    login_response = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": phone, "password": password},
    )
    assert login_response.status_code == 200
    if hasattr(client, "cookies"):
        client.cookies.clear()
    return auth_headers(login_response.json()["data"]["accessToken"])


def register_student_auth_headers(client: object, *, exam_category_code: str = "SCIENCE_ENGINEERING", joint_exam_group_code: str = "SCIENCE_ENGINEERING_3") -> dict[str, str]:
    suffix = str(int(uuid4().hex[:10], 16))[-8:].zfill(8)
    phone = f"139{suffix}"
    sms_response = client.post(
        "/api/question-bank/auth/sms-code",
        json={"phone": phone, "purpose": "register"},
    )
    assert sms_response.status_code == 200
    sms_code = sms_response.json()["data"]["devCode"]

    register_response = client.post(
        "/api/question-bank/auth/register",
        json={
            "phone": phone,
            "smsCode": sms_code,
            "password": "Study123",
            "role": "student",
            "name": "段位测试学生",
            "examCategoryCode": exam_category_code,
            "jointExamGroupCode": joint_exam_group_code,
            "vocationalMajor": "计算机类",
            "prepStage": "基础阶段",
            "employeeNo": "",
        },
    )
    assert register_response.status_code == 200

    login_response = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": phone, "password": "Study123"},
    )
    assert login_response.status_code == 200
    if hasattr(client, "cookies"):
        client.cookies.clear()
    return auth_headers(login_response.json()["data"]["accessToken"])


def required_permissions(raw: str) -> list[str]:
    return [item.strip() for item in str(raw or "").split(",") if item.strip()]


def permission_marker_visible(raw: str, permissions: list[str], mode: str = "all") -> bool:
    required = required_permissions(raw)
    if not required:
        return True
    permission_set = {str(item).strip() for item in permissions if str(item).strip()}
    normalized_mode = str(mode or "all").strip().lower()
    if normalized_mode == "any":
        return any(key in permission_set for key in required)
    return all(key in permission_set for key in required)


TEACHER_MENU_PERMISSION_RULES = (
    ("题库管理", "question:manage"),
    ("学生账号", "student:manage"),
    ("知识点三级树", "question:manage"),
    ("内容体系字典", "paper:manage"),
    ("试卷管理", "paper:manage"),
    ("学情管理", "analytics:view"),
)


def parse_page_bootstrap(raw: str) -> dict[str, object]:
    payload = json.loads(raw)
    assert isinstance(payload, dict)
    assert "route" in payload
    assert "viewKey" in payload
    assert "pageTitle" in payload
    return payload


def seed_student_hot_state(
    service: object,
    student_user_id: str,
    *,
    points: int | None = None,
    title: str | None = None,
    unlocked_titles: list[str] | None = None,
    check_in_dates: list[str] | None = None,
    ai_quota: dict[str, object] | None = None,
    exam_session: dict[str, object] | None = None,
    daily_progress: dict[str, dict[str, object]] | None = None,
    points_ledger: list[dict[str, object]] | None = None,
) -> None:
    repo = service.repository
    current = repo.get_student_profile_state(student_user_id) or {}
    now_iso = "2026-03-23T00:00:00Z"
    repo.upsert_student_profile_state(
        {
            "studentUserId": student_user_id,
            "examCategoryCode": str(current.get("examCategoryCode", "SCIENCE_ENGINEERING")).strip() or "SCIENCE_ENGINEERING",
            "jointExamGroupCode": str(current.get("jointExamGroupCode", "SCIENCE_ENGINEERING_3")).strip()
            or "SCIENCE_ENGINEERING_3",
            "points": int(points if points is not None else current.get("points", 0) or 0),
            "title": str(title if title is not None else current.get("title", "备考新星")).strip() or "备考新星",
            "unlockedTitles": list(unlocked_titles if unlocked_titles is not None else current.get("unlockedTitles", ["备考新星"])),
            "checkInDates": list(check_in_dates if check_in_dates is not None else current.get("checkInDates", [])),
            "aiQuota": dict(ai_quota if ai_quota is not None else current.get("aiQuota", {"dailyLimit": 20, "usedCount": 0, "quotaDate": ""})),
            "examSession": dict(exam_session if exam_session is not None else current.get("examSession", {"answeredCount": 0, "elapsedSec": 0, "updateTime": ""})),
            "createTime": str(current.get("createTime", now_iso)).strip() or now_iso,
            "updateTime": now_iso,
        }
    )
    for progress_date, row in (daily_progress or {}).items():
        repo.upsert_student_daily_progress(
            {
                "studentUserId": student_user_id,
                "progressDate": progress_date,
                "checkInCount": int(row.get("checkInCount", 0) or 0),
                "practiceAnswers": int(row.get("practiceAnswers", 0) or 0),
                "papersCompleted": int(row.get("papersCompleted", 0) or 0),
                "wrongBookReviewed": int(row.get("wrongBookReviewed", 0) or 0),
                "rewardedKeys": list(row.get("rewardedKeys", [])) if isinstance(row.get("rewardedKeys", []), list) else [],
                "createTime": str(row.get("createTime", progress_date)).strip() or progress_date,
                "updateTime": str(row.get("updateTime", progress_date)).strip() or progress_date,
            }
        )
    for item in points_ledger or []:
        repo.upsert_student_points_ledger(
            {
                "studentUserId": student_user_id,
                "eventKey": str(item.get("eventKey", "")).strip(),
                "reason": str(item.get("reason", "")).strip(),
                "points": int(item.get("points", 0) or 0),
                "createTime": str(item.get("createTime", now_iso)).strip() or now_iso,
                "updateTime": str(item.get("updateTime", item.get("createTime", now_iso))).strip() or now_iso,
                "extJson": dict(item),
            }
        )


def visible_teacher_menu_labels_by_permissions(html: str, permissions: list[str]) -> list[str]:
    try:
        payload = json.loads(html)
    except json.JSONDecodeError:
        payload = None
    if isinstance(payload, dict) and "viewKey" in payload:
        return [
            label
            for label, permission in TEACHER_MENU_PERMISSION_RULES
            if permission_marker_visible(permission, permissions)
        ]

    rows = re.findall(
        r'<a\b(?=[^>]*\bclass="[^"]*\bmenu-item\b[^"]*")(?=[^>]*\bdata-required-permission="([^"]+)")[^>]*>(.*?)</a>',
        html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    visible: list[str] = []
    for required, raw_label in rows:
        if not permission_marker_visible(required, permissions):
            continue
        label = re.sub(r"<[^>]+>", "", raw_label).strip()
        if label:
            visible.append(label)
    return visible


def test_auth_register_password_login_me_and_logout(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository

    sms_response = client.post(
        "/api/question-bank/auth/sms-code",
        json={"phone": "13800000011", "purpose": "register"},
    )
    assert sms_response.status_code == 200
    sms_code = sms_response.json()["data"]["devCode"]
    assert re.fullmatch(r"\d{6}", sms_code)

    register_response = client.post(
        "/api/question-bank/auth/register",
        json={
            "phone": "13800000011",
            "smsCode": sms_code,
            "password": "Study123",
            "role": "student",
            "name": "注册学生",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "vocationalMajor": "计算机类",
            "prepStage": "基础阶段",
            "employeeNo": "",
        },
    )
    assert register_response.status_code == 200
    register_data = register_response.json()["data"]
    assert register_data["role"] == "student"
    assert register_data["status"] == "ENABLED"
    assert register_data["auditRequired"] is False

    login_response = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800000011", "password": "Study123"},
    )
    assert login_response.status_code == 200
    login_data = login_response.json()["data"]
    token = login_data["accessToken"]
    assert login_data["tokenType"] == "Bearer"
    assert login_data["role"] == "student"
    assert login_data["userId"] == register_data["userId"]
    assert "qbAccessTok" "en=" in login_response.headers.get("set-cookie", "")
    assert "qbCsrfTok" "en=" in login_response.headers.get("set-cookie", "")

    me_response = client.get("/api/question-bank/auth/me", headers=auth_headers(token))
    assert me_response.status_code == 200
    me_data = me_response.json()["data"]
    assert me_data["userId"] == register_data["userId"]
    assert me_data["role"] == "student"
    assert me_data["phone"] == "13800000011"
    assert me_data["examCategoryCode"] == "SCIENCE_ENGINEERING"
    assert me_data["jointExamGroupCode"] == "SCIENCE_ENGINEERING_3"

    registered_user = repo.get_user_by_id(register_data["userId"])
    assert registered_user is not None
    registered_user_ext = json.loads(registered_user["extJson"])
    assert "examCategoryCode" not in registered_user_ext
    assert "jointExamGroupCode" not in registered_user_ext
    registered_state = repo.get_student_profile_state(register_data["userId"])
    assert registered_state is not None
    assert registered_state["examCategoryCode"] == "SCIENCE_ENGINEERING"
    assert registered_state["jointExamGroupCode"] == "SCIENCE_ENGINEERING_3"
    assert registered_state["vocationalMajor"] == "计算机类"
    assert registered_state["prepStage"] == "基础阶段"

    logout_response = client.post("/api/question-bank/auth/logout", headers=auth_headers(token))
    assert logout_response.status_code == 200
    assert logout_response.json()["data"]["loggedOut"] is True

    expired_me = client.get("/api/question-bank/auth/me", headers=auth_headers(token))
    assert expired_me.status_code == 401
    assert expired_me.json()["code"] == "AUTH_UNAUTHORIZED"


def test_auth_login_by_sms_and_reset_password(tmp_path: Path):
    client = make_client(tmp_path)

    login_sms = client.post(
        "/api/question-bank/auth/sms-code",
        json={"phone": "13800000004", "purpose": "login"},
    )
    assert login_sms.status_code == 200
    login_code = login_sms.json()["data"]["devCode"]

    login_response = client.post(
        "/api/question-bank/auth/login/sms",
        json={"phone": "13800000004", "smsCode": login_code},
    )
    assert login_response.status_code == 200
    assert login_response.json()["data"]["role"] == "teacher"

    reset_sms = client.post(
        "/api/question-bank/auth/sms-code",
        json={"phone": "13800000002", "purpose": "reset_password"},
    )
    assert reset_sms.status_code == 200
    reset_code = reset_sms.json()["data"]["devCode"]

    reset_response = client.post(
        "/api/question-bank/auth/password/reset",
        json={"phone": "13800000002", "smsCode": reset_code, "newPassword": "Teach1234"},
    )
    assert reset_response.status_code == 200
    assert reset_response.json()["data"]["userId"] == "teacher-001"

    old_password_login = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800000002", "password": "seed-password-teacher-001"},
    )
    assert old_password_login.status_code == 403
    assert old_password_login.json()["code"] == "QUESTION_FORBIDDEN"

    new_password_login = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800000002", "password": "Teach1234"},
    )
    assert new_password_login.status_code == 200
    assert new_password_login.json()["data"]["userId"] == "teacher-001"


def test_super_admin_cookie_login_allows_admin_console_and_logout_clears_cookie(tmp_path: Path):
    client = make_client(tmp_path)
    login_response = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800000001", "password": "seed-password-admin-001"},
    )
    assert login_response.status_code == 200
    assert "qbAccessTok" "en=" in login_response.headers.get("set-cookie", "")
    assert "qbCsrfTok" "en=" in login_response.headers.get("set-cookie", "")

    cookie_console = client.get("/api/question-bank/admin/console")
    assert cookie_console.status_code == 200

    logout_response = client.post("/api/question-bank/auth/logout")
    assert logout_response.status_code == 200
    assert "qbAccessTok" "en=" in logout_response.headers.get("set-cookie", "")
    assert "qbCsrfTok" "en=" in logout_response.headers.get("set-cookie", "")
    assert "Max-Age=0" in logout_response.headers.get("set-cookie", "")

    after_logout_console = client.get("/api/question-bank/admin/console")
    assert after_logout_console.status_code == 401
    assert "登录态" in after_logout_console.json()["message"]


def test_super_admin_cookie_session_requires_csrf_for_admin_write(tmp_path: Path):
    client = make_client(tmp_path)
    login_response = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800000001", "password": "seed-password-admin-001"},
    )
    assert login_response.status_code == 200
    csrf_token = client.cookies.get("qbCsrfToken")
    assert csrf_token

    without_csrf = client.post(
        "/api/question-bank/admin/settings",
        json={
            "platformName": "专升本 ALL AI",
            "defaultExamMinutes": 120,
            "dailyCheckInPoints": 10,
            "practiceRewardThreshold": 10,
            "practiceRewardPoints": 20,
            "paperRewardPoints": 30,
            "wrongBookRewardThreshold": 5,
            "wrongBookRewardPoints": 15,
            "aiDailyLimit": 20,
        },
    )
    assert without_csrf.status_code == 403
    assert "安全校验失败" in without_csrf.json()["message"]

    with_csrf = client.post(
        "/api/question-bank/admin/settings",
        headers={"X-CSRF-Token": csrf_token},
        json={
            "platformName": "专升本 ALL AI",
            "defaultExamMinutes": 120,
            "dailyCheckInPoints": 10,
            "practiceRewardThreshold": 10,
            "practiceRewardPoints": 20,
            "paperRewardPoints": 30,
            "wrongBookRewardThreshold": 5,
            "wrongBookRewardPoints": 15,
            "aiDailyLimit": 20,
        },
    )
    assert with_csrf.status_code == 200


def test_password_login_rate_limit_blocks_repeated_failures(tmp_path: Path):
    client = make_client(tmp_path)
    for _ in range(5):
        failed = client.post(
            "/api/question-bank/auth/login/password",
            json={"phone": "13800009999", "password": "wrong-password"},
        )
        assert failed.status_code == 403
        assert failed.json()["message"] == "账号或密码错误。"

    blocked = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800009999", "password": "wrong-password"},
    )
    assert blocked.status_code == 403
    assert "登录尝试过于频繁" in blocked.json()["message"]


def test_init_db_migrates_plaintext_password_to_hashed_storage(tmp_path: Path):
    db_path = tmp_path / "question_bank.db"
    client = make_client(tmp_path)
    assert client.get("/api/question-bank/subjects", headers=teacher_headers()).status_code == 200

    with sqlite3.connect(db_path) as connection:
        connection.execute('UPDATE "user" SET password = ? WHERE id = ?', ("legacy-password-001", "teacher-001"))
        migrate_legacy_passwords(connection)
        connection.commit()

    migrated_login = client.post(
        "/api/question-bank/auth/login/password",
        json={"phone": "13800000002", "password": "legacy-password-001"},
    )
    assert migrated_login.status_code == 200

    with sqlite3.connect(db_path) as connection:
        row = connection.execute('SELECT password FROM "user" WHERE id = ?', ("teacher-001",)).fetchone()
    assert row is not None
    assert str(row[0]).startswith("sha256$")


def test_get_connection_enables_wal_normal_and_cache_size(tmp_path: Path):
    db_path = tmp_path / "question_bank.db"
    client = make_client(tmp_path)
    assert client.get("/api/question-bank/subjects", headers=teacher_headers()).status_code == 200

    with get_connection(db_path) as connection:
        journal_mode = connection.execute("PRAGMA journal_mode;").fetchone()[0]
        synchronous = connection.execute("PRAGMA synchronous;").fetchone()[0]
        cache_size = connection.execute("PRAGMA cache_size;").fetchone()[0]

    assert str(journal_mode).lower() == "wal"
    assert int(synchronous) == 1
    assert int(cache_size) == -10000


def test_student_profile_backfill_only_migrates_fat_legacy_snapshot_and_seeds_thin_defaults(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository
    fat_student_id = "student-001"
    thin_student_id = "student-002"
    fat_record = repo.get_student_question_bank("question-seed-001", fat_student_id)
    thin_record = repo.get_student_question_bank("question-seed-005", thin_student_id)
    assert fat_record is not None
    assert thin_record is not None

    fat_ext = json.loads(fat_record["extJson"])
    fat_ext["studentProfile"] = {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
        "checkInDates": ["2026-03-23"],
        "points": 66,
        "title": "连刷达人",
        "unlockedTitles": ["备考新星", "连刷达人"],
        "dailyProgress": {
            "2026-03-23": {
                "checkInCount": 1,
                "practiceAnswers": 10,
                "papersCompleted": 0,
                "wrongBookReviewed": 5,
                "rewardedKeys": ["practiceReward", "wrongBookReward"],
            }
        },
        "pointsLedger": [
            {
                "eventKey": "fat-backfill:1",
                "reason": "旧积分账本",
                "points": 20,
                "createTime": "2026-03-23T08:00:00Z",
            }
        ],
        "aiQuota": {"dailyLimit": 24, "usedCount": 2, "quotaDate": "2026-03-23"},
        "examSession": {"answeredCount": 7, "elapsedSec": 320, "updateTime": "2026-03-23T09:00:00Z"},
    }

    with get_connection(repo.db_path) as connection:
        connection.execute("DELETE FROM student_profile_state WHERE studentUserId IN (?, ?)", (fat_student_id, thin_student_id))
        connection.execute("DELETE FROM student_daily_progress WHERE studentUserId = ?", (fat_student_id,))
        connection.execute("DELETE FROM student_points_ledger WHERE studentUserId = ?", (fat_student_id,))
        connection.execute(
            """
            UPDATE student_question_record
            SET extJson = ?
            WHERE studentUserId = ? AND questionId = ?
            """,
            (json.dumps(fat_ext, ensure_ascii=False), fat_student_id, fat_record["questionId"]),
        )
        backfill_student_profile_hot_state_tables(connection)
        ensure_student_profile_state_defaults(connection)
        connection.commit()

    fat_state = repo.get_student_profile_state(fat_student_id)
    assert fat_state is not None
    assert fat_state["points"] == 66
    assert fat_state["title"] == "连刷达人"
    assert fat_state["checkInDates"] == ["2026-03-23"]
    assert fat_state["aiQuota"] == {"dailyLimit": 24, "usedCount": 2, "quotaDate": "2026-03-23"}
    assert fat_state["examSession"] == {"answeredCount": 7, "elapsedSec": 320, "updateTime": "2026-03-23T09:00:00Z"}
    fat_progress = repo.get_student_daily_progress(fat_student_id, "2026-03-23")
    assert fat_progress is not None
    assert fat_progress["practiceAnswers"] == 10
    assert fat_progress["rewardedKeys"] == ["practiceReward", "wrongBookReward"]
    fat_ledger = repo.list_student_points_ledger(fat_student_id)
    assert any(item["eventKey"] == "fat-backfill:1" for item in fat_ledger)

    refreshed_fat_record = repo.get_student_question_bank(fat_record["questionId"], fat_student_id)
    assert refreshed_fat_record is not None
    refreshed_fat_row = repo.get_student_question_record_row(fat_student_id, fat_record["questionId"])
    assert refreshed_fat_row is not None
    assert refreshed_fat_row["profileAnchorFlag"] == 1
    refreshed_fat_ext = json.loads(refreshed_fat_record["extJson"])
    assert dict(refreshed_fat_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }

    thin_state = repo.get_student_profile_state(thin_student_id)
    assert thin_state is not None
    assert thin_state["points"] == 0
    assert thin_state["title"] == "备考新星"
    assert thin_state["unlockedTitles"] == ["备考新星"]
    assert thin_state["checkInDates"] == []
    assert thin_state["aiQuota"] == {"dailyLimit": 20, "usedCount": 0, "quotaDate": ""}
    assert thin_state["examSession"] == {"answeredCount": 0, "elapsedSec": 0, "updateTime": ""}
    assert repo.list_student_points_ledger(thin_student_id) == []
    assert repo.get_student_daily_progress(thin_student_id, "2026-03-23") is None


def test_student_profile_backfill_merges_into_default_formal_rows(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository
    fat_student_id = "student-001"
    fat_record = repo.get_student_question_bank("question-seed-001", fat_student_id)
    assert fat_record is not None

    fat_ext = json.loads(fat_record["extJson"])
    fat_ext["studentProfile"] = {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
        "checkInDates": ["2026-03-23"],
        "points": 66,
        "title": "连刷达人",
        "unlockedTitles": ["备考新星", "连刷达人"],
        "dailyProgress": {
            "2026-03-23": {
                "checkInCount": 1,
                "practiceAnswers": 10,
                "papersCompleted": 0,
                "wrongBookReviewed": 5,
                "rewardedKeys": ["practiceReward", "wrongBookReward"],
            }
        },
        "pointsLedger": [
            {
                "eventKey": "fat-merge:1",
                "reason": "旧积分账本",
                "points": 20,
                "createTime": "2026-03-23T08:00:00Z",
            }
        ],
        "aiQuota": {"dailyLimit": 24, "usedCount": 2, "quotaDate": "2026-03-23"},
        "examSession": {"answeredCount": 7, "elapsedSec": 320, "updateTime": "2026-03-23T09:00:00Z"},
    }

    repo.upsert_student_profile_state(
        {
            "studentUserId": fat_student_id,
            "points": 0,
            "title": "备考新星",
            "unlockedTitles": ["备考新星"],
            "checkInDates": [],
            "aiQuota": {"dailyLimit": 20, "usedCount": 0, "quotaDate": ""},
            "examSession": {"answeredCount": 0, "elapsedSec": 0, "updateTime": ""},
            "createTime": "2026-03-17T00:00:00Z",
            "updateTime": "2026-03-17T00:00:00Z",
        }
    )
    repo.upsert_student_daily_progress(
        {
            "studentUserId": fat_student_id,
            "progressDate": "2026-03-23",
            "checkInCount": 0,
            "practiceAnswers": 6,
            "papersCompleted": 0,
            "wrongBookReviewed": 0,
            "rewardedKeys": ["practiceReward"],
            "createTime": "2026-03-23",
            "updateTime": "2026-03-23",
        }
    )

    with get_connection(repo.db_path) as connection:
        connection.execute("DELETE FROM student_points_ledger WHERE studentUserId = ?", (fat_student_id,))
        connection.execute(
            """
            UPDATE student_question_record
            SET extJson = ?
            WHERE studentUserId = ? AND questionId = ?
            """,
            (json.dumps(fat_ext, ensure_ascii=False), fat_student_id, fat_record["questionId"]),
        )
        backfill_student_profile_hot_state_tables(connection)
        ensure_student_profile_state_defaults(connection)
        connection.commit()

    fat_state = repo.get_student_profile_state(fat_student_id)
    assert fat_state is not None
    assert fat_state["points"] == 66
    assert fat_state["title"] == "连刷达人"
    assert fat_state["unlockedTitles"] == ["备考新星", "连刷达人"]
    assert fat_state["checkInDates"] == ["2026-03-23"]
    assert fat_state["aiQuota"] == {"dailyLimit": 24, "usedCount": 2, "quotaDate": "2026-03-23"}
    assert fat_state["examSession"] == {"answeredCount": 7, "elapsedSec": 320, "updateTime": "2026-03-23T09:00:00Z"}

    fat_progress = repo.get_student_daily_progress(fat_student_id, "2026-03-23")
    assert fat_progress is not None
    assert fat_progress["checkInCount"] == 1
    assert fat_progress["practiceAnswers"] == 10
    assert fat_progress["wrongBookReviewed"] == 5
    assert fat_progress["rewardedKeys"] == ["practiceReward", "wrongBookReward"]

    fat_ledger = repo.list_student_points_ledger(fat_student_id)
    assert any(item["eventKey"] == "fat-merge:1" for item in fat_ledger)

    refreshed_fat_record = repo.get_student_question_bank(fat_record["questionId"], fat_student_id)
    assert refreshed_fat_record is not None
    refreshed_fat_ext = json.loads(refreshed_fat_record["extJson"])
    assert dict(refreshed_fat_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }


def test_list_student_records_by_user_is_not_truncated_to_1000(tmp_path: Path):
    db_path = tmp_path / "question_bank.db"
    client = make_client(tmp_path)
    repository = client.app.state.service.repository

    with sqlite3.connect(db_path) as connection:
        connection.row_factory = sqlite3.Row
        base_question = connection.execute("SELECT * FROM question ORDER BY id LIMIT 1").fetchone()
        assert base_question is not None
        create_time = str(base_question["createTime"])
        update_time = str(base_question["updateTime"])
        ext_json = str(base_question["extJson"])
        question_rows = []
        student_record_rows = []
        for index in range(1005):
            question_id = f"perf-record-{index:04d}"
            question_rows.append(
                (
                    question_id,
                    str(base_question["knowledgeId"]),
                    str(base_question["userId"]),
                    str(base_question["type"]),
                    f"批量性能回归题 {index}",
                    str(base_question["optionsJson"]),
                    str(base_question["answer"]),
                    str(base_question["status"]),
                    ext_json,
                    create_time,
                    update_time,
                )
            )
            student_record_rows.append(
                (
                    f"student-record-{index:04d}",
                    "student-001",
                    question_id,
                    "ACTIVE",
                    f"2026-03-20T00:{index % 60:02d}:00Z",
                    "",
                    1 if index % 2 == 0 else 0,
                    1,
                    1 if index % 2 == 0 else 0,
                    0 if index % 2 == 0 else 1,
                    30 + (index % 5),
                    "CHAPTER_PRACTICE",
                    "",
                    0,
                    0,
                    json.dumps(
                        {
                            "chapterPractice": {
                                "submittedAt": f"2026-03-20T00:{index % 60:02d}:00Z",
                                "answerDurationSec": 30 + (index % 5),
                                "isCorrect": index % 2 == 0,
                            }
                        },
                        ensure_ascii=False,
                    ),
                    create_time,
                    update_time,
                )
            )

        connection.executemany(
            """
            INSERT INTO question (
              id, knowledgeId, userId, type, stem, optionsJson,
              answer, status, extJson, createTime, updateTime
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            question_rows,
        )

        connection.executemany(
            """
            INSERT INTO student_question_record (
              id, studentUserId, questionId, status, lastSubmittedAt, lastAnswer, lastIsCorrect,
              answerCount, correctCount, wrongCount, totalAnswerDurationSec,
              latestSourceType, latestPaperId, wrongBookFlag, personalBankFlag,
              extJson, createTime, updateTime
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            student_record_rows,
        )
        connection.commit()

    items = repository.list_student_records_by_user("student-001")
    perf_ids = {str(item["id"]) for item in items if str(item["id"]).startswith("perf-record-")}
    assert len(perf_ids) == 1005
    assert perf_ids == {f"perf-record-{index:04d}" for index in range(1005)}


def test_get_student_profile_record_row_prefers_profile_anchor_record(tmp_path: Path):
    client = make_client(tmp_path)
    repository = client.app.state.service.repository

    non_profile_row = repository.get_student_question_record_row("student-001", "question-seed-008")
    assert non_profile_row is not None
    assert non_profile_row["profileAnchorFlag"] == 0
    with get_connection(repository.db_path) as connection:
        connection.execute(
            """
            UPDATE student_question_record
            SET updateTime = ?
            WHERE studentUserId = ? AND questionId = ?
            """,
            ("2099-01-01T00:00:00Z", "student-001", "question-seed-008"),
        )
        connection.commit()

    profile_row = repository.get_student_profile_record_row("student-001")
    assert profile_row is not None
    assert profile_row["questionId"] == "question-seed-001"
    assert profile_row["profileAnchorFlag"] == 1
    profile_ext = json.loads(profile_row["extJson"])
    assert dict(profile_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }


def test_list_questions_uses_fixed_envelope_and_pagination(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/api/question-bank/questions", headers=teacher_headers())
    assert response.status_code == 200
    body = response.json()
    assert body["code"] == "OK"
    assert body["message"] == "success"
    assert set(body["data"].keys()) == {"items", "page", "size", "total"}


def test_question_list_supports_keyword_filter_and_update_time_sort(tmp_path: Path):
    client = make_client(tmp_path)
    first = payload()
    first["title"] = "排序测试题目-A"
    first["content"] = "排序测试题目-A"
    first["analysis"] = "排序测试解析-A"
    second = payload()
    second["title"] = "排序测试题目-B"
    second["content"] = "排序测试题目-B"
    second["analysis"] = "排序测试解析-B"
    created_first = client.post("/api/question-bank/questions", headers=teacher_headers(), json=first)
    assert created_first.status_code == 200
    first_id = created_first.json()["data"]["id"]
    created_second = client.post("/api/question-bank/questions", headers=teacher_headers(), json=second)
    assert created_second.status_code == 200
    second_id = created_second.json()["data"]["id"]

    created_first_data = created_first.json()["data"]
    update_payload = {
        "title": "排序测试题目-A（更新后）",
        "content": "排序测试题目-A（更新后）",
        "type": str(created_first_data.get("type", "single_choice")),
        "knowledgePoints": [str(created_first_data.get("knowledgeId", "")).strip()],
        "options": json.loads(str(created_first_data.get("optionsJson", "[]")) or "[]"),
        "answer": str(created_first_data.get("answer", "")).strip() or "B",
        "status": "DRAFT",
        "analysis": "排序测试解析-A（更新后）",
    }
    updated = client.put(f"/api/question-bank/questions/{first_id}", headers=teacher_headers(), json=update_payload)
    assert updated.status_code == 200

    listed = client.get(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        params={"keyword": "排序测试", "userId": "teacher-001"},
    )
    assert listed.status_code == 200
    ids = [item["id"] for item in listed.json()["data"]["items"] if item["id"] in {first_id, second_id}]
    assert ids[:2] == [first_id, second_id]


def test_knowledge_page_and_tree_use_fixed_envelope(tmp_path: Path):
    client = make_client(tmp_path)
    page = client.get("/teacher/knowledge", params={"role": "teacher", "userId": "teacher-001"})
    assert page.status_code == 200
    page_payload = parse_page_bootstrap(page.text)
    assert page_payload["route"] == "/teacher/knowledge"
    assert page_payload["viewKey"] == "teacher-knowledge"
    assert page_payload["pageTitle"] == "知识点三级树"
    assert page_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}
    assert page_payload["questionStatuses"]
    assert page_payload["questionTypes"]

    tree = client.get("/api/question-bank/knowledge/tree", headers=teacher_headers())
    assert tree.status_code == 200
    body = tree.json()
    assert body["code"] == "OK"
    assert body["message"] == "success"
    assert isinstance(body["data"], dict)
    assert set(body["data"].keys()) == {"nodes", "links"}
    assert isinstance(body["data"]["nodes"], list)
    assert isinstance(body["data"]["links"], list)
    assert body["data"]["nodes"]
    first_node = body["data"]["nodes"][0]
    assert {"id", "label", "mastery", "size", "x", "y"}.issubset(set(first_node.keys()))
    assert 0.0 <= float(first_node["mastery"]) <= 1.0
    assert int(first_node["size"]) >= 1
    assert any(link["type"] == "parent" for link in body["data"]["links"])


def test_knowledge_tree_graph_includes_parent_and_prerequisite_links(tmp_path: Path):
    client = make_client(tmp_path)
    prerequisite = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": POLITICS_SECTION_ID,
            "name": "图结构前置节点",
            "sort": 60,
            "status": "ENABLED",
            "extJson": {"weight": "HIGH"},
        },
    )
    assert prerequisite.status_code == 200
    prerequisite_id = prerequisite.json()["data"]["id"]

    target = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": POLITICS_SECTION_ID,
            "name": "图结构目标节点",
            "sort": 70,
            "status": "ENABLED",
            "extJson": {"weight": "HIGH", "prerequisites": [prerequisite_id]},
        },
    )
    assert target.status_code == 200
    target_id = target.json()["data"]["id"]

    tree = client.get("/api/question-bank/knowledge/tree", headers=teacher_headers())
    assert tree.status_code == 200
    links = tree.json()["data"]["links"]
    assert any(link == {"source": POLITICS_SECTION_ID, "target": target_id, "type": "parent"} for link in links)
    assert any(link == {"source": prerequisite_id, "target": target_id, "type": "prerequisite"} for link in links)


def test_knowledge_tree_supports_subject_scope_sql_filter(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get(
        "/api/question-bank/knowledge/tree",
        headers=teacher_headers(),
        params={
            "policyVersion": "HB_ZSB_2026",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "subjectCode": "INFO_TECH_INTRO",
        },
    )
    assert response.status_code == 200
    data = response.json()["data"]
    labels = {str(item.get("label", "")) for item in data.get("nodes", []) if isinstance(item, dict)}
    assert "信息技术概论" in labels
    assert "计算机基础知识" in labels
    assert "计算机系统组成" in labels
    assert "政治" not in labels


def _build_outline_docx_bytes(lines: list[str]) -> bytes:
    document = Document()
    for line in lines:
        document.add_paragraph(str(line))
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_parse_knowledge_graph_from_word_requires_scope(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/knowledge-graph/parse-from-word",
        headers=teacher_headers(),
        data={},
        files={
            "file": (
                "knowledge-outline.docx",
                _build_outline_docx_bytes(["第一章 信息系统基础", "1. 信息系统概述"]),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert response.status_code == 422
    assert "请先通过三级联动选择科目" in str(response.json().get("message", ""))


def test_parse_knowledge_graph_from_word_creates_highlight_nodes(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/knowledge-graph/parse-from-word",
        headers=teacher_headers(),
        data={
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "subjectCode": "INFO_TECH_INTRO",
        },
        files={
            "file": (
                "knowledge-outline.docx",
                _build_outline_docx_bytes(
                    [
                        "第一章 信息系统基础",
                        "1. 信息系统概述",
                        "2. 数据与信息",
                        "第二章 程序设计基础",
                        "1. 顺序结构",
                        "2. 分支结构",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert isinstance(data.get("createdNodeIds"), list)
    assert isinstance(data.get("recognizedNodeIds"), list)
    assert len(data["recognizedNodeIds"]) >= len(data["createdNodeIds"])
    assert str(data.get("scope", {}).get("subject_code", "")) == "INFO_TECH_INTRO"
    parser_report = data.get("parserReport", {})
    assert str(parser_report.get("parserMode", "")) in {"heuristic", "llm"}

    graph = client.get(
        "/api/question-bank/knowledge/tree",
        headers=teacher_headers(),
        params={
            "policyVersion": "HB_ZSB_2026",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "subjectCode": "INFO_TECH_INTRO",
        },
    )
    assert graph.status_code == 200
    labels = {
        str(item.get("label", ""))
        for item in graph.json()["data"].get("nodes", [])
        if isinstance(item, dict)
    }
    assert "信息系统基础" in labels
    assert "程序设计基础" in labels
    assert "信息系统概述" in labels
    assert "分支结构" in labels


def test_parse_knowledge_graph_from_word_matches_l4_and_l5_semantic_pool(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/knowledge-graph/parse-from-word",
        headers=teacher_headers(),
        data={
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "subjectCode": "INFO_TECH_INTRO",
        },
        files={
            "file": (
                "knowledge-outline.docx",
                _build_outline_docx_bytes(
                    [
                        "第一章 程序设计基础",
                        "1. 程序设计基础",
                        "2. 顺序结构",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert response.status_code == 200
    parser_report = response.json()["data"]["parserReport"]
    semantic_pool = parser_report.get("semanticPool", [])
    semantic_levels = {int(item.get("level", 0) or 0) for item in parser_report.get("semanticPool", [])}
    assert 4 in semantic_levels
    assert 5 in semantic_levels
    svc = client.app.state.service
    assert int(svc._best_outline_semantic_candidate("算法与程序设计", semantic_pool)["level"]) == 4
    assert int(svc._best_outline_semantic_candidate("了解区块链关键技术", semantic_pool)["level"]) == 5


def test_build_knowledge_outline_llm_prompt_includes_knowledge_points_list(tmp_path: Path):
    client = make_client(tmp_path)
    svc = client.app.state.service
    scope_filters = svc._resolve_knowledge_graph_scope_filters(
        actor=Actor(role="teacher", user_id="teacher-001"),
        exam_category_code="SCIENCE_ENGINEERING",
        joint_exam_group_code="SCIENCE_ENGINEERING_3",
        subject_code="INFO_TECH_INTRO",
        policy_version="HB_ZSB_2026",
    )
    knowledge_points_list = svc._load_subject_knowledge_points_list(scope_filters)
    assert knowledge_points_list

    prompt = svc._build_knowledge_outline_llm_prompt(
        "第一章 信息系统基础\n1. 信息系统概述",
        "INFO_TECH_INTRO",
        knowledge_points_list,
    )
    assert "候选知识点清单" in prompt
    assert "语义匹配" in prompt
    assert "module_code" in prompt
    first_module_code = str(knowledge_points_list[0].get("module_code", "")).strip()
    assert first_module_code
    assert first_module_code in prompt


def test_upsert_outline_prefers_existing_id_from_module_code(tmp_path: Path):
    client = make_client(tmp_path)
    svc = client.app.state.service
    scope_filters = svc._resolve_knowledge_graph_scope_filters(
        actor=Actor(role="teacher", user_id="teacher-001"),
        exam_category_code="SCIENCE_ENGINEERING",
        joint_exam_group_code="SCIENCE_ENGINEERING_3",
        subject_code="INFO_TECH_INTRO",
        policy_version="HB_ZSB_2026",
    )
    subject_name = svc._resolve_subject_display_name("INFO_TECH_INTRO", "SCIENCE_ENGINEERING_3")
    root_id, _ = svc._ensure_subject_root_node(scope_filters, subject_name)
    chapter_id, _ = svc._ensure_scoped_knowledge_node(
        parent_id=root_id,
        name="语义挂载去重章节",
        scope_filters=scope_filters,
        status="ENABLED",
    )
    existing_point_id, _ = svc._ensure_scoped_knowledge_node(
        parent_id=chapter_id,
        name="旧版考点名称",
        scope_filters=scope_filters,
        status="ENABLED",
    )
    before_children = svc.repository.list_knowledge_children(chapter_id, "")
    knowledge_points_list = svc._load_subject_knowledge_points_list(scope_filters)

    created_node_ids, recognized_node_ids = svc._upsert_outline_to_knowledge_graph(
        chapter_rows=[
            {
                "chapter": "语义挂载去重章节",
                "points": [
                    {
                        "point": "新版考点名称",
                        "module_code": existing_point_id,
                    }
                ],
            }
        ],
        scope_filters=scope_filters,
        subject_name=subject_name,
        knowledge_points_list=knowledge_points_list,
    )
    after_children = svc.repository.list_knowledge_children(chapter_id, "")

    assert existing_point_id in recognized_node_ids
    assert existing_point_id not in created_node_ids
    assert len(after_children) == len(before_children)
    assert len([item for item in after_children if str(item.get("id", "")).strip() == existing_point_id]) == 1


def test_knowledge_update_prerequisites_supports_connect_flow(tmp_path: Path):
    client = make_client(tmp_path)
    prerequisite = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": POLITICS_SECTION_ID,
            "name": "连线前置节点",
            "sort": 80,
            "status": "ENABLED",
            "extJson": {"weight": "HIGH"},
        },
    )
    assert prerequisite.status_code == 200
    prerequisite_id = prerequisite.json()["data"]["id"]

    target = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": POLITICS_SECTION_ID,
            "name": "连线目标节点",
            "sort": 90,
            "status": "ENABLED",
            "extJson": {"weight": "HIGH"},
        },
    )
    assert target.status_code == 200
    target_id = target.json()["data"]["id"]

    updated = client.post(
        f"/api/question-bank/knowledge/{target_id}/prerequisites",
        headers=teacher_headers(),
        json={"sourceId": prerequisite_id},
    )
    assert updated.status_code == 200
    ext_json = json.loads(updated.json()["data"]["extJson"])
    assert prerequisite_id in ext_json.get("prerequisites", [])

    tree = client.get("/api/question-bank/knowledge/tree", headers=teacher_headers())
    assert tree.status_code == 200
    links = tree.json()["data"]["links"]
    assert any(link == {"source": prerequisite_id, "target": target_id, "type": "prerequisite"} for link in links)


def test_knowledge_layout_save_persists_graph_coordinates(tmp_path: Path):
    client = make_client(tmp_path)
    target_id = POLITICS_POINT_ID

    saved = client.post(
        "/api/question-bank/knowledge/layout",
        headers=teacher_headers(),
        json={"nodes": [{"id": target_id, "x": 321.5, "y": 188.25}]},
    )
    assert saved.status_code == 200
    assert saved.json()["data"]["updatedCount"] == 1
    assert saved.json()["data"]["updatedIds"] == [target_id]

    tree = client.get("/api/question-bank/knowledge/tree", headers=teacher_headers())
    assert tree.status_code == 200
    node = next((item for item in tree.json()["data"]["nodes"] if item["id"] == target_id), None)
    assert node is not None
    assert float(node["x"]) == 321.5
    assert float(node["y"]) == 188.25

    detail = client.get(f"/api/question-bank/knowledge/{target_id}", headers=teacher_headers())
    assert detail.status_code == 200
    ext_json = json.loads(detail.json()["data"]["extJson"])
    layout = ext_json.get("graphLayout", {})
    assert float(layout.get("x", 0)) == 321.5
    assert float(layout.get("y", 0)) == 188.25


def test_student_can_read_knowledge_tree_for_ai_progress_graph(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/api/question-bank/knowledge/tree", headers=student_headers())
    assert response.status_code == 200
    data = response.json()["data"]
    assert isinstance(data.get("nodes"), list)
    assert isinstance(data.get("links"), list)


def test_knowledge_create_update_children_and_delete(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json=knowledge_payload(POLITICS_SECTION_ID),
    )
    assert created.status_code == 200
    item = created.json()["data"]
    assert tuple(item.keys()) == ("id", "parentId", "name", "sort", "status", "extJson", "createTime", "updateTime")
    assert item["parentId"] == POLITICS_SECTION_ID

    children = client.get(
        "/api/question-bank/knowledge/children",
        headers=teacher_headers(),
        params={"parentId": POLITICS_SECTION_ID},
    )
    assert children.status_code == 200
    assert any(child["id"] == item["id"] for child in children.json()["data"])

    updated = client.put(
        f"/api/question-bank/knowledge/{item['id']}",
        headers=teacher_headers(),
        json={
            "id": item["id"],
            "parentId": POLITICS_SECTION_ID,
            "name": "已更新知识点",
            "sort": 40,
            "status": "DISABLED",
            "extJson": item["extJson"],
            "createTime": item["createTime"],
        },
    )
    assert updated.status_code == 200
    assert updated.json()["data"]["name"] == "已更新知识点"
    assert updated.json()["data"]["status"] == "DISABLED"

    deleted = client.delete(f"/api/question-bank/knowledge/{item['id']}", headers=teacher_headers())
    assert deleted.status_code == 200
    assert deleted.json()["data"]["id"] == item["id"]


def test_knowledge_allows_deeper_levels_and_blocks_duplicate_sibling(tmp_path: Path):
    client = make_client(tmp_path)

    duplicate = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": POLITICS_SECTION_ID,
            "name": "导论",
            "sort": 50,
            "status": "ENABLED",
            "extJson": json.dumps({}, ensure_ascii=False),
        },
    )
    assert duplicate.status_code == 422
    assert duplicate.json()["code"] == "QUESTION_VALIDATION_FAILED"

    fourth = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": POLITICS_CHAPTER_ID,
            "name": "第四层节点",
            "sort": 10,
            "status": "ENABLED",
            "extJson": json.dumps({}, ensure_ascii=False),
        },
    )
    assert fourth.status_code == 200
    fourth_item = fourth.json()["data"]
    assert fourth_item["parentId"] == POLITICS_CHAPTER_ID
    fourth_ext = fourth_item["extJson"] if isinstance(fourth_item["extJson"], dict) else json.loads(fourth_item["extJson"])
    assert int(fourth_ext.get("level", 0)) == 5

    fifth = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": fourth_item["id"],
            "name": "第五层节点",
            "sort": 10,
            "status": "ENABLED",
            "extJson": json.dumps({}, ensure_ascii=False),
        },
    )
    assert fifth.status_code == 422
    assert fifth.json()["code"] == "QUESTION_VALIDATION_FAILED"
    assert "L1-L5" in fifth.json()["message"]


def test_knowledge_delete_rejects_parent_with_children(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.delete(f"/api/question-bank/knowledge/{POLITICS_SECTION_ID}", headers=teacher_headers())
    assert response.status_code == 422
    assert response.json()["code"] == "QUESTION_VALIDATION_FAILED"


def test_knowledge_move_supports_up_and_down_sorting(tmp_path: Path):
    client = make_client(tmp_path)
    before = client.get(
        "/api/question-bank/knowledge/children",
        headers=teacher_headers(),
        params={"parentId": POLITICS_SECTION_ID},
    )
    assert before.status_code == 200
    before_items = before.json()["data"]
    assert len(before_items) >= 2
    moving_id = before_items[1]["id"]

    move_up = client.post(f"/api/question-bank/knowledge/{moving_id}/sort/up", headers=teacher_headers())
    assert move_up.status_code == 200
    assert move_up.json()["data"]["id"] == moving_id

    after_up = client.get(
        "/api/question-bank/knowledge/children",
        headers=teacher_headers(),
        params={"parentId": POLITICS_SECTION_ID},
    )
    assert after_up.status_code == 200
    after_up_items = after_up.json()["data"]
    assert after_up_items[0]["id"] == moving_id

    move_down = client.post(f"/api/question-bank/knowledge/{moving_id}/sort/down", headers=teacher_headers())
    assert move_down.status_code == 200
    after_down = client.get(
        "/api/question-bank/knowledge/children",
        headers=teacher_headers(),
        params={"parentId": POLITICS_SECTION_ID},
    )
    assert after_down.status_code == 200
    assert after_down.json()["data"][1]["id"] == moving_id


def test_question_fields_are_aligned_in_create_response(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=payload(),
    )
    assert response.status_code == 200
    question = response.json()["data"]
    assert tuple(question.keys()) == QUESTION_FIELDS


def test_question_create_populates_chapter_and_point_codes_in_ext_json(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=payload(),
    )
    assert response.status_code == 200
    question = response.json()["data"]
    ext_json = question["extJson"] if isinstance(question["extJson"], dict) else json.loads(question["extJson"])
    chapter_code = str(ext_json.get("chapterCode", "")).strip()
    point_code = str(ext_json.get("pointCode", "")).strip()
    assert re.fullmatch(r"CH_\d{3}", chapter_code)
    assert re.fullmatch(r"PT_\d{3}_\d{3}", point_code)


def test_list_questions_supports_chapter_code_and_point_code_filters(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=payload(),
    )
    assert created.status_code == 200
    created_question = created.json()["data"]
    created_id = str(created_question["id"])
    ext_json = created_question["extJson"] if isinstance(created_question["extJson"], dict) else json.loads(created_question["extJson"])
    chapter_code = str(ext_json.get("chapterCode", "")).strip()
    point_code = str(ext_json.get("pointCode", "")).strip()
    assert chapter_code
    assert point_code

    chapter_filtered = client.get(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        params={"chapterCode": chapter_code},
    )
    assert chapter_filtered.status_code == 200
    chapter_items = chapter_filtered.json()["data"]["items"]
    assert any(str(item.get("id", "")) == created_id for item in chapter_items)

    point_filtered = client.get(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        params={"pointCode": point_code},
    )
    assert point_filtered.status_code == 200
    point_items = point_filtered.json()["data"]["items"]
    assert any(str(item.get("id", "")) == created_id for item in point_items)


def test_list_questions_supports_shared_professional_subject_exam_scope(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        params={
            "examCategoryCode": "LITERATURE",
            "jointExamGroupCode": "LITERATURE_1",
            "subjectCode": "ARTS_HISTORY_FOUNDATION",
        },
    )
    assert response.status_code == 200
    payload = response.json()["data"]
    items = payload["items"]
    item_ids = {str(item.get("id", "")).strip() for item in items}

    assert payload["total"] >= 1
    assert "question-seed-std-arts-history-foundation-001" in item_ids


def test_teacher_cannot_create_other_owners_question(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        json=payload(user_id="teacher-999"),
    )
    assert response.status_code == 403
    assert response.json()["code"] == "QUESTION_FORBIDDEN"


def test_question_create_requires_scope_labels(tmp_path: Path):
    client = make_client(tmp_path)
    question_payload = payload()
    question_payload.pop("examCategoryCode", None)

    response = client.post("/api/question-bank/questions", headers=teacher_headers(), json=question_payload)

    assert response.status_code == 422
    assert "examCategoryCode" in str(response.json().get("detail", ""))


def test_question_create_rejects_mismatched_exam_category_and_joint_group(tmp_path: Path):
    client = make_client(tmp_path)
    question_payload = payload()
    question_payload["examCategoryCode"] = "MANAGEMENT"
    question_payload["jointExamGroupCode"] = "SCIENCE_ENGINEERING_3"
    question_payload["subjectCode"] = "INFO_TECH_INTRO"

    response = client.post("/api/question-bank/questions", headers=teacher_headers(), json=question_payload)

    assert response.status_code == 422
    detail_text = str(response.json().get("detail", ""))
    normalized_detail = (
        detail_text.replace("joint_exam_group_code", "jointExamGroupCode").replace("exam_category_code", "examCategoryCode")
    )
    assert "jointExamGroupCode 与 examCategoryCode 不匹配" in normalized_detail


def test_student_is_forbidden_for_question_bank(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get(
        "/api/question-bank/questions",
        headers={"X-Role": "student", "X-User-Id": "student-001"},
    )
    assert response.status_code == 403
    assert response.json()["code"] == "QUESTION_FORBIDDEN"


def test_status_flow_requires_cross_review_and_teacher_publish(tmp_path: Path):
    client = make_client(tmp_path)
    create_response = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        json=payload(user_id="teacher-001"),
    )
    question_id = create_response.json()["data"]["id"]

    qa_response = client.post(
        f"/api/question-bank/questions/{question_id}/status/QA_IN_PROGRESS",
        headers=teacher_headers("teacher-001"),
    )
    assert qa_response.status_code == 200

    same_teacher_review = client.post(
        f"/api/question-bank/questions/{question_id}/status/REVIEW_PENDING",
        headers=teacher_headers("teacher-001"),
    )
    assert same_teacher_review.status_code == 403

    cross_review = client.post(
        f"/api/question-bank/questions/{question_id}/status/REVIEW_PENDING",
        headers=teacher_headers("teacher-002"),
    )
    assert cross_review.status_code == 200

    owner_publish = client.post(
        f"/api/question-bank/questions/{question_id}/status/PUBLISHED",
        headers=teacher_headers("teacher-001"),
    )
    assert owner_publish.status_code == 403

    publish_response = client.post(
        f"/api/question-bank/questions/{question_id}/status/PUBLISHED",
        headers=teacher_headers("teacher-002"),
    )
    assert publish_response.status_code == 200
    assert publish_response.json()["data"]["status"] == "PUBLISHED"


def test_review_pending_question_is_visible_to_cross_teacher_pool(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        json=payload(user_id="teacher-001"),
    )
    assert created.status_code == 200
    question_id = created.json()["data"]["id"]

    qa_response = client.post(
        f"/api/question-bank/questions/{question_id}/status/QA_IN_PROGRESS",
        headers=teacher_headers("teacher-001"),
    )
    assert qa_response.status_code == 200

    review_response = client.post(
        f"/api/question-bank/questions/{question_id}/status/REVIEW_PENDING",
        headers=teacher_headers("teacher-002"),
    )
    assert review_response.status_code == 200

    cross_teacher_list = client.get("/api/question-bank/questions", headers=teacher_headers("teacher-002"))
    assert cross_teacher_list.status_code == 200
    ids = [item["id"] for item in cross_teacher_list.json()["data"]["items"]]
    assert question_id in ids


def test_draft_question_is_visible_to_cross_teacher_pool(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        json=payload(user_id="teacher-001"),
    )
    assert created.status_code == 200
    question_id = created.json()["data"]["id"]

    cross_teacher_list = client.get("/api/question-bank/questions", headers=teacher_headers("teacher-002"))
    assert cross_teacher_list.status_code == 200
    ids = [item["id"] for item in cross_teacher_list.json()["data"]["items"]]
    assert question_id in ids


def test_published_question_is_visible_to_cross_teacher_pool(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        json=payload(user_id="teacher-001"),
    )
    assert created.status_code == 200
    question_id = created.json()["data"]["id"]

    qa_response = client.post(
        f"/api/question-bank/questions/{question_id}/status/QA_IN_PROGRESS",
        headers=teacher_headers("teacher-001"),
    )
    assert qa_response.status_code == 200

    review_response = client.post(
        f"/api/question-bank/questions/{question_id}/status/REVIEW_PENDING",
        headers=teacher_headers("teacher-002"),
    )
    assert review_response.status_code == 200

    publish_response = client.post(
        f"/api/question-bank/questions/{question_id}/status/PUBLISHED",
        headers=teacher_headers("teacher-002"),
    )
    assert publish_response.status_code == 200

    cross_teacher_list = client.get("/api/question-bank/questions", headers=teacher_headers("teacher-002"))
    assert cross_teacher_list.status_code == 200
    ids = [item["id"] for item in cross_teacher_list.json()["data"]["items"]]
    assert question_id in ids


def test_review_records_keep_summary_and_create_time_trace(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        json=payload(user_id="teacher-001"),
    )
    assert created.status_code == 200
    question_id = created.json()["data"]["id"]

    to_qa = client.post(
        f"/api/question-bank/questions/{question_id}/status/QA_IN_PROGRESS",
        headers=teacher_headers("teacher-001"),
    )
    assert to_qa.status_code == 200
    to_review = client.post(
        f"/api/question-bank/questions/{question_id}/status/REVIEW_PENDING",
        headers=teacher_headers("teacher-002"),
    )
    assert to_review.status_code == 200

    detail = client.get(f"/api/question-bank/questions/{question_id}", headers=teacher_headers("teacher-001"))
    assert detail.status_code == 200
    ext = json.loads(detail.json()["data"]["extJson"])
    summary = ext.get("reviewSummary", {})
    assert int(summary.get("reviewCount", 0)) == 2
    assert summary.get("latestStatus") == "REVIEW_PENDING"
    assert summary.get("latestReviewerUserId") == "teacher-002"
    assert summary.get("latestReviewedAt")

    reviews = client.get(f"/api/question-bank/questions/{question_id}/reviews", headers=teacher_headers("teacher-001"))
    assert reviews.status_code == 200
    rows = reviews.json()["data"]
    assert len(rows) >= 2
    assert rows[0]["createTime"]
    assert rows[0]["status"] == "REVIEW_PENDING"


def test_question_update_recomputes_ext_json_metadata_when_knowledge_changes(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=payload(),
    )
    assert created.status_code == 200
    question = created.json()["data"]

    updated = client.put(
        f"/api/question-bank/questions/{question['id']}",
        headers=teacher_headers(),
        json={"knowledgePoints": [ENGLISH_POINT_ID]},
    )
    assert updated.status_code == 200
    ext = json.loads(updated.json()["data"]["extJson"])
    assert ext["subjectId"] == "subject-english"
    assert ext["chapter"] == "词汇知识"
    assert ext["knowledgeTags"] == ["词汇知识", "掌握单词拼写和音节结构知识，准确认读单词中的字母或字母组合等"]


def test_question_update_recomputes_public_content_tags_on_cross_subject_change(tmp_path: Path):
    client = make_client(tmp_path)
    create_payload = payload()
    create_payload["moduleCode"] = "POLITICS_CUSTOM_001"
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=create_payload,
    )
    assert created.status_code == 200
    question = created.json()["data"]

    updated = client.put(
        f"/api/question-bank/questions/{question['id']}",
        headers=teacher_headers(),
        json={"knowledgePoints": [ENGLISH_POINT_ID]},
    )
    assert updated.status_code == 200
    ext = json.loads(updated.json()["data"]["extJson"])
    assert ext["subjectId"] == "subject-english"
    assert ext["subjectCode"] == "ENGLISH"
    assert ext["subjectType"] == "PUBLIC"
    assert ext["jointExamGroupCode"] == ""
    assert ext["examCategoryCode"] == ""
    assert ext["applicableGroups"] == all_joint_exam_group_codes()
    assert ext["moduleCode"] == ""


def test_question_update_clears_stale_professional_content_tags_when_cross_subject_changes(tmp_path: Path):
    client = make_client(tmp_path)
    create_payload = payload()
    create_payload["moduleCode"] = "POLITICS_CUSTOM_001"
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=create_payload,
    )
    assert created.status_code == 200
    question = created.json()["data"]

    updated = client.put(
        f"/api/question-bank/questions/{question['id']}",
        headers=teacher_headers(),
        json={"knowledgePoints": [INFO_TECH_POINT_ID]},
    )
    assert updated.status_code == 200
    ext = json.loads(updated.json()["data"]["extJson"])
    assert ext["subjectId"] == "subject-computer"
    assert ext["chapter"] == "计算机系统组成"
    assert ext["knowledgeTags"] == [
        "计算机系统组成",
        "掌握计算机硬件系统组成，掌握冯·诺依曼计算机体系结构。掌握微型计算机五大部件（运算器、控制器、存储器、输入/输出设备）的功能，理解微型计算机的工作原理和工作过程",
    ]
    assert ext["subjectCode"] == ""
    assert ext["subjectType"] == ""
    assert ext["jointExamGroupCode"] == ""
    assert ext["examCategoryCode"] == ""
    assert ext["applicableGroups"] == []
    assert ext["moduleCode"] == ""


def test_question_update_keeps_valid_professional_content_tags_when_explicitly_provided(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=payload(),
    )
    assert created.status_code == 200
    question = created.json()["data"]
    update_ext = json.loads(question["extJson"])
    update_ext["subjectCode"] = "INFO_TECH_INTRO"
    update_ext["subjectType"] = "PROFESSIONAL_1"
    update_ext["jointExamGroupCode"] = "SCIENCE_ENGINEERING_3"
    update_ext["moduleCode"] = "INFO_TECH_INTRO_CUSTOM_001"
    update_ext["examCategoryCode"] = "SCIENCE_ENGINEERING"

    updated = client.put(
        f"/api/question-bank/questions/{question['id']}",
        headers=teacher_headers(),
        json={"knowledgePoints": [INFO_TECH_POINT_ID], "extJson": update_ext},
    )
    assert updated.status_code == 200
    ext = json.loads(updated.json()["data"]["extJson"])
    assert ext["subjectCode"] == "INFO_TECH_INTRO"
    assert ext["subjectType"] == "PROFESSIONAL_1"
    assert ext["jointExamGroupCode"] == "SCIENCE_ENGINEERING_3"
    assert ext["examCategoryCode"] == "SCIENCE_ENGINEERING"
    assert ext["applicableGroups"] == ["SCIENCE_ENGINEERING_3", "SCIENCE_ENGINEERING_4"]
    assert ext["moduleCode"] == "INFO_TECH_INTRO_CUSTOM_001"


def test_question_update_rejects_invalid_professional_content_tags(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=payload(),
    )
    assert created.status_code == 200
    question = created.json()["data"]
    update_ext = json.loads(question["extJson"])
    update_ext["subjectCode"] = "INFO_TECH_INTRO"
    update_ext["subjectType"] = "PROFESSIONAL_2"
    update_ext["jointExamGroupCode"] = "SCIENCE_ENGINEERING_3"
    update_ext["examCategoryCode"] = "SCIENCE_ENGINEERING"

    updated = client.put(
        f"/api/question-bank/questions/{question['id']}",
        headers=teacher_headers(),
        json={"knowledgePoints": [INFO_TECH_POINT_ID], "extJson": update_ext},
    )
    assert updated.status_code == 422
    assert updated.json()["code"] == "QUESTION_VALIDATION_FAILED"
    assert "subjectCode、subjectType 与 jointExamGroupCode 不匹配" in updated.json()["message"]


def test_template_import_uses_same_contract(tmp_path: Path):
    client = make_client(tmp_path)
    content = "\n".join(
        [
            "【题型】single_choice",
            "【难度】medium",
            "【题干】下列哪一项体现实践观点？",
            "【选项】A.重视经验|B.脱离实践",
            "【答案】A",
            "【解析】实践决定认识。",
            "【知识点】实践与认识",
        ]
    )
    response = client.post(
        "/api/question-bank/imports/template",
        headers=teacher_headers(),
        files={"file": ("questions.txt", io.BytesIO(content.encode("utf-8")), "text/plain")},
        data={"knowledgeId": POLITICS_POINT_ID},
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["imported"] == 1
    assert data["failed"] == 0


def test_template_import_decodes_utf16_text_file(tmp_path: Path):
    client = make_client(tmp_path)
    unique_stem = "\u7f16\u7801\u56de\u5f52UTF16\u9898\u5e72\u6d4b\u8bd5"
    content = "\n".join(
        [
            "\u3010\u9898\u578b\u3011single_choice",
            "\u3010\u96be\u5ea6\u3011medium",
            f"\u3010\u9898\u5e72\u3011{unique_stem}",
            "\u3010\u9009\u9879\u3011A.\u6b63\u786e|B.\u9519\u8bef",
            "\u3010\u7b54\u6848\u3011A",
            "\u3010\u89e3\u6790\u3011UTF16 \u4e0a\u4f20\u7f16\u7801\u89e3\u6790\u9a8c\u8bc1",
            "\u3010\u77e5\u8bc6\u70b9\u3011\u5b9e\u8df5\u4e0e\u8ba4\u8bc6",
        ]
    )
    response = client.post(
        "/api/question-bank/imports/template",
        headers=teacher_headers(),
        files={"file": ("questions.txt", io.BytesIO(content.encode("utf-16")), "text/plain")},
        data={"knowledgeId": POLITICS_POINT_ID},
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["imported"] == 1
    assert data["failed"] == 0

    listed = client.get(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        params={"keyword": unique_stem, "userId": "teacher-001"},
    )
    assert listed.status_code == 200
    items = listed.json()["data"]["items"]
    assert any(str(item.get("stem", "")).strip() == unique_stem for item in items)


def test_template_import_example_is_served_by_api(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/api/question-bank/imports/template/example", headers=teacher_headers())
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["format"] == "txt"
    assert data["fileName"] == "question-batch-template.txt"
    assert "【题型】single_choice" in data["content"]
    assert "【题型】subjective" in data["content"]
    assert "---" in data["content"]
    assert "f(x)=x^2+1" in data["content"]


def test_parse_word_content_extracts_formula_text_from_docx_xml() -> None:
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
  <w:body>
    <w:p>
      <w:r><w:t>题干：</w:t></w:r>
      <m:oMath>
        <m:r><m:t>f(x)=x^2+1</m:t></m:r>
      </m:oMath>
    </w:p>
    <w:p>
      <w:r><w:t>This is a continuation line.</w:t></w:r>
    </w:p>
  </w:body>
</w:document>
"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml"
    ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    relationships = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1"
    Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
    Target="word/document.xml"/>
</Relationships>
"""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", relationships)
        archive.writestr("word/document.xml", document_xml)

    extracted = service_shared.parse_word_content(buffer.getvalue())

    assert "f(x)=x^2+1" in extracted
    assert "This is a continuation line." in extracted


def test_template_preview_returns_question_contract_and_block_errors(tmp_path: Path):
    client = make_client(tmp_path)
    content = "\n---\n".join(
        [
            "\n".join(
                [
                    "【题型】single_choice",
                    "【难度】medium",
                    "【题干】下列哪一项体现实践观点？",
                    "【选项】A.重视经验|B.来源于实践",
                    "【答案】B",
                    "【解析】实践决定认识。",
                    "【知识点】实践与认识",
                ]
            ),
            "\n".join(
                [
                    "【题型】single_choice",
                    "【难度】medium",
                    "【题干】缺少解析的错误题目",
                    "【选项】A.选项一|B.选项二",
                    "【答案】A",
                    "【知识点】实践与认识",
                ]
            ),
        ]
    )
    response = client.post(
        "/api/question-bank/imports/template/preview",
        headers=teacher_headers(),
        files={"file": ("questions.txt", io.BytesIO(content.encode("utf-8")), "text/plain")},
        data={"knowledgeId": POLITICS_POINT_ID},
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["validCount"] == 1
    assert data["invalidCount"] == 1
    assert tuple(data["items"][0].keys()) == QUESTION_FIELDS
    assert data["errorLogFileName"].endswith(".log")
    assert "question-bank import error log" in data["errorLog"]
    assert "errorCount=1" in data["errorLog"]
    ext = json.loads(data["items"][0]["extJson"])
    assert ext["subjectId"] == "subject-politics"
    assert ext["chapter"] == "导论"
    assert "第 2 题校对失败" in data["errors"][0]


def test_template_import_supports_selected_preview_indexes(tmp_path: Path):
    client = make_client(tmp_path)
    knowledge = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json={
            "id": "",
            "parentId": POLITICS_SECTION_ID,
            "name": "逐题勾选导入测试",
            "sort": 60,
            "status": "ENABLED",
            "extJson": json.dumps({}, ensure_ascii=False),
        },
    )
    assert knowledge.status_code == 200
    knowledge_id = knowledge.json()["data"]["id"]
    content = "\n---\n".join(
        [
            "\n".join(
                [
                    "【题型】single_choice",
                    "【难度】medium",
                    "【题干】逐题勾选导入题目一",
                    "【选项】A.选项一|B.选项二",
                    "【答案】A",
                    "【解析】解析一。",
                    "【知识点】逐题勾选导入测试",
                ]
            ),
            "\n".join(
                [
                    "【题型】single_choice",
                    "【难度】hard",
                    "【题干】逐题勾选导入题目二",
                    "【选项】A.选项甲|B.选项乙",
                    "【答案】B",
                    "【解析】解析二。",
                    "【知识点】逐题勾选导入测试",
                ]
            ),
        ]
    )
    response = client.post(
        "/api/question-bank/imports/template",
        headers=teacher_headers(),
        files={"file": ("questions.txt", io.BytesIO(content.encode("utf-8")), "text/plain")},
        data={"knowledgeId": knowledge_id, "selectedIndexes": "1"},
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["imported"] == 1
    assert data["failed"] == 0

    question_list = client.get(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        params={"knowledgeId": knowledge_id, "userId": "teacher-001"},
    )
    assert question_list.status_code == 200
    items = question_list.json()["data"]["items"]
    assert len(items) == 1
    assert items[0]["stem"] == "逐题勾选导入题目二"


def test_template_import_rejects_out_of_range_selected_preview_indexes(tmp_path: Path):
    client = make_client(tmp_path)
    content = "\n".join(
        [
            "【题型】single_choice",
            "【难度】medium",
            "【题干】越界勾选导入题目",
            "【选项】A.选项一|B.选项二",
            "【答案】A",
            "【解析】解析内容。",
            "【知识点】实践与认识",
        ]
    )
    response = client.post(
        "/api/question-bank/imports/template",
        headers=teacher_headers(),
        files={"file": ("questions.txt", io.BytesIO(content.encode("utf-8")), "text/plain")},
        data={"knowledgeId": POLITICS_POINT_ID, "selectedIndexes": "5"},
    )
    assert response.status_code == 422
    assert response.json()["code"] == "QUESTION_VALIDATION_FAILED"


def test_template_import_reports_block_error_when_knowledge_tags_exceed_limit(tmp_path: Path):
    client = make_client(tmp_path)
    content = "\n".join(
        [
            "【题型】single_choice",
            "【难度】medium",
            "【题干】哪项体现实践？",
            "【选项】A.来源于实践|B.脱离实践",
            "【答案】A",
            "【解析】实践决定认识。",
            "【知识点】知识点1,知识点2,知识点3,知识点4",
        ]
    )
    response = client.post(
        "/api/question-bank/imports/template",
        headers=teacher_headers(),
        files={"file": ("questions.txt", io.BytesIO(content.encode("utf-8")), "text/plain")},
        data={"knowledgeId": POLITICS_POINT_ID},
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["imported"] == 0
    assert data["failed"] == 1
    assert "模板中的【知识点】最多支持 3 个" in data["errors"][0]
    assert data["errorLogFileName"].endswith(".log")
    assert "question-bank import error log" in data["errorLog"]
    assert "errorCount=1" in data["errorLog"]


def test_index_page_keeps_fixed_layout_copy(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/teacher/questions", params={"role": "teacher", "userId": "teacher-001"})
    assert response.status_code == 200
    page_payload = parse_page_bootstrap(response.text)
    assert page_payload["route"] == "/teacher/questions"
    assert page_payload["viewKey"] == "teacher-questions"
    assert page_payload["pageTitle"] == "题库管理"
    assert page_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}
    assert page_payload["questionStatuses"]
    assert page_payload["questionTypes"]


def test_index_page_script_ids_match_template_ids(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/teacher/questions", params={"role": "teacher", "userId": "teacher-001"})
    assert response.status_code == 200
    page_payload = parse_page_bootstrap(response.text)
    assert page_payload["viewKey"] == "teacher-questions"
    view = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "views" / "Teacher" / "QuestionManagement.vue").read_text(encoding="utf-8")
    assert "BaseFilterPanel" in view
    assert "QuestionUpload" in view
    assert "knowledgeId" in view
    assert "userId" in view
    assert "type" in view
    assert "stem" in view
    assert "optionsJson" in view
    assert "answer" in view
    assert "status" in view
    assert "extJsonObject" in view
    assert "createTime" in view
    assert "updateTime" in view


def test_index_page_uses_dynamic_filter_selects_and_drawers(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/teacher/questions", params={"role": "teacher", "userId": "teacher-001"})
    assert response.status_code == 200
    page_payload = parse_page_bootstrap(response.text)
    assert page_payload["viewKey"] == "teacher-questions"
    view = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "views" / "Teacher" / "QuestionManagement.vue").read_text(encoding="utf-8")
    assert "BaseFilterPanel" in view
    assert "filterModel" in view
    assert "detailDrawerVisible" in view
    assert "editDialogVisible" in view
    assert "ElMessageBox.prompt" in view
    assert "QuestionUpload" in view


def test_index_page_navigation_links_use_clean_urls(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/teacher/questions", params={"role": "teacher", "userId": "teacher-001"})
    assert response.status_code == 200
    page_payload = parse_page_bootstrap(response.text)
    assert page_payload["route"] == "/teacher/questions"
    routes = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "router" / "teacherRoutes.js").read_text(encoding="utf-8")
    assert "path: '/teacher'" in routes
    assert "path: '/admin'" in routes
    assert "path: '/messages'" in routes
    assert "path: 'questions'" in routes
    assert "path: 'knowledge'" in routes
    assert "path: 'content-system'" in routes
    assert "path: 'papers'" in routes
    assert "path: 'analytics'" in routes
    assert "?role=" not in routes


def test_teacher_questions_page_reads_identity_from_cookie_without_query(tmp_path: Path):
    client = make_client(tmp_path)
    client.cookies.set("qbRole", "teacher")
    client.cookies.set("qbUserId", "teacher-002")
    response = client.get("/teacher/questions")
    assert response.status_code == 200
    page_payload = parse_page_bootstrap(response.text)
    assert page_payload["actor"] == {"role": "teacher", "userId": "teacher-002"}


def test_teacher_question_script_aligns_baseline_and_workflow_rules() -> None:
    view = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "views" / "Teacher" / "QuestionManagement.vue").read_text(encoding="utf-8")
    routes = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "router" / "teacherRoutes.js").read_text(encoding="utf-8")
    assert "fetchContentBaseline" in view
    assert "row.latestStatus === 'QA_IN_PROGRESS'" in view
    assert "label: '提审'" in view
    assert "label: '通过'" in view
    assert "label: '驳回'" in view
    assert "targetStatus: 'PUBLISHED'" in view
    assert "targetStatus: 'REJECTED'" in view
    assert "requiredPermissions: ['question:manage']" in routes
    assert "requiredPermissions: ['analytics:view']" in routes


def test_portal_page_keeps_three_role_copy(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/", params={"role": "teacher", "userId": "teacher-001"})
    assert response.status_code == 200
    page_payload = parse_page_bootstrap(response.text)
    assert page_payload["route"] == "/"
    assert page_payload["viewKey"] == "portal"
    assert page_payload["pageTitle"] == "三角色门户"
    assert page_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}


def test_portal_page_removes_non_current_role_direct_entries(tmp_path: Path):
    client = make_client(tmp_path)
    teacher_view = client.get("/", params={"role": "teacher", "userId": "teacher-001"})
    assert teacher_view.status_code == 200
    page_payload = parse_page_bootstrap(teacher_view.text)
    assert page_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}


def test_portal_script_super_admin_entrance_checks_login_state() -> None:
    script = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "main.js").read_text(encoding="utf-8")
    entry_type = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "utils" / "entryType.js").read_text(encoding="utf-8")
    assert "fetchManagementProfile" in script
    assert "resolveEntryTypeFromLocation" in script
    assert "normalizedPathname.startsWith('/admin/')" in entry_type
    assert "normalizedPathname.startsWith('/teacher/')" in entry_type
    assert "normalizedPathname.startsWith('/student/')" in entry_type


def test_login_script_supports_next_redirect_for_super_admin() -> None:
    login_view = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "views" / "Auth" / "Login.vue").read_text(encoding="utf-8")
    dev_login = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "utils" / "devLogin.js").read_text(encoding="utf-8")
    assert "URLSearchParams(window.location.search)" in login_view
    assert "resolveLocationRedirectPath" in login_view
    assert "window.location.assign(nextPath)" in login_view
    assert "homePath: '/admin/home'" in dev_login


def test_login_page_keeps_portal_and_messages_navigation(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/login")
    assert response.status_code == 200
    page_payload = parse_page_bootstrap(response.text)
    assert page_payload["route"] == "/login"
    assert page_payload["viewKey"] == "login"
    assert page_payload["pageTitle"] == "注册登录"
    assert page_payload["actor"] is None
    assert page_payload["questionStatuses"] == []
    assert page_payload["questionTypes"] == []


def test_frontend_scripts_do_not_use_native_dialog_apis() -> None:
    frontend_dir = Path(__file__).resolve().parents[1] / "frontend" / "src"
    for path in frontend_dir.rglob("*"):
        if path.suffix not in {".js", ".ts", ".vue"}:
            continue
        text = path.read_text(encoding="utf-8")
        assert "window.confirm(" not in text
        assert "window.prompt(" not in text
        assert "window.alert(" not in text


def test_ui_kit_applies_advanced_filter_to_generic_filter_forms() -> None:
    component = (
        Path(__file__).resolve().parents[1]
        / "frontend"
        / "src"
        / "components"
        / "common"
        / "BaseFilterPanel.vue"
    ).read_text(encoding="utf-8")
    assert "const localFilters = reactive({" in component
    assert "keyword: ''" in component
    assert "examCategoryCode: ''" in component
    assert "jointExamGroupCode: ''" in component
    assert "subjectCode: ''" in component
    assert "emit('search', emitFilterModel())" in component
    assert "emit('reset')" in component


def test_role_portals_are_strictly_isolated(tmp_path: Path):
    client = make_client(tmp_path)

    super_admin_to_teacher = client.get("/teacher/questions", headers=super_admin_headers(), follow_redirects=False)
    assert super_admin_to_teacher.status_code == 302
    assert super_admin_to_teacher.headers["location"] == "/admin/home"
    super_admin_to_student = client.get("/student/home", headers=super_admin_headers(), follow_redirects=False)
    assert super_admin_to_student.status_code == 200
    teacher_page = client.get("/admin/control-center", headers=teacher_headers(), follow_redirects=False)
    assert teacher_page.status_code == 302
    assert teacher_page.headers["location"] == "/login"
    student_to_teacher = client.get("/teacher/questions", headers=student_headers(), follow_redirects=False)
    assert student_to_teacher.status_code == 302
    assert student_to_teacher.headers["location"] == "/student/home?role=student&userId=student-001"
    student_page = client.get("/admin/control-center", headers=student_headers(), follow_redirects=False)
    assert student_page.status_code == 302
    assert student_page.headers["location"] == "/login"


def test_super_admin_can_use_student_analysis_and_practice_apis(tmp_path: Path):
    client = make_client(tmp_path)

    analysis_page = client.get("/student/analysis", headers=super_admin_headers())
    assert analysis_page.status_code == 200
    analysis_tasks_page = client.get("/student/analysis/tasks", headers=super_admin_headers())
    assert analysis_tasks_page.status_code == 200

    knowledge_tree = client.get(
        "/api/knowledge-tree",
        headers=super_admin_headers(),
        params={"subjectCode": "POLITICS"},
    )
    assert knowledge_tree.status_code == 200
    assert knowledge_tree.json()["code"] == "OK"

    practice_questions = client.get(
        "/api/question-bank/student/practice/questions",
        headers=super_admin_headers(),
        params={"subjectCode": "POLITICS", "page": 1, "size": 5},
    )
    assert practice_questions.status_code == 200
    assert practice_questions.json()["code"] == "OK"
    query_super_admin = client.get("/admin/control-center", params={"role": "super_admin", "userId": "admin-001"}, follow_redirects=False)
    assert query_super_admin.status_code == 302
    assert query_super_admin.headers["location"] == "/login"

    assert client.get("/api/question-bank/questions", headers=super_admin_headers()).status_code == 403
    assert client.get("/api/question-bank/admin/console", headers=teacher_headers()).status_code == 401
    assert client.get("/api/question-bank/admin/users", headers=student_headers()).status_code == 401


def test_knowledge_tree_response_allows_wrong_count_field(tmp_path: Path):
    client = make_client(tmp_path)

    knowledge_tree = client.get(
        "/api/knowledge-tree",
        headers=student_headers(),
        params={"subjectCode": "POLITICS"},
    )

    assert knowledge_tree.status_code == 200
    payload = knowledge_tree.json()["data"]
    assert payload["nodes"]
    assert all("wrongCount" in item for item in payload["nodes"])
    assert all("shortLabel" in item for item in payload["nodes"])
    assert all("fullLabel" in item for item in payload["nodes"])


def test_knowledge_tree_supports_shared_professional_subject_scope_filter(tmp_path: Path):
    client = make_client(tmp_path)

    knowledge_tree = client.get(
        "/api/knowledge-tree",
        headers=teacher_headers("teacher-001"),
        params={
            "subjectCode": "ARTS_HISTORY_FOUNDATION",
            "examCategoryCode": "LITERATURE",
            "jointExamGroupCode": "LITERATURE_1",
        },
    )

    assert knowledge_tree.status_code == 200
    payload = knowledge_tree.json()["data"]
    assert len(payload["nodes"]) >= 1


def test_init_db_backfills_l5_diagnostic_nodes_for_sparse_subjects(tmp_path: Path):
    client = make_client(tmp_path)

    with get_connection(client.app.state.db_path) as connection:
        art_chapter_total = connection.execute(
            """
            SELECT COUNT(*) AS total
            FROM knowledge
            WHERE COALESCE(json_extract(extJson, '$.subjectCode'), '') = 'ART_INTRODUCTION'
              AND COALESCE(json_extract(extJson, '$.level'), 0) = 4
            """
        ).fetchone()["total"]
        art_l5_total = connection.execute(
            """
            SELECT COUNT(*) AS total
            FROM knowledge
            WHERE COALESCE(json_extract(extJson, '$.subjectCode'), '') = 'ART_INTRODUCTION'
              AND COALESCE(json_extract(extJson, '$.level'), 0) = 5
            """
        ).fetchone()["total"]
        synthetic_total = connection.execute(
            """
            SELECT COUNT(*) AS total
            FROM knowledge
            WHERE id LIKE 'ART_INTRODUCTION-%-diag-l5'
            """
        ).fetchone()["total"]

    assert art_chapter_total > 0
    assert art_l5_total >= art_chapter_total
    assert synthetic_total > 0


def test_init_db_backfills_minimum_published_demo_questions_per_subject(tmp_path: Path):
    client = make_client(tmp_path)

    with get_connection(client.app.state.db_path) as connection:
        counts = {
            row["subject"]: row["total"]
            for row in connection.execute(
                """
                SELECT COALESCE(json_extract(extJson, '$.subjectCode'), '') AS subject,
                       COUNT(*) AS total
                FROM question
                WHERE status = 'PUBLISHED'
                GROUP BY subject
                """
            ).fetchall()
        }

    assert counts["ADVANCED_MATH_1"] >= 4
    assert counts["ENGLISH"] >= 4
    assert counts["ART_INTRODUCTION"] >= 4


def test_analytics_page_contains_keyword_filter_for_advanced_panel(tmp_path: Path):
    client = make_client(tmp_path)
    page = client.get("/teacher/analytics", params={"role": "teacher", "userId": "teacher-001"})
    assert page.status_code == 200
    page_payload = parse_page_bootstrap(page.text)
    assert page_payload["viewKey"] == "teacher-analytics"
    component = (
        Path(__file__).resolve().parents[1]
        / "frontend"
        / "src"
        / "views"
        / "Teacher"
        / "Analytics.vue"
    ).read_text(encoding="utf-8")
    assert "const selectedClassId = computed(() => String(route.query.classId || '').trim())" in component
    assert "const selectedStudentUserId = computed(() => String(route.query.studentUserId || '').trim())" in component
    assert "const raw = route.query.subjectCodes || route.query.subjectCode || ''" in component
    assert 'placeholder="选择班级"' in component
    assert 'placeholder="选择学生"' in component
    assert 'placeholder="选择 L3 / L4 / L5 知识路径"' in component
    assert "@change=\"handleKnowledgePathChange\"" in component


def test_student_practice_endpoint_returns_question_contract(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get(
        "/api/question-bank/student/practice/questions?subjectId=subject-politics&chapter=导论",
        headers=student_headers(),
    )
    assert response.status_code == 200
    items = response.json()["data"]["items"]
    assert items
    assert tuple(items[0].keys()) == QUESTION_FIELDS


def test_student_practice_chapter_statuses_follow_unlock_rule(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get(
        "/api/question-bank/student/practice/chapters?subjectId=subject-politics",
        headers=student_headers(),
    )
    assert response.status_code == 200
    items = response.json()["data"]
    assert [item["chapter"] for item in items] == ["导论", "世界的物质性及其发展规律"]
    assert items[0]["isUnlocked"] is True
    assert items[0]["statusLabel"] in {"已解锁", "正在闯关"}
    assert items[1]["isUnlocked"] is False
    assert items[1]["statusLabel"] == "未解锁"


def test_student_practice_filters_locked_chapters(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get(
        "/api/question-bank/student/practice/questions?subjectId=subject-politics",
        headers=student_headers(),
    )
    assert response.status_code == 200
    items = response.json()["data"]["items"]
    ids = [item["id"] for item in items]
    assert "question-seed-001" in ids
    assert "question-seed-002" in ids
    assert "question-seed-003" not in ids


def test_student_submit_practice_rejects_locked_chapter(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/student/practice/questions/question-seed-003/submit",
        headers=student_headers(),
        json={"answer": "A", "elapsedSec": 20},
    )
    assert response.status_code == 422
    assert response.json()["code"] == "QUESTION_VALIDATION_FAILED"
    assert "章节尚未解锁" in response.json()["message"]


def test_student_submit_practice_updates_student_state(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "B", "elapsedSec": 26},
    )
    assert response.status_code == 200
    question = response.json()["data"]
    ext = json.loads(question["extJson"])
    assert ext["studentState"]["chapterPractice"]["lastAnswer"] == "B"
    assert ext["studentState"]["chapterPractice"]["isCorrect"] is True


def test_student_submit_practice_timeout_marks_answer_wrong(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "B", "elapsedSec": 120},
    )
    assert response.status_code == 200
    ext = json.loads(response.json()["data"]["extJson"])
    chapter_practice = ext["studentState"]["chapterPractice"]
    assert chapter_practice["lastAnswer"] == "B"
    assert chapter_practice["timeLimitSec"] == 60
    assert chapter_practice["isTimeout"] is True
    assert chapter_practice["isCorrect"] is False


def test_student_submit_practice_rejects_negative_elapsed_sec(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "B", "elapsedSec": -1},
    )
    assert response.status_code == 422


def test_challenge_points_repeated_correct_answers_count_each_submit(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)
    summary_before = client.get(
        "/api/question-bank/student/challenge-points",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert summary_before.status_code == 200
    baseline_total = int(summary_before.json()["data"]["total"])

    first_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "B", "elapsedSec": 25},
    )
    assert first_submit.status_code == 200
    first_data = first_submit.json()["data"]
    assert first_data["isCorrect"] is True
    assert first_data["challengePointGranted"] is True
    assert int(first_data["challengePointDelta"]) == 1

    second_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "B", "elapsedSec": 27},
    )
    assert second_submit.status_code == 200
    second_data = second_submit.json()["data"]
    assert second_data["isCorrect"] is True
    assert second_data["challengePointGranted"] is True
    assert int(second_data["challengePointDelta"]) == 1

    summary_after = client.get(
        "/api/question-bank/student/challenge-points",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert summary_after.status_code == 200
    data = summary_after.json()["data"]
    assert int(data["total"]) == baseline_total + 2
    assert int(data["todayDelta"]) >= 2

def test_challenge_points_replayed_submit_with_same_attempt_id_is_deduplicated(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)
    replay_key = "challenge-attempt-001"

    first_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "B", "elapsedSec": 22, "attemptKey": replay_key},
    )
    assert first_submit.status_code == 200
    first_data = first_submit.json()["data"]
    assert first_data["isCorrect"] is True
    assert first_data["challengePointGranted"] is True
    assert int(first_data["challengePointDelta"]) == 1

    replay_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "B", "elapsedSec": 22, "attemptKey": replay_key},
    )
    assert replay_submit.status_code == 200
    replay_data = replay_submit.json()["data"]
    assert replay_data["isCorrect"] is True
    assert replay_data["challengePointGranted"] is False
    assert int(replay_data["challengePointDelta"]) == 0

    new_attempt_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "B", "elapsedSec": 23, "attemptKey": "challenge-attempt-002"},
    )
    assert new_attempt_submit.status_code == 200
    new_attempt_data = new_attempt_submit.json()["data"]
    assert new_attempt_data["isCorrect"] is True
    assert new_attempt_data["challengePointGranted"] is True
    assert int(new_attempt_data["challengePointDelta"]) == 1


def test_challenge_point_level_thresholds_follow_fixed_breakpoints(tmp_path: Path):
    client = make_client(tmp_path)
    svc = client.app.state.service

    assert svc._build_challenge_point_level(0)["levelName"] == "刷题青铜"
    assert svc._build_challenge_point_level(999)["levelName"] == "刷题青铜"
    assert svc._build_challenge_point_level(1000)["levelName"] == "刷题白银"
    assert svc._build_challenge_point_level(2499)["levelName"] == "刷题白银"
    assert svc._build_challenge_point_level(2500)["levelName"] == "刷题黄金"
    assert svc._build_challenge_point_level(4499)["levelName"] == "刷题黄金"
    assert svc._build_challenge_point_level(4500)["levelName"] == "刷题铂金"
    assert svc._build_challenge_point_level(6499)["levelName"] == "刷题铂金"
    assert svc._build_challenge_point_level(6500)["levelName"] == "刷题钻石"
    assert svc._build_challenge_point_level(7999)["levelName"] == "刷题钻石"
    assert svc._build_challenge_point_level(8000)["levelName"] == "刷题星耀"
    assert svc._build_challenge_point_level(9299)["levelName"] == "刷题星耀"
    assert svc._build_challenge_point_level(9300)["levelName"] == "荣耀王者"
    assert svc._build_challenge_point_level(9799)["levelName"] == "荣耀王者"
    assert svc._build_challenge_point_level(9800)["levelName"] == "传奇王者"
    assert svc._build_challenge_point_level(10000)["levelName"] == "传奇王者"
    assert svc._build_challenge_point_level(10000)["isTopLevel"] is True


def test_check_in_growth_points_do_not_change_challenge_points(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)

    summary_before = client.get(
        "/api/question-bank/student/challenge-points",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert summary_before.status_code == 200
    total_before = int(summary_before.json()["data"]["total"])

    check_in = client.post("/api/question-bank/student/check-in", headers=headers)
    assert check_in.status_code == 200
    assert int(check_in.json()["data"]["points"]) >= 1

    summary_after = client.get(
        "/api/question-bank/student/challenge-points",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert summary_after.status_code == 200
    assert int(summary_after.json()["data"]["total"]) == total_before


def test_challenge_point_summary_includes_correct_submit_counters(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)

    submit_response = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "B", "elapsedSec": 24},
    )
    assert submit_response.status_code == 200

    summary = client.get(
        "/api/question-bank/student/challenge-points",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert summary.status_code == 200
    data = summary.json()["data"]
    assert "correctSubmitCount" in data
    assert "todayCorrectSubmitCount" in data
    assert int(data["correctSubmitCount"]) == int(data["total"])
    assert int(data["todayCorrectSubmitCount"]) == int(data["todayDelta"])


def test_paper_management_manual_create_and_export(tmp_path: Path):
    client = make_client(tmp_path)
    create_response = client.post(
        "/api/question-bank/papers/manual",
        headers=teacher_headers(),
        json={
            "paperName": "新建章节卷",
            "subjectId": "subject-politics",
            "paperType": "chapter",
            "paperStatus": "PUBLISHED",
            "durationMinutes": 20,
            "totalScore": 20,
            "visibleToStudents": True,
            "questionIds": ["question-seed-001", "question-seed-002"],
        },
    )
    assert create_response.status_code == 200
    paper_id = create_response.json()["data"]["paperId"]

    list_response = client.get(
        f"/api/question-bank/papers/questions?paperId={paper_id}",
        headers=teacher_headers(),
    )
    assert list_response.status_code == 200
    items = list_response.json()["data"]["items"]
    assert len(items) == 2
    assert tuple(items[0].keys()) == QUESTION_FIELDS

    export_response = client.get(f"/api/question-bank/papers/{paper_id}/export", headers=teacher_headers())
    assert export_response.status_code == 200
    assert "试卷ID" in export_response.json()["data"]["content"]


def test_manual_paper_supports_question_scores(tmp_path: Path):
    client = make_client(tmp_path)
    create_response = client.post(
        "/api/question-bank/papers/manual",
        headers=teacher_headers(),
        json={
            "paperName": "自定义分值手动卷",
            "subjectId": "subject-politics",
            "paperType": "chapter",
            "paperStatus": "DRAFT",
            "durationMinutes": 30,
            "totalScore": 12,
            "visibleToStudents": False,
            "publishClassIds": ["class-a", "class-b", "class-a"],
            "questionIds": ["question-seed-001", "question-seed-002"],
            "questionScores": {
                "question-seed-001": 7,
                "question-seed-002": 5,
            },
        },
    )
    assert create_response.status_code == 200
    paper_id = create_response.json()["data"]["paperId"]

    list_response = client.get(
        f"/api/question-bank/papers/questions?paperId={paper_id}",
        headers=teacher_headers(),
    )
    assert list_response.status_code == 200
    rows = list_response.json()["data"]["items"]
    assert len(rows) == 2
    for row in rows:
        ext = json.loads(row["extJson"])
        binding = next(item for item in ext.get("paperBindings", []) if item["paperId"] == paper_id)
        score_map = binding.get("questionScoreMap", {})
        assert score_map == {"question-seed-001": 7, "question-seed-002": 5}
        assert binding.get("subjectId") == "subject-politics"
        assert binding.get("targetClassIds") == ["class-a", "class-b"]


def test_manual_paper_supports_joint_group_scope_without_subject(tmp_path: Path):
    client = make_client(tmp_path)
    create_response = client.post(
        "/api/question-bank/papers/manual",
        headers=teacher_headers(),
        json={
            "paperName": "专业组手动组卷",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "paperType": "chapter",
            "paperStatus": "DRAFT",
            "durationMinutes": 30,
            "totalScore": 10,
            "visibleToStudents": False,
            "questionIds": ["question-seed-001", "question-seed-002"],
        },
    )
    assert create_response.status_code == 200
    paper_id = str(create_response.json()["data"]["paperId"]).strip()
    assert paper_id

    list_response = client.get(
        "/api/question-bank/papers/questions",
        headers=teacher_headers(),
        params={
            "paperId": paper_id,
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "page": 1,
            "size": 20,
        },
    )
    assert list_response.status_code == 200
    rows = list_response.json()["data"]["items"]
    assert len(rows) == 2
    for row in rows:
        ext = json.loads(row["extJson"])
        binding = next(item for item in ext.get("paperBindings", []) if item["paperId"] == paper_id)
        assert str(binding.get("examCategoryCode", "")).strip() == "SCIENCE_ENGINEERING"
        assert str(binding.get("jointExamGroupCode", "")).strip() == "SCIENCE_ENGINEERING_3"
        assert str(binding.get("subjectCode", "")).strip() == ""
        assert str(binding.get("subjectId", "")).strip() == ""


def test_teacher_can_get_my_classes(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/api/user/my-classes", headers=teacher_headers())
    assert response.status_code == 200
    rows = response.json()["data"]
    assert isinstance(rows, list)
    assert rows
    first = rows[0]
    assert "class_id" in first
    assert "class_name" in first


def test_teacher_can_read_paper_target_weights_profile(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/api/question-bank/papers/target-weights", headers=teacher_headers())
    assert response.status_code == 200
    data = response.json()["data"]
    assert isinstance(data["selectedVersionId"], str)
    assert isinstance(data["selectedVersionName"], str)
    assert isinstance(data["knowledgeWeights"], list)
    assert isinstance(data["targetWeightMap"], dict)
    assert data["knowledgeWeights"]

    first_knowledge_id = str(data["knowledgeWeights"][0]["knowledgeId"])
    filtered_response = client.get(
        f"/api/question-bank/papers/target-weights?knowledgeIds={first_knowledge_id}",
        headers=teacher_headers(),
    )
    assert filtered_response.status_code == 200
    filtered_data = filtered_response.json()["data"]
    assert len(filtered_data["knowledgeWeights"]) == 1
    assert filtered_data["knowledgeWeights"][0]["knowledgeId"] == first_knowledge_id
    assert first_knowledge_id in filtered_data["targetWeightMap"]


def test_question_delete_supports_undo_restore(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post("/api/question-bank/questions", headers=teacher_headers(), json=payload())
    assert created.status_code == 200
    question_id = created.json()["data"]["id"]

    deleted = client.delete(f"/api/question-bank/questions/{question_id}", headers=teacher_headers())
    assert deleted.status_code == 200
    delete_data = deleted.json()["data"]
    assert delete_data["id"] == question_id
    assert delete_data["undoSnapshotId"].startswith("undo-")

    missing = client.get(f"/api/question-bank/questions/{question_id}", headers=teacher_headers())
    assert missing.status_code == 404

    restored = client.post(
        f"/api/question-bank/questions/deleted/{delete_data['undoSnapshotId']}/restore",
        headers=teacher_headers(),
    )
    assert restored.status_code == 200
    assert restored.json()["data"]["id"] == question_id

    detail = client.get(f"/api/question-bank/questions/{question_id}", headers=teacher_headers())
    assert detail.status_code == 200


def test_paper_delete_supports_undo_restore(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/papers/manual",
        headers=teacher_headers(),
        json={
            "paperName": "可撤销删除试卷",
            "subjectId": "subject-politics",
            "paperType": "chapter",
            "paperStatus": "PUBLISHED",
            "durationMinutes": 20,
            "totalScore": 20,
            "visibleToStudents": False,
            "questionIds": ["question-seed-001", "question-seed-002"],
        },
    )
    assert created.status_code == 200
    paper_id = created.json()["data"]["paperId"]

    deleted = client.delete(f"/api/question-bank/papers/{paper_id}", headers=teacher_headers())
    assert deleted.status_code == 200
    delete_data = deleted.json()["data"]
    assert delete_data["paperId"] == paper_id
    assert delete_data["undoSnapshotId"].startswith("undo-")

    after_delete = client.get(
        f"/api/question-bank/papers/questions?paperId={paper_id}",
        headers=teacher_headers(),
    )
    assert after_delete.status_code == 200
    assert after_delete.json()["data"]["items"] == []

    restored = client.post(
        f"/api/question-bank/papers/deleted/{delete_data['undoSnapshotId']}/restore",
        headers=teacher_headers(),
    )
    assert restored.status_code == 200
    assert restored.json()["data"]["paperId"] == paper_id

    after_restore = client.get(
        f"/api/question-bank/papers/questions?paperId={paper_id}",
        headers=teacher_headers(),
    )
    assert after_restore.status_code == 200
    assert len(after_restore.json()["data"]["items"]) == 2


def test_paper_status_transition_follows_state_machine(tmp_path: Path):
    client = make_client(tmp_path)
    create_response = client.post(
        "/api/question-bank/papers/manual",
        headers=teacher_headers(),
        json={
            "paperName": "状态机校验卷",
            "subjectId": "subject-politics",
            "paperType": "chapter",
            "paperStatus": "DRAFT",
            "durationMinutes": 20,
            "totalScore": 20,
            "visibleToStudents": False,
            "questionIds": ["question-seed-001", "question-seed-002"],
        },
    )
    assert create_response.status_code == 200
    paper_id = create_response.json()["data"]["paperId"]

    direct_publish = client.post(f"/api/question-bank/papers/{paper_id}/status/PUBLISHED", headers=teacher_headers())
    assert direct_publish.status_code == 422
    assert direct_publish.json()["code"] == "QUESTION_INVALID_STATUS"

    to_review = client.post(f"/api/question-bank/papers/{paper_id}/status/REVIEW_PENDING", headers=teacher_headers())
    assert to_review.status_code == 200

    published = client.post(f"/api/question-bank/papers/{paper_id}/status/PUBLISHED", headers=teacher_headers())
    assert published.status_code == 200

    offlined = client.post(f"/api/question-bank/papers/{paper_id}/status/OFFLINE", headers=teacher_headers())
    assert offlined.status_code == 200


def test_student_can_list_and_submit_published_paper(tmp_path: Path):
    client = make_client(tmp_path)
    paper_list = client.get("/api/question-bank/student/papers/questions", headers=student_headers())
    assert paper_list.status_code == 200
    items = paper_list.json()["data"]["items"]
    assert items

    response = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 40, "marked": True},
            ],
            "totalElapsedSec": 70,
        },
    )
    assert response.status_code == 200
    report = response.json()["data"]
    assert report["paperId"] == "paper-demo-001"
    assert report["score"] == report["totalScore"]


def test_student_submit_paper_counts_unanswered_as_wrong(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
            ],
            "totalElapsedSec": 30,
        },
    )
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["totalScore"] == 20
    assert data["score"] == 10
    assert "question-seed-002" in data["wrongQuestionIds"]


def test_student_submit_paper_creates_subjective_ai_marking_task(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/student/papers/paper-english-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-004", "answer": "B", "elapsedSec": 35, "marked": False},
                {
                    "questionId": "question-seed-005",
                    "answer": "终身学习有助于持续成长并提升适应能力。",
                    "elapsedSec": 120,
                    "marked": True,
                },
            ],
            "totalElapsedSec": 155,
        },
    )
    assert response.status_code == 200
    report = response.json()["data"]
    assert report["reportId"].startswith("paper-report-")
    assert report["pendingSubjectiveCount"] == 1
    task_ids = report["pendingSubjectiveTaskIds"]
    assert len(task_ids) == 1

    task_list = client.get(
        "/api/question-bank/tasks?page=1&size=20&questionId=question-seed-005",
        headers=student_headers(),
    )
    assert task_list.status_code == 200
    items = task_list.json()["data"]["items"]
    assert any(item["id"] == task_ids[0] and item["type"] == "AI_MARKING" for item in items)

    reports_before = client.get("/api/question-bank/student/papers/reports?page=1&size=5", headers=student_headers())
    assert reports_before.status_code == 200
    first_report = reports_before.json()["data"]["items"][0]
    assert first_report["reportId"] == report["reportId"]
    assert first_report["paperId"] == "paper-english-001"
    assert first_report["subjectiveMarking"]["total"] == 1
    assert first_report["subjectiveMarking"]["pendingCount"] >= 0

    task = poll_task(client, str(task_ids[0]), student_headers())
    assert task["status"] == "COMPLETED"

    reports_after = client.get("/api/question-bank/student/papers/reports?page=1&size=5", headers=student_headers())
    assert reports_after.status_code == 200
    refreshed_report = reports_after.json()["data"]["items"][0]
    assert refreshed_report["paperId"] == "paper-english-001"
    assert refreshed_report["pendingSubjectiveCount"] == 0
    assert refreshed_report["subjectiveMarking"]["pendingCount"] == 0
    assert refreshed_report["subjectiveMarking"]["averageScore"] >= 0


def test_student_submit_same_paper_same_day_is_idempotent(tmp_path: Path):
    client = make_client(tmp_path)
    payload = {
        "answers": [
            {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
            {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 40, "marked": True},
        ],
        "totalElapsedSec": 70,
    }
    first = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json=payload,
    )
    assert first.status_code == 200
    second = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json=payload,
    )
    assert second.status_code == 200
    first_data = first.json()["data"]
    second_data = second.json()["data"]
    assert second_data["paperId"] == first_data["paperId"]
    assert second_data["score"] == first_data["score"]
    assert second_data["submittedAt"] == first_data["submittedAt"]


def test_student_daily_simulation_limit_is_scoped_by_subject(tmp_path: Path):
    client = make_client(tmp_path)
    first = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 40, "marked": False},
            ],
            "totalElapsedSec": 70,
        },
    )
    assert first.status_code == 200

    create_second = client.post(
        "/api/question-bank/papers/manual",
        headers=teacher_headers(),
        json={
            "paperName": "政治第二套章节卷",
            "subjectId": "subject-politics",
            "paperType": "chapter",
            "paperStatus": "PUBLISHED",
            "durationMinutes": 20,
            "totalScore": 20,
            "visibleToStudents": True,
            "questionIds": ["question-seed-001", "question-seed-002"],
        },
    )
    assert create_second.status_code == 200
    second_paper_id = create_second.json()["data"]["paperId"]

    second_submit = client.post(
        f"/api/question-bank/student/papers/{second_paper_id}/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 40, "marked": False},
            ],
            "totalElapsedSec": 70,
        },
    )
    assert second_submit.status_code == 422
    assert second_submit.json()["code"] == "QUESTION_VALIDATION_FAILED"
    assert "每个科目每日最多完成 1 次全真模拟考试" in second_submit.json()["message"]


def test_student_can_list_recent_paper_reports(tmp_path: Path):
    client = make_client(tmp_path)
    submit = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 40, "marked": True},
            ],
            "totalElapsedSec": 70,
        },
    )
    assert submit.status_code == 200

    reports = client.get("/api/question-bank/student/papers/reports?page=1&size=10", headers=student_headers())
    assert reports.status_code == 200
    body = reports.json()["data"]
    assert body["items"]
    report = body["items"][0]
    assert report["reportId"].startswith("paper-report-")
    assert report["paperId"] == "paper-demo-001"
    assert report["subjectId"] == "subject-politics"
    assert report["score"] == 20
    assert report["pendingSubjectiveCount"] == 0
    assert report["subjectiveMarking"]["total"] == 0


def test_student_can_view_paper_report_detail(tmp_path: Path):
    client = make_client(tmp_path)
    submit = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 35, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 45, "marked": True},
            ],
            "totalElapsedSec": 80,
        },
    )
    assert submit.status_code == 200
    report_id = submit.json()["data"]["reportId"]

    detail = client.get(f"/api/question-bank/student/papers/reports/{report_id}", headers=student_headers())
    assert detail.status_code == 200
    data = detail.json()["data"]
    assert data["reportId"] == report_id
    assert data["paperId"] == "paper-demo-001"
    assert len(data["typeAccuracy"]) >= 1
    assert len(data["questionResults"]) == 2
    assert data["summary"]["questionCount"] == 2
    assert data["summary"]["markedCount"] == 1
    assert any(item["questionId"] == "question-seed-002" and item["marked"] is True for item in data["questionResults"])


def test_student_paper_report_history_keeps_multiple_records_for_same_paper(tmp_path: Path):
    client = make_client(tmp_path)
    submit = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 30, "marked": False},
            ],
            "totalElapsedSec": 60,
        },
    )
    assert submit.status_code == 200

    client.app.state.service._save_paper_report(  # noqa: SLF001
        "student-001",
        "paper-demo-001",
        {
            "reportId": "paper-report-history-seed",
            "paperId": "paper-demo-001",
            "subjectId": "subject-politics",
            "subjectIds": ["subject-politics"],
            "score": 10,
            "totalScore": 20,
            "scoreRate": 0.5,
            "totalElapsedSec": 120,
            "submittedAt": "2026-03-16T08:00:00Z",
            "typeAccuracy": [{"type": "single_choice", "accuracy": 0.5, "correct": 1, "total": 2}],
            "wrongQuestionIds": ["question-seed-002"],
            "pendingSubjectiveTaskIds": [],
            "reportDetail": {"typeAccuracy": [], "questionResults": []},
        },
    )

    reports = client.get("/api/question-bank/student/papers/reports?page=1&size=20", headers=student_headers())
    assert reports.status_code == 200
    items = [item for item in reports.json()["data"]["items"] if item["paperId"] == "paper-demo-001"]
    assert len(items) >= 2


def test_analytics_summary_supports_student_scope_and_teacher_records(tmp_path: Path):
    client = make_client(tmp_path)

    student_summary = client.get("/api/question-bank/analytics/summary", headers=student_headers())
    assert student_summary.status_code == 200
    student_data = student_summary.json()["data"]
    assert "mastery" in student_data
    assert isinstance(student_data["mastery"], list)
    assert all(item["studentUserId"] == "student-001" for item in student_data["mastery"])

    summary = client.get("/api/question-bank/analytics/summary", headers=teacher_headers())
    assert summary.status_code == 200
    data = summary.json()["data"]
    assert "mastery" in data
    assert isinstance(data["mastery"], list)

    records = client.get("/api/question-bank/analytics/records", headers=teacher_headers())
    assert records.status_code == 200
    items = records.json()["data"]["items"]
    assert items
    ext = json.loads(items[0]["extJson"])
    assert "analytics" in ext


def test_analytics_records_total_is_stable_across_page_sizes(tmp_path: Path):
    client = make_client(tmp_path)

    size_one = client.get(
        "/api/question-bank/analytics/records",
        headers=teacher_headers(),
        params={"page": 1, "size": 1},
    )
    assert size_one.status_code == 200
    size_one_data = size_one.json()["data"]
    assert size_one_data["size"] == 1
    assert len(size_one_data["items"]) <= 1

    size_two = client.get(
        "/api/question-bank/analytics/records",
        headers=teacher_headers(),
        params={"page": 1, "size": 2},
    )
    assert size_two.status_code == 200
    size_two_data = size_two.json()["data"]
    assert size_two_data["size"] == 2
    assert len(size_two_data["items"]) <= 2
    assert size_one_data["total"] == size_two_data["total"]


def test_generate_adaptive_practice_returns_unique_question_ids(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service

    result = service.generate_adaptive_practice({"count": 6}, Actor(role="student", user_id="student-001"))
    question_ids = result.get("questionIds", [])

    assert isinstance(question_ids, list)
    assert 0 < len(question_ids) <= 6
    assert len(question_ids) == len(set(question_ids))

    selected_questions = service.repository.list_questions_by_ids(question_ids)
    assert len(selected_questions) == len(question_ids)
    assert all(str(item.get("status", "")) == "PUBLISHED" for item in selected_questions)


def test_generate_adaptive_practice_supports_preferred_knowledge_id(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service

    result = service.generate_adaptive_practice(
        {"count": 6, "knowledgeId": INFO_TECH_POINT_ID},
        Actor(role="student", user_id="student-001"),
    )
    question_ids = result.get("questionIds", [])

    assert isinstance(question_ids, list)
    assert question_ids
    selected_questions = service.repository.list_questions_by_ids(question_ids)
    assert selected_questions
    assert any(str(item.get("knowledgeId", "")).strip() == INFO_TECH_POINT_ID for item in selected_questions)


def test_adaptive_gap_allocation_prefers_weakest_knowledge(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service

    target_knowledge_ids = [
        ADVANCED_MATH_POINT_ID,
        ENGLISH_POINT_ID,
        POLITICS_POINT_ID,
    ]
    mastery_snapshot = [
        {"knowledgeId": ADVANCED_MATH_POINT_ID, "mastery": 0.3435},
        {"knowledgeId": ENGLISH_POINT_ID, "mastery": 0.8718},
        {"knowledgeId": POLITICS_POINT_ID, "mastery": 0.8938},
    ]

    quota_map = service._allocate_adaptive_question_counts(target_knowledge_ids, mastery_snapshot, 10)

    assert sum(quota_map.values()) == 10
    assert quota_map[ADVANCED_MATH_POINT_ID] == 8
    assert quota_map[ENGLISH_POINT_ID] == 1
    assert quota_map[POLITICS_POINT_ID] == 1


def test_adaptive_question_selection_respects_gap_quota(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service

    target_knowledge_ids = ["knowledge-math", "knowledge-english", "knowledge-politics"]
    quota_map = {"knowledge-math": 8, "knowledge-english": 1, "knowledge-politics": 1}
    difficulty_cycle = ["medium", "easy", "hard"]

    def build_question(question_id: str, knowledge_id: str, difficulty: str) -> dict[str, str]:
        return {
            "id": question_id,
            "knowledgeId": knowledge_id,
            "extJson": json.dumps({"difficulty": difficulty}),
        }

    question_pool = [
        *[build_question(f"math-{index:02d}", "knowledge-math", "medium") for index in range(1, 11)],
        *[build_question(f"english-{index:02d}", "knowledge-english", "easy") for index in range(1, 6)],
        *[build_question(f"politics-{index:02d}", "knowledge-politics", "hard") for index in range(1, 6)],
    ]

    selected_ids = service._select_adaptive_question_ids(
        question_pool,
        target_knowledge_ids,
        difficulty_cycle,
        10,
        quota_map,
    )
    selected_questions = [item for item in question_pool if item["id"] in set(selected_ids)]
    selected_count_by_knowledge: dict[str, int] = {}
    for question in selected_questions:
        knowledge_id = str(question.get("knowledgeId", ""))
        selected_count_by_knowledge[knowledge_id] = selected_count_by_knowledge.get(knowledge_id, 0) + 1

    assert len(selected_ids) == 10
    assert len(selected_ids) == len(set(selected_ids))
    assert selected_count_by_knowledge.get("knowledge-math", 0) >= 8
    assert selected_count_by_knowledge.get("knowledge-english", 0) >= 1
    assert selected_count_by_knowledge.get("knowledge-politics", 0) >= 1


def test_new_pages_keep_fixed_layout_copy(tmp_path: Path):
    client = make_client(tmp_path)
    teacher_home_page = client.get("/teacher/home", params={"role": "teacher", "userId": "teacher-001"})
    assert teacher_home_page.status_code == 200
    teacher_home_payload = parse_page_bootstrap(teacher_home_page.text)
    assert teacher_home_payload["route"] == "/teacher/home"
    assert teacher_home_payload["viewKey"] == "teacher-home"
    assert teacher_home_payload["pageTitle"] == "教师工作台"
    assert teacher_home_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}

    home_page = client.get("/student/home", params={"role": "student", "userId": "student-001"})
    assert home_page.status_code == 200
    home_payload = parse_page_bootstrap(home_page.text)
    assert home_payload["route"] == "/student/home"
    assert home_payload["viewKey"] == "student-home"
    assert home_payload["pageTitle"] == "专属学习台"
    assert home_payload["actor"] == {"role": "student", "userId": "student-001"}

    student_page = client.get("/student/practice", params={"role": "student", "userId": "student-001"})
    assert student_page.status_code == 200
    practice_payload = parse_page_bootstrap(student_page.text)
    assert practice_payload["route"] == "/student/practice"
    assert practice_payload["viewKey"] == "student-practice"
    assert practice_payload["pageTitle"] == "学生端刷题页"
    assert practice_payload["actor"] == {"role": "student", "userId": "student-001"}

    chapter_page = client.get("/student/practice/chapter", params={"role": "student", "userId": "student-001"})
    assert chapter_page.status_code == 200
    chapter_payload = parse_page_bootstrap(chapter_page.text)
    assert chapter_payload["route"] == "/student/practice/chapter"
    assert chapter_payload["viewKey"] == "student-practice"
    assert chapter_payload["pageTitle"] == "章节闯关"
    assert chapter_payload["actor"] == {"role": "student", "userId": "student-001"}

    free_page = client.get("/student/practice/free", params={"role": "student", "userId": "student-001"})
    assert free_page.status_code == 200
    free_payload = parse_page_bootstrap(free_page.text)
    assert free_payload["route"] == "/student/practice/free"
    assert free_payload["viewKey"] == "student-practice"
    assert free_payload["pageTitle"] == "自由练习"
    assert free_payload["actor"] == {"role": "student", "userId": "student-001"}

    mock_page = client.get("/student/practice/mock", params={"role": "student", "userId": "student-001"})
    assert mock_page.status_code == 200
    mock_payload = parse_page_bootstrap(mock_page.text)
    assert mock_payload["route"] == "/student/practice/mock"
    assert mock_payload["viewKey"] == "student-practice"
    assert mock_payload["pageTitle"] == "模拟考试"
    assert mock_payload["actor"] == {"role": "student", "userId": "student-001"}

    wrong_book_page = client.get("/student/wrong-book", params={"role": "student", "userId": "student-001"}, follow_redirects=False)
    assert wrong_book_page.status_code == 307
    assert wrong_book_page.headers["location"].endswith("/student/question-bank/repair?role=student&userId=student-001")

    personal_bank_page = client.get("/student/personal-bank", params={"role": "student", "userId": "student-001"}, follow_redirects=False)
    assert personal_bank_page.status_code == 307
    assert personal_bank_page.headers["location"].endswith("/student/question-bank/archive?role=student&userId=student-001")

    question_bank_guide_page = client.get("/student/question-bank/guide", params={"role": "student", "userId": "student-001"})
    assert question_bank_guide_page.status_code == 200
    question_bank_guide_payload = parse_page_bootstrap(question_bank_guide_page.text)
    assert question_bank_guide_payload["route"] == "/student/question-bank/guide"
    assert question_bank_guide_payload["viewKey"] == "student-question-bank-guide"
    assert question_bank_guide_payload["pageTitle"] == "使用文档"
    assert question_bank_guide_payload["actor"] == {"role": "student", "userId": "student-001"}

    analysis_page = client.get("/student/analysis", params={"role": "student", "userId": "student-001"})
    assert analysis_page.status_code == 200
    analysis_payload = parse_page_bootstrap(analysis_page.text)
    assert analysis_payload["route"] == "/student/analysis"
    assert analysis_payload["viewKey"] == "student-analysis"
    assert analysis_payload["pageTitle"] == "知识诊断"
    assert analysis_payload["actor"] == {"role": "student", "userId": "student-001"}

    analysis_overview_page = client.get("/student/analysis/overview", params={"role": "student", "userId": "student-001"})
    assert analysis_overview_page.status_code == 200
    analysis_overview_payload = parse_page_bootstrap(analysis_overview_page.text)
    assert analysis_overview_payload["route"] == "/student/analysis/overview"
    assert analysis_overview_payload["viewKey"] == "student-analysis"
    assert analysis_overview_payload["pageTitle"] == "知识诊断"
    assert analysis_overview_payload["actor"] == {"role": "student", "userId": "student-001"}

    analysis_tasks_page = client.get("/student/analysis/tasks", params={"role": "student", "userId": "student-001"})
    assert analysis_tasks_page.status_code == 200
    analysis_tasks_payload = parse_page_bootstrap(analysis_tasks_page.text)
    assert analysis_tasks_payload["route"] == "/student/analysis/tasks"
    assert analysis_tasks_payload["viewKey"] == "student-analysis"
    assert analysis_tasks_payload["pageTitle"] == "知识诊断今日任务"
    assert analysis_tasks_payload["actor"] == {"role": "student", "userId": "student-001"}

    analysis_points_page = client.get("/student/analysis/points", params={"role": "student", "userId": "student-001"})
    assert analysis_points_page.status_code == 200
    analysis_points_payload = parse_page_bootstrap(analysis_points_page.text)
    assert analysis_points_payload["route"] == "/student/analysis/points"
    assert analysis_points_payload["viewKey"] == "student-analysis"
    assert analysis_points_payload["pageTitle"] == "练习积分"
    assert analysis_points_payload["actor"] == {"role": "student", "userId": "student-001"}

    legacy_points_page = client.get("/student/points", params={"role": "student", "userId": "student-001"}, follow_redirects=False)
    assert legacy_points_page.status_code == 307
    assert legacy_points_page.headers["location"].endswith("/student/analysis/points?role=student&userId=student-001")

    paper_page = client.get("/teacher/papers", params={"role": "teacher", "userId": "teacher-001"})
    assert paper_page.status_code == 200
    paper_payload = parse_page_bootstrap(paper_page.text)
    assert paper_payload["route"] == "/teacher/papers"
    assert paper_payload["viewKey"] == "teacher-papers"
    assert paper_payload["pageTitle"] == "试卷管理"
    assert paper_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}

    analytics_page = client.get("/teacher/analytics", params={"role": "teacher", "userId": "teacher-001"})
    assert analytics_page.status_code == 200
    analytics_payload = parse_page_bootstrap(analytics_page.text)
    assert analytics_payload["route"] == "/teacher/analytics"
    assert analytics_payload["viewKey"] == "teacher-analytics"
    assert analytics_payload["pageTitle"] == "学情页"
    assert analytics_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}

    import_history_page = client.get(
        "/teacher/import-history",
        params={"role": "teacher", "userId": "teacher-001"},
        follow_redirects=False,
    )
    assert import_history_page.status_code == 307
    assert import_history_page.headers["location"].endswith("/teacher/questions?role=teacher&userId=teacher-001#import-history")

    import_history_detail_page = client.get(
        "/teacher/import-history/task/task-demo-001",
        params={"role": "teacher", "userId": "teacher-001"},
    )
    assert import_history_detail_page.status_code == 200
    import_history_detail_payload = parse_page_bootstrap(import_history_detail_page.text)
    assert import_history_detail_payload["route"] == "/teacher/import-history/task/task-demo-001"
    assert import_history_detail_payload["viewKey"] == "teacher-import-history-detail"
    assert import_history_detail_payload["pageTitle"] == "导入任务详情"
    assert import_history_detail_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}

    teacher_guide_page = client.get("/teacher/guide", params={"role": "teacher", "userId": "teacher-001"})
    assert teacher_guide_page.status_code == 200
    teacher_guide_payload = parse_page_bootstrap(teacher_guide_page.text)
    assert teacher_guide_payload["route"] == "/teacher/guide"
    assert teacher_guide_payload["viewKey"] == "teacher-guide"
    assert teacher_guide_payload["pageTitle"] == "使用文档"
    assert teacher_guide_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}


def test_teacher_with_student_manage_permission_can_manage_student_accounts(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    bind_teacher_scope = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-001",
            "role": "teacher",
            "name": "教师A",
            "mobile": "13800000002",
            "enabled": True,
            "permissions": ["question:manage", "paper:manage", "analytics:view", "student:manage", "message:send"],
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert bind_teacher_scope.status_code == 200

    teacher_headers_auth = teacher_auth_headers(client)
    student_account_page = client.get("/teacher/student-accounts", headers=teacher_headers_auth)
    assert student_account_page.status_code == 200
    page_payload = parse_page_bootstrap(student_account_page.text)
    assert page_payload["route"] == "/teacher/student-accounts"
    assert page_payload["viewKey"] == "teacher-student-accounts"
    assert page_payload["pageTitle"] == "学生账号开通"
    assert page_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}

    scoped_students = client.get(
        "/api/question-bank/admin/users?page=1&size=50&role=student",
        headers=teacher_headers_auth,
    )
    assert scoped_students.status_code == 200
    scoped_items = scoped_students.json()["data"]["items"]
    assert scoped_items
    assert {str(item["role"]) for item in scoped_items} == {"student"}
    assert {str(item["jointExamGroupCode"]) for item in scoped_items} == {"SCIENCE_ENGINEERING_3"}

    create_student = client.post(
        "/api/question-bank/admin/users",
        headers=teacher_headers_auth,
        json={
            "userId": "student-scope-001",
            "role": "student",
            "name": "理工新生",
            "mobile": "13800000081",
            "enabled": True,
            "permissions": [],
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "vocationalMajor": "计算机类",
            "prepStage": "基础阶段",
        },
    )
    assert create_student.status_code == 200

    reject_cross_scope = client.post(
        "/api/question-bank/admin/users",
        headers=teacher_headers_auth,
        json={
            "userId": "student-cross-001",
            "role": "student",
            "name": "跨组学生",
            "mobile": "13800000082",
            "enabled": True,
            "permissions": [],
            "examCategoryCode": "LITERATURE",
            "jointExamGroupCode": "LITERATURE_11",
            "vocationalMajor": "汉语言",
            "prepStage": "基础阶段",
        },
    )
    assert reject_cross_scope.status_code == 403
    assert any(
        keyword in reject_cross_scope.json()["message"]
        for keyword in ("学科门类", "联考专业组")
    )

    export_students = client.get(
        "/api/question-bank/admin/students/export?format=csv",
        headers=teacher_headers_auth,
    )
    assert export_students.status_code == 200
    export_content = export_students.json()["data"]["content"]
    assert "SCIENCE_ENGINEERING_3" in export_content
    assert "LITERATURE_11" not in export_content


def test_student_dashboard_profile_and_check_in(tmp_path: Path):
    client = make_client(tmp_path)

    def assert_profile_groups_are_scoped(dashboard_data: dict[str, object]) -> None:
        selected_exam_category_code = str(dashboard_data.get("examCategoryCode", ""))
        available_categories = dashboard_data.get("availableExamCategories", [])
        category_group_codes = {
            str(item.get("examCategoryCode", "")): {
                str(group.get("jointExamGroupCode", ""))
                for group in item.get("jointExamGroups", [])
                if isinstance(group, dict)
            }
            for item in available_categories
            if isinstance(item, dict)
        }
        available_joint_exam_groups = dashboard_data.get("availableJointExamGroups", [])
        returned_group_codes = {
            str(item.get("jointExamGroupCode", ""))
            for item in available_joint_exam_groups
            if isinstance(item, dict)
        }
        assert returned_group_codes <= category_group_codes.get(selected_exam_category_code, set())
        assert all(
            str(item.get("examCategoryCode", "")) == selected_exam_category_code
            for item in available_joint_exam_groups
            if isinstance(item, dict)
        )

    dashboard = client.get("/api/question-bank/student/dashboard", headers=student_headers())
    assert dashboard.status_code == 200
    data = dashboard.json()["data"]
    assert_profile_groups_are_scoped(data)
    assert len(data["coreSubjects"]) == 4
    assert "aiQuota" in data
    assert set(data["aiQuota"].keys()) == {"dailyLimit", "usedCount"}
    assert "studentState" in data
    assert "chapterPracticeTree" in data
    assert isinstance(data["chapterPracticeTree"], list)
    assert "checkInDates" not in data
    assert len(data["dailyTasks"]) >= 3
    first_task = data["dailyTasks"][0]
    assert {"taskKey", "actionPath", "actionQuery", "actionLabel", "isDone"}.issubset(first_task.keys())
    assert first_task["actionPath"].startswith("/student/")
    assert isinstance(first_task["actionQuery"], dict)
    assert first_task["actionPath"] == "/student/practice/chapter"
    assert first_task["actionQuery"]["module"] == "chapter"

    profile = client.post(
        "/api/question-bank/student/profile",
        headers=student_headers(),
        json={"examCategoryCode": "LITERATURE", "jointExamGroupCode": "LITERATURE_11"},
    )
    assert profile.status_code == 200
    assert profile.json()["data"]["examCategoryCode"] == "LITERATURE"
    assert profile.json()["data"]["jointExamGroupCode"] == "LITERATURE_11"
    assert_profile_groups_are_scoped(profile.json()["data"])

    check_in = client.post("/api/question-bank/student/check-in", headers=student_headers())
    assert check_in.status_code == 200
    assert check_in.json()["data"]["points"] >= 2


def test_student_profile_rejects_joint_exam_group_outside_selected_exam_category(tmp_path: Path):
    client = make_client(tmp_path)
    profile = client.post(
        "/api/question-bank/student/profile",
        headers=student_headers(),
        json={"examCategoryCode": "LITERATURE", "jointExamGroupCode": "SCIENCE_ENGINEERING_3"},
    )
    assert profile.status_code == 422
    assert "jointExamGroupCode 与 examCategoryCode 不匹配。" in profile.json()["message"]


def test_student_analysis_tasks_page_contract_aligns_global_tasks_with_subject_weakness_tree(tmp_path: Path):
    client = make_client(tmp_path)

    politics_page = client.get(
        "/student/analysis/tasks",
        params={"role": "student", "userId": "student-001", "subjectCode": "POLITICS"},
    )
    english_page = client.get(
        "/student/analysis/tasks",
        params={"role": "student", "userId": "student-001", "subjectCode": "ENGLISH"},
    )
    assert politics_page.status_code == 200
    assert english_page.status_code == 200

    politics_payload = parse_page_bootstrap(politics_page.text)
    english_payload = parse_page_bootstrap(english_page.text)
    assert politics_payload["route"] == "/student/analysis/tasks"
    assert english_payload["route"] == "/student/analysis/tasks"
    assert politics_payload["viewKey"] == "student-analysis"
    assert english_payload["viewKey"] == "student-analysis"
    assert politics_payload["pageTitle"] == "知识诊断今日任务"
    assert english_payload["pageTitle"] == "知识诊断今日任务"
    assert politics_payload["actor"] == {"role": "student", "userId": "student-001"}
    assert english_payload["actor"] == {"role": "student", "userId": "student-001"}

    dashboard = client.get("/api/question-bank/student/dashboard", headers=student_headers())
    assert dashboard.status_code == 200
    dashboard_data = dashboard.json()["data"]
    assert isinstance(dashboard_data.get("dailyTasks"), list)
    assert len(dashboard_data["dailyTasks"]) >= 3
    assert all(
        {"taskKey", "actionPath", "actionQuery", "actionLabel", "isDone"}.issubset(task.keys())
        for task in dashboard_data["dailyTasks"]
        if isinstance(task, dict)
    )
    wrong_book_task = next(task for task in dashboard_data["dailyTasks"] if task.get("taskKey") == "wrongBookReward")
    assert wrong_book_task["actionPath"] == "/student/question-bank/repair"
    assert wrong_book_task["actionQuery"] == {}

    task_list = client.get("/api/question-bank/tasks?page=1&size=20", headers=student_headers())
    assert task_list.status_code == 200
    task_list_data = task_list.json()["data"]
    assert isinstance(task_list_data.get("items"), list)
    assert "total" in task_list_data
    assert "aiQuota" in task_list_data
    assert set(task_list_data["aiQuota"].keys()) == {"dailyLimit", "usedCount"}

    politics_tree = client.get(
        "/api/question-bank/knowledge/tree",
        headers=student_headers(),
        params={"subjectCode": "POLITICS"},
    )
    english_tree = client.get(
        "/api/question-bank/knowledge/tree",
        headers=student_headers(),
        params={"subjectCode": "ENGLISH"},
    )
    assert politics_tree.status_code == 200
    assert english_tree.status_code == 200

    politics_node_ids = {str(item.get("id", "")) for item in politics_tree.json()["data"]["nodes"]}
    english_node_ids = {str(item.get("id", "")) for item in english_tree.json()["data"]["nodes"]}
    assert POLITICS_POINT_ID in politics_node_ids
    assert ENGLISH_POINT_ID not in politics_node_ids
    assert ENGLISH_POINT_ID in english_node_ids
    assert POLITICS_POINT_ID not in english_node_ids


def test_exam_task_student_ids_prefers_formal_profile_scope_for_assignment_class_meta(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)
    repo = client.app.state.service.repository

    profile_saved = client.post(
        "/api/question-bank/student/profile",
        headers=student_headers("student-001"),
        json={"examCategoryCode": "LITERATURE", "jointExamGroupCode": "LITERATURE_11"},
    )
    assert profile_saved.status_code == 200

    user = repo.get_user_by_id("student-001")
    assert user is not None
    user_ext = json.loads(user["extJson"])
    user_ext["jointExamGroupCode"] = "SCIENCE_ENGINEERING_3"
    with get_connection(repo.db_path) as connection:
        connection.execute(
            """
            UPDATE user
            SET extJson = ?
            WHERE id = ?
            """,
            (json.dumps(user_ext, ensure_ascii=False), "student-001"),
        )
        for question_id in ("question-seed-001", "question-seed-002", "question-seed-std-politics-001"):
            row = connection.execute("SELECT extJson FROM question WHERE id = ?", (question_id,)).fetchone()
            if row:
                ext_json = json.loads(row[0])
                if "chapterCode" not in ext_json and "chapter_code" in ext_json:
                    ext_json["chapterCode"] = ext_json.get("chapter_code", "")
                if "pointCode" not in ext_json and "point_code" in ext_json:
                    ext_json["pointCode"] = ext_json.get("point_code", "")
                connection.execute(
                    "UPDATE question SET extJson = ? WHERE id = ?",
                    (json.dumps(ext_json, ensure_ascii=False), question_id),
                )
        connection.commit()

    created = client.post(
        "/api/question-bank/exam-tasks",
        headers=admin_headers,
        json={
            "taskName": "正式画像班级归属任务",
            "taskType": "CHAPTER",
            "subjectCode": "POLITICS",
            "examCategoryCode": "LITERATURE",
            "jointExamGroupCode": "LITERATURE_11",
            "studentIds": ["student-001"],
            "targetQuestionCount": 5,
            "status": "PUBLISHED",
        },
    )
    assert created.status_code == 200
    task_id = created.json()["data"]["id"]

    detail = client.get(f"/api/question-bank/exam-tasks/{task_id}", headers=admin_headers)
    assert detail.status_code == 200
    assignments = detail.json()["data"]["assignments"]
    assert len(assignments) == 1
    assert assignments[0]["classId"] == "LITERATURE_11"
    assert assignments[0]["className"]


def test_teacher_qa_thread_prefers_formal_profile_scope_when_user_ext_is_stale(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository

    profile_saved = client.post(
        "/api/question-bank/student/profile",
        headers=student_headers("student-001"),
        json={"examCategoryCode": "LITERATURE", "jointExamGroupCode": "LITERATURE_11"},
    )
    assert profile_saved.status_code == 200

    user = repo.get_user_by_id("student-001")
    assert user is not None
    user_ext = json.loads(user["extJson"])
    user_ext["examCategoryCode"] = "SCIENCE_ENGINEERING"
    user_ext["jointExamGroupCode"] = "SCIENCE_ENGINEERING_3"
    user_ext["vocationalMajor"] = "旧专业"
    with get_connection(repo.db_path) as connection:
        connection.execute(
            """
            UPDATE user
            SET extJson = ?
            WHERE id = ?
            """,
            (json.dumps(user_ext, ensure_ascii=False), "student-001"),
        )
        connection.commit()

    created = client.post(
        "/api/question-bank/messages/teacher-qa/threads",
        headers=student_headers("student-001"),
        data={
            "subjectCode": "POLITICS",
            "title": "正式画像答疑范围",
            "content": "确认老师答疑线程写入的是正式画像范围。",
        },
    )
    assert created.status_code == 200
    thread = created.json()["data"]
    assert thread["examCategoryCode"] == "LITERATURE"
    assert thread["jointExamGroupCode"] == "LITERATURE_11"
    assert thread["vocationalMajor"] != "旧专业"


def test_content_baseline_api_and_page_are_aligned(tmp_path: Path):
    client = make_client(tmp_path)

    baseline = client.get("/api/question-bank/content/baseline", headers=teacher_headers())
    assert baseline.status_code == 200
    data = baseline.json()["data"]
    assert data["policyVersionCode"] == "HB_ZSB_2026"
    assert len(data["examCategories"]) == 10
    joint_group_count = sum(len(item.get("jointExamGroups", [])) for item in data["examCategories"])
    assert joint_group_count == 40
    all_majors = "、".join(
        str(group.get("majorListText", ""))
        for category in data["examCategories"]
        for group in category.get("jointExamGroups", [])
    )
    assert "水产养殖学" not in all_majors
    assert data["examCategories"][2]["jointExamGroups"]
    assert data["examCategories"][2]["jointExamGroups"][0]["subjects"]

    page = client.get("/teacher/content-system", params={"role": "teacher", "userId": "teacher-001"})
    assert page.status_code == 200
    page_payload = parse_page_bootstrap(page.text)
    assert page_payload["route"] == "/teacher/content-system"
    assert page_payload["viewKey"] == "teacher-content-system"
    assert page_payload["pageTitle"] == "内容体系字典"
    assert page_payload["actor"] == {"role": "teacher", "userId": "teacher-001"}


def test_professional_tree_api_returns_recursive_hierarchy_with_public_subjects(tmp_path: Path):
    client = make_client(tmp_path)

    response = client.get("/api/question-bank/professional-tree", headers=teacher_headers())
    assert response.status_code == 200
    exam_categories = response.json()["data"]
    assert len(exam_categories) == 10
    total_groups = sum(len(item.get("children", [])) for item in exam_categories)
    assert total_groups == 40

    for category in exam_categories:
        groups = category.get("children", [])
        assert groups
        category_subject_codes = {
            str(subject.get("code", ""))
            for group in groups
            for subject in group.get("children", [])
        }
        assert "POLITICS" in category_subject_codes
        assert "ENGLISH" in category_subject_codes
        for group in groups:
            group_subjects = group.get("children", [])
            assert group_subjects
            for subject in group_subjects:
                assert str(subject.get("subjectType", "")) in {"PUBLIC", "PROFESSIONAL"}
                if str(subject.get("subjectType", "")) == "PUBLIC":
                    assert str(subject.get("subjectSlot", "")) == "PUBLIC"
                else:
                    assert str(subject.get("subjectSlot", "")).startswith("PROFESSIONAL")

    science_engineering = next(item for item in exam_categories if item.get("code") == "SCIENCE_ENGINEERING")
    science_group_3 = next(item for item in science_engineering["children"] if item.get("code") == "SCIENCE_ENGINEERING_3")
    science_subject_names = [str(subject.get("name", "")) for subject in science_group_3["children"]]
    assert "信息技术概论" in science_subject_names
    assert "高等数学（一）" in science_subject_names


def test_content_filters_and_student_scope_follow_joint_exam_group(tmp_path: Path):
    client = make_client(tmp_path)

    teacher_list = client.get(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-002"),
        params={
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_4",
            "subjectCode": "INFO_TECH_INTRO",
        },
    )
    assert teacher_list.status_code == 200
    teacher_items = teacher_list.json()["data"]["items"]
    assert [item["id"] for item in teacher_items] == ["question-seed-008"]

    visible_for_science = client.get(
        "/api/question-bank/student/practice/questions",
        headers=student_headers(),
        params={"subjectCode": "INFO_TECH_INTRO", "module": "free"},
    )
    assert visible_for_science.status_code == 200
    visible_ids = [item["id"] for item in visible_for_science.json()["data"]["items"]]
    assert "question-seed-008" in visible_ids

    profile = client.post(
        "/api/question-bank/student/profile",
        headers=student_headers(),
        json={"examCategoryCode": "LITERATURE", "jointExamGroupCode": "LITERATURE_11"},
    )
    assert profile.status_code == 200

    hidden_for_literature = client.get(
        "/api/question-bank/student/practice/questions",
        headers=student_headers(),
        params={"subjectCode": "INFO_TECH_INTRO", "module": "free"},
    )
    assert hidden_for_literature.status_code == 200
    assert hidden_for_literature.json()["data"]["items"] == []


def test_wrong_book_and_personalized_paper_generation(tmp_path: Path):
    client = make_client(tmp_path)
    wrong_book = client.get("/api/question-bank/student/wrong-book/questions", headers=student_headers())
    assert wrong_book.status_code == 200
    items = wrong_book.json()["data"]["items"]
    assert items
    assert tuple(items[0].keys()) == QUESTION_FIELDS

    paper = client.post("/api/question-bank/student/wrong-book/papers", headers=student_headers())
    assert paper.status_code == 200
    paper_id = paper.json()["data"]["paperId"]

    available = client.get("/api/question-bank/student/papers/questions", headers=student_headers())
    assert available.status_code == 200
    assert any(paper_id in item["extJson"] for item in available.json()["data"]["items"])


def test_student_practice_and_wrong_book_support_chapter_code_and_point_code_filters(tmp_path: Path):
    client = make_client(tmp_path)

    question_payload = payload()
    question_payload["status"] = "PUBLISHED"
    question_payload["examCategoryCode"] = "SCIENCE_ENGINEERING"
    question_payload["jointExamGroupCode"] = "SCIENCE_ENGINEERING_3"
    question_payload["subjectCode"] = "INFO_TECH_INTRO"
    question_payload["knowledgePoints"] = [INFO_TECH_POINT_ID]
    question_payload["extJson"] = {"subjectId": "subject-computer"}
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers(),
        json=question_payload,
    )
    assert created.status_code == 200
    practice_target = created.json()["data"]
    practice_ext = practice_target["extJson"] if isinstance(practice_target["extJson"], dict) else json.loads(practice_target["extJson"])
    chapter_code = str(practice_ext.get("chapterCode", "")).strip()
    assert chapter_code

    filtered_practice = client.get(
        "/api/question-bank/student/practice/questions",
        headers=student_headers(),
        params={
            "subjectCode": str(practice_ext.get("subjectCode", "")).strip(),
            "chapterCode": chapter_code,
            "module": "free",
        },
    )
    assert filtered_practice.status_code == 200
    filtered_practice_items = filtered_practice.json()["data"]["items"]
    assert filtered_practice_items
    assert practice_target["id"] in [item["id"] for item in filtered_practice_items]
    assert all(json.loads(item["extJson"]).get("chapterCode") == chapter_code for item in filtered_practice_items)

    wrong_submit = client.post(
        f"/api/question-bank/student/practice/questions/{practice_target['id']}/submit",
        headers=student_headers(),
        json={"answer": "A", "elapsedSec": 18, "sourceType": "FREE_PRACTICE"},
    )
    assert wrong_submit.status_code == 200

    wrong_book = client.get(
        "/api/question-bank/student/wrong-book/questions",
        headers=student_headers(),
        params={"subjectCode": str(practice_ext.get("subjectCode", "")).strip()},
    )
    assert wrong_book.status_code == 200
    wrong_items = wrong_book.json()["data"]["items"]
    assert wrong_items
    wrong_target = next(item for item in wrong_items if item["id"] == practice_target["id"])
    wrong_ext = json.loads(wrong_target["extJson"])
    point_code = str(wrong_ext.get("pointCode", "")).strip()
    assert point_code

    filtered_wrong_book = client.get(
        "/api/question-bank/student/wrong-book/questions",
        headers=student_headers(),
        params={
            "subjectCode": str(wrong_ext.get("subjectCode", "")).strip(),
            "pointCode": point_code,
        },
    )
    assert filtered_wrong_book.status_code == 200
    filtered_wrong_items = filtered_wrong_book.json()["data"]["items"]
    assert filtered_wrong_items
    assert wrong_target["id"] in [item["id"] for item in filtered_wrong_items]
    assert all(json.loads(item["extJson"]).get("pointCode") == point_code for item in filtered_wrong_items)

    filtered_wrong_summary = client.get(
        "/api/question-bank/student/wrong-book/summary",
        headers=student_headers(),
        params={
            "subjectCode": str(wrong_ext.get("subjectCode", "")).strip(),
            "pointCode": point_code,
        },
    )
    assert filtered_wrong_summary.status_code == 200
    filtered_summary_items = filtered_wrong_summary.json()["data"]["questionInsights"]
    assert filtered_summary_items
    assert wrong_target["id"] in [item["questionId"] for item in filtered_summary_items]
    assert all(str(item.get("pointCode", "")).strip() == point_code for item in filtered_summary_items)

    knowledge_tree = client.get(
        "/api/question-bank/knowledge/tree",
        headers=student_headers(),
        params={"subjectCode": str(wrong_ext.get("subjectCode", "")).strip()},
    )
    assert knowledge_tree.status_code == 200
    tree_data = knowledge_tree.json()["data"]
    node_map = {str(item["id"]): item for item in tree_data["nodes"]}
    parent_by_id = {
        str(link["target"]): str(link["source"])
        for link in tree_data["links"]
        if str(link.get("type", "")) == "parent"
    }
    path_cursor = str(wrong_target.get("knowledgeId", "")).strip()
    l3_ancestor_id = ""
    while path_cursor:
        current_node = node_map.get(path_cursor, {})
        if int(current_node.get("level", 0) or 0) == 3:
            l3_ancestor_id = path_cursor
            break
        path_cursor = parent_by_id.get(path_cursor, "")
    assert l3_ancestor_id

    filtered_wrong_book_by_path = client.get(
        "/api/question-bank/student/wrong-book/questions",
        headers=student_headers(),
        params={
            "subjectCode": str(wrong_ext.get("subjectCode", "")).strip(),
            "knowledgePathNodeId": l3_ancestor_id,
        },
    )
    assert filtered_wrong_book_by_path.status_code == 200
    filtered_wrong_path_items = filtered_wrong_book_by_path.json()["data"]["items"]
    assert filtered_wrong_path_items
    assert wrong_target["id"] in [item["id"] for item in filtered_wrong_path_items]


def test_wrong_book_reasoned_paper_generation(tmp_path: Path):
    client = make_client(tmp_path)
    wrong_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "A", "elapsedSec": 30},
    )
    assert wrong_submit.status_code == 200

    reasoned = client.post(
        "/api/question-bank/student/wrong-book/papers/reasoned?reasonCode=KNOWLEDGE_GAP&questionCount=10",
        headers=student_headers(),
    )
    assert reasoned.status_code == 200
    data = reasoned.json()["data"]
    assert data["paperId"].startswith("paper-reason-")
    assert data["reasonCode"] == "KNOWLEDGE_GAP"


def test_wrong_book_personalized_paper_generation_respects_subject_scope(tmp_path: Path):
    client = make_client(tmp_path)
    paper = client.post(
        "/api/question-bank/student/wrong-book/papers",
        headers=student_headers(),
        params={"subjectCode": "POLITICS"},
    )
    assert paper.status_code == 200
    data = paper.json()["data"]
    paper_id = data["paperId"]

    paper_questions = client.get(
        f"/api/question-bank/student/papers/{paper_id}/questions",
        headers=student_headers(),
    )
    assert paper_questions.status_code == 200
    paper_items = paper_questions.json()["data"]["items"]
    assert paper_items
    assert all(json.loads(item["extJson"]).get("subjectCode") == "POLITICS" for item in paper_items)
    assert len(data["questionIds"]) == len(paper_items)

    available = client.get("/api/question-bank/student/papers/questions", headers=student_headers())
    assert available.status_code == 200
    assert any(data["paperId"] in item["extJson"] for item in available.json()["data"]["items"])


def test_student_error_book_summary_enforces_joint_group_and_returns_repair_signals(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository
    record = repo.get_student_question_bank("question-seed-002", "student-001")
    assert record is not None
    record_ext = json.loads(record["extJson"])
    record_ext["wrongBook"] = {
        "isCollected": True,
        "wrongCount": 4,
        "collectedAt": "2026-03-15T08:00:00Z",
        "reasonStats": [
            {"reasonCode": "KNOWLEDGE_GAP", "reasonLabel": "知识点掌握不牢", "count": 4},
        ],
    }
    repo.upsert_student_question_bank(
        {
            **record,
            "extJson": json.dumps(record_ext, ensure_ascii=False),
        }
    )

    forbidden_summary = client.get(
        "/api/student/error-book/summary",
        headers={**student_headers(), "X-Joint-Group": "MANAGEMENT_1"},
        params={"subjectCode": "POLITICS"},
    )
    assert forbidden_summary.status_code == 403

    summary = client.get(
        "/api/student/error-book/summary",
        headers={**student_headers(), "X-Joint-Group": "SCIENCE_ENGINEERING_3"},
        params={"subjectCode": "POLITICS"},
    )
    assert summary.status_code == 200
    data = summary.json()["data"]
    assert data["policyVersionCode"] == "HB_ZSB_2026"
    assert data["jointExamGroupCode"] == "SCIENCE_ENGINEERING_3"
    assert data["selectedSubjectCode"] == "POLITICS"
    assert data["currentSubject"]["subjectCode"] == "POLITICS"
    assert data["currentSubject"]["errorCoverageRate"] >= 0
    assert data["currentSubject"]["knowledgeCoverageRate"] >= 0
    assert data["currentSubject"]["chapterInducerSuggestions"]
    assert data["currentSubject"]["aiSuggestions"]["repairSuggestions"]
    assert data["currentSubject"]["aiSuggestions"]["reviewWarnings"]
    first_question = next(item for item in data["questionInsights"] if item["subjectCode"] == "POLITICS")
    assert "reviewStatusLabel" in first_question
    assert "jointGroupAverageAccuracy" in first_question
    assert "benchmarkStatusText" in first_question
    assert "benchmarkTagType" in first_question
    assert "showBenchmarkRiskBadge" in first_question
    assert "benchmarkRiskBadgeText" in first_question
    assert "errorInducerLabel" in first_question


def test_repository_persists_hot_state_tables(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository

    seeded_row = repo.get_student_question_record_row("student-001", "question-seed-001")
    assert seeded_row is not None
    assert seeded_row["studentUserId"] == "student-001"
    assert seeded_row["questionId"] == "question-seed-001"

    record = repo.get_student_question_bank("question-seed-002", "student-001")
    assert record is not None
    record_ext = json.loads(record["extJson"])
    record_ext["chapterPractice"] = {
        "lastAnswer": "A",
        "isCorrect": False,
        "answerDurationSec": 88,
        "submitCount": 3,
        "correctCount": 1,
        "submittedAt": "2026-03-22T10:00:00Z",
    }
    record_ext["wrongBook"] = {
        "isCollected": True,
        "wrongCount": 2,
        "reviewCount": 3,
        "postWrongAttemptCount": 4,
        "postWrongCorrectCount": 1,
        "collectedAt": "2026-03-20T09:00:00Z",
        "reviewedAt": "2026-03-22T10:10:00Z",
        "lastReasonCode": "KNOWLEDGE_GAP",
        "lastReasonLabel": "知识点掌握不牢",
    }
    record_ext["personalBank"] = {"isCollected": True, "sourceType": "HARVESTED_ARCHIVE"}
    repo.upsert_student_question_bank({**record, "extJson": json.dumps(record_ext, ensure_ascii=False)})

    student_row = repo.get_student_question_record_row("student-001", "question-seed-002")
    assert student_row is not None
    assert student_row["lastAnswer"] == "A"
    assert student_row["answerCount"] == 3
    assert student_row["correctCount"] == 1
    assert student_row["wrongCount"] == 2
    assert student_row["wrongBookFlag"] == 1
    assert student_row["wrongBookReviewCount"] == 3
    assert student_row["wrongBookPostWrongAttemptCount"] == 4
    assert student_row["wrongBookPostWrongCorrectCount"] == 1
    assert student_row["wrongBookCollectedAt"] == "2026-03-20T09:00:00Z"
    assert student_row["wrongBookReviewedAt"] == "2026-03-22T10:10:00Z"
    assert student_row["wrongBookLastReasonCode"] == "KNOWLEDGE_GAP"
    assert student_row["personalBankFlag"] == 1
    assert student_row["latestSourceType"] == "HARVESTED_ARCHIVE"


def test_student_profile_hot_state_prefers_formal_tables(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)
    me = client.get("/api/question-bank/auth/me", headers=headers)
    assert me.status_code == 200
    student_user_id = me.json()["data"]["userId"]
    service = client.app.state.service
    repo = service.repository

    record, record_ext, profile = service._load_student_profile_bundle(student_user_id)
    assert dict(record_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }
    progress_date = "2026-03-23"
    ledger_item = {
        "eventKey": "practiceReward:2026-03-23",
        "reason": "完成章节刷题10道",
        "points": 20,
        "createTime": "2026-03-23T09:30:00Z",
    }
    seed_student_hot_state(
        service,
        student_user_id,
        points=88,
        title="连刷达人",
        unlocked_titles=["备考新星", "连刷达人"],
        check_in_dates=[progress_date],
        ai_quota={"dailyLimit": 25, "usedCount": 3, "quotaDate": progress_date},
        exam_session={"answeredCount": 12, "elapsedSec": 640, "updateTime": "2026-03-23T09:45:00Z"},
        daily_progress={
            progress_date: {
                "checkInCount": 1,
                "practiceAnswers": 10,
                "papersCompleted": 1,
                "wrongBookReviewed": 5,
                "rewardedKeys": ["practiceReward", "paperReward"],
            }
        },
        points_ledger=[ledger_item],
    )

    daily_progress_row = repo.get_student_daily_progress(student_user_id, progress_date)
    assert daily_progress_row is not None
    assert daily_progress_row["checkInCount"] == 1
    assert daily_progress_row["practiceAnswers"] == 10
    assert daily_progress_row["papersCompleted"] == 1
    assert daily_progress_row["wrongBookReviewed"] == 5
    assert daily_progress_row["rewardedKeys"] == ["practiceReward", "paperReward"]

    ledger_rows = repo.list_student_points_ledger(student_user_id)
    assert any(item["eventKey"] == ledger_item["eventKey"] for item in ledger_rows)
    profile_state_row = repo.get_student_profile_state(student_user_id)
    assert profile_state_row is not None
    assert profile_state_row["checkInDates"] == [progress_date]
    assert profile_state_row["points"] == 88
    assert profile_state_row["title"] == "连刷达人"
    assert profile_state_row["unlockedTitles"] == ["备考新星", "连刷达人"]
    assert profile_state_row["aiQuota"] == {"dailyLimit": 25, "usedCount": 3, "quotaDate": progress_date}
    assert profile_state_row["examSession"] == {"answeredCount": 12, "elapsedSec": 640, "updateTime": "2026-03-23T09:45:00Z"}

    record = repo.get_student_question_bank(record["questionId"], student_user_id)
    assert record is not None
    cleared_ext = json.loads(record["extJson"])
    cleared_profile = dict(cleared_ext.get("studentProfile", {}))
    assert set(cleared_profile.keys()) == {"examCategoryCode", "jointExamGroupCode"}
    cleared_profile["checkInDates"] = []
    cleared_profile["points"] = 0
    cleared_profile["title"] = "备考新星"
    cleared_profile["unlockedTitles"] = ["备考新星"]
    cleared_profile["dailyProgress"] = {}
    cleared_profile["pointsLedger"] = []
    cleared_profile["aiQuota"] = {"dailyLimit": 20, "usedCount": 0, "quotaDate": ""}
    cleared_profile["examSession"] = {"answeredCount": 0, "elapsedSec": 0, "updateTime": ""}
    cleared_ext["studentProfile"] = cleared_profile
    repo.upsert_student_question_bank({**record, "extJson": json.dumps(cleared_ext, ensure_ascii=False)})
    normalized_immediately = repo.get_student_question_bank(record["questionId"], student_user_id)
    assert normalized_immediately is not None
    normalized_immediately_ext = json.loads(normalized_immediately["extJson"])
    assert dict(normalized_immediately_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }

    _, _, reloaded_profile = service._load_student_profile_bundle(student_user_id)
    assert reloaded_profile["checkInDates"] == [progress_date]
    assert reloaded_profile["points"] == 88
    assert reloaded_profile["title"] == "连刷达人"
    assert reloaded_profile["unlockedTitles"] == ["备考新星", "连刷达人"]
    assert reloaded_profile["dailyProgress"][progress_date]["practiceAnswers"] == 10
    assert reloaded_profile["dailyProgress"][progress_date]["rewardedKeys"] == ["practiceReward", "paperReward"]
    assert any(item["eventKey"] == ledger_item["eventKey"] for item in reloaded_profile["pointsLedger"])
    assert reloaded_profile["aiQuota"] == {"dailyLimit": 25, "usedCount": 3, "quotaDate": progress_date}
    assert reloaded_profile["examSession"] == {"answeredCount": 12, "elapsedSec": 640, "updateTime": "2026-03-23T09:45:00Z"}
    normalized_record = repo.get_student_question_bank(record["questionId"], student_user_id)
    assert normalized_record is not None
    normalized_ext = json.loads(normalized_record["extJson"])
    assert dict(normalized_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }


def test_hot_state_endpoints_keep_legacy_student_profile_snapshot_thin(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)
    me = client.get("/api/question-bank/auth/me", headers=headers)
    assert me.status_code == 200
    student_user_id = me.json()["data"]["userId"]
    service = client.app.state.service

    record, _, _ = service._load_student_profile_bundle(student_user_id)  # noqa: SLF001
    check_in = client.post("/api/question-bank/student/check-in", headers=headers)
    assert check_in.status_code == 200

    submit_session = client.post(
        "/api/question-bank/student/submit",
        headers=headers,
        json={"answeredCount": 7, "elapsedSec": 320},
    )
    assert submit_session.status_code == 200
    assert submit_session.json()["data"]["answeredCount"] == 7

    ai_tutor = client.post(
        "/api/question-bank/student/practice/questions/question-seed-005/ai-tutor",
        headers=headers,
        json={"prompt": "继续讲解这道题", "promptImageUrl": ""},
    )
    assert ai_tutor.status_code == 200

    refreshed_record = service.repository.get_student_question_bank(record["questionId"], student_user_id)
    assert refreshed_record is not None
    refreshed_ext = json.loads(refreshed_record["extJson"])
    assert dict(refreshed_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }

    _, _, reloaded_profile = service._load_student_profile_bundle(student_user_id)  # noqa: SLF001
    assert reloaded_profile["checkInDates"]
    assert reloaded_profile["points"] >= check_in.json()["data"]["points"]
    assert reloaded_profile["aiQuota"]["usedCount"] == 1
    assert reloaded_profile["examSession"]["answeredCount"] == 7
    assert reloaded_profile["examSession"]["elapsedSec"] == 320


def test_save_student_profile_keeps_hot_state_and_only_updates_cold_snapshot(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)
    me = client.get("/api/question-bank/auth/me", headers=headers)
    assert me.status_code == 200
    student_user_id = me.json()["data"]["userId"]
    service = client.app.state.service
    repo = service.repository

    record, record_ext, profile = service._load_student_profile_bundle(student_user_id)
    today = datetime.now(timezone.utc).date().isoformat()
    seed_student_hot_state(
        service,
        student_user_id,
        points=66,
        title="连刷达人",
        unlocked_titles=["备考新星", "连刷达人"],
        check_in_dates=[today],
        ai_quota={"dailyLimit": 20, "usedCount": 2, "quotaDate": today},
        exam_session={"answeredCount": 5, "elapsedSec": 180, "updateTime": f"{today}T10:00:00Z"},
    )

    saved = client.post(
        "/api/question-bank/student/profile",
        headers=headers,
        json={"examCategoryCode": "LITERATURE", "jointExamGroupCode": "LITERATURE_11"},
    )
    assert saved.status_code == 200
    data = saved.json()["data"]
    assert data["examCategoryCode"] == "LITERATURE"
    assert data["jointExamGroupCode"] == "LITERATURE_11"
    assert data["points"] == 66
    assert data["title"] == "连刷达人"
    assert data["aiQuota"]["usedCount"] == 2
    assert data["examSession"]["answeredCount"] == 5

    refreshed_record = repo.get_student_question_bank(record["questionId"], student_user_id)
    assert refreshed_record is not None
    refreshed_ext = json.loads(refreshed_record["extJson"])
    assert dict(refreshed_ext.get("studentProfile", {})) == {
        "examCategoryCode": "LITERATURE",
        "jointExamGroupCode": "LITERATURE_11",
    }

    profile_state = repo.get_student_profile_state(student_user_id)
    assert profile_state is not None
    assert profile_state["examCategoryCode"] == "LITERATURE"
    assert profile_state["jointExamGroupCode"] == "LITERATURE_11"
    assert profile_state["points"] == 66
    assert profile_state["checkInDates"] == [today]
    assert profile_state["aiQuota"] == {"dailyLimit": 20, "usedCount": 2, "quotaDate": today}
    assert profile_state["examSession"] == {"answeredCount": 5, "elapsedSec": 180, "updateTime": f"{today}T10:00:00Z"}

    refreshed_me = client.get("/api/question-bank/auth/me", headers=headers)
    assert refreshed_me.status_code == 200
    refreshed_me_data = refreshed_me.json()["data"]
    assert refreshed_me_data["examCategoryCode"] == "LITERATURE"
    assert refreshed_me_data["jointExamGroupCode"] == "LITERATURE_11"

    managed_user = service._get_managed_user(student_user_id)  # noqa: SLF001
    assert managed_user is not None
    assert managed_user["examCategoryCode"] == "LITERATURE"
    assert managed_user["jointExamGroupCode"] == "LITERATURE_11"

    persisted_state = service._load_system_state()  # noqa: SLF001
    persisted_managed_user = next(item for item in persisted_state["managedUsers"] if item["userId"] == student_user_id)
    assert persisted_managed_user["examCategoryCode"] == ""
    assert persisted_managed_user["jointExamGroupCode"] == ""

    teacher_students = client.get("/api/question-bank/teacher/error-book/students", headers=teacher_headers())
    assert teacher_students.status_code == 200
    teacher_student_row = next(item for item in teacher_students.json()["data"] if item["studentUserId"] == student_user_id)
    assert teacher_student_row["examCategoryCode"] == "LITERATURE"
    assert teacher_student_row["jointExamGroupCode"] == "LITERATURE_11"

    admin_headers = super_admin_auth_headers(client)
    raw_state = service._load_system_state()  # noqa: SLF001
    raw_managed_user = next(item for item in raw_state["managedUsers"] if item["userId"] == student_user_id)
    raw_managed_user["examCategoryCode"] = "SCIENCE_ENGINEERING"
    raw_managed_user["jointExamGroupCode"] = "SCIENCE_ENGINEERING_3"
    service._save_system_state(raw_state)  # noqa: SLF001

    refreshed_raw_state = service._load_system_state()  # noqa: SLF001
    refreshed_raw_managed_user = next(item for item in refreshed_raw_state["managedUsers"] if item["userId"] == student_user_id)
    assert refreshed_raw_managed_user["examCategoryCode"] == ""
    assert refreshed_raw_managed_user["jointExamGroupCode"] == ""
    refreshed_user = repo.get_user_by_id(student_user_id)
    assert refreshed_user is not None
    refreshed_user_ext = json.loads(refreshed_user["extJson"])
    assert "examCategoryCode" not in refreshed_user_ext
    assert "jointExamGroupCode" not in refreshed_user_ext

    export_students = client.get("/api/question-bank/admin/students/export?format=csv", headers=admin_headers)
    assert export_students.status_code == 200
    export_line = next(
        line for line in export_students.json()["data"]["content"].splitlines() if line.startswith(f"{student_user_id},")
    )
    assert ",LITERATURE,LITERATURE_11," in export_line

    send_title = f"正式画像分群-{student_user_id}"
    sent = client.post(
        "/api/question-bank/messages/send",
        headers=admin_headers,
        json={
            "targetMode": "cohort",
            "examCategoryCode": "LITERATURE",
            "jointExamGroupCode": "LITERATURE_11",
            "category": "SYSTEM_NOTICE",
            "title": send_title,
            "content": "分群应命中正式画像学生。",
        },
    )
    assert sent.status_code == 200
    assert sent.json()["data"]["sentCount"] >= 1

    student_messages = client.get(
        "/api/question-bank/messages?page=1&size=50",
        headers=student_headers(student_user_id),
    )
    assert student_messages.status_code == 200
    assert any(item["title"] == send_title for item in student_messages.json()["data"]["items"])


def test_seeded_student_profile_snapshot_is_thin_after_init(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository
    record = repo.get_student_question_bank("question-seed-001", "student-001")
    assert record is not None
    record_ext = json.loads(record["extJson"])
    assert dict(record_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }


def test_load_student_profile_bundle_seeds_formal_cold_selection_from_legacy_snapshot(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service
    repo = service.repository
    student_user_id = "student-001"
    record = repo.get_student_question_bank("question-seed-001", student_user_id)
    assert record is not None

    record_ext = json.loads(record["extJson"])
    record_ext["studentProfile"] = {
        "examCategoryCode": "LITERATURE",
        "jointExamGroupCode": "LITERATURE_11",
    }
    repo.upsert_student_question_bank({**record, "extJson": json.dumps(record_ext, ensure_ascii=False)})
    with get_connection(repo.db_path) as connection:
        connection.execute("DELETE FROM student_profile_state WHERE studentUserId = ?", (student_user_id,))
        connection.commit()

    _, _, profile = service._load_student_profile_bundle(student_user_id)  # noqa: SLF001
    assert profile["examCategoryCode"] == "LITERATURE"
    assert profile["jointExamGroupCode"] == "LITERATURE_11"

    profile_state = repo.get_student_profile_state(student_user_id)
    assert profile_state is not None
    assert profile_state["examCategoryCode"] == "LITERATURE"
    assert profile_state["jointExamGroupCode"] == "LITERATURE_11"


def test_legacy_daily_progress_and_points_ledger_do_not_override_formal_state(tmp_path: Path):
    client = make_client(tmp_path)
    headers = register_student_auth_headers(client)
    me = client.get("/api/question-bank/auth/me", headers=headers)
    assert me.status_code == 200
    student_user_id = me.json()["data"]["userId"]
    service = client.app.state.service
    repo = service.repository

    record, record_ext, profile = service._load_student_profile_bundle(student_user_id)
    progress_date = "2026-03-23"
    seed_student_hot_state(
        service,
        student_user_id,
        daily_progress={
            progress_date: {
                "checkInCount": 1,
                "practiceAnswers": 6,
                "papersCompleted": 0,
                "wrongBookReviewed": 2,
                "rewardedKeys": ["practiceReward"],
            }
        },
        points_ledger=[
            {
                "eventKey": "legacy-guard:1",
                "reason": "正式表数据",
                "points": 9,
                "createTime": "2026-03-23T08:00:00Z",
            }
        ],
    )

    record = repo.get_student_question_bank(record["questionId"], student_user_id)
    assert record is not None
    forged_ext = json.loads(record["extJson"])
    forged_ext["studentProfile"] = {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
        "dailyProgress": {
            progress_date: {
                "checkInCount": 99,
                "practiceAnswers": 999,
                "papersCompleted": 9,
                "wrongBookReviewed": 9,
                "rewardedKeys": ["forgedReward"],
            }
        },
        "pointsLedger": [
            {
                "eventKey": "forged:1",
                "reason": "伪造旧账本",
                "points": 999,
                "createTime": "2026-03-23T09:00:00Z",
            }
        ],
    }
    with get_connection(repo.db_path) as connection:
        connection.execute(
            """
            UPDATE student_question_record
            SET extJson = ?
            WHERE studentUserId = ? AND questionId = ?
            """,
            (json.dumps(forged_ext, ensure_ascii=False), student_user_id, record["questionId"]),
        )
        connection.commit()

    _, _, reloaded_profile = service._load_student_profile_bundle(student_user_id)
    assert reloaded_profile["dailyProgress"][progress_date]["practiceAnswers"] == 6
    assert reloaded_profile["dailyProgress"][progress_date]["rewardedKeys"] == ["practiceReward"]
    assert any(item["eventKey"] == "legacy-guard:1" for item in reloaded_profile["pointsLedger"])
    assert all(item["eventKey"] != "forged:1" for item in reloaded_profile["pointsLedger"])


def test_wrong_book_list_prefers_formal_hot_state_columns_when_extjson_is_stale(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository
    record = repo.get_student_question_bank("question-seed-008", "student-001")
    assert record is not None
    record_ext = json.loads(record["extJson"])
    record_ext["wrongBook"] = {}
    repo.upsert_student_question_bank({**record, "extJson": json.dumps(record_ext, ensure_ascii=False)})

    with get_connection(repo.db_path) as connection:
        connection.execute(
            """
            UPDATE student_question_record
            SET wrongBookFlag = 1,
                wrongCount = 5,
                wrongBookReviewCount = 2,
                wrongBookCollectedAt = '2026-03-20T09:00:00Z',
                wrongBookReviewedAt = '2026-03-22T08:30:00Z',
                wrongBookLastReasonCode = 'KNOWLEDGE_GAP',
                wrongBookLastReasonLabel = '知识点掌握不牢'
            WHERE studentUserId = ? AND questionId = ?
            """,
            ("student-001", "question-seed-008"),
        )
        connection.commit()

    wrong_book = client.get(
        "/api/question-bank/student/wrong-book/questions",
        headers=student_headers(),
    )
    assert wrong_book.status_code == 200
    target = next(item for item in wrong_book.json()["data"]["items"] if item["id"] == "question-seed-008")
    target_ext = json.loads(target["extJson"])
    wrong_state = target_ext["studentState"]["wrongBook"]
    assert wrong_state["isCollected"] is True
    assert wrong_state["wrongCount"] == 5
    assert wrong_state["reviewCount"] == 2
    assert wrong_state["reviewedAt"] == "2026-03-22T08:30:00Z"

    summary = client.get(
        "/api/student/error-book/summary",
        headers=student_headers(),
    )
    assert summary.status_code == 200
    target_insight = next(item for item in summary.json()["data"]["questionInsights"] if item["questionId"] == "question-seed-008")
    assert target_insight["reviewCount"] == 2

    repo.upsert_paper_report_payload(
        {
            "studentUserId": "student-001",
            "reportId": "paper-report-test-001",
            "paperId": "paper-demo-001",
            "subjectId": "POLITICS",
            "subjectIds": ["POLITICS"],
            "score": 82,
            "totalScore": 100,
            "scoreRate": 0.82,
            "totalElapsedSec": 1200,
            "submittedAt": "2026-03-22T09:00:00Z",
            "pendingSubjectiveTaskIds": ["task-001"],
            "reportDetail": {"typeAccuracy": [], "questionResults": []},
        }
    )
    repo.upsert_message_send_history_payload(
        {
            "traceId": "trace-test-001",
            "scheduleId": "",
            "senderUserId": "teacher-001",
            "targetMode": "userIds",
            "targetCount": 2,
            "sentCount": 2,
            "category": "SYSTEM_NOTICE",
            "title": "开课提醒",
            "content": "请准时上课",
            "sendAt": "2026-03-22T09:05:00Z",
            "status": "SENT",
            "createTime": "2026-03-22T09:05:00Z",
        }
    )

    paper_row = repo.get_paper_report_row("paper-report-test-001")
    assert paper_row is not None
    assert paper_row["studentUserId"] == "student-001"
    assert paper_row["paperId"] == "paper-demo-001"
    assert paper_row["pendingSubjectiveCount"] == 1

    history_row = repo.get_message_send_history_row("trace-test-001")
    assert history_row is not None
    assert history_row["senderUserId"] == "teacher-001"
    assert history_row["targetCount"] == 2
    assert history_row["status"] == "SENT"


def test_student_paper_report_read_path_prefers_formal_table_when_legacy_json_missing(tmp_path: Path):
    client = make_client(tmp_path)
    submit = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 35, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 45, "marked": True},
            ],
            "totalElapsedSec": 80,
        },
    )
    assert submit.status_code == 200
    report_id = submit.json()["data"]["reportId"]
    system_record = client.app.state.service.repository.get_student_question_bank("question-seed-001", "__system__")
    assert system_record is not None
    assert "paperReports" not in json.loads(system_record["extJson"])

    reports = client.get("/api/question-bank/student/papers/reports?page=1&size=10", headers=student_headers())
    assert reports.status_code == 200
    assert any(item["reportId"] == report_id for item in reports.json()["data"]["items"])

    detail = client.get(f"/api/question-bank/student/papers/reports/{report_id}", headers=student_headers())
    assert detail.status_code == 200
    assert detail.json()["data"]["reportId"] == report_id


def test_message_send_history_read_path_prefers_formal_table_when_legacy_json_missing(tmp_path: Path):
    client = make_client(tmp_path)
    sent = client.post(
        "/api/question-bank/messages/send",
        headers=teacher_headers(),
        json={
            "targetMode": "userIds",
            "userIds": ["student-001"],
            "category": "SYSTEM_NOTICE",
            "title": "正式表优先",
            "content": "发送历史应优先读取正式表。",
        },
    )
    assert sent.status_code == 200
    trace_id = sent.json()["data"]["traceId"]
    system_record = client.app.state.service.repository.get_student_question_bank("question-seed-001", "__system__")
    assert system_record is not None
    assert "messageSendHistory" not in json.loads(system_record["extJson"])

    history = client.get("/api/question-bank/messages/send-history?page=1&size=20", headers=teacher_headers())
    assert history.status_code == 200
    assert any(item["traceId"] == trace_id for item in history.json()["data"]["items"])


def test_repository_analytics_rollups_prefer_formal_table_when_legacy_json_missing(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository
    record = repo.get_student_question_bank("question-seed-001", "student-001")
    assert record is not None
    record_ext = json.loads(record["extJson"])
    record_ext["chapterPractice"] = {
        "lastAnswer": "B",
        "isCorrect": True,
        "answerDurationSec": 66,
        "submitCount": 1,
        "correctCount": 1,
        "submittedAt": "2026-03-22T11:00:00Z",
    }
    repo.upsert_student_question_bank({**record, "extJson": json.dumps(record_ext, ensure_ascii=False)})
    student_row = repo.get_user_by_id("student-001")
    assert student_row is not None
    assert "studentRecords" not in json.loads(student_row["extJson"])

    rows = repo.aggregate_student_analytics_rollups({"studentUserId": "student-001", "subjectId": "", "chapter": "", "paperId": ""})
    assert any(item["rowType"] == "student_subject" and item["studentUserId"] == "student-001" for item in rows)


def test_student_error_book_similar_questions_and_word_export(tmp_path: Path):
    client = make_client(tmp_path)
    headers = {**student_headers(), "X-Joint-Group": "SCIENCE_ENGINEERING_3"}
    service = client.app.state.service
    repo = service.repository
    user = repo.get_user_by_id("student-001")
    assert user is not None
    user_ext = json.loads(user["extJson"])
    user_ext["examCategoryCode"] = ""
    user_ext["jointExamGroupCode"] = ""
    with get_connection(repo.db_path) as connection:
        connection.execute(
            """
            UPDATE user
            SET extJson = ?
            WHERE id = ?
            """,
            (json.dumps(user_ext, ensure_ascii=False), "student-001"),
        )
        for question_id in ("question-seed-001", "question-seed-002", "question-seed-std-politics-001"):
            row = connection.execute("SELECT extJson FROM question WHERE id = ?", (question_id,)).fetchone()
            if row:
                ext_json = json.loads(row[0])
                if "chapterCode" not in ext_json and "chapter_code" in ext_json:
                    ext_json["chapterCode"] = ext_json.get("chapter_code", "")
                if "pointCode" not in ext_json and "point_code" in ext_json:
                    ext_json["pointCode"] = ext_json.get("point_code", "")
                connection.execute(
                    "UPDATE question SET extJson = ? WHERE id = ?",
                    (json.dumps(ext_json, ensure_ascii=False), question_id),
                )
        connection.commit()

    wrong_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "A", "elapsedSec": 25},
    )
    assert wrong_submit.status_code == 200

    similar = client.get(
        "/api/question-bank/student/wrong-book/questions/question-seed-002/similar",
        headers=headers,
    )
    assert similar.status_code == 200
    similar_data = similar.json()["data"]
    assert similar_data["pointCode"] == "PT_001_001"
    assert len(similar_data["items"]) >= 1
    assert all(json.loads(item["extJson"]).get("pointCode") == "PT_001_001" for item in similar_data["items"])

    export = client.post(
        "/api/question-bank/student/wrong-book/exports/word",
        headers=headers,
        json={"subjectCode": "POLITICS", "questionIds": ["question-seed-001", "question-seed-002"]},
    )
    assert export.status_code == 200
    export_data = export.json()["data"]
    assert export_data["format"] == "docx"
    assert export_data["fileName"].endswith(".docx")
    assert export_data["questionCount"] == 2

    document = Document(io.BytesIO(base64.b64decode(export_data["contentBase64"])))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "SMART REPAIR WORKBOOK" in paragraph_text
    assert "个人线下提分卷" in paragraph_text
    assert "错题修复中心线下背诵版" in paragraph_text
    assert "目录导览" in paragraph_text
    assert "章节索引" in paragraph_text
    assert "错因统计汇总" in paragraph_text
    assert "章节错因分布" in paragraph_text
    assert "CHAPTER" in paragraph_text
    assert "本章重点：" in paragraph_text
    assert "本章错因 TOP3：" in paragraph_text
    assert "本章推荐复习顺序：" in paragraph_text
    assert "本章完成目标：" in paragraph_text
    assert "本章自测通过标准：" in paragraph_text
    assert "预计完成时长：" in paragraph_text
    assert "复习建议：" in paragraph_text
    assert "错因标签：" in paragraph_text
    assert "背诵默写区" in paragraph_text
    assert "答案附录" in paragraph_text
    assert "附录章节：" in paragraph_text
    assert "见文末附录" in paragraph_text
    assert "难度角标：" in paragraph_text
    assert "本卷结构：目录导览 -> 题目正文 -> 答案附录" in paragraph_text
    assert "返回原题页：" in paragraph_text
    assert "标签：掌握度" in paragraph_text

    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "学生姓名" in table_text
    assert "理工考生" in table_text
    assert "SCIENCE_ENGINEERING_3" in table_text
    assert "学科门类" in table_text
    assert "专业方向" in table_text
    assert "计算机类" in table_text
    assert "强化阶段" in table_text
    assert "导出时间" in table_text
    assert "本卷结构" in table_text
    assert "本章小结" in table_text
    assert "错因标签" in table_text
    assert "累计次数" in table_text
    assert "占比" in table_text
    assert "等级提示" in table_text
    assert "风险分层" in table_text
    assert "章节" in table_text
    assert "错因分布" in table_text
    assert "章节主错因" in table_text
    assert "章节提示" in table_text
    assert "主错因" in table_text
    assert "出现次数" in table_text
    assert "次级错因" in table_text
    assert "专业画像摘要：" in paragraph_text
    assert "学情一句话结论：" in paragraph_text
    assert "老师寄语：" in paragraph_text

    header_text = "\n".join(paragraph.text for paragraph in document.sections[0].header.paragraphs)
    assert "智能学情修复中心" in header_text
    footer_xml = document.sections[0].footer.paragraphs[0]._element.xml
    assert "PAGE" in footer_xml
    document_xml = document._element.body.xml
    assert "PAGEREF qb_q_1" in document_xml
    assert 'w:hyperlink' in document_xml
    assert 'w:anchor="qb_q_1"' in document_xml
    assert "回到正文" in document_xml

    page_break_count = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            page_break_count += run._element.xml.count('w:type="page"')
    assert page_break_count >= 4


def test_student_wrong_reason_is_persisted_after_wrong_answer(tmp_path: Path):
    client = make_client(tmp_path)
    submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "A", "elapsedSec": 30},
    )
    assert submit.status_code == 200
    ext = json.loads(submit.json()["data"]["extJson"])
    wrong_book = ext["studentState"]["wrongBook"]
    assert wrong_book["isCollected"] is True
    assert wrong_book["lastReasonCode"] in {"KNOWLEDGE_GAP", "TIMEOUT", "BLANK", "EXPRESSION"}
    assert wrong_book["reasonStats"]
    assert wrong_book["reasonStats"][0]["count"] >= 1


def test_student_error_book_review_status_can_reach_mastered(tmp_path: Path):
    client = make_client(tmp_path)
    headers = {**student_headers(), "X-Joint-Group": "SCIENCE_ENGINEERING_3"}

    first_wrong = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=headers,
        json={"answer": "A", "elapsedSec": 25},
    )
    assert first_wrong.status_code == 200

    for _ in range(3):
        correct = client.post(
            "/api/question-bank/student/practice/questions/question-seed-001/submit",
            headers=headers,
            json={"answer": "B", "elapsedSec": 20},
        )
        assert correct.status_code == 200

    summary = client.get(
        "/api/student/error-book/summary",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert summary.status_code == 200
    data = summary.json()["data"]
    target = next(item for item in data["questionInsights"] if item["questionId"] == "question-seed-001")
    assert target["reviewStatusLabel"] == "已斩获"
    assert target["reviewAccuracyRate"] >= 0.8

    archived = client.post(
        "/api/question-bank/student/wrong-book/archive-harvested",
        headers=headers,
        json={"questionIds": ["question-seed-001"], "subjectCode": "POLITICS"},
    )
    assert archived.status_code == 200
    assert archived.json()["data"]["archivedCount"] == 1

    refreshed_summary = client.get(
        "/api/student/error-book/summary",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert refreshed_summary.status_code == 200
    refreshed_items = refreshed_summary.json()["data"]["questionInsights"]
    assert not any(item["questionId"] == "question-seed-001" for item in refreshed_items)

    personal_bank = client.get(
        "/api/question-bank/student/practice/questions",
        headers=headers,
        params={"onlyPersonalBank": "true", "subjectCode": "POLITICS"},
    )
    assert personal_bank.status_code == 200
    archived_item = next(
        item
        for item in personal_bank.json()["data"]["items"]
        if item["id"] == "question-seed-001"
    )
    archived_ext = json.loads(archived_item["extJson"])
    assert archived_ext["studentState"]["personalBank"]["sourceLabel"] == "已斩获归档"
    assert archived_ext["studentState"]["wrongBook"]["archivedAt"]

    restored = client.post(
        "/api/question-bank/student/wrong-book/restore-archived",
        headers=headers,
        json={"questionIds": ["question-seed-001"]},
    )
    assert restored.status_code == 200
    assert restored.json()["data"]["restoredCount"] == 1

    restored_summary = client.get(
        "/api/student/error-book/summary",
        headers=headers,
        params={"subjectCode": "POLITICS"},
    )
    assert restored_summary.status_code == 200
    assert any(item["questionId"] == "question-seed-001" for item in restored_summary.json()["data"]["questionInsights"])


def test_teacher_error_book_center_can_select_student_analyze_and_export(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service
    repo = service.repository
    with get_connection(repo.db_path) as connection:
        for question_id in ("question-seed-001", "question-seed-002", "question-seed-std-politics-001"):
            row = connection.execute("SELECT extJson FROM question WHERE id = ?", (question_id,)).fetchone()
            if row:
                ext_json = json.loads(row[0])
                if "chapterCode" not in ext_json and "chapter_code" in ext_json:
                    ext_json["chapterCode"] = ext_json.get("chapter_code", "")
                if "pointCode" not in ext_json and "point_code" in ext_json:
                    ext_json["pointCode"] = ext_json.get("point_code", "")
                connection.execute(
                    "UPDATE question SET extJson = ? WHERE id = ?",
                    (json.dumps(ext_json, ensure_ascii=False), question_id),
                )
        connection.commit()

    wrong_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "A", "elapsedSec": 25},
    )
    assert wrong_submit.status_code == 200

    students = client.get("/api/question-bank/teacher/error-book/students", headers=teacher_headers())
    assert students.status_code == 200
    student_rows = students.json()["data"]
    assert any(item["studentUserId"] == "student-001" for item in student_rows)

    summary = client.get(
        "/api/question-bank/teacher/error-book/summary",
        headers=teacher_headers(),
        params={"userId": "student-001", "subjectCode": "POLITICS"},
    )
    assert summary.status_code == 200
    summary_data = summary.json()["data"]
    assert summary_data["studentMeta"]["studentUserId"] == "student-001"
    assert summary_data["currentSubject"]["subjectCode"] == "POLITICS"

    multi_subject_rows = client.get(
        "/api/question-bank/teacher/error-book/questions",
        headers=teacher_headers(),
        params={"userId": "student-001", "subjectCodes": "POLITICS,INFO_TECH_INTRO"},
    )
    assert multi_subject_rows.status_code == 200
    items = multi_subject_rows.json()["data"]["items"]
    assert items
    ext_subject_codes = {str(json.loads(item["extJson"]).get("subjectCode", "")).strip() for item in items}
    assert ext_subject_codes.issubset({"POLITICS", "INFO_TECH_INTRO"})

    similar = client.get(
        "/api/question-bank/teacher/error-book/questions/question-seed-002/similar",
        headers=teacher_headers(),
        params={"userId": "student-001"},
    )
    assert similar.status_code == 200

    export = client.post(
        "/api/question-bank/teacher/error-book/exports/word",
        headers=teacher_headers(),
        json={"studentUserId": "student-001", "subjectCodes": "POLITICS,INFO_TECH_INTRO"},
    )
    assert export.status_code == 200
    export_data = export.json()["data"]
    assert export_data["format"] == "docx"
    assert export_data["studentUserId"] == "student-001"
    export_bytes = base64.b64decode(export_data["contentBase64"])
    with zipfile.ZipFile(io.BytesIO(export_bytes)) as export_archive:
        export_styles_xml = export_archive.read("word/styles.xml").decode("utf-8")
    assert 'w:eastAsia="zh-CN"' in export_styles_xml
    assert "Microsoft YaHei" in export_styles_xml

    class_overview = client.get(
        "/api/question-bank/teacher/error-book/class-overview",
        headers=teacher_headers(),
        params={"classId": "SCIENCE_ENGINEERING_3", "subjectCodes": "POLITICS,INFO_TECH_INTRO"},
    )
    assert class_overview.status_code == 200
    class_data = class_overview.json()["data"]
    assert class_data["classMeta"]["classId"] == "SCIENCE_ENGINEERING_3"
    assert class_data["summaryCards"]["studentCount"] >= 1

    batch_export = client.post(
        "/api/question-bank/teacher/error-book/exports/word",
        headers=teacher_headers(),
        json={"studentUserIds": ["student-001", "student-002"], "subjectCodes": "POLITICS"},
    )
    assert batch_export.status_code == 200
    batch_data = batch_export.json()["data"]
    assert batch_data["format"] == "zip"
    assert batch_data["fileName"].endswith(".zip")
    assert len(batch_data["items"]) == 2

    class_report = client.post(
        "/api/question-bank/teacher/error-book/class-exports/report",
        headers=teacher_headers(),
        json={"classId": "SCIENCE_ENGINEERING_3", "subjectCodes": "POLITICS,INFO_TECH_INTRO"},
    )
    assert class_report.status_code == 200
    class_report_data = class_report.json()["data"]
    assert class_report_data["format"] == "docx"
    assert class_report_data["fileName"].endswith(".docx")
    class_report_document = Document(io.BytesIO(base64.b64decode(class_report_data["contentBase64"])))
    class_report_text = "\n".join(paragraph.text for paragraph in class_report_document.paragraphs)
    assert "首页摘要结论：" in class_report_text
    assert "风险等级结论卡：" in class_report_text
    assert "班级一句话建议卡：" in class_report_text
    assert "班级错因统计图表化文本" in class_report_text
    assert "章节错因图表化文本" in class_report_text
    class_report_bytes = base64.b64decode(class_report_data["contentBase64"])
    with zipfile.ZipFile(io.BytesIO(class_report_bytes)) as class_report_archive:
        class_report_styles_xml = class_report_archive.read("word/styles.xml").decode("utf-8")
    assert 'w:eastAsia="zh-CN"' in class_report_styles_xml
    assert "Microsoft YaHei" in class_report_styles_xml

    class_package = client.post(
        "/api/question-bank/teacher/error-book/class-exports/package",
        headers=teacher_headers(),
        json={"classId": "SCIENCE_ENGINEERING_3", "subjectCodes": "POLITICS"},
    )
    assert class_package.status_code == 200
    class_package_data = class_package.json()["data"]
    assert class_package_data["format"] == "zip"
    assert class_package_data["fileName"].endswith(".zip")
    export_item_names = [item["studentName"] for item in class_package_data["items"]]
    assert export_item_names == sorted(export_item_names)
    zip_bytes = base64.b64decode(class_package_data["contentBase64"])
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as archive:
        archive_names = archive.namelist()
        readme_text = archive.read("README.txt").decode("utf-8")
    assert any(name.endswith(".docx") and "class-error-book-report" in name for name in archive_names)
    assert any(name.endswith(".docx") and "理工考生-student-001" in name for name in archive_names)
    assert "教师端全班分析包 README" in readme_text
    assert "使用说明:" in readme_text
    assert "推荐阅读顺序:" in readme_text
    assert "目录树说明:" in readme_text
    assert "  |-" in readme_text


def test_ai_marking_and_ai_tutor_are_persisted_in_ext_json(tmp_path: Path):
    client = make_client(tmp_path)
    marking = client.post(
        "/api/question-bank/student/practice/questions/question-seed-005/ai-marking",
        headers=student_headers("student-002"),
        json={"answer": "终身学习可以提升适应能力并促进持续成长。", "answerImageUrl": ""},
    )
    assert marking.status_code == 200
    marking_task = marking.json()["data"]
    assert tuple(marking_task.keys()) == TASK_FIELDS
    assert marking_task["type"] == "AI_MARKING"

    tutor = client.post(
        "/api/question-bank/student/practice/questions/question-seed-005/ai-tutor",
        headers=student_headers("student-002"),
        json={"prompt": "这道题的采分点是什么？", "promptImageUrl": ""},
    )
    assert tutor.status_code == 200
    tutor_task = tutor.json()["data"]
    assert tuple(tutor_task.keys()) == TASK_FIELDS
    assert tutor_task["type"] == "AI_TUTOR"

    completed_marking = poll_task(client, str(marking_task["id"]), student_headers("student-002"))
    completed_tutor = poll_task(client, str(tutor_task["id"]), student_headers("student-002"))
    assert completed_marking["status"] == "COMPLETED"
    assert completed_tutor["status"] == "COMPLETED"

    task_list = client.get("/api/question-bank/tasks?page=1&size=10", headers=student_headers("student-002"))
    assert task_list.status_code == 200
    ids = [item["id"] for item in task_list.json()["data"]["items"]]
    assert marking_task["id"] in ids
    assert tutor_task["id"] in ids

    wrong_book = client.get("/api/question-bank/student/wrong-book/questions", headers=student_headers("student-002"))
    assert wrong_book.status_code == 200
    ext = json.loads(wrong_book.json()["data"]["items"][0]["extJson"])
    assert "aiMarking" in ext["studentState"]
    assert "aiTutor" in ext["studentState"]


def test_task_visibility_is_limited_to_owner_or_super_admin(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)
    task = client.post(
        "/api/question-bank/student/practice/questions/question-seed-005/ai-tutor",
        headers=student_headers("student-002"),
        json={"prompt": "帮我拆一下答题步骤", "promptImageUrl": ""},
    )
    assert task.status_code == 200
    task_id = task.json()["data"]["id"]

    forbidden = client.get(f"/api/question-bank/tasks/{task_id}", headers=student_headers("student-001"))
    assert forbidden.status_code == 403
    assert forbidden.json()["code"] == "TASK_FORBIDDEN"

    admin = client.get(f"/api/question-bank/tasks/{task_id}", headers=admin_headers)
    assert admin.status_code == 200
    assert admin.json()["data"]["id"] == task_id


def test_task_owner_can_cancel_running_or_queued_task(tmp_path: Path):
    client = make_client(tmp_path)
    task = client.post(
        "/api/question-bank/student/practice/questions/question-seed-005/ai-tutor",
        headers=student_headers("student-002"),
        json={"prompt": "帮我先拆一下答题步骤", "promptImageUrl": ""},
    )
    assert task.status_code == 200
    task_id = task.json()["data"]["id"]

    cancelled = client.post(f"/api/question-bank/tasks/{task_id}/cancel", headers=student_headers("student-002"))
    assert cancelled.status_code == 200
    assert cancelled.json()["data"]["status"] == "CANCELLED"
    assert cancelled.json()["data"]["progress"] == 100

    detail = client.get(f"/api/question-bank/tasks/{task_id}", headers=student_headers("student-002"))
    assert detail.status_code == 200
    assert detail.json()["data"]["status"] == "CANCELLED"

    second_cancel = client.post(f"/api/question-bank/tasks/{task_id}/cancel", headers=student_headers("student-002"))
    assert second_cancel.status_code == 422
    assert second_cancel.json()["code"] == "TASK_VALIDATION_FAILED"


def test_admin_console_settings_and_user_directory(tmp_path: Path):
    client = make_client(tmp_path)
    unauthorized = client.get("/api/question-bank/admin/console", headers=super_admin_headers())
    assert unauthorized.status_code == 401
    assert "登录态" in unauthorized.json()["message"]
    unauthorized_home = client.get(
        "/admin/home",
        params={"role": "super_admin", "userId": "admin-001"},
        follow_redirects=False,
    )
    assert unauthorized_home.status_code == 302
    assert unauthorized_home.headers["location"] == "/login"
    unauthorized_page = client.get(
        "/admin/control-center",
        params={"role": "super_admin", "userId": "admin-001"},
        follow_redirects=False,
    )
    assert unauthorized_page.status_code == 302
    assert unauthorized_page.headers["location"] == "/login"
    admin_headers = super_admin_auth_headers(client)

    home_page = client.get("/admin/home", headers=admin_headers)
    assert home_page.status_code == 200
    home_payload = parse_page_bootstrap(home_page.text)
    assert home_payload["route"] == "/admin/home"
    assert home_payload["viewKey"] == "admin-home"
    assert home_payload["pageTitle"] == "管理驾驶舱"
    assert home_payload["actor"] == {"role": "super_admin", "userId": "admin-001"}
    assert isinstance(home_payload["csrfToken"], str) and home_payload["csrfToken"]

    page = client.get("/admin/control-center", headers=admin_headers)
    assert page.status_code == 200
    page_payload = parse_page_bootstrap(page.text)
    assert page_payload["route"] == "/admin/control-center"
    assert page_payload["viewKey"] == "admin-control-center"
    assert page_payload["pageTitle"] == "超管控制台"
    assert page_payload["actor"] == {"role": "super_admin", "userId": "admin-001"}
    assert isinstance(page_payload["csrfToken"], str) and page_payload["csrfToken"]

    console = client.get("/api/question-bank/admin/console", headers=admin_headers)
    assert console.status_code == 200
    data = console.json()["data"]
    assert "systemSettings" in data
    assert "managedUsers" in data
    assert data["managedUsersPagination"]["total"] >= len(data["managedUsers"])

    settings = client.post(
        "/api/question-bank/admin/settings",
        headers=admin_headers,
        json={
            "platformName": "专升本ALL AI Pro",
            "defaultExamMinutes": 150,
            "dailyCheckInPoints": 12,
            "practiceRewardThreshold": 10,
            "practiceRewardPoints": 20,
            "paperRewardPoints": 30,
            "wrongBookRewardThreshold": 5,
            "wrongBookRewardPoints": 15,
            "aiDailyLimit": 25,
        },
    )
    assert settings.status_code == 200
    assert settings.json()["data"]["platformName"] == "专升本ALL AI Pro"

    saved_user = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "student-009",
            "role": "student",
            "name": "新考生",
            "mobile": "13800000019",
            "enabled": True,
            "permissions": [],
            "examCategoryCode": "MANAGEMENT",
            "jointExamGroupCode": "MANAGEMENT_1",
            "vocationalMajor": "财经商贸类",
            "prepStage": "基础阶段",
        },
    )
    assert saved_user.status_code == 200

    imported = client.post(
        "/api/question-bank/admin/students/import",
        headers=admin_headers,
        json={
            "csvText": "\n".join(
                [
                    "userId,name,mobile,examCategoryCode,jointExamGroupCode,vocationalMajor,prepStage",
                    "student-010,批量考生,13800000020,LITERATURE,LITERATURE_11,语言类,冲刺阶段",
                ]
            )
        },
    )
    assert imported.status_code == 200
    assert imported.json()["data"]["imported"] == 1

    users = client.get("/api/question-bank/admin/users?role=student&page=1&size=20", headers=admin_headers)
    assert users.status_code == 200
    payload = users.json()["data"]
    assert payload["page"] == 1
    assert payload["size"] == 20
    assert payload["total"] >= 2
    user_ids = [item["userId"] for item in payload["items"]]
    assert "student-009" in user_ids
    assert "student-010" in user_ids


def test_admin_student_import_supports_quoted_csv_and_row_error_details(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    imported = client.post(
        "/api/question-bank/admin/students/import",
        headers=admin_headers,
        json={
            "csvText": "\n".join(
                [
                    "userId,name,mobile,examCategoryCode,jointExamGroupCode,vocationalMajor,prepStage",
                    "student-021,\"批量,考生\",13800000021,MANAGEMENT,MANAGEMENT_1,\"财经,商贸类\",基础阶段",
                    "student-022,缺列考生,13800000022,MANAGEMENT",
                ]
            )
        },
    )
    assert imported.status_code == 200
    result = imported.json()["data"]
    assert result["imported"] == 1
    assert result["failed"] == 1
    assert result["errorDetails"][0]["errorCode"] == "FIELD_COUNT_INVALID"

    users = client.get(
        "/api/question-bank/admin/users?keyword=student-021&page=1&size=10",
        headers=admin_headers,
    )
    assert users.status_code == 200
    items = users.json()["data"]["items"]
    assert len(items) == 1
    assert items[0]["name"] == "批量,考生"


def test_admin_student_directory_sync_keeps_snapshot_thin_and_hot_state_stable(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)
    service = client.app.state.service
    repo = service.repository

    created = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "student-031",
            "role": "student",
            "name": "目录迁移考生",
            "mobile": "13800000031",
            "enabled": True,
            "permissions": [],
            "examCategoryCode": "MANAGEMENT",
            "jointExamGroupCode": "MANAGEMENT_1",
            "vocationalMajor": "财经商贸类",
            "prepStage": "基础阶段",
        },
    )
    assert created.status_code == 200

    records = repo.list_student_records_by_user("student-031")
    assert len(records) == 1
    created_ext = json.loads(records[0]["extJson"])
    student_record = created_ext.get("studentRecord", {})
    record_ext = json.loads(student_record.get("extJson", "{}"))
    assert student_record["formalState"]["profileAnchorFlag"] == 1
    assert dict(record_ext.get("studentProfile", {})) == {
        "examCategoryCode": "MANAGEMENT",
        "jointExamGroupCode": "MANAGEMENT_1",
    }
    created_state = repo.get_student_profile_state("student-031")
    assert created_state is not None
    assert created_state["examCategoryCode"] == "MANAGEMENT"
    assert created_state["jointExamGroupCode"] == "MANAGEMENT_1"
    assert created_state["vocationalMajor"] == "财经商贸类"
    assert created_state["prepStage"] == "基础阶段"
    assert created_state["points"] == 0
    assert created_state["title"] == "备考新星"
    assert created_state["aiQuota"] == {"dailyLimit": 20, "usedCount": 0, "quotaDate": ""}
    created_user = repo.get_user_by_id("student-031")
    assert created_user is not None
    created_user_ext = json.loads(created_user["extJson"])
    assert "examCategoryCode" not in created_user_ext
    assert "jointExamGroupCode" not in created_user_ext
    created_state_snapshot = service._load_system_state()  # noqa: SLF001
    created_raw_managed_user = next(item for item in created_state_snapshot["managedUsers"] if item["userId"] == "student-031")
    assert created_raw_managed_user["examCategoryCode"] == ""
    assert created_raw_managed_user["jointExamGroupCode"] == ""

    seed_student_hot_state(
        service,
        "student-031",
        points=88,
        title="连刷达人",
        unlocked_titles=["备考新星", "连刷达人"],
        check_in_dates=["2026-03-23"],
        ai_quota={"dailyLimit": 24, "usedCount": 3, "quotaDate": "2026-03-23"},
        exam_session={"answeredCount": 9, "elapsedSec": 420, "updateTime": "2026-03-23T10:00:00Z"},
    )

    updated = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "student-031",
            "role": "student",
            "name": "目录迁移考生",
            "mobile": "13800000031",
            "enabled": True,
            "permissions": [],
            "examCategoryCode": "LITERATURE",
            "jointExamGroupCode": "LITERATURE_11",
            "vocationalMajor": "语言类",
            "prepStage": "冲刺阶段",
        },
    )
    assert updated.status_code == 200

    updated_records = repo.list_student_records_by_user("student-031")
    assert len(updated_records) == 1
    updated_ext = json.loads(updated_records[0]["extJson"])
    updated_record = updated_ext.get("studentRecord", {})
    updated_record_ext = json.loads(updated_record.get("extJson", "{}"))
    assert updated_record["formalState"]["profileAnchorFlag"] == 1
    assert dict(updated_record_ext.get("studentProfile", {})) == {
        "examCategoryCode": "LITERATURE",
        "jointExamGroupCode": "LITERATURE_11",
    }
    _, _, reloaded_profile = service._load_student_profile_bundle("student-031")  # noqa: SLF001
    updated_state = repo.get_student_profile_state("student-031")
    assert updated_state is not None
    assert updated_state["examCategoryCode"] == "LITERATURE"
    assert updated_state["jointExamGroupCode"] == "LITERATURE_11"
    assert updated_state["vocationalMajor"] == "语言类"
    assert updated_state["prepStage"] == "冲刺阶段"
    assert reloaded_profile["points"] == 88
    assert reloaded_profile["title"] == "连刷达人"
    assert reloaded_profile["checkInDates"] == ["2026-03-23"]
    assert reloaded_profile["aiQuota"] == {"dailyLimit": 24, "usedCount": 3, "quotaDate": "2026-03-23"}
    assert reloaded_profile["examSession"] == {"answeredCount": 9, "elapsedSec": 420, "updateTime": "2026-03-23T10:00:00Z"}
    updated_user = repo.get_user_by_id("student-031")
    assert updated_user is not None
    updated_user_ext = json.loads(updated_user["extJson"])
    assert "examCategoryCode" not in updated_user_ext
    assert "jointExamGroupCode" not in updated_user_ext
    updated_state_snapshot = service._load_system_state()  # noqa: SLF001
    updated_raw_managed_user = next(item for item in updated_state_snapshot["managedUsers"] if item["userId"] == "student-031")
    assert updated_raw_managed_user["examCategoryCode"] == ""
    assert updated_raw_managed_user["jointExamGroupCode"] == ""

    imported = client.post(
        "/api/question-bank/admin/students/import",
        headers=admin_headers,
        json={
            "csvText": "\n".join(
                [
                    "userId,name,mobile,examCategoryCode,jointExamGroupCode,vocationalMajor,prepStage",
                    "student-032,批量同步考生,13800000032,SCIENCE_ENGINEERING,SCIENCE_ENGINEERING_3,计算机类,强化阶段",
                ]
            )
        },
    )
    assert imported.status_code == 200
    assert imported.json()["data"]["imported"] == 1

    imported_records = repo.list_student_records_by_user("student-032")
    assert len(imported_records) == 1
    imported_ext = json.loads(imported_records[0]["extJson"])
    imported_student_record = imported_ext.get("studentRecord", {})
    imported_record_ext = json.loads(imported_student_record.get("extJson", "{}"))
    assert imported_student_record["formalState"]["profileAnchorFlag"] == 1
    assert dict(imported_record_ext.get("studentProfile", {})) == {
        "examCategoryCode": "SCIENCE_ENGINEERING",
        "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
    }
    imported_state = repo.get_student_profile_state("student-032")
    assert imported_state is not None
    assert imported_state["examCategoryCode"] == "SCIENCE_ENGINEERING"
    assert imported_state["jointExamGroupCode"] == "SCIENCE_ENGINEERING_3"
    assert imported_state["vocationalMajor"] == "计算机类"
    assert imported_state["prepStage"] == "强化阶段"
    assert imported_state["points"] == 0
    assert imported_state["title"] == "备考新星"
    assert imported_state["unlockedTitles"] == ["备考新星"]
    imported_user = repo.get_user_by_id("student-032")
    assert imported_user is not None
    imported_user_ext = json.loads(imported_user["extJson"])
    assert "examCategoryCode" not in imported_user_ext
    assert "jointExamGroupCode" not in imported_user_ext
    imported_state_snapshot = service._load_system_state()  # noqa: SLF001
    imported_raw_managed_user = next(item for item in imported_state_snapshot["managedUsers"] if item["userId"] == "student-032")
    assert imported_raw_managed_user["examCategoryCode"] == ""
    assert imported_raw_managed_user["jointExamGroupCode"] == ""


def test_system_state_save_rejects_stale_version(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service

    stale_state = service._load_system_state()
    latest_state = service._load_system_state()
    latest_state["systemSettings"]["platformName"] = "并发更新版本A"
    service._save_system_state(latest_state)

    stale_state["systemSettings"]["platformName"] = "并发更新版本B"
    with pytest.raises(Exception) as exc_info:
        service._save_system_state(stale_state)
    assert "系统状态已被其他管理员更新" in str(exc_info.value)


def test_messages_center_send_read_and_settings(tmp_path: Path):
    client = make_client(tmp_path)

    page = client.get("/messages", params={"role": "student", "userId": "student-001"})
    assert page.status_code == 200
    page_payload = parse_page_bootstrap(page.text)
    assert page_payload["route"] == "/messages"
    assert page_payload["viewKey"] == "messages"
    assert page_payload["pageTitle"] == "消息中心"
    assert page_payload["actor"] == {"role": "student", "userId": "student-001"}

    sent = client.post(
        "/api/question-bank/messages/send",
        headers=teacher_headers(),
        json={
            "userIds": ["student-001"],
            "category": "STUDY_REMINDER",
            "title": "学习提醒",
            "content": "请尽快完成今日刷题。",
        },
    )
    assert sent.status_code == 200
    assert sent.json()["data"]["sentCount"] == 1

    settings = client.post(
        "/api/question-bank/messages/settings",
        headers=student_headers(),
        json={
            "allowAiTutor": True,
            "allowSystemNotice": True,
            "allowReviewNotice": True,
            "allowStudyReminder": True,
            "allowWeeklyReport": False,
            "allowPointsNotice": True,
        },
    )
    assert settings.status_code == 200

    messages = client.get("/api/question-bank/messages?readStatus=unread", headers=student_headers())
    assert messages.status_code == 200
    items = messages.json()["data"]["items"]
    assert items
    assert items[0]["title"] == "学习提醒"

    marked = client.post(f"/api/question-bank/messages/{items[0]['messageId']}/read", headers=student_headers())
    assert marked.status_code == 200

    unread = client.get("/api/question-bank/messages?readStatus=unread", headers=student_headers())
    assert unread.status_code == 200
    assert all(item["messageId"] != items[0]["messageId"] for item in unread.json()["data"]["items"])


def test_messages_center_supports_batch_mark_read(tmp_path: Path):
    client = make_client(tmp_path)

    sent = client.post(
        "/api/question-bank/messages/send",
        headers=teacher_headers(),
        json={
            "userIds": ["student-001"],
            "category": "STUDY_REMINDER",
            "title": "批量已读测试",
            "content": "批量已读测试消息。",
        },
    )
    assert sent.status_code == 200

    unread_before = client.get("/api/question-bank/messages?readStatus=unread&page=1&size=20", headers=student_headers())
    assert unread_before.status_code == 200
    unread_items = unread_before.json()["data"]["items"]
    assert unread_items
    message_ids = [item["messageId"] for item in unread_items[:2]]

    batch_mark = client.post(
        "/api/question-bank/messages/read/batch",
        headers=student_headers(),
        json={"messageIds": message_ids},
    )
    assert batch_mark.status_code == 200
    assert batch_mark.json()["data"]["markedCount"] >= 1

    unread_after = client.get("/api/question-bank/messages?readStatus=unread&page=1&size=20", headers=student_headers())
    assert unread_after.status_code == 200
    remained_ids = [item["messageId"] for item in unread_after.json()["data"]["items"]]
    assert all(message_id not in remained_ids for message_id in message_ids)


def test_student_growth_personal_bank_and_wrong_book_review(tmp_path: Path):
    client = make_client(tmp_path)

    toggle = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/personal-bank",
        headers=student_headers(),
        json={"isCollected": True},
    )
    assert toggle.status_code == 200

    personal_bank = client.get(
        "/api/question-bank/student/practice/questions?onlyPersonalBank=true",
        headers=student_headers(),
    )
    assert personal_bank.status_code == 200
    ids = [item["id"] for item in personal_bank.json()["data"]["items"]]
    assert "question-seed-001" in ids

    review = client.post(
        "/api/question-bank/student/wrong-book/questions/question-seed-008/review",
        headers=student_headers(),
    )
    assert review.status_code == 200
    assert review.json()["data"]["reviewCount"] >= 1

    for _ in range(10):
        answer = client.post(
            "/api/question-bank/student/practice/questions/question-seed-001/submit",
            headers=student_headers(),
            json={"answer": "B", "elapsedSec": 30},
        )
        assert answer.status_code == 200

    dashboard = client.get("/api/question-bank/student/dashboard", headers=student_headers())
    assert dashboard.status_code == 200
    data = dashboard.json()["data"]
    assert data["personalBankCount"] >= 1
    assert any(item["reason"] == "完成章节刷题10道" for item in data["recentPointsLedger"])


def test_wrong_book_review_does_not_upgrade_status_without_correct_retry(tmp_path: Path):
    client = make_client(tmp_path)
    headers = student_headers()

    baseline_summary = client.get(
        "/api/student/error-book/summary",
        headers=headers,
    )
    assert baseline_summary.status_code == 200
    baseline_items = baseline_summary.json()["data"]["questionInsights"]
    target_before = next(
        item
        for item in baseline_items
        if item["reviewStatusLabel"] == "生疏"
    )
    assert target_before["reviewStatusLabel"] == "生疏"
    question_id = target_before["questionId"]
    subject_code = target_before["subjectCode"]

    for _ in range(2):
        review = client.post(
            f"/api/question-bank/student/wrong-book/questions/{question_id}/review",
            headers=headers,
        )
        assert review.status_code == 200

    refreshed_summary = client.get(
        "/api/student/error-book/summary",
        headers=headers,
        params={"subjectCode": subject_code},
    )
    assert refreshed_summary.status_code == 200
    refreshed_items = refreshed_summary.json()["data"]["questionInsights"]
    target_after = next(
        item
        for item in refreshed_items
        if item["questionId"] == question_id
    )
    assert target_after["reviewCount"] >= 2
    assert target_after["reviewStatusLabel"] == "生疏"


def seed_personal_bank_question(
    repo: object,
    *,
    question_id: str,
    stem: str,
    archived_at: str = "",
    is_archived: bool = False,
    source_type: str = "",
    source_label: str = "",
) -> None:
    source_question = repo.get_question("question-seed-001")
    assert source_question is not None
    repo.create_question(
        {
            **source_question,
            "id": question_id,
            "stem": stem,
            "status": "PUBLISHED",
            "createTime": "2026-03-23T10:00:00Z",
            "updateTime": "2026-03-23T10:00:00Z",
        }
    )
    repo.upsert_student_question_bank(
        {
            "id": "",
            "questionId": question_id,
            "studentUserId": "student-001",
            "status": "ACTIVE",
            "extJson": json.dumps(
                {
                    "personalBank": {
                        "isCollected": True,
                        "collectedAt": "2026-03-23T10:00:00Z",
                        "sourceType": source_type,
                        "sourceLabel": source_label,
                    },
                    "wrongBook": {
                        "isArchived": is_archived,
                        "archivedAt": archived_at,
                    },
                },
                ensure_ascii=False,
            ),
        }
    )


def test_personal_bank_independent_page_summary_and_export(tmp_path: Path):
    client = make_client(tmp_path)

    toggle = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/personal-bank",
        headers=student_headers(),
        json={"isCollected": True},
    )
    assert toggle.status_code == 200

    subjects = client.get("/api/question-bank/subjects", headers=student_headers())
    assert subjects.status_code == 200

    summary = client.get("/api/question-bank/student/personal-bank/summary", headers=student_headers())
    assert summary.status_code == 200
    assert summary.json()["data"]["totalCount"] >= 1

    listed = client.get("/api/question-bank/student/personal-bank/questions", headers=student_headers())
    assert listed.status_code == 200
    ids = [item["id"] for item in listed.json()["data"]["items"]]
    assert "question-seed-001" in ids

    exported = client.get("/api/question-bank/student/personal-bank/export?format=csv", headers=student_headers())
    assert exported.status_code == 200
    assert exported.json()["data"]["format"] == "csv"
    assert "questionId,subjectId,chapter" in exported.json()["data"]["content"]


def test_personal_bank_summary_and_export_are_not_truncated_at_500(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository

    for index in range(505):
        seed_personal_bank_question(
            repo,
            question_id=f"question-personal-bank-{index:03d}",
            stem=f"沉淀题库全量导出题 {index}",
        )

    summary = client.get(
        "/api/question-bank/student/personal-bank/summary",
        headers=student_headers(),
        params={"keyword": "沉淀题库全量导出题"},
    )
    assert summary.status_code == 200
    assert summary.json()["data"]["totalCount"] == 505

    exported = client.get(
        "/api/question-bank/student/personal-bank/export",
        headers=student_headers(),
        params={"format": "csv", "keyword": "沉淀题库全量导出题"},
    )
    assert exported.status_code == 200
    content_lines = exported.json()["data"]["content"].splitlines()
    assert len(content_lines) == 506
    assert "question-personal-bank-504" in exported.json()["data"]["content"]


def test_personal_bank_filters_align_for_archive_window_and_question_ids(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository

    seed_personal_bank_question(
        repo,
        question_id="question-personal-archive-recent",
        stem="沉淀题库筛选联动题-最近归档",
        archived_at="2026-03-22T10:00:00Z",
        is_archived=True,
        source_type="HARVESTED_ARCHIVE",
        source_label="已斩获归档",
    )
    seed_personal_bank_question(
        repo,
        question_id="question-personal-archive-old",
        stem="沉淀题库筛选联动题-历史归档",
        archived_at="2026-01-10T10:00:00Z",
        is_archived=True,
        source_type="HARVESTED_ARCHIVE",
        source_label="已斩获归档",
    )
    seed_personal_bank_question(
        repo,
        question_id="question-personal-active",
        stem="沉淀题库筛选联动题-未归档",
    )

    recent_summary = client.get(
        "/api/question-bank/student/personal-bank/summary",
        headers=student_headers(),
        params={"keyword": "沉淀题库筛选联动题", "archiveWindow": "LAST_7_DAYS"},
    )
    assert recent_summary.status_code == 200
    recent_data = recent_summary.json()["data"]
    assert recent_data["totalCount"] == 1
    assert recent_data["archivedCount"] == 1

    archived_list = client.get(
        "/api/question-bank/student/personal-bank/questions",
        headers=student_headers(),
        params={"keyword": "沉淀题库筛选联动题", "archiveWindow": "ARCHIVED"},
    )
    assert archived_list.status_code == 200
    archived_ids = [item["id"] for item in archived_list.json()["data"]["items"]]
    assert archived_ids == ["question-personal-archive-old", "question-personal-archive-recent"]

    filtered_list = client.get(
        "/api/question-bank/student/personal-bank/questions",
        headers=student_headers(),
        params={
            "questionIds": "question-personal-archive-recent,question-personal-active",
            "keyword": "沉淀题库筛选联动题",
        },
    )
    assert filtered_list.status_code == 200
    filtered_ids = [item["id"] for item in filtered_list.json()["data"]["items"]]
    assert filtered_ids == ["question-personal-active", "question-personal-archive-recent"]

    exported = client.get(
        "/api/question-bank/student/personal-bank/export",
        headers=student_headers(),
        params={"format": "csv", "keyword": "沉淀题库筛选联动题", "archiveWindow": "EARLIER"},
    )
    assert exported.status_code == 200
    exported_content = exported.json()["data"]["content"]
    assert "question-personal-archive-old" in exported_content
    assert "question-personal-archive-recent" not in exported_content
    assert "question-personal-active" not in exported_content


def test_personal_bank_summary_contains_review_plan(tmp_path: Path):
    client = make_client(tmp_path)
    toggle = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/personal-bank",
        headers=student_headers(),
        json={"isCollected": True},
    )
    assert toggle.status_code == 200

    wrong_submit = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "A", "elapsedSec": 30},
    )
    assert wrong_submit.status_code == 200

    summary = client.get("/api/question-bank/student/personal-bank/summary", headers=student_headers())
    assert summary.status_code == 200
    data = summary.json()["data"]
    assert "reviewPlan" in data
    assert {item["planKey"] for item in data["reviewPlan"]} == {"todayDue", "highWrong", "unanswered"}
    due_plan = next(item for item in data["reviewPlan"] if item["planKey"] == "todayDue")
    assert "question-seed-001" in due_plan["questionIds"]
    assert due_plan["planId"].startswith("student-review-plan::student-001::todayDue")
    assert due_plan["status"] in {"PENDING", "IN_PROGRESS", "COMPLETED"}


def test_personal_bank_review_plan_endpoints_support_start_and_complete(tmp_path: Path):
    client = make_client(tmp_path)

    for question_id in ("question-seed-001", "question-seed-002"):
        toggle = client.post(
            f"/api/question-bank/student/practice/questions/{question_id}/personal-bank",
            headers=student_headers(),
            json={"isCollected": True},
        )
        assert toggle.status_code == 200

    wrong_answers = {
        "question-seed-001": "A",
        "question-seed-002": "A",
    }
    for question_id, answer in wrong_answers.items():
        wrong_submit = client.post(
            f"/api/question-bank/student/practice/questions/{question_id}/submit",
            headers=student_headers(),
            json={"answer": answer, "elapsedSec": 30},
        )
        assert wrong_submit.status_code == 200

    listed = client.get("/api/question-bank/student/personal-bank/review-plans", headers=student_headers())
    assert listed.status_code == 200
    rows = listed.json()["data"]
    due_plan = next(item for item in rows if item["planKey"] == "todayDue")
    assert due_plan["questionCount"] >= 2
    assert {"question-seed-001", "question-seed-002"}.issubset(set(due_plan["questionIds"]))

    detail = client.get(
        f"/api/question-bank/student/personal-bank/review-plans/{due_plan['planId']}",
        headers=student_headers(),
        params={"questionIds": "question-seed-001"},
    )
    assert detail.status_code == 200
    detail_data = detail.json()["data"]
    assert detail_data["questionCount"] == 1
    assert detail_data["questionIds"] == ["question-seed-001"]
    assert [item["questionId"] for item in detail_data["items"]] == ["question-seed-001"]

    started = client.post(
        f"/api/question-bank/student/personal-bank/review-plans/{due_plan['planId']}/start",
        headers=student_headers(),
        params={"questionIds": "question-seed-001"},
    )
    assert started.status_code == 200
    started_data = started.json()["data"]
    assert started_data["status"] == "IN_PROGRESS"
    assert started_data["questionCount"] == 1
    assert started_data["questionIds"] == ["question-seed-001"]

    completed = client.post(
        f"/api/question-bank/student/personal-bank/review-plans/{due_plan['planId']}/questions/question-seed-001/complete",
        headers=student_headers(),
        params={"questionIds": "question-seed-001"},
    )
    assert completed.status_code == 200
    completed_data = completed.json()["data"]
    assert completed_data["status"] == "IN_PROGRESS"
    assert completed_data["questionCount"] == 1
    assert completed_data["questionIds"] == ["question-seed-001"]
    completed_item = completed_data["items"][0]
    assert completed_item["questionId"] == "question-seed-001"
    assert completed_item["status"] == "COMPLETED"


def test_personal_bank_toggle_updates_formal_columns(tmp_path: Path):
    client = make_client(tmp_path)

    toggle = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/personal-bank",
        headers=student_headers(),
        json={"isCollected": True},
    )
    assert toggle.status_code == 200

    repo = client.app.state.service.repository
    record = repo.get_student_question_bank("question-seed-001", "student-001")
    assert record is not None
    assert int(record["personalBankFlag"]) == 1
    assert str(record["personalBankCollectedAt"]).strip()


def test_paper_templates_overview_and_export_formats(tmp_path: Path):
    client = make_client(tmp_path)

    template = client.post(
        "/api/question-bank/papers/templates",
        headers=teacher_headers(),
        json={
            "templateName": "政治模板",
            "paperType": "simulation",
            "subjectId": "subject-politics",
            "chapter": "马克思主义哲学",
            "difficulty": "medium",
            "totalScore": 30,
            "durationMinutes": 45,
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "typeRules": [
                {"type": "single_choice", "count": 1, "questionScore": 10},
                {"type": "multiple_choice", "count": 1, "questionScore": 20},
            ],
        },
    )
    assert template.status_code == 200

    templates = client.get("/api/question-bank/papers/templates", headers=teacher_headers())
    assert templates.status_code == 200
    assert templates.json()["data"][0]["templateName"] == "政治模板"

    submit = client.post(
        "/api/question-bank/student/papers/paper-demo-001/submit",
        headers=student_headers(),
        json={
            "answers": [
                {"questionId": "question-seed-001", "answer": "B", "elapsedSec": 30, "marked": False},
                {"questionId": "question-seed-002", "answer": "ABC", "elapsedSec": 40, "marked": True},
            ],
            "totalElapsedSec": 70,
        },
    )
    assert submit.status_code == 200

    overview = client.get("/api/question-bank/papers/overview", headers=teacher_headers())
    assert overview.status_code == 200
    assert any(item["attemptCount"] >= 1 for item in overview.json()["data"])

    exported = client.get("/api/question-bank/papers/paper-demo-001/export?format=pdf", headers=teacher_headers())
    assert exported.status_code == 200
    assert exported.json()["data"]["format"] == "pdf"


def test_paper_template_supports_joint_group_scope_without_subject(tmp_path: Path):
    client = make_client(tmp_path)

    template = client.post(
        "/api/question-bank/papers/templates",
        headers=teacher_headers(),
        json={
            "templateName": "专业组模板",
            "paperType": "simulation",
            "chapter": "理工组综合训练",
            "difficulty": "medium",
            "totalScore": 20,
            "durationMinutes": 45,
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "typeRules": [
                {"type": "single_choice", "count": 2, "questionScore": 10},
            ],
        },
    )
    assert template.status_code == 200
    template_data = template.json()["data"]
    assert str(template_data.get("examCategoryCode", "")).strip() == "SCIENCE_ENGINEERING"
    assert str(template_data.get("jointExamGroupCode", "")).strip() == "SCIENCE_ENGINEERING_3"
    assert str(template_data.get("subjectCode", "")).strip() == ""
    assert str(template_data.get("subjectId", "")).strip() == ""

    templates = client.get("/api/question-bank/papers/templates", headers=teacher_headers())
    assert templates.status_code == 200
    row = next(item for item in templates.json()["data"] if str(item.get("templateName", "")) == "专业组模板")
    assert str(row.get("jointExamGroupCode", "")).strip() == "SCIENCE_ENGINEERING_3"
    assert str(row.get("subjectCode", "")).strip() == ""


def test_ai_generate_paper_returns_paper_id_and_updates_overview(tmp_path: Path):
    client = make_client(tmp_path)

    generated = client.post(
        "/api/question-bank/papers/ai-generate",
        headers=teacher_headers(),
        json={
            "subjectId": "subject-politics",
            "classIds": [],
            "totalCount": 10,
            "difficulty": 3,
            "knowledgeScope": [],
        },
    )
    assert generated.status_code == 200
    generated_data = generated.json()["data"]
    paper_id = str(generated_data.get("paperId", "") or generated_data.get("paper_id", "")).strip()
    assert paper_id

    overview = client.get("/api/question-bank/papers/overview", headers=teacher_headers())
    assert overview.status_code == 200
    assert any(str(item.get("paperId", "")).strip() == paper_id for item in overview.json()["data"])


def test_ai_generate_paper_supports_joint_group_scope_without_subject(tmp_path: Path):
    client = make_client(tmp_path)

    generated = client.post(
        "/api/question-bank/papers/ai-generate",
        headers=teacher_headers(),
        json={
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "classIds": [],
            "totalCount": 10,
            "difficulty": 3,
            "knowledgeScope": [],
        },
    )
    assert generated.status_code == 200
    generated_data = generated.json()["data"]
    paper_id = str(generated_data.get("paperId", "") or generated_data.get("paper_id", "")).strip()
    assert paper_id
    assert str(generated_data.get("joint_exam_group_code", "")).strip() == "SCIENCE_ENGINEERING_3"
    assert str(generated_data.get("subject_code", "")).strip() == ""
    assert str(generated_data.get("subject_id", "")).strip() == ""

    listed = client.get(
        "/api/question-bank/papers/questions",
        headers=teacher_headers(),
        params={
            "paperId": paper_id,
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "page": 1,
            "size": 20,
        },
    )
    assert listed.status_code == 200
    items = listed.json()["data"]["items"]
    assert items
    ext_json = json.loads(items[0]["extJson"])
    paper_bindings = ext_json.get("paperBindings", [])
    matched_binding = next((item for item in paper_bindings if str(item.get("paperId", "")).strip() == paper_id), None)
    assert matched_binding is not None
    assert str(matched_binding.get("jointExamGroupCode", "")).strip() == "SCIENCE_ENGINEERING_3"
    assert str(matched_binding.get("subjectCode", "")).strip() == ""


def test_save_auto_paper_persists_scope_binding_fields(tmp_path: Path):
    client = make_client(tmp_path)

    saved = client.post(
        "/api/question-bank/papers/auto",
        headers=teacher_headers(),
        json={
            "paperName": "自动组卷-作用域字段落库",
            "subjectId": "subject-politics",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "subjectCode": "POLITICS",
            "paperType": "chapter",
            "paperStatus": "DRAFT",
            "durationMinutes": 45,
            "totalScore": 10,
            "visibleToStudents": True,
            "typeRules": [
                {"type": "single_choice", "count": 1, "questionScore": 10},
            ],
        },
    )
    assert saved.status_code == 200
    paper_id = str(saved.json()["data"]["paperId"]).strip()
    assert paper_id

    listed = client.get(
        "/api/question-bank/papers/questions",
        headers=teacher_headers(),
        params={
            "paperId": paper_id,
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "subjectCode": "POLITICS",
            "page": 1,
            "size": 20,
        },
    )
    assert listed.status_code == 200
    items = listed.json()["data"]["items"]
    assert isinstance(items, list) and items
    ext_json = json.loads(items[0]["extJson"])
    paper_bindings = ext_json.get("paperBindings", [])
    matched_binding = next((item for item in paper_bindings if str(item.get("paperId", "")).strip() == paper_id), None)
    assert matched_binding is not None
    assert str(matched_binding.get("examCategoryCode", "")).strip() == "SCIENCE_ENGINEERING"
    assert str(matched_binding.get("jointExamGroupCode", "")).strip() == "SCIENCE_ENGINEERING_3"
    assert str(matched_binding.get("subjectCode", "")).strip() == "POLITICS"


def test_save_auto_paper_supports_joint_group_scope_without_subject(tmp_path: Path):
    client = make_client(tmp_path)

    saved = client.post(
        "/api/question-bank/papers/auto",
        headers=teacher_headers(),
        json={
            "paperName": "自动组卷-专业组模式",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "paperType": "chapter",
            "paperStatus": "DRAFT",
            "durationMinutes": 45,
            "totalScore": 10,
            "visibleToStudents": True,
            "typeRules": [
                {"type": "single_choice", "count": 1, "questionScore": 10},
            ],
        },
    )
    assert saved.status_code == 200
    paper_id = str(saved.json()["data"]["paperId"]).strip()
    assert paper_id

    listed = client.get(
        "/api/question-bank/papers/questions",
        headers=teacher_headers(),
        params={
            "paperId": paper_id,
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "page": 1,
            "size": 20,
        },
    )
    assert listed.status_code == 200
    items = listed.json()["data"]["items"]
    assert items
    ext_json = json.loads(items[0]["extJson"])
    paper_bindings = ext_json.get("paperBindings", [])
    matched_binding = next((item for item in paper_bindings if str(item.get("paperId", "")).strip() == paper_id), None)
    assert matched_binding is not None
    assert str(matched_binding.get("jointExamGroupCode", "")).strip() == "SCIENCE_ENGINEERING_3"
    assert str(matched_binding.get("subjectCode", "")).strip() == ""
    assert str(matched_binding.get("subjectId", "")).strip() == ""


def test_save_auto_paper_rejects_partial_professional_scope(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.post(
        "/api/question-bank/papers/auto",
        headers=teacher_headers(),
        json={
            "paperName": "自动组卷-缺少专业组",
            "subjectId": "subject-politics",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "paperType": "chapter",
            "paperStatus": "DRAFT",
            "durationMinutes": 45,
            "totalScore": 10,
            "visibleToStudents": True,
            "typeRules": [
                {"type": "single_choice", "count": 1, "questionScore": 10},
            ],
        },
    )
    assert response.status_code == 422
    assert "examCategoryCode、jointExamGroupCode 需同时提供，subjectCode 可选。" in str(response.json().get("detail", ""))


def test_paper_questions_support_keyword_filter_and_template_undo(tmp_path: Path):
    client = make_client(tmp_path)

    question_payload = payload()
    question_payload["title"] = "组卷关键词唯一题干-静态资源优化"
    question_payload["content"] = "组卷关键词唯一题干-静态资源优化"
    created = client.post("/api/question-bank/questions", headers=teacher_headers(), json=question_payload)
    assert created.status_code == 200
    question_id = created.json()["data"]["id"]

    review_pending = client.post(
        f"/api/question-bank/questions/{question_id}/status/REVIEW_PENDING",
        headers=teacher_headers(),
        json={},
    )
    assert review_pending.status_code == 200

    published = client.post(
        f"/api/question-bank/questions/{question_id}/status/PUBLISHED",
        headers=teacher_headers("teacher-002"),
        json={},
    )
    assert published.status_code == 200

    filtered = client.get(
        "/api/question-bank/papers/questions",
        headers=teacher_headers(),
        params={"keyword": "静态资源优化", "page": 1, "size": 20},
    )
    assert filtered.status_code == 200
    assert any(item["id"] == question_id for item in filtered.json()["data"]["items"])

    template = client.post(
        "/api/question-bank/papers/templates",
        headers=teacher_headers(),
        json={
            "templateName": "模板撤销测试",
            "paperType": "simulation",
            "subjectId": "subject-politics",
            "chapter": "马克思主义哲学",
            "difficulty": "medium",
            "totalScore": 30,
            "durationMinutes": 45,
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "jointExamGroupCode": "SCIENCE_ENGINEERING_3",
            "typeRules": [
                {"type": "single_choice", "count": 1, "questionScore": 10},
                {"type": "multiple_choice", "count": 1, "questionScore": 20},
            ],
        },
    )
    assert template.status_code == 200
    template_id = template.json()["data"]["templateId"]

    deleted = client.delete(f"/api/question-bank/papers/templates/{template_id}", headers=teacher_headers())
    assert deleted.status_code == 200
    undo_snapshot_id = deleted.json()["data"]["undoSnapshotId"]

    restored = client.post(
        f"/api/question-bank/papers/templates/deleted/{undo_snapshot_id}/restore",
        headers=teacher_headers(),
    )
    assert restored.status_code == 200
    assert restored.json()["data"]["templateId"] == template_id


def test_analytics_enhanced_summary_and_export(tmp_path: Path):
    client = make_client(tmp_path)

    client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "B", "elapsedSec": 30},
    )

    summary = client.get(
        "/api/question-bank/analytics/summary",
        headers=teacher_headers(),
        params={"startDate": "2026-01-01", "endDate": "2026-12-31"},
    )
    assert summary.status_code == 200
    data = summary.json()["data"]
    assert "weakKnowledgeTags" in data
    assert "lowActivityStudents" in data
    assert "aiReport" in data
    assert "coverageRate" in data
    assert "studentRankings" in data
    assert "averageAnswerDurationSec" in data
    assert "subjectCoverage" in data
    assert isinstance(data["subjectCoverage"], list)
    assert data["subjectCoverage"]
    assert all("knowledgeCoverageRate" in item for item in data["subjectCoverage"])
    assert any(
        item["coveredPointCount"] >= 1 and item["totalPointCount"] >= item["coveredPointCount"]
        for item in data["subjectCoverage"]
    )

    exported = client.get(
        "/api/question-bank/analytics/export",
        headers=teacher_headers(),
        params={"startDate": "2026-01-01", "endDate": "2026-12-31", "format": "pdf"},
    )
    assert exported.status_code == 200
    assert exported.json()["data"]["format"] == "pdf"


def test_list_all_student_records_uses_temp_table_when_ref_pool_exceeds_threshold(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository

    submit_one = client.post(
        "/api/question-bank/student/practice/questions/question-seed-001/submit",
        headers=student_headers(),
        json={"answer": "B", "elapsedSec": 20},
    )
    assert submit_one.status_code == 200
    submit_two = client.post(
        "/api/question-bank/student/practice/questions/question-seed-002/submit",
        headers=student_headers(),
        json={"answer": "A", "elapsedSec": 24},
    )
    assert submit_two.status_code == 200

    called = {"temp_table": False}
    original = repo._load_question_rows_for_student_records_with_temp_table  # noqa: SLF001

    def tracking_loader(connection: sqlite3.Connection, question_ids: list[str], policy_version: str):
        called["temp_table"] = True
        return original(connection, question_ids, policy_version)

    monkeypatch.setattr(repository_module, "STUDENT_RECORD_TEMP_TABLE_THRESHOLD", 1)
    monkeypatch.setattr(repo, "_load_question_rows_for_student_records_with_temp_table", tracking_loader)  # noqa: SLF001

    items = repo.list_all_student_records(
        "student-001",
        {
            "subjectId": "",
            "chapter": "",
            "paperId": "",
            "policyVersionCode": "HB_ZSB_2026",
        },
    )

    assert items
    assert called["temp_table"] is True


def test_knowledge_delete_supports_undo_restore(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/knowledge",
        headers=teacher_headers(),
        json=knowledge_payload(POLITICS_SECTION_ID),
    )
    assert created.status_code == 200
    knowledge_id = created.json()["data"]["id"]

    deleted = client.delete(f"/api/question-bank/knowledge/{knowledge_id}", headers=teacher_headers())
    assert deleted.status_code == 200
    undo_snapshot_id = deleted.json()["data"]["undoSnapshotId"]
    assert undo_snapshot_id.startswith("undo-")

    restored = client.post(
        f"/api/question-bank/knowledge/deleted/{undo_snapshot_id}/restore",
        headers=teacher_headers(),
    )
    assert restored.status_code == 200
    assert restored.json()["data"]["id"] == knowledge_id


def test_question_batch_delete_restore_and_reject_reason(tmp_path: Path):
    client = make_client(tmp_path)
    created_ids = []
    for index in range(2):
        question_payload = payload()
        question_payload["title"] = f"批量操作题目-{index}"
        question_payload["content"] = f"批量操作题目-{index}"
        question_payload["status"] = "DRAFT"
        created = client.post("/api/question-bank/questions", headers=teacher_headers(), json=question_payload)
        assert created.status_code == 200
        created_ids.append(created.json()["data"]["id"])

    batch_deleted = client.post(
        "/api/question-bank/questions/delete/batch",
        headers=teacher_headers(),
        json={"questionIds": created_ids},
    )
    assert batch_deleted.status_code == 200
    undo_snapshot_id = batch_deleted.json()["data"]["undoSnapshotId"]
    assert undo_snapshot_id.startswith("undo-")

    restored = client.post(
        f"/api/question-bank/questions/deleted/batch/{undo_snapshot_id}/restore",
        headers=teacher_headers(),
    )
    assert restored.status_code == 200
    assert restored.json()["data"]["restoredCount"] == 2

    to_reject = created_ids[0]
    to_review_pending = client.post(
        f"/api/question-bank/questions/{to_reject}/status/REVIEW_PENDING",
        headers=teacher_headers(),
        json={},
    )
    assert to_review_pending.status_code == 200

    reject_without_reason = client.post(
        "/api/question-bank/questions/status/batch",
        headers=teacher_headers("teacher-002"),
        json={"questionIds": [to_reject], "targetStatus": "REJECTED"},
    )
    assert reject_without_reason.status_code == 422

    reject_with_reason = client.post(
        "/api/question-bank/questions/status/batch",
        headers=teacher_headers("teacher-002"),
        json={"questionIds": [to_reject], "targetStatus": "REJECTED", "reason": "解析不完整"},
    )
    assert reject_with_reason.status_code == 200
    assert reject_with_reason.json()["data"]["updatedCount"] == 1


def test_messages_support_cohort_schedule_history_and_recall(tmp_path: Path):
    client = make_client(tmp_path)

    sent = client.post(
        "/api/question-bank/messages/send",
        headers=teacher_headers(),
        json={
            "targetMode": "cohort",
            "examCategoryCode": "SCIENCE_ENGINEERING",
            "category": "STUDY_REMINDER",
            "title": "分群提醒",
            "content": "请完成本周刷题计划。",
        },
    )
    assert sent.status_code == 200
    trace_id = sent.json()["data"]["traceId"]
    assert trace_id.startswith("msg-send-")
    assert sent.json()["data"]["sentCount"] >= 1

    history = client.get("/api/question-bank/messages/send-history?page=1&size=20", headers=teacher_headers())
    assert history.status_code == 200
    assert any(item["traceId"] == trace_id for item in history.json()["data"]["items"])

    recalled = client.post(f"/api/question-bank/messages/send-history/{trace_id}/recall", headers=teacher_headers())
    assert recalled.status_code == 200
    assert recalled.json()["data"]["recalledCount"] >= 1

    scheduled = client.post(
        "/api/question-bank/messages/send",
        headers=teacher_headers(),
        json={
            "targetMode": "userIds",
            "userIds": ["student-001"],
            "category": "SYSTEM_NOTICE",
            "title": "定时消息",
            "content": "这是一条定时消息。",
            "sendAt": "2099-01-01T00:00:00",
        },
    )
    assert scheduled.status_code == 200
    assert scheduled.json()["data"]["scheduledCount"] == 1


def test_admin_created_teacher_account_can_access_teacher_side_and_respects_permissions(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    created_teacher = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-qa-001",
            "role": "teacher",
            "name": "联调教师",
            "mobile": "13800000039",
            "enabled": True,
            "permissions": ["question:manage", "paper:manage", "analytics:view"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert created_teacher.status_code == 200

    teacher_page = client.get("/teacher/questions", headers=teacher_headers("teacher-qa-001"))
    assert teacher_page.status_code == 200
    assert "题库管理" in teacher_page.text

    teacher_list = client.get("/api/question-bank/questions", headers=teacher_headers("teacher-qa-001"))
    assert teacher_list.status_code == 200

    teacher_create = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-qa-001"),
        json=payload(user_id="teacher-qa-001"),
    )
    assert teacher_create.status_code == 200
    assert teacher_create.json()["data"]["userId"] == "teacher-qa-001"

    downgraded_teacher = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-qa-001",
            "role": "teacher",
            "name": "联调教师",
            "mobile": "13800000039",
            "enabled": True,
            "permissions": ["paper:manage"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert downgraded_teacher.status_code == 200

    forbidden_page = client.get("/teacher/questions", headers=teacher_headers("teacher-qa-001"))
    assert forbidden_page.status_code == 403
    forbidden_api = client.get("/api/question-bank/questions", headers=teacher_headers("teacher-qa-001"))
    assert forbidden_api.status_code == 403
    assert "question:manage" in forbidden_api.json()["message"]

    paper_allowed = client.get("/api/question-bank/papers/overview", headers=teacher_headers("teacher-qa-001"))
    assert paper_allowed.status_code == 200


def test_permission_validation_and_import_export_guards(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    restricted_user = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-009",
            "role": "teacher",
            "name": "受限老师",
            "mobile": "13800000029",
            "enabled": True,
            "permissions": ["paper:manage"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert restricted_user.status_code == 200

    forbidden = client.get("/api/question-bank/questions", headers=teacher_headers("teacher-009"))
    assert forbidden.status_code == 403
    assert "question:manage" in forbidden.json()["message"]

    disabled_user = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-010",
            "role": "teacher",
            "name": "停用老师",
            "mobile": "13800000030",
            "enabled": False,
            "permissions": ["question:manage"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert disabled_user.status_code == 200

    disabled = client.get("/api/question-bank/questions", headers=teacher_headers("teacher-010"))
    assert disabled.status_code == 403
    assert "已停用" in disabled.json()["message"]

    template_example = client.get("/api/question-bank/imports/template/example", headers=teacher_headers())
    assert template_example.status_code == 200
    assert "【题干】" in template_example.json()["data"]["content"]

    student_export = client.get("/api/question-bank/admin/students/export?format=csv", headers=admin_headers)
    assert student_export.status_code == 200
    assert "userId,name,mobile" in student_export.json()["data"]["content"]

    bad_dates = client.get(
        "/api/question-bank/analytics/summary",
        headers=teacher_headers(),
        params={"startDate": "2026-12-31", "endDate": "2026-01-01"},
    )
    assert bad_dates.status_code == 422

    bad_format = client.get("/api/question-bank/papers/paper-demo-001/export?format=xlsx", headers=teacher_headers())
    assert bad_format.status_code == 422


def test_permission_change_updates_auth_me_and_teacher_menu_visibility(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    created_teacher = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-ui-001",
            "role": "teacher",
            "name": "菜单显隐老师",
            "mobile": "13800000049",
            "enabled": True,
            "permissions": ["question:manage", "paper:manage", "analytics:view", "message:send"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert created_teacher.status_code == 200

    full_me = client.get("/api/question-bank/auth/me", headers=teacher_headers("teacher-ui-001"))
    assert full_me.status_code == 200
    full_permissions = full_me.json()["data"]["permissions"]
    assert set(full_permissions) == {"question:manage", "paper:manage", "analytics:view", "message:send"}

    full_menu_page = client.get("/teacher/papers", headers=teacher_headers("teacher-ui-001"))
    assert full_menu_page.status_code == 200
    full_visible = visible_teacher_menu_labels_by_permissions(full_menu_page.text, full_permissions)
    assert {"题库管理", "知识点三级树", "内容体系字典", "试卷管理", "学情管理"}.issubset(set(full_visible))

    downgraded_teacher = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-ui-001",
            "role": "teacher",
            "name": "菜单显隐老师",
            "mobile": "13800000049",
            "enabled": True,
            "permissions": ["paper:manage"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert downgraded_teacher.status_code == 200

    downgraded_me = client.get("/api/question-bank/auth/me", headers=teacher_headers("teacher-ui-001"))
    assert downgraded_me.status_code == 200
    downgraded_permissions = downgraded_me.json()["data"]["permissions"]
    assert downgraded_permissions == ["paper:manage"]

    downgraded_menu_page = client.get("/teacher/papers", headers=teacher_headers("teacher-ui-001"))
    assert downgraded_menu_page.status_code == 200
    downgraded_visible = visible_teacher_menu_labels_by_permissions(downgraded_menu_page.text, downgraded_permissions)
    assert "内容体系字典" in downgraded_visible
    assert "试卷管理" in downgraded_visible
    assert "题库管理" not in downgraded_visible
    assert "知识点三级树" not in downgraded_visible
    assert "学情管理" not in downgraded_visible

    forbidden_questions_page = client.get("/teacher/questions", headers=teacher_headers("teacher-ui-001"))
    assert forbidden_questions_page.status_code == 403
    forbidden_analytics_page = client.get("/teacher/analytics", headers=teacher_headers("teacher-ui-001"))
    assert forbidden_analytics_page.status_code == 403


def test_permission_change_hides_message_send_ui_marker_and_blocks_send_api(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    created_teacher = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-ui-002",
            "role": "teacher",
            "name": "消息显隐老师",
            "mobile": "13800000050",
            "enabled": True,
            "permissions": ["question:manage", "paper:manage", "analytics:view", "message:send"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert created_teacher.status_code == 200

    messages_page = client.get("/messages", headers=teacher_headers("teacher-ui-002"))
    assert messages_page.status_code == 200
    page_payload = parse_page_bootstrap(messages_page.text)
    assert page_payload["route"] == "/messages"
    assert page_payload["viewKey"] == "messages"
    assert "message:send" in page_payload["permissions"]

    full_me = client.get("/api/question-bank/auth/me", headers=teacher_headers("teacher-ui-002"))
    assert full_me.status_code == 200
    assert "message:send" in full_me.json()["data"]["permissions"]
    assert permission_marker_visible("message:send", full_me.json()["data"]["permissions"])

    downgraded_teacher = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-ui-002",
            "role": "teacher",
            "name": "消息显隐老师",
            "mobile": "13800000050",
            "enabled": True,
            "permissions": ["paper:manage"],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert downgraded_teacher.status_code == 200

    downgraded_me = client.get("/api/question-bank/auth/me", headers=teacher_headers("teacher-ui-002"))
    assert downgraded_me.status_code == 200
    downgraded_permissions = downgraded_me.json()["data"]["permissions"]
    assert downgraded_permissions == ["paper:manage"]
    assert not permission_marker_visible("message:send", downgraded_permissions)

    forbidden_send = client.post(
        "/api/question-bank/messages/send",
        headers=teacher_headers("teacher-ui-002"),
        json={
            "targetMode": "userIds",
            "userIds": ["student-001"],
            "category": "SYSTEM_NOTICE",
            "title": "无权限发送",
            "content": "该请求应被拦截",
        },
    )
    assert forbidden_send.status_code == 403
    assert "message:send" in forbidden_send.json()["message"]


def test_teacher_post_tags_can_drive_permission_template_and_require_recruit_scope(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    missing_scope = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-post-001",
            "role": "teacher",
            "name": "岗位老师",
            "mobile": "13800000090",
            "enabled": True,
            "permissions": [],
            "postTags": ["recruit"],
            "managedStudentIds": [],
            "managedJointExamGroupCodes": [],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert missing_scope.status_code == 422
    assert "数据范围" in missing_scope.json()["message"]

    created = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-post-001",
            "role": "teacher",
            "name": "岗位老师",
            "mobile": "13800000090",
            "enabled": True,
            "permissions": [],
            "postTags": ["recruit", "teach"],
            "managedStudentIds": ["student-001"],
            "managedJointExamGroupCodes": [],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert created.status_code == 200
    created_data = created.json()["data"]
    assert set(created_data["permissions"]) == {"student:manage", "analytics:view", "message:send", "question:manage", "paper:manage"}
    assert set(created_data["postTags"]) == {"recruit", "teach"}
    assert created_data["managedStudentIds"] == ["student-001"]

    me_response = client.get("/api/question-bank/auth/me", headers=teacher_headers("teacher-post-001"))
    assert me_response.status_code == 200
    me_data = me_response.json()["data"]
    assert set(me_data["postTags"]) == {"recruit", "teach"}
    assert me_data["managedStudentIds"] == ["student-001"]


def test_recruit_teacher_scope_filters_student_directory_by_managed_student_ids(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)
    recruit_teacher_login_cred = "seed-password-teacher-recruit-001"

    created = client.post(
        "/api/question-bank/admin/users",
        headers=admin_headers,
        json={
            "userId": "teacher-recruit-001",
            "role": "teacher",
            "name": "招生老师",
            "mobile": "13800000091",
            "enabled": True,
            "permissions": [],
            "postTags": ["recruit"],
            "managedStudentIds": ["student-001"],
            "managedJointExamGroupCodes": [],
            "examCategoryCode": "",
            "jointExamGroupCode": "",
            "vocationalMajor": "",
            "prepStage": "",
        },
    )
    assert created.status_code == 200

    scoped_students = client.get(
        "/api/question-bank/admin/users?page=1&size=20&role=student",
        headers=teacher_auth_headers(
            client,
            phone="13800000091",
            password=recruit_teacher_login_cred,
        ),
    )
    assert scoped_students.status_code == 200
    scoped_items = scoped_students.json()["data"]["items"]
    assert len(scoped_items) == 1
    assert scoped_items[0]["userId"] == "student-001"

def test_unauthorized_access(tmp_path: Path):
    client = make_client(tmp_path)

    teacher_admin_console = client.get("/api/question-bank/admin/console", headers=teacher_headers())
    assert teacher_admin_console.status_code == 401

    student_admin_users = client.get("/api/question-bank/admin/users?page=1&size=10", headers=student_headers())
    assert student_admin_users.status_code == 401

    cross_category_access = client.get(
        "/api/question-bank/student/practice/questions?examCategoryCode=ART",
        headers=student_headers("student-001"),
    )
    assert cross_category_access.status_code == 403
    assert cross_category_access.json()["code"] == "QUESTION_FORBIDDEN"


def test_peer_review_logic(tmp_path: Path):
    client = make_client(tmp_path)
    created = client.post(
        "/api/question-bank/questions",
        headers=teacher_headers("teacher-001"),
        json=payload(user_id="teacher-001"),
    )
    assert created.status_code == 200
    question = created.json()["data"]
    question_id = question["id"]

    to_qa = client.post(
        f"/api/question-bank/questions/{question_id}/status/QA_IN_PROGRESS",
        headers=teacher_headers("teacher-001"),
    )
    assert to_qa.status_code == 200
    qa_question = to_qa.json()["data"]

    owner_submit_review = client.put(
        f"/api/question-bank/questions/{question_id}",
        headers=teacher_headers("teacher-001"),
        json={"status": "REVIEW_PENDING"},
    )
    assert owner_submit_review.status_code == 403

    peer_submit_review = client.put(
        f"/api/question-bank/questions/{question_id}",
        headers=teacher_headers("teacher-002"),
        json={"status": "REVIEW_PENDING"},
    )
    assert peer_submit_review.status_code == 200

    detail = client.get(f"/api/question-bank/questions/{question_id}", headers=teacher_headers("teacher-001"))
    assert detail.status_code == 200
    ext_json = json.loads(detail.json()["data"]["extJson"])
    review_records = ext_json.get("reviewRecords", [])
    assert any(
        str(item.get("status", "")) == "REVIEW_PENDING"
        and str(item.get("reviewerId", "")) == "teacher-002"
        and str(item.get("timestamp", "")).strip()
        for item in review_records
        if isinstance(item, dict)
    )


def test_dashboard_filtering(tmp_path: Path):
    client = make_client(tmp_path)

    dashboard = client.get("/api/question-bank/student/dashboard", headers=student_headers("student-001"))
    assert dashboard.status_code == 200
    data = dashboard.json()["data"]
    assert data["examCategoryCode"] == "SCIENCE_ENGINEERING"
    onboarding = data.get("onboarding", {})
    assert onboarding.get("completed") is False
    assert onboarding.get("subscriptionActive") is False
    assert onboarding.get("quickDiagnosisCompleted") is False
    subject_codes = {str(item.get("subjectCode", "")) for item in data.get("coreSubjects", []) if isinstance(item, dict)}
    assert "ART_INTRODUCTION" not in subject_codes
    assert "FINE_ARTS_COMPREHENSIVE" not in subject_codes

    service = client.app.state.service
    repo = service.repository
    record = repo.get_student_question_bank("question-seed-001", "student-001")
    assert record is not None
    record_ext = json.loads(record["extJson"])
    record_ext["studentProfile"] = {"examCategoryCode": "", "jointExamGroupCode": ""}
    with get_connection(repo.db_path) as connection:
        connection.execute(
            """
            UPDATE student_question_record
            SET extJson = ?
            WHERE studentUserId = ? AND questionId = ?
            """,
            (json.dumps(record_ext, ensure_ascii=False), "student-001", "question-seed-001"),
        )
        connection.commit()

    incomplete = client.get("/api/question-bank/student/dashboard", headers=student_headers("student-001"))
    assert incomplete.status_code == 200
    incomplete_data = incomplete.json()["data"]
    assert incomplete_data["examCategoryCode"] == "SCIENCE_ENGINEERING"
    assert incomplete_data["jointExamGroupCode"] == "SCIENCE_ENGINEERING_3"


def test_ai_quota_limit(tmp_path: Path):
    client = make_client(tmp_path)
    service = client.app.state.service
    today = datetime.now(timezone.utc).date().isoformat()
    seed_student_hot_state(
        service,
        "student-002",
        ai_quota={"dailyLimit": 20, "usedCount": 0, "quotaDate": today},
    )

    for index in range(20):
        response = client.post(
            "/api/question-bank/student/practice/questions/question-seed-005/ai-tutor",
            headers=student_headers("student-002"),
            json={"prompt": f"第{index + 1}次 AI 追问", "promptImageUrl": ""},
        )
        assert response.status_code == 200

    tasks = client.get("/api/question-bank/tasks?page=1&size=50", headers=student_headers("student-002"))
    assert tasks.status_code == 200
    ai_quota = tasks.json()["data"]["aiQuota"]
    assert ai_quota["dailyLimit"] == 20
    assert ai_quota["usedCount"] == 20

    exhausted = client.post(
        "/api/question-bank/student/practice/questions/question-seed-005/ai-tutor",
        headers=student_headers("student-002"),
        json={"prompt": "第21次 AI 追问", "promptImageUrl": ""},
    )
    assert exhausted.status_code == 422
    assert exhausted.json()["code"] == "QUESTION_VALIDATION_FAILED"
    assert "配额已用完" in exhausted.json()["message"]


def _seed_redeem_code_for_test(
    repo: object,
    *,
    code: str,
    plan_code: str = "AI_SCORE_BOOST_30D",
    expires_at: str = "2099-12-31T00:00:00Z",
) -> str:
    batch_id = f"redeem-batch-{uuid4().hex[:8]}"
    now_iso = "2026-03-26T00:00:00Z"
    repo.create_redeem_code_batch(
        {
            "id": batch_id,
            "batchCode": f"BATCH-{uuid4().hex[:8].upper()}",
            "batchName": "测试兑换批次",
            "channelCode": "TEST",
            "planCode": plan_code,
            "totalCount": 10,
            "usedCount": 0,
            "expiresAt": expires_at,
            "status": "ACTIVE",
            "createdByUserId": "admin-001",
            "extJson": {},
            "createTime": now_iso,
            "updateTime": now_iso,
        }
    )
    repo.create_redeem_code(
        {
            "id": f"redeem-code-{uuid4().hex[:8]}",
            "batchId": batch_id,
            "code": code,
            "planCode": plan_code,
            "status": "UNUSED",
            "expiresAt": expires_at,
            "usedByUserId": "",
            "usedAt": "",
            "sourceOrderId": "",
            "extJson": {},
            "createTime": now_iso,
            "updateTime": now_iso,
        }
    )
    return batch_id


def test_student_subscription_status_default_inactive(tmp_path: Path):
    client = make_client(tmp_path)
    response = client.get("/api/question-bank/student/subscription/status", headers=student_headers())
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["subscription"]["status"] == "INACTIVE"
    assert data["subscription"]["subscriptionActive"] is False
    assert data["subscriptionActive"] is False


def test_student_subscription_redeem_once_limit(tmp_path: Path):
    client = make_client(tmp_path)
    repo = client.app.state.service.repository
    _seed_redeem_code_for_test(repo, code="TEST-REDEEM-001")
    _seed_redeem_code_for_test(repo, code="TEST-REDEEM-002")

    first = client.post(
        "/api/question-bank/student/subscription/redeem",
        headers=student_headers(),
        json={"code": "TEST-REDEEM-001"},
    )
    assert first.status_code == 200
    first_data = first.json()["data"]
    assert first_data["subscription"]["status"] == "ACTIVE"
    assert first_data["subscription"]["subscriptionActive"] is True
    dashboard_after_redeem = client.get("/api/question-bank/student/dashboard", headers=student_headers())
    assert dashboard_after_redeem.status_code == 200
    onboarding = dashboard_after_redeem.json()["data"]["onboarding"]
    assert onboarding["completed"] is True
    assert onboarding["subscriptionActive"] is True
    assert onboarding["quickDiagnosisCompleted"] is False

    second = client.post(
        "/api/question-bank/student/subscription/redeem",
        headers=student_headers(),
        json={"code": "TEST-REDEEM-002"},
    )
    assert second.status_code == 422
    assert second.json()["code"] == "QUESTION_VALIDATION_FAILED"
    assert "仅可兑换一次" in second.json()["message"]


def test_student_subscription_mock_order_confirm_idempotent(tmp_path: Path):
    client = make_client(tmp_path)

    created = client.post(
        "/api/question-bank/student/subscription/mock-orders",
        headers=student_headers(),
        json={"planCode": "AI_SCORE_BOOST_30D"},
    )
    assert created.status_code == 200
    order_id = str(created.json()["data"]["order"]["orderId"])
    assert order_id

    first_confirm = client.post(
        f"/api/question-bank/student/subscription/mock-orders/{order_id}/confirm",
        headers=student_headers(),
        json={"transactionNo": "MOCK-TXN-001"},
    )
    assert first_confirm.status_code == 200
    first_data = first_confirm.json()["data"]
    assert first_data["idempotent"] is False
    assert first_data["order"]["status"] == "PAID"
    first_end_time = str(first_data["subscription"]["endTime"])
    assert first_data["subscription"]["status"] == "ACTIVE"

    second_confirm = client.post(
        f"/api/question-bank/student/subscription/mock-orders/{order_id}/confirm",
        headers=student_headers(),
        json={"transactionNo": "MOCK-TXN-001"},
    )
    assert second_confirm.status_code == 200
    second_data = second_confirm.json()["data"]
    assert second_data["idempotent"] is True
    assert second_data["order"]["status"] == "PAID"
    assert str(second_data["subscription"]["endTime"]) == first_end_time


def test_teacher_forbidden_but_super_admin_can_access_student_subscription_and_diagnosis_endpoints(tmp_path: Path):
    client = make_client(tmp_path)
    teacher_token_headers = teacher_auth_headers(client)
    admin_headers = super_admin_auth_headers(client)

    teacher_status_response = client.get(
        "/api/question-bank/student/subscription/status",
        headers=teacher_token_headers,
    )
    assert teacher_status_response.status_code == 403

    teacher_diagnosis_start_response = client.post(
        "/api/question-bank/student/diagnosis/quick/start",
        headers=teacher_token_headers,
        json={"questionCount": 3},
    )
    assert teacher_diagnosis_start_response.status_code == 403

    teacher_redeem_response = client.post(
        "/api/question-bank/student/subscription/redeem",
        headers=teacher_token_headers,
        json={"code": "TEST-REDEEM-FORBIDDEN"},
    )
    assert teacher_redeem_response.status_code == 403

    admin_status_response = client.get(
        "/api/question-bank/student/subscription/status",
        headers=admin_headers,
    )
    assert admin_status_response.status_code == 200

    admin_diagnosis_start_response = client.post(
        "/api/question-bank/student/diagnosis/quick/start",
        headers=admin_headers,
        json={"questionCount": 3},
    )
    assert admin_diagnosis_start_response.status_code == 200


def test_admin_can_create_redeem_code_batch(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)
    response = client.post(
        "/api/question-bank/admin/redeem-code/batches",
        headers=admin_headers,
        json={
            "batchName": "春季拉新",
            "totalCount": 3,
            "planCode": "AI_SCORE_BOOST_30D",
            "channelCode": "SPRING",
            "expiresAt": "2099-12-31T00:00:00Z",
            "codePrefix": "SPRING26",
        },
    )
    assert response.status_code == 200
    payload = response.json()["data"]
    batch = payload["batch"]
    codes = payload["codes"]
    assert batch["batchName"] == "春季拉新"
    assert batch["totalCount"] == 3
    assert batch["usedCount"] == 0
    assert batch["unusedCount"] == 3
    assert len(codes) == 3


def test_non_super_admin_cannot_create_redeem_code_batch(tmp_path: Path):
    client = make_client(tmp_path)
    teacher_token_headers = teacher_auth_headers(client)
    response = client.post(
        "/api/question-bank/admin/redeem-code/batches",
        headers=teacher_token_headers,
        json={
            "batchName": "无权限批次",
            "totalCount": 1,
            "planCode": "AI_SCORE_BOOST_30D",
            "channelCode": "TEST",
            "expiresAt": "2099-12-31T00:00:00Z",
            "codePrefix": "NOAUTH",
        },
    )
    assert response.status_code == 403


def test_admin_list_redeem_code_batches_contains_usage_counters(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    created = client.post(
        "/api/question-bank/admin/redeem-code/batches",
        headers=admin_headers,
        json={
            "batchName": "统计批次",
            "totalCount": 2,
            "planCode": "AI_SCORE_BOOST_30D",
            "channelCode": "STATS",
            "expiresAt": "2099-12-31T00:00:00Z",
            "codePrefix": "STATS",
        },
    )
    assert created.status_code == 200
    created_payload = created.json()["data"]
    created_batch = created_payload["batch"]
    redeem_code = str(created_payload["codes"][0])

    redeem_response = client.post(
        "/api/question-bank/student/subscription/redeem",
        headers=student_headers(),
        json={"code": redeem_code},
    )
    assert redeem_response.status_code == 200

    listed = client.get(
        "/api/question-bank/admin/redeem-code/batches?page=1&size=20&status=ACTIVE",
        headers=admin_headers,
    )
    assert listed.status_code == 200
    page_payload = listed.json()["data"]
    items = page_payload["items"]
    matched = [item for item in items if str(item.get("id", "")) == str(created_batch["id"])]
    assert matched
    first = matched[0]
    assert first["usedCount"] == 1
    assert first["unusedCount"] == 1
    assert first["usageRate"] == 0.5


def test_admin_conversion_overview_returns_expected_counters(tmp_path: Path):
    client = make_client(tmp_path)
    admin_headers = super_admin_auth_headers(client)

    created = client.post(
        "/api/question-bank/admin/redeem-code/batches",
        headers=admin_headers,
        json={
            "batchName": "转化总览批次",
            "totalCount": 1,
            "planCode": "AI_SCORE_BOOST_30D",
            "channelCode": "OVERVIEW",
            "expiresAt": "2099-12-31T00:00:00Z",
            "codePrefix": "OVERVIEW",
        },
    )
    assert created.status_code == 200
    redeem_code = str(created.json()["data"]["codes"][0])

    redeem_response = client.post(
        "/api/question-bank/student/subscription/redeem",
        headers=student_headers(),
        json={"code": redeem_code},
    )
    assert redeem_response.status_code == 200

    order_response = client.post(
        "/api/question-bank/student/subscription/mock-orders",
        headers=student_headers(),
        json={"planCode": "AI_SCORE_BOOST_30D"},
    )
    assert order_response.status_code == 200
    order_id = str(order_response.json()["data"]["order"]["orderId"])
    assert order_id

    confirm_response = client.post(
        f"/api/question-bank/student/subscription/mock-orders/{order_id}/confirm",
        headers=student_headers(),
        json={"transactionNo": "MOCK-TXN-OVERVIEW-001"},
    )
    assert confirm_response.status_code == 200

    overview_response = client.get(
        "/api/question-bank/admin/conversion/overview?startDate=2000-01-01&endDate=2099-12-31",
        headers=admin_headers,
    )
    assert overview_response.status_code == 200
    data = overview_response.json()["data"]
    summary = data["summary"]
    assert summary["redeemSubmitCount"] >= 1
    assert summary["redeemSuccessCount"] >= 1
    assert summary["mockOrderCreatedCount"] >= 1
    assert summary["mockPaymentSuccessCount"] >= 1
    assert summary["subscriptionActivatedCount"] >= 2
    counters = {str(item.get("eventType", "")): int(item.get("eventCount", 0) or 0) for item in data["eventTypeCounters"]}
    assert counters.get("REDEEM_SUCCESS", 0) >= 1
    assert counters.get("MOCK_PAYMENT_SUCCESS", 0) >= 1


def test_student_quick_diagnosis_start_submit_and_idempotent(tmp_path: Path):
    client = make_client(tmp_path)

    start_response = client.post(
        "/api/question-bank/student/diagnosis/quick/start",
        headers=student_headers(),
        json={"questionCount": 3},
    )
    assert start_response.status_code == 200
    start_data = start_response.json()["data"]
    session_id = str(start_data["sessionId"])
    question_ids = list(start_data["questionIds"])
    questions_preview = start_data.get("questions", [])
    assert session_id
    assert len(question_ids) == 3
    assert isinstance(questions_preview, list)
    assert len(questions_preview) == 3
    for item in questions_preview:
        assert str(item.get("questionId", "")).strip() in question_ids
        assert "answer" not in item

    repo = client.app.state.service.repository
    question_rows = {str(row["id"]): row for row in repo.list_questions_by_ids(question_ids)}
    submit_payload = {
        "answers": [
            {
                "questionId": question_id,
                "answer": str(question_rows[question_id]["answer"]),
                "elapsedSec": 12,
            }
            for question_id in question_ids
        ],
        "sourceType": "ONBOARDING",
    }
    submit_response = client.post(
        f"/api/question-bank/student/diagnosis/quick/{session_id}/submit",
        headers=student_headers(),
        json=submit_payload,
    )
    assert submit_response.status_code == 200
    submit_data = submit_response.json()["data"]
    assert submit_data["idempotent"] is False
    assert submit_data["status"] == "COMPLETED"
    assert submit_data["answeredCount"] == 3
    assert submit_data["correctCount"] == 3
    assert submit_data["accuracy"] == 1.0

    idempotent_response = client.post(
        f"/api/question-bank/student/diagnosis/quick/{session_id}/submit",
        headers=student_headers(),
        json=submit_payload,
    )
    assert idempotent_response.status_code == 200
    idempotent_data = idempotent_response.json()["data"]
    assert idempotent_data["idempotent"] is True
    assert idempotent_data["status"] == "COMPLETED"
    assert idempotent_data["answeredCount"] == 3
    dashboard_response = client.get("/api/question-bank/student/dashboard", headers=student_headers())
    assert dashboard_response.status_code == 200
    onboarding = dashboard_response.json()["data"]["onboarding"]
    assert onboarding["completed"] is True
    assert onboarding["quickDiagnosisCompleted"] is True
    assert onboarding["subscriptionActive"] is False
    assert onboarding["latestQuickDiagnosisSession"]["status"] == "COMPLETED"
    assert onboarding["latestQuickDiagnosisSession"]["sessionId"] == session_id

    admin_headers = super_admin_auth_headers(client)
    overview_response = client.get(
        "/api/question-bank/admin/conversion/overview?startDate=2000-01-01&endDate=2099-12-31",
        headers=admin_headers,
    )
    assert overview_response.status_code == 200
    summary = overview_response.json()["data"]["summary"]
    assert summary["quickDiagnosisStartCount"] >= 1
    assert summary["quickDiagnosisCompleteCount"] >= 1


def test_teacher_cannot_access_student_quick_diagnosis_api(tmp_path: Path):
    client = make_client(tmp_path)
    start_response = client.post(
        "/api/question-bank/student/diagnosis/quick/start",
        headers=teacher_headers(),
        json={"questionCount": 3},
    )
    assert start_response.status_code == 403
